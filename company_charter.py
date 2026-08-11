"""
Agentic Company Charter generator. Fills output/company_charters/
Company_Charter_TEMPLATE_WebSourced.docx *in place* (python-docx, preserving
every style/color/table already baked into the template) with real,
sourced facts about one MahaRERA project -- replacing the manual
Node/docx-js scripts used for the first (Pranami Bliss) charter.

Reuses category_data already fetched by main.py, the documents already
downloaded to documents_dir, and -- if available -- the deep_research.json
already produced for this same project (its macro_market/micro_market/
promoter_external findings feed Area Intelligence and Corporate Identity
directly, so this pass doesn't re-research what deep_research.py already
confirmed).

Guardrails, same philosophy as deep_research.py:
  - Every sourced field carries its source; claims materially affecting the
    corporate-identity / litigation / FSI sections are independently
    re-verified (reusing deep_research._verify_claim) before being trusted.
  - Anything that can't be confirmed goes in `gaps`, verbatim, never
    silently filled in or approximated as fact.
  - Distances are estimated via web_search (no live browser/Maps API here)
    -- the template's own Methodology Note says so explicitly, matching its
    existing "state plainly what's confirmed vs approximated" philosophy.

Requires: ANTHROPIC_API_KEY (same as deep_research.py), and a local
Tesseract OCR install for scanned/no-text PDFs (falls back to a plain
"[OCR unavailable]" marker per-document if Tesseract isn't found, rather
than failing the whole run).

Optional precision upgrade for the Distances table: set
COMPANY_CHARTER_USE_MAPS_SCRAPE=1 to have _refine_distances_with_maps()
launch a headless Playwright browser per landmark and read the real
driving route off Google Maps, replacing the model's web_search estimate.
Off by default -- verified against the live site, but it scrapes Google's
consumer UI rather than their paid Distance Matrix/Routes API, so it may
not comply with Google's Terms of Service and can break without warning
if Google changes their page. Always falls back to the existing estimate
on any failure.

    python company_charter.py <REG_NO>
    COMPANY_CHARTER_USE_MAPS_SCRAPE=1 python company_charter.py <REG_NO>
"""

import argparse
import concurrent.futures
import copy
import json
import os
import re
import shutil
import sys
import time
from datetime import datetime

import fitz  # PyMuPDF
import pytesseract
import requests
from PIL import Image

import config
import deep_research
import finalize_report
import run_archive

MODEL = deep_research.MODEL

# pytesseract shells out to the tesseract binary by name -- if it's installed
# but not on PATH (confirmed live: this was the actual state on this machine,
# silently degrading every scanned document to "[OCR unavailable for this
# page]" rather than raising anything visible), it fails silently into that
# same per-page marker in _extract_document_text below. TESSERACT_CMD lets a
# different machine/OS override this; the two hardcoded fallbacks cover the
# default Windows installer locations.
if not shutil.which("tesseract"):
    for _candidate in (
        os.environ.get("TESSERACT_CMD"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ):
        if _candidate and os.path.exists(_candidate):
            pytesseract.pytesseract.tesseract_cmd = _candidate
            break
TEMPLATE_PATH = os.path.join(config.OUTPUT_ROOT, "company_charters", "Company_Charter_TEMPLATE_WebSourced.docx")
CHARTER_OUTPUT_DIR = os.path.join(config.OUTPUT_ROOT, "company_charters")

# --- CLAUDE.md loaders -----------------------------------------------------
# CLAUDE.md is auto-loaded by Claude Code for interactive/coding sessions,
# but this pipeline also calls the Claude API directly and unattended (via
# _run_charter_pass and the citation-completeness judge below) -- no rules file
# is auto-loaded there, so the three sections are read explicitly and routed to
# exactly the calls each one is scoped for. Section A never leaves this process
# as API content; Section B goes into every charter content-generating/verifying
# call; Section C only into external-doc-variant calls.
#
# The rules live in rules.md, not CLAUDE.md. They were split out once CLAUDE.md
# was rewritten as the pipeline's operational map: the rules are ~250 lines and
# two thirds of them are injected into API calls, so keeping them in the file
# every Claude Code session auto-loads made that file too long to be read
# reliably. CLAUDE.md now describes the flow and points here for the content.
_RULES_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rules.md")
_RULES_SECTION_RE = re.compile(
    r"--- Section {marker}:.*?---\n(.*?)(?=\n--- Section |\Z)", re.DOTALL
)


def _read_rules_section(marker: str) -> str:
    with open(_RULES_PATH, "r", encoding="utf-8") as f:
        text = f.read()
    pattern = re.compile(_RULES_SECTION_RE.pattern.format(marker=marker), re.DOTALL)
    m = pattern.search(text)
    if not m:
        raise RuntimeError(f"rules.md Section {marker} not found -- has the file been restructured?")
    body = m.group(1).strip()
    if not body:
        raise RuntimeError(f"rules.md Section {marker} is empty")
    return body


def _coding_time_notes() -> str:
    """rules.md Section A -- coding-time rules for Claude Code / human
    sessions only. Documentation use only; NEVER pass this into any runtime
    API call."""
    return _read_rules_section("A")


def _common_content_rules() -> str:
    """Section B -- prepended to the system prompt of every charter
    content-generating/verifying Claude API call, both doc variants."""
    return _read_rules_section("B")


def _external_citation_rule() -> str:
    """Section C -- appended only when the call is building or checking
    doc_variant == "external" content."""
    return _read_rules_section("C")


def _charter_system_blocks(*, external: bool, extra: str = "") -> list:
    """Assembles the `system` param for a charter-specific Claude API call as
    a list of text blocks: Section B first (marked as a cacheable prompt
    prefix -- identical across every call and every project), then Section C
    only when `external`, then any call-specific instructions. Passing a
    list (rather than a plain string) to deep_research._run_agentic_pass's
    `system` param is what the Anthropic API needs to actually cache the
    Section B prefix."""
    blocks = [{"type": "text", "text": _common_content_rules(), "cache_control": {"type": "ephemeral"}}]
    if external:
        blocks.append({"type": "text", "text": _external_citation_rule()})
    if extra:
        blocks.append({"type": "text", "text": extra})
    return blocks

# Document labels worth extracting full text from (case-insensitive substring
# match against the manifest's `label`/`saved_filename`) -- everything else
# is listed in the Document Library table by name only, with an explicit
# "not opened this pass" reason, matching the template's own allowance for
# that (see paragraph 63/64 in the template).
_HIGH_PRIORITY_DOC_KEYWORDS = (
    "title", "legal", "encumbrance", "form b", "declaration", "ioa", "iod",
    "sanction", "approval", "layout", "plan", "allotment", "agreement",
)
_MAX_CHARS_PER_DOC = 6000
_MAX_TOTAL_DOC_CHARS = 30000


def _extract_document_text(path: str) -> str:
    """Native text first; falls back to Tesseract OCR per page with no
    extractable text. Never raises -- a broken/unreadable file or a missing
    Tesseract install yields a plain marker instead of failing the run."""
    try:
        parts = []
        with fitz.open(path) as doc:
            for page in doc:
                text = page.get_text().strip()
                if text:
                    parts.append(text)
                    continue
                try:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    parts.append(pytesseract.image_to_string(img))
                except Exception:
                    parts.append("[OCR unavailable for this page]")
        return "\n".join(parts).strip()
    except Exception as e:
        return f"[Could not read this document: {e}]"


# ---------------------------------------------------------------------------
# Complaint-order outcome extraction. complaints.json has always carried
# orderDmsRefNo/orderFileName per complaint -- this session's own document-
# grounding work is what surfaced that these were never actually downloaded
# (see api_client.download_complaint_orders, which fetches them via the same
# DMS mechanism already used for project documents). Once downloaded, this
# reuses _extract_document_text (the OCR path this same session fixed) to
# read each order and classifies its outcome via a small, named set of
# keywords MahaRERA orders commonly use -- checked in priority order, most
# specific first, since e.g. "disposed of as withdrawn" should classify as
# "withdrawn" rather than the more generic "disposed_other" match. Anything
# that matches none of them is reported as not determinable, never guessed.
# ---------------------------------------------------------------------------

_ORDER_OUTCOME_PATTERNS = [
    ("withdrawn", re.compile(r"disposed?\s+(?:of|off)\s+as\s+withdrawn", re.I)),
    ("settled_by_conciliation", re.compile(r"\bconciliat(?:ion|ed)\b", re.I)),
    ("dismissed", re.compile(r"\b(?:complaint|appeal)\s+(?:is|stands)?\s*dismissed\b|\bdismissed\b", re.I)),
    ("allowed", re.compile(r"\b(?:complaint|appeal)\s+(?:is|stands)?\s*allowed\b", re.I)),
    ("disposed_other", re.compile(r"\bdisposed?\s+(?:of|off)\b", re.I)),
]


def _extract_complaint_outcome(document_text: str) -> str:
    """Best-effort classification of a downloaded complaint-order PDF's
    outcome from its extracted text. Returns "not determinable from
    extracted text" rather than guessing if nothing matches -- this is a
    heuristic label meant to summarize a large complaint count at a glance,
    not a substitute for a human reading the actual order for anything that
    matters to a real decision."""
    if not document_text or document_text.startswith("[Could not read") or "[OCR unavailable" in document_text:
        return "not determinable -- document could not be read"
    for label, pattern in _ORDER_OUTCOME_PATTERNS:
        if pattern.search(document_text):
            return label
    return "not determinable from extracted text"


def summarize_complaint_outcomes(complaint_orders_manifest: list, documents_dir: str) -> dict:
    """For each successfully downloaded/reused complaint-order PDF, reads it
    and classifies its outcome. Returns {outcome_counts: {label: count},
    per_complaint: [{complaint_registration_no, complaint_id, outcome,
    order_filename}]} -- both the aggregate breakdown for a summary
    paragraph and the per-complaint detail for anyone who wants to trace a
    specific outcome back to its actual order file."""
    outcome_counts = {}
    per_complaint = []
    for entry in complaint_orders_manifest:
        if entry.get("status") not in ("downloaded", "reused") or not entry.get("saved_filename"):
            continue
        path = os.path.join(documents_dir, entry["saved_filename"])
        text = _extract_document_text(path)
        outcome = _extract_complaint_outcome(text)
        outcome_counts[outcome] = outcome_counts.get(outcome, 0) + 1
        per_complaint.append({
            "complaint_registration_no": entry.get("complaint_registration_no"),
            "complaint_id": entry.get("complaint_id"),
            "outcome": outcome,
            "order_filename": entry["saved_filename"],
        })
    return {"outcome_counts": outcome_counts, "per_complaint": per_complaint}


# ---------------------------------------------------------------------------
# Entity-name normalisation. Scraped/filed names arrive in inconsistent
# shapes -- confirmed live in a single project's own data: "SUNDEEP  PODDAR"
# (double space, ALL CAPS) alongside "Vimanam Realty LLP" (title case) and
# "SAHEJTA REALITY LLP" (ALL CAPS) in the same list. Rendering those
# side by side in one table reads as sloppy, so every name that reaches a
# table goes through here first.
#
# Deliberately conservative: an input that is NOT effectively all-caps is
# left alone apart from whitespace collapsing, so a deliberately-mixed name
# ("A plus Architects plus Planners") is never re-cased into something its
# owner doesn't use. Only genuinely ALL-CAPS input gets title-cased, then a
# small token map restores the forms that title-casing gets wrong.
# ---------------------------------------------------------------------------

_NAME_TOKEN_FIXES = {
    "LLP": "LLP", "LLC": "LLC", "PLC": "PLC", "HUF": "HUF", "OPC": "OPC",
    "JV": "JV", "AOP": "AOP", "PVT": "Pvt", "PVT.": "Pvt.", "LTD": "Ltd",
    "LTD.": "Ltd.", "CO": "Co", "CO.": "Co.", "AND": "and", "OF": "of",
    "THE": "the",
}


def _normalise_entity_name(name: str) -> str:
    """Collapses whitespace in a person/company name, and title-cases it if
    (and only if) it arrived effectively ALL CAPS. Returns "" for a missing
    name rather than "None"."""
    cleaned = re.sub(r"\s+", " ", str(name or "").strip())
    if not cleaned:
        return ""

    letters = [ch for ch in cleaned if ch.isalpha()]
    if letters and not all(ch.isupper() for ch in letters):
        return cleaned  # already mixed case -- the owner's own styling, leave it

    words = []
    for i, word in enumerate(cleaned.split(" ")):
        fixed = _NAME_TOKEN_FIXES.get(word.upper())
        if fixed is not None:
            # A lowercase connector ("and"/"of"/"the") never leads a name.
            words.append(fixed.capitalize() if i == 0 and fixed.islower() else fixed)
        else:
            words.append(word.capitalize())
    return " ".join(words)


# ---------------------------------------------------------------------------
# LLP-aware terminology -- an LLP has partners (or "designated partners"), not
# directors; a company has directors. Confirmed live as a real document bug:
# Trimity Realty LLP's Charter was reading "director" throughout despite
# ZaubaCorp's own Designation field correctly saying "Designated Partner" --
# the underlying facts were right, only the surrounding template prose
# ignored entity type. Both charter_report.py and charter_research_prep.py
# call these rather than each keeping their own copy.
# ---------------------------------------------------------------------------

def _is_llp(facts: dict) -> bool:
    profile = facts.get("company_profile_check") or {}
    corp = facts.get("corporate_identity") or {}
    text = " ".join([
        str(profile.get("class_of_company") or ""),
        str((corp.get("organization_type") or {}).get("value") or ""),
    ]).lower()
    return "limited liability partnership" in text or re.search(r"\bllp\b", text) is not None


def _is_partnership(facts: dict) -> bool:
    """Broader than _is_llp: also true for a plain (unincorporated)
    Partnership Firm under the Indian Partnership Act, 1932 -- confirmed
    live as a real gap for IRA Homes, a promoter with no CIN or LLPIN at
    all (no MCA registration of any kind), whose partners are still
    "partners", never "directors", even though it isn't an LLP and
    "Designated Partner" (an LLP-Act-specific title) doesn't apply to
    them either. Use this for the plain director/partner noun; keep
    _is_llp for the "designated partner" label specifically, which is
    only ever correct for a true LLP."""
    if _is_llp(facts):
        return True
    profile = facts.get("company_profile_check") or {}
    corp = facts.get("corporate_identity") or {}
    text = " ".join([
        str(profile.get("class_of_company") or ""),
        str((corp.get("organization_type") or {}).get("value") or ""),
    ]).lower()
    return "partnership firm" in text or bool(re.search(r"\bpartnership\b", text))


def _role_word(facts: dict, count: int = 1, designated: bool = False) -> str:
    """The correct noun for a person who sits on this entity's board/
    partnership, singular or plural per `count`. `designated=True` asks for
    "designated partner" specifically (only meaningful when it's a true
    LLP; ignored for a plain partnership firm or a company, which just
    gets "partner" or "director" respectively)."""
    if _is_partnership(facts):
        base = "designated partner" if (designated and _is_llp(facts)) else "partner"
    else:
        base = "director"
    return base if count == 1 else base + "s"


# ---------------------------------------------------------------------------
# Professional team of record -- code-computed from MahaRERA's own structured
# `professionals` category, NOT model-authored.
#
# Why this exists: the model was writing local_planning.professionals_of_record
# from the DOCUMENTS it was handed and reporting, verbatim on a real project,
# "Engineer and CA firms of record were listed in MahaRERA's professionals
# category but not individually named in the documents reviewed this pass" --
# while that very category, already fetched and passed to it, named all five
# (architect, engineer, chartered accountant, and two "Other" firms) with
# their registration numbers. Same class of thing as document_library and
# promoter_portfolio: a deterministic list should never be re-derived or
# paraphrased by a model when it can be read directly.
#
# Registration numbers are reported EXACTLY as filed, never cleaned up or
# dropped for looking wrong -- one real project's architect CoA number is
# literally "a", and that filing-quality signal belongs in front of a reader
# rather than being quietly sanitised away.
# ---------------------------------------------------------------------------

# professionalTypeName -> the field that carries that role's registration id.
_PROFESSIONAL_REGISTRATION_FIELDS = {
    "Architect": ("architectCoARegistrationNo", "CoA registration"),
    "Engineer": ("engineerLicenseNo", "Engineer licence"),
    "Chartered Accountant": ("caIcaiMembershipNo", "ICAI membership"),
    "Real Estate Agent": ("realEstateAgentReraRegNo", "RERA registration"),
}


def summarize_professionals(category_data: dict) -> list:
    """Returns [{role, name, is_individual, registration_label,
    registration_number}] -- one entry per professional on MahaRERA's record
    for this project, in the order the portal lists them. Returns [] when the
    category is absent or empty (never a fabricated entry), so a caller can
    render an honest "none on record" line instead."""
    professionals = category_data.get("professionals")
    if not isinstance(professionals, list):
        return []

    team = []
    for entry in professionals:
        if not isinstance(entry, dict):
            continue
        role = str(entry.get("professionalTypeName") or "").strip() or "Not stated"
        name = _normalise_entity_name(
            entry.get("entityCompanyName") or entry.get("profileName") or entry.get("firstName") or ""
        )
        if not name:
            continue

        reg_field, reg_label = _PROFESSIONAL_REGISTRATION_FIELDS.get(role, (None, None))
        reg_number = str(entry.get(reg_field) or "").strip() if reg_field else ""

        team.append({
            "role": role,
            "name": name,
            "is_individual": str(entry.get("professionalPersonalTypeName") or "").strip().lower().startswith("person"),
            "registration_label": reg_label or "",
            "registration_number": reg_number,
        })
    return team


def _select_documents_for_extraction(documents_manifest: list, documents_dir: str) -> tuple[dict, list]:
    """Returns (extracted_text_by_label, all_labels_with_status) -- the
    former feeds the Claude prompt (capped total size), the latter fills the
    Document Library table for every document regardless of whether it was
    opened."""
    extracted = {}
    all_labels = []
    total_chars = 0

    for entry in documents_manifest:
        label = entry.get("label") or entry.get("saved_filename") or "Unnamed document"
        available = entry.get("status") in ("downloaded", "reused") and entry.get("saved_filename")
        if not available:
            all_labels.append({"document_name": label, "status": f"Not available ({entry.get('status')})"})
            continue

        if label in extracted:
            # Multiple manifest entries can share the exact same generic
            # label (e.g. MahaRERA's own catch-all "Other -- Legal" bucket
            # covering many distinct sale deeds/orders) -- extracting a
            # second one would silently overwrite the first in `extracted`
            # (same dict key) while still burning the shared char budget,
            # starving out later, distinctly-labeled, higher-value documents
            # (a real bug, confirmed live: it silently dropped the Title
            # Report and Form B for a project whose library happened to have
            # a dozen same-labeled documents ahead of them). Keep only the
            # first document under a given label; skip re-extracting the rest.
            all_labels.append({"document_name": label, "status": "Not opened this pass -- another document already opened under this same shared label"})
            continue

        is_priority = any(k in label.lower() for k in _HIGH_PRIORITY_DOC_KEYWORDS) or any(
            k in (entry["saved_filename"] or "").lower() for k in _HIGH_PRIORITY_DOC_KEYWORDS
        )
        if is_priority and total_chars < _MAX_TOTAL_DOC_CHARS:
            path = os.path.join(documents_dir, entry["saved_filename"])
            text = _extract_document_text(path)[:_MAX_CHARS_PER_DOC]
            extracted[label] = text
            total_chars += len(text)
            all_labels.append({"document_name": label, "status": "Opened directly this pass"})
        else:
            reason = "not among the identified high-priority legal/regulatory document types" if is_priority else "not a high-priority legal/regulatory document type for this pass"
            all_labels.append({"document_name": label, "status": f"Not opened this pass -- {reason}"})

    return extracted, all_labels


_FIELD_WITH_SOURCE = {
    "type": "object",
    "properties": {"value": {"type": "string"}, "source": {"type": "string"}},
    "required": ["value", "source"],
}
_PLAIN_FIELD = {"type": "string"}
_INT_FIELD = {"type": "integer"}

_CHARTER_FACTS_SCHEMA = {
    "type": "object",
    "properties": {
        "methodology_note": _PLAIN_FIELD,
        "executive_summary": _PLAIN_FIELD,
        "land_identification": {
            "type": "object",
            "properties": {k: _FIELD_WITH_SOURCE for k in (
                "survey_cts_plot_numbers", "village_locality", "mandal_taluka_district",
                "pincode", "total_gross_area", "area_affected", "net_area",
            )},
        },
        "corporate_identity": {
            "type": "object",
            "properties": {k: _FIELD_WITH_SOURCE for k in (
                "promoter_name", "organization_type", "cin_llpin", "registered_office_main",
                "registered_office_board_resolution", "registered_office_planning_stage",
                "authorized_signatory", "partners_directors", "landowner_investor",
            )},
        },
        "address_discrepancy_note": _PLAIN_FIELD,
        "corporate_registry_cross_check": _PLAIN_FIELD,
        "litigation_status": _FIELD_WITH_SOURCE,
        "location_coordinates_note": _PLAIN_FIELD,
        "neighbourhood": {
            "type": "object",
            "properties": {k: _PLAIN_FIELD for k in ("east", "west", "north", "south")},
        },
        "distances": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "landmark": {"type": "string"},
                    "distance_time": {"type": "string"},
                    "route_note": {"type": "string"},
                },
                "required": ["landmark", "distance_time", "route_note"],
            },
        },
        "connectivity": {
            "type": "object",
            "properties": {k: _PLAIN_FIELD for k in ("road", "rail", "metro", "air")},
        },
        "social_infrastructure": _PLAIN_FIELD,
        "fsi_governing_framework": _PLAIN_FIELD,
        "fsi_interpretation": _PLAIN_FIELD,
        "fsi_metrics": {
            "type": "object",
            "properties": {
                **{k: _PLAIN_FIELD for k in (
                    "net_land_area", "approved_bua", "sanctioned_bua", "mortgage_area", "implied_fsi",
                )},
                # {value, source} rather than plain text: MahaRERA's own API exposes NO lender/bank
                # field anywhere (confirmed by checking every one of the 9 raw category payloads --
                # projects, partners, professionals, sro_details, past_experiences, documents,
                # complaints, appeals, spocs -- zero matches for bank/lender/mortgage/finance/loan in
                # any of them). The only place a lender's identity ever appears is free text inside a
                # downloaded title/encumbrance/Form-B document, so this needs the same independent
                # re-verification path as land_identification/corporate_identity, which a bare string
                # wouldn't get.
                "mortgage_lender": _FIELD_WITH_SOURCE,
            },
        },
        "rules_statutory": {
            "type": "object",
            "properties": {k: _PLAIN_FIELD for k in (
                "governing_act", "planning_approval_sequence", "allotment_mechanics",
            )},
        },
        "rera_compliance": {
            "type": "object",
            "properties": {k: _PLAIN_FIELD for k in (
                "registration_summary", "collection_account", "escrow_subaccounts",
                "litigations_complaints_appeals", "statutory_declaration", "construction_progress",
            )},
        },
        "local_planning": {
            "type": "object",
            "properties": {k: _PLAIN_FIELD for k in (
                "authority_of_record", "project_type", "professionals_of_record",
            )},
        },
        "micro_market_overview": _PLAIN_FIELD,
        "comparables": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "project": {"type": "string"}, "configuration": {"type": "string"},
                    "pricing": {"type": "string"}, "source": {"type": "string"},
                    "distance_km": {"type": "string", "description": "Approximate straight-line or driving distance from the subject project, as a bare number with NO 'km' suffix -- e.g. '3.8', not '3.8 km'"},
                },
                "required": ["project", "configuration", "pricing", "source", "distance_km"],
            },
        },
        "area_intelligence_trend": _PLAIN_FIELD,
        "rera_core_fields": {
            "type": "object",
            "properties": {
                **{k: _PLAIN_FIELD for k in (
                    "project_name", "registration_number", "promoter_name", "authority",
                    "plan_approval_number", "project_status", "approved_date",
                    "proposed_completion_date", "project_type", "litigations_per_record",
                    "promoter_land_owner_investor", "collection_bank_account", "total_building_units",
                )},
                # Structured companions to the prose fields above, so downstream scoring/KPI-card
                # code can read the real number directly instead of regexing free text (see the
                # _SYSTEM_PROMPT rule on these -- a project's exact phrasing shouldn't determine
                # whether its own clean record silently drops out of the Developer Score).
                "total_complaints_count": _INT_FIELD,
                "total_appeals_count": _INT_FIELD,
                "units_total": _INT_FIELD,
                "units_sold": _INT_FIELD,
                "completion_date_current": _PLAIN_FIELD,
                "completion_date_original": _PLAIN_FIELD,
            },
        },
        # Feeds the code-computed Developer Score's "track record years" criterion --
        # deliberately its own top-level object, not nested under rera_core_fields, since
        # MahaRERA's own record never carries this (it's promoter/group history, not a
        # RERA-filed fact). Omit this object entirely if it can't be confirmed -- see the
        # _SYSTEM_PROMPT rule below -- rather than guessing a number.
        "developer_track_record": {
            "type": "object",
            "properties": {
                "years_in_industry": _INT_FIELD,
                "years_in_industry_basis": _PLAIN_FIELD,
            },
        },
        "unit_summary_note": _PLAIN_FIELD,
        "blocks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "block_wing": {"type": "string"}, "floors": {"type": "string"},
                    "config": {"type": "string"}, "units_counted": {"type": "string"},
                    "note": {"type": "string"},
                },
                "required": ["block_wing", "floors", "config", "units_counted", "note"],
            },
        },
        "documents_reviewed_note": _PLAIN_FIELD,
        "documents_absent_note": _PLAIN_FIELD,
        "gaps": {"type": "array", "items": {"type": "string"}},
        "sources": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "label": {"type": "string"},
                    "ref": {"type": "string"},
                    "topic": {"type": "string", "description": "Short category this source backs. Used both to check whether a material topic is backed by more than one independent source, AND to automatically render an inline citation next to any paragraph on that topic -- so use EXACTLY one of these known values whenever it applies (a value outside this list still gets a source-count check, but will NOT get an automatic inline citation anywhere): 'land_title', 'corporate_identity', 'litigation', 'pricing', 'market_trend', 'distance', 'project_registration' (any fact drawn from MahaRERA's own project record -- rera_core_fields, rules_statutory, local_planning, neighbourhood, unit_summary_note all draw from this), 'legal_documents', 'reputation', 'credit_rating', 'insolvency_status', 'company_profile', 'group_companies'."},
                    "published_date": {"type": "string", "description": "The date the underlying information was published or last updated, as YYYY-MM-DD if the source discloses one, or the literal string 'unknown' if it does not (e.g. a live database record with no publish date shown) -- never guessed or invented."},
                    "accessed_date": {"type": "string", "description": "The date this specific source was actually retrieved/read during this research pass, as YYYY-MM-DD."},
                },
                "required": ["label", "ref", "topic", "published_date", "accessed_date"],
            },
        },
    },
    "required": [
        "methodology_note", "executive_summary", "land_identification", "corporate_identity",
        "address_discrepancy_note", "corporate_registry_cross_check", "litigation_status",
        "location_coordinates_note", "neighbourhood", "distances", "connectivity",
        "social_infrastructure", "fsi_governing_framework", "fsi_interpretation", "fsi_metrics",
        "rules_statutory", "rera_compliance", "local_planning", "micro_market_overview",
        "comparables", "area_intelligence_trend", "rera_core_fields", "unit_summary_note",
        "blocks", "documents_reviewed_note", "documents_absent_note", "gaps", "sources",
    ],
}

_CHARTER_JSON_SHAPE = json.dumps(_CHARTER_FACTS_SCHEMA)

_SYSTEM_PROMPT = f"""You are producing the factual content for a real-estate Company Charter \
document (a due-diligence reference, not marketing material) for one MahaRERA-registered \
project, following a fixed template structure.

Rules:
- Every field must be either a real, sourced fact or an explicit statement that it could not \
be confirmed (put those in `gaps`, do not leave a field blank or guess).
- `rera_core_fields.total_complaints_count`, `total_appeals_count`, `units_total`, and \
`units_sold` must be the literal integers your own `litigations_per_record`/`total_building_units` \
prose describes -- not a restatement of the prose, the actual number a downstream scoring step \
can read directly. `completion_date_current` is the current proposed completion date as \
YYYY-MM-DD; `completion_date_original` is the ORIGINAL declared date as YYYY-MM-DD if it has ever \
been extended/pushed back, or the exact literal empty string "" if it has never been extended \
(same as current) -- never omit these because the prose already covers it, and never guess a \
number you can't confirm (state the uncertainty in the prose field instead, and leave the count \
fields out of your JSON entirely for that one project rather than invent a number).
- Distances/routes: you do NOT have a live Maps browsing tool here -- use web_search to find \
driving distances/times and state plainly in `location_coordinates_note` that these are \
web-search estimates, not live Maps-verified routes (this template explicitly wants that kind \
of honesty about what's confirmed vs approximated). Each `distances[].landmark` MUST be an \
actual named place (e.g. "Chhatrapati Shivaji Maharaj International Airport"), never a generic \
category label like "Nearest airport" -- a later step may look up the precise route for the \
named place you give.
- Do not conflate similarly-named companies -- cross-check any CIN/LLPIN claim against the \
promoter's registered legal name and address from the provided RERA data before accepting it.
- For the FSI section, show the computation (net area, approved/sanctioned BUA -> implied FSI) \
using the actual figures given, and state plainly whether this is your interpretation or a \
figure confirmed by an actual sanctioned-FSI certificate.
- The standing gap about promoter shareholding splits/personal net worth is permanent and \
already in the template -- do not attempt to research it, and do not repeat it in your own \
`gaps` list.
- Comparable projects (`comparables`) must be GENUINELY INDEPENDENT developments located within \
approximately 3-5 km of the subject project -- not the subject's own adjacent phase or the same \
complex under a different marketing name (e.g. do not count a project's own "Phase 2" or an \
immediately-adjacent same-developer extension as a comparable; it is the same product, not a \
market comparable). State the approximate distance for each comparable in `distance_km`. If no \
genuinely independent comparable exists within that radius, say so in `gaps` rather than \
substituting a same-complex neighbor or a locality-wide average as if it satisfied the radius.
- Use web_search to find how many years this promoter (or its parent group, if this is a \
group-affiliated SPV incorporated recently for one project) has actually been active in real- \
estate development, and fill `developer_track_record.years_in_industry` (a plain integer) with \
`years_in_industry_basis` stating what start event you counted from (e.g. "group's first \
recorded project launch in 1998, per the group's own corporate website") and its source. If you \
cannot confirm this from any source, omit `developer_track_record` entirely rather than guessing \
a number or assuming the SPV's own incorporation date equals the group's real experience.
- Every entry in `sources` needs a `topic` (see the exact known values listed on that field), a \
`published_date` (YYYY-MM-DD if the source states one, or literally "unknown" if it does not -- \
never invent one), and an `accessed_date` (YYYY-MM-DD, when you actually looked at it this pass). \
These feed a later, code-computed confidence score -- an omitted or fabricated date is worse than \
an honest "unknown".
- A later, code-only step renders an inline citation next to several prose fields (rules_statutory, \
local_planning, unit_summary_note, neighbourhood, connectivity, micro_market_overview, \
area_intelligence_trend, location_coordinates_note) by matching that field's likely topic against \
`sources[].topic`. It does NOT read your prose to figure out what backs it -- if a fact in one of \
those fields came from a source, tag that source with the matching known topic value (see the field \
description on `sources[].topic`) or it will render with no citation at all, even though you did \
have a source. The `_FIELD_WITH_SOURCE` fields (land_identification, corporate_identity, \
litigation_status, fsi_metrics.mortgage_lender) work differently: their own `source` value is shown \
directly, so just make sure it's a real, specific document name/URL, not a vague description.
- Use web_search as many times as you need across multiple turns. Once you have everything \
you need, your FINAL reply must be ONLY a single raw JSON object -- no prose, no markdown \
code fences, nothing before or after it -- matching exactly this JSON Schema: \
{_CHARTER_JSON_SHAPE}"""


def _run_charter_pass(user_prompt: str) -> dict:
    # Delegates to deep_research's tool-runner helper (iterates the
    # BetaToolRunner to its final BetaMessage and parses the raw-JSON reply)
    # rather than duplicating that logic here. `facts` produced here is
    # shared by both Internal and External renders, so only Section B
    # (doc-variant-agnostic) is injected -- never Section C.
    system = _charter_system_blocks(external=False, extra=_SYSTEM_PROMPT)
    return deep_research._run_agentic_pass(user_prompt, system, label="charter_pass")


_MAPS_SCRAPE_ENV_VAR = "COMPANY_CHARTER_USE_MAPS_SCRAPE"
_MAPS_ROUTE_PATTERN = re.compile(
    r"Copy link\n.*?\n((?:\d+\s*hr\s*)?\d+\s*min)\n([\d.]+\s*km)\nvia ([^\n]+)"
)
_MAPS_MIN_TOKEN_OVERLAP = 0.5


def _tokenize(s: str) -> set:
    return {w for w in re.findall(r"[a-z0-9]+", s.lower()) if len(w) > 2}


def _destination_plausibly_resolved(requested: str, resolved: str) -> bool:
    """Sanity check against Maps silently mis-resolving a bad/ambiguous
    destination to an unrelated nearby point instead of failing outright --
    confirmed live: a nonsense query like "Not a real place asdkjaskdj12345"
    still returns a route, just to a random unrelated address, with zero
    word overlap with what was asked for. A genuine match (even lightly
    reworded/expanded by Maps, e.g. a city name appended) retains most of
    the requested landmark's significant words."""
    requested_tokens = _tokenize(requested)
    if not requested_tokens:
        return True  # nothing meaningful to check against
    overlap = requested_tokens & _tokenize(resolved)
    return len(overlap) / len(requested_tokens) >= _MAPS_MIN_TOKEN_OVERLAP


def _lookup_maps_distance(origin: str, destination: str) -> dict | None:
    """Launches a headless Playwright browser, opens Google Maps driving
    directions from origin to destination, and reads the rendered
    duration/distance/route off the page. Returns None on ANY failure
    (selector miss, timeout, network, or Maps resolving `destination` to an
    unrelated place -- see _destination_plausibly_resolved) -- callers must
    fall back to the model's own web_search estimate rather than treat None
    as an error.

    Verified against the live site (short in-city and long inter-city
    routes, plus a deliberately bogus destination to confirm the
    mis-resolution check actually fires) before shipping, but this scrapes
    Google Maps' consumer web UI rather than their Distance Matrix/Routes
    API -- it can break without warning if Google changes their page, and
    likely doesn't comply with Google's Terms of Service (which is exactly
    why they sell an API for this). Off by default -- see
    _MAPS_SCRAPE_ENV_VAR."""
    from urllib.parse import quote

    from playwright.sync_api import sync_playwright

    url = f"https://www.google.com/maps/dir/{quote(origin)}/{quote(destination)}"
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=20000)
            page.get_by_role("radio", name="Driving").click(timeout=10000)
            page.wait_for_timeout(4000)
            dest_box = page.get_by_role("combobox").nth(1)
            resolved_destination = dest_box.locator("input, textarea").first.input_value()
            text = page.inner_text("body")
            browser.close()
        if not _destination_plausibly_resolved(destination, resolved_destination):
            return None
        match = _MAPS_ROUTE_PATTERN.search(text)
        if not match:
            return None
        duration, distance, route = match.groups()
        return {"duration": duration, "distance": distance, "route": route}
    except Exception:
        return None


def _refine_distances_with_maps(facts: dict, origin: str) -> dict:
    """Opt-in (COMPANY_CHARTER_USE_MAPS_SCRAPE=1) precision upgrade: replaces
    each distance entry's web_search-estimated distance_time/route_note with
    a live-scraped Google Maps driving route when the lookup succeeds,
    leaving the model's own estimate untouched (not silently dropped) when
    it doesn't -- so a scrape failure degrades to exactly today's behavior
    rather than blanking the field."""
    if os.environ.get(_MAPS_SCRAPE_ENV_VAR) != "1":
        return facts

    for entry in facts.get("distances", []):
        landmark = entry.get("landmark")
        if not landmark:
            continue
        result = _lookup_maps_distance(origin, landmark)
        if result:
            entry["distance_time"] = f"{result['duration']} / {result['distance']}"
            entry["route_note"] = f"via {result['route']} (live Google Maps driving route, scraped this run)"
    return facts


def _verify_one_field(field_name: str, field, gaps: list, stats: dict) -> None:
    """Shared check applied to every {value, source}-shaped field below --
    only claims sourced from a live URL are re-verified (a claim sourced
    from a downloaded primary document is already as good as it gets; there
    is no independent web check to run against it -- see
    _check_document_grounding for that category instead). A claim that was
    actually checked and failed is demoted into `gaps` with its value
    replaced, never left silently asserted. A claim whose check could not
    even run (no ANTHROPIC_API_KEY, network failure, rate limit --
    deep_research._verify_claim's "verification_error" status) is a
    different, weaker finding than "checked and failed": it is NOT proof the
    claim is wrong, so the value is left untouched, but the fact that no
    independent check actually happened is still surfaced as an explicit gap
    -- this is what lets that absence show up in the document on its own,
    rather than depending on whoever is running this pass to remember and
    manually re-check it themselves.

    `stats` is mutated in place ({"attempted", "confirmed", "failed",
    "could_not_run"}) so _compute_documentation_confidence_score can report a
    real Independent Verification Rate instead of parsing gap text."""
    if not isinstance(field, dict) or not field.get("value") or not field.get("source"):
        return
    if not field["source"].startswith("http"):
        return
    stats["attempted"] += 1
    verdict = deep_research._verify_claim(f"{field_name}: {field['value']}", field["source"], label="material_claim_verify")
    status = verdict.get("status")
    if status == "confirmed":
        stats["confirmed"] += 1
        return
    if status == "verification_error":
        stats["could_not_run"] += 1
        gaps.append(
            f"{field_name}: {field['value']} (NOT independently re-verified this pass -- "
            f"{verdict.get('reason', 'verification error')} -- treat as unconfirmed, not as checked-and-passed)"
        )
        return
    stats["failed"] += 1
    gaps.append(f"{field_name}: {field['value']} (verification: {verdict.get('reason', 'unconfirmed')})")
    field["value"] = "Not independently confirmed -- see gaps"


def _verify_material_claims(facts: dict) -> dict:
    """Re-checks the highest-stakes sourced claims (corporate identity, land
    identification, litigation, mortgage lender) the same way
    deep_research._verify_claim does. Unconfirmed ones move to `gaps`
    rather than staying asserted."""
    gaps = list(facts.get("gaps", []))
    stats = {"attempted": 0, "confirmed": 0, "failed": 0, "could_not_run": 0}
    for section_name in ("land_identification", "corporate_identity", "fsi_metrics"):
        for field_name, field in list(facts.get(section_name, {}).items()):
            _verify_one_field(field_name, field, gaps, stats)
    _verify_one_field("litigation_status", facts.get("litigation_status"), gaps, stats)
    facts["gaps"] = gaps
    facts.setdefault("_verification_stats", {})["url"] = stats
    return facts


# ---------------------------------------------------------------------------
# Document-grounding check -- a second, DIFFERENT failure mode from the web
# re-verification above. _verify_one_field only ever re-checks claims sourced
# from a live URL; a claim sourced from a downloaded primary document (Form
# B, the title report, MahaRERA's own raw JSON) has no independent web check
# to run against it -- by design, that document IS the primary source. But
# "no independent check exists" is not the same as "cannot be checked at
# all": if a name, code, or figure was simply misread or mistranscribed out
# of that document while building `facts`, nothing above would ever catch it.
# This check is mechanical (plain string matching, no API call, no network,
# works with or without ANTHROPIC_API_KEY) -- it asks only "do this field's
# distinguishing, verbatim-checkable details actually appear somewhere in the
# text of the document(s) it cites?", not "is this a good summary of them" --
# so it is deliberately narrow: it flags a name/code that doesn't appear
# anywhere in its cited source at all (the exact class of error just found:
# a wrong director name), not a paraphrase or a computed/aggregated number
# that legitimately won't appear as a literal string in the source.
# ---------------------------------------------------------------------------

_PROPER_NOUN_RUN_RE = re.compile(r"\b([A-Z][a-zA-Z']+(?:\s+[A-Z][a-zA-Z']+)+)\b")
# Mixed letter+digit codes (CIN, survey/plot numbers, ...) -- deliberately
# excludes pure-digit tokens (dates, counts) since those are frequently
# computed/aggregated rather than copied verbatim from the source text, and
# flagging them would be a false alarm, not a real transcription signal.
_ALPHANUMERIC_CODE_RE = re.compile(r"\b(?=[A-Za-z0-9]*[A-Za-z])(?=[A-Za-z0-9]*\d)[A-Za-z0-9/\-]{5,}\b")


def _tokenize_significant(value: str) -> set:
    tokens = {m.group(1) for m in _PROPER_NOUN_RUN_RE.finditer(value)}
    tokens |= {m.group(0) for m in _ALPHANUMERIC_CODE_RE.finditer(value)}
    return tokens


def _document_grounding_text(source: str, extracted_docs: dict, category_data: dict, documents_manifest: list) -> str:
    """Best-effort combined text of whatever this field's `source` string
    actually references: any extracted-document label OR saved_filename
    mentioned by name (both are checked -- a `source` string commonly names
    the raw filename, e.g. "Mamurdi_Supplemental title report.pdf", while
    extracted_docs is keyed by the manifest's often-different display label,
    e.g. "Legal Title report"; matching on label alone silently missed the
    real extracted text for every field authored this way, confirmed live),
    plus any `raw/<category>.json` mention resolved against that category's
    own already-fetched payload. Returns "" if the source doesn't resolve to
    anything we actually have text for (e.g. a document that wasn't opened
    this pass) -- callers must treat that as "nothing to check", not a
    failure."""
    source_lower = source.lower()
    matched_labels = {label for label in extracted_docs if label.lower() in source_lower}
    for entry in documents_manifest:
        label = entry.get("label") or entry.get("saved_filename") or "Unnamed document"
        filename = entry.get("saved_filename") or ""
        if label in extracted_docs and filename and filename.lower() in source_lower:
            matched_labels.add(label)
    combined = [extracted_docs[label] for label in matched_labels]
    for category, data in category_data.items():
        if f"raw/{category}.json" in source or f"raw\\{category}.json" in source:
            combined.append(json.dumps(data or {}))
    return "\n".join(combined)


_DOCUMENT_CLAIM_VERIFY_SYSTEM_PROMPT = """Re-check one specific claim against the actual extracted text of \
the local document it cites. This is NOT a web claim -- you are given the document's own text directly \
below; do not use web_search, judge only from the text provided (a faithful paraphrase or a computed/ \
summarized statement should still be judged "confirmed" if the text supports it -- this is a meaning \
check, not an exact-string match; that cheaper check already ran and failed, which is why you're being \
asked).

Your FINAL reply must be ONLY a single raw JSON object -- no prose, no markdown code fences -- matching \
exactly this shape: {"status": "confirmed" | "unsupported" | "stale", "reason": "one sentence"}"""


_GROK_ENV_VAR = "XAI_API_KEY"
_GROK_MODEL = "grok-4.5"
_GROK_BASE_URL = "https://api.x.ai/v1"


def _parse_json_text(text: str) -> dict:
    """Strips a leading/trailing markdown code fence if present, then parses
    raw JSON -- shared by every "final reply must be ONLY a JSON object"
    verifier prompt in this file, regardless of which provider answered."""
    text = text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:]
        text = text.strip()
    return json.loads(text)


def _verify_document_claim_via_grok(claim: str, document_text: str) -> dict:
    """Fallback path for _verify_document_claim when Claude can't run the
    check (missing/invalid ANTHROPIC_API_KEY, rate limit, etc.) but an xAI
    key is available -- genuine redundancy against one provider being down
    or unconfigured, not a way to make this check API-free (it still needs
    a working key, just not necessarily this one). Uses the OpenAI-compatible
    SDK already installed in this environment, pointed at xAI's endpoint, per
    xAI's own documented pattern (client.responses.create with an `input`
    list, not client.chat.completions.create) -- response-text extraction
    tries the documented output_text convenience property first and falls
    back to walking `response.output` directly, since that exact shape
    wasn't independently confirmed against a live call (no XAI_API_KEY was
    available in this environment to test against)."""
    from openai import OpenAI

    client = OpenAI(api_key=os.environ[_GROK_ENV_VAR], base_url=_GROK_BASE_URL)
    prompt = f"Claim: {claim}\n\nExtracted document text:\n{document_text[:_MAX_CHARS_PER_DOC]}"
    response = client.responses.create(
        model=_GROK_MODEL,
        input=[
            {"role": "system", "content": _DOCUMENT_CLAIM_VERIFY_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
    )
    text = getattr(response, "output_text", None)
    if not text:
        text = response.output[0].content[0].text
    result = _parse_json_text(text)
    if result.get("status") not in ("confirmed", "unsupported", "stale"):
        return {"status": "unsupported", "reason": "verifier returned an unrecognized status"}
    return result


def _verify_document_claim(claim: str, document_text: str) -> dict:
    """Same {status, reason} shape as deep_research._verify_claim, but the
    model is handed the document's own extracted text directly instead of
    being told to search the web for a URL. This is strictly an escalation
    on top of _check_document_grounding's free mechanical check, called only
    for a field that check couldn't confirm -- it costs a real API call and
    is bounded by the same extraction/OCR quality as the mechanical check: it
    cannot confirm something the extracted text genuinely doesn't contain,
    only recognize a paraphrase of something the text DOES contain -- exactly
    the gap a literal substring match leaves open.

    Tries Claude first (the project's primary provider); if that fails for
    any reason and an xAI key is configured, falls back to Grok before
    giving up -- redundancy against one provider being down or unconfigured,
    not a way to avoid needing a working key at all. Only returns
    "verification_error" if neither provider could complete the check."""
    prompt = f"Claim: {claim}\n\nExtracted document text:\n{document_text[:_MAX_CHARS_PER_DOC]}"
    try:
        result = deep_research._run_agentic_pass(prompt, _DOCUMENT_CLAIM_VERIFY_SYSTEM_PROMPT, label="document_grounding_verify")
        if result.get("status") not in ("confirmed", "unsupported", "stale"):
            return {"status": "unsupported", "reason": "verifier returned an unrecognized status"}
        return result
    except Exception as claude_error:
        if not os.environ.get(_GROK_ENV_VAR):
            return {"status": "verification_error", "reason": f"verification could not run: {claude_error}"}
        try:
            return _verify_document_claim_via_grok(claim, document_text)
        except Exception as grok_error:
            return {
                "status": "verification_error",
                "reason": f"verification could not run on either provider -- Claude: {claude_error}; Grok: {grok_error}",
            }


def _check_document_grounding(facts: dict, extracted_docs: dict, category_data: dict, documents_manifest: list) -> dict:
    """Flags, as an explicit gap, a {value, source}-shaped field whose
    distinguishing details don't appear anywhere in the text of the local
    document(s)/raw-JSON it cites. Only runs against non-URL sources;
    anything starting with "http" is already covered by
    _verify_material_claims.

    A field that fails the free mechanical check gets one more chance: an
    LLM call (_verify_document_claim) that judges the claim against the same
    extracted text, catching a faithful paraphrase the substring match can't
    recognize. If that second check confirms it, no gap is added at all --
    the mechanical flag was a false alarm, not an error. If it also fails,
    the field is demoted the same way _verify_material_claims demotes a
    failed web claim: value replaced, explicit gap recorded -- we now have a
    real, checked failure, not just a heuristic flag. Only a genuine
    "verification_error" (the LLM check couldn't even be attempted) falls
    back to the original soft "manual re-check" gap without touching the
    value, since a heuristic flag alone is not proof of an error."""
    gaps = list(facts.get("gaps", []))
    stats = {"attempted": 0, "confirmed": 0, "failed": 0, "could_not_run": 0}
    fields_to_check = [
        (f"{section}.{name}", field)
        for section in ("land_identification", "corporate_identity", "fsi_metrics")
        for name, field in facts.get(section, {}).items()
    ]
    fields_to_check.append(("litigation_status", facts.get("litigation_status")))

    for field_name, field in fields_to_check:
        if not isinstance(field, dict) or not field.get("value") or not field.get("source"):
            continue
        if field["source"].startswith("http"):
            continue
        significant = _tokenize_significant(field["value"])
        if not significant:
            continue  # nothing verbatim-checkable here (pure narrative/summary text)
        grounding_text = _document_grounding_text(field["source"], extracted_docs, category_data, documents_manifest)
        if not grounding_text:
            continue  # cited document wasn't actually opened this pass -- nothing to check against
        stats["attempted"] += 1
        found = {t for t in significant if t.lower() in grounding_text.lower()}
        if found:
            stats["confirmed"] += 1
            continue

        verdict = _verify_document_claim(f"{field_name}: {field['value']}", grounding_text)
        status = verdict.get("status")
        if status == "confirmed":
            stats["confirmed"] += 1
            continue  # mechanical flag was a false alarm -- a faithful paraphrase, not an error
        sample = ", ".join(list(significant)[:3])
        if status == "verification_error":
            stats["could_not_run"] += 1
            gaps.append(
                f"{field_name}: none of its distinguishing details ({sample}) were found verbatim in "
                f"the text of its cited source document, and a second-opinion check could not run "
                f"({verdict.get('reason', 'verification error')}); flagged for a manual re-check "
                f"rather than trusted on transcription alone."
            )
            continue
        stats["failed"] += 1
        gaps.append(
            f"{field_name}: {field['value']} (neither a literal match nor an independent re-check against "
            f"its cited source document's text supported this claim -- {verdict.get('reason', 'unconfirmed')})"
        )
        field["value"] = "Not independently confirmed -- see gaps"
    facts["gaps"] = gaps
    facts.setdefault("_verification_stats", {})["document"] = stats
    return facts


# ---------------------------------------------------------------------------
# Credit-rating lookup (Phase 2) -- confirmed live against ICRA's real public
# rating database: a "Search by entity" box on icra.in backs onto a genuine
# AJAX endpoint (POST /Rating/GetRatingCompanys, requires the
# X-Requested-With: XMLHttpRequest header or it 404s -- an ASP.NET MVC
# AJAX-only-action quirk, not a real access restriction) returning
# {id, label} matches; the resulting /Rating/RatingDetails page is fully
# server-rendered (a plain requests.get() gets the real content, no JS
# execution needed) and its first table is always "Instrument | Ratings".
# Live-tested end to end against "Godrej Properties Limited" -- returned
# real, current ratings (ICRA A1+ commercial paper, AA+ (Stable) long-term).
#
# As of this pass, ICRA is no longer the only agency checked -- see
# _lookup_infomerics_rating below, added after ICRA-only coverage produced
# a real, misleading "not found" for a promoter's sister entity that was
# actually rated (by Infomerics, not ICRA). CRISIL's equivalent
# company-rating search still wasn't confirmed reverse-engineerable in the
# time spent on it (its site exposes an SME-specific autocomplete on a
# different product, not the main corporate-rating search) -- rather than
# build against an unconfirmed mechanism, coverage stays scoped to what
# actually works. CRISIL/CARE/India Ratings coverage can be added later if
# a working search endpoint is found for any of them.
#
# Matches ONLY on an exact (case-insensitive) name match against each
# agency's own company list -- never a fuzzy "probably the same company"
# guess, since misattributing a rating to the wrong legal entity would be
# a serious factual error in a due-diligence document. A promoter/SPV
# having no public rating from either agency is the NORMAL case, not a red
# flag: both only rate developers that sought a public rating (typically
# larger, listed, or NCD-issuing entities) -- most MahaRERA promoters are
# too small/private to ever be rated. Callers who also want to check a
# distinct parent/group entity (e.g. "Godrej Properties Limited" for
# promoter "Godrej Skyline Developers Limited") must call this a second
# time with that name explicitly and label the result as the PARENT's
# rating, never as if it were the promoter's own -- this function does not
# guess a parent itself.
# ---------------------------------------------------------------------------

_ICRA_SEARCH_URL = "https://www.icra.in/Rating/GetRatingCompanys"
_ICRA_DETAIL_URL = "https://www.icra.in/Rating/RatingDetails"


def _icra_search_companies(name: str) -> list:
    resp = requests.post(
        _ICRA_SEARCH_URL,
        data={"term": name},
        headers={"X-Requested-With": "XMLHttpRequest"},
        timeout=config.REQUEST_TIMEOUT,
    )
    resp.raise_for_status()
    return resp.json()


def _icra_fetch_rating_detail(company_id: str, company_name: str) -> dict:
    from urllib.parse import quote

    from bs4 import BeautifulSoup

    url = f"{_ICRA_DETAIL_URL}?CompanyId={company_id}&CompanyName={quote(company_name)}"
    resp = requests.get(url, timeout=config.REQUEST_TIMEOUT)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")

    instruments = []
    for table in soup.find_all("table"):
        headers = [th.get_text(strip=True) for th in table.find_all("th")]
        if headers[:2] == ["Instrument", "Ratings"]:
            for row in table.find_all("tr")[1:]:
                cells = [td.get_text(strip=True) for td in row.find_all("td")]
                if len(cells) >= 2 and cells[0]:
                    instruments.append({"instrument": cells[0], "rating": cells[1]})
            break

    return {"instruments": instruments, "url": url}


def _lookup_icra_rating(company_name: str) -> dict:
    try:
        matches = _icra_search_companies(company_name)
    except (requests.RequestException, ValueError) as e:
        return {"found": False, "note": f"ICRA lookup could not run this pass: {e}"}

    exact = next(
        (m for m in matches if str(m.get("label", "")).strip().lower() == company_name.strip().lower()),
        None,
    )
    if not exact:
        return {"found": False, "note": "No public ICRA rating found for this exact legal entity name."}

    try:
        detail = _icra_fetch_rating_detail(exact["id"], exact["label"])
    except requests.RequestException as e:
        return {"found": False, "note": f"Found a matching ICRA entity ({exact['label']}) but could not fetch its rating detail this pass: {e}"}

    return {
        "found": True,
        "agency": "ICRA",
        "company_name": exact["label"],
        "instruments": detail["instruments"],
        "url": detail["url"],
    }


# ---------------------------------------------------------------------------
# Infomerics credit-rating lookup -- confirmed live against Infomerics'
# real public rating database (infomerics.com), added specifically because
# ICRA-only coverage produced a misleading "not found" for a real MahaRERA
# promoter's sister entity (Pranami Estates Private Limited) that turned
# out to be actively rated by Infomerics, not ICRA -- a materially
# different result (a live downgrade trajectory: BBB- in 2021 down to
# BB (INC), Negative outlook, by March 2026) that ICRA-only checking would
# have silently missed and reported as "no rating found."
#
# Mechanism: the site's header search box (id="search") is a plain
# client-side filter over a public autocomplete API --
# cms.infomerics.com/api/companies/autocomplete?query=<name> -- returning
# [{CompanyName, slug, documentId}, ...]; the matched `slug` then fetches
# cms.infomerics.com/api/companies/<slug> for the CURRENT instrument
# rating(s) (instrument, amount, rating, outlook, date). Both are plain,
# unauthenticated JSON GETs -- no CAPTCHA, no session, confirmed with a
# bare requests.get() (no custom headers needed, unlike ICRA's
# X-Requested-With quirk). Only the CURRENT rating is fetched here, not
# the full historical rationale table also shown on the site's own
# press-release page -- that would need a second, not-yet-identified
# endpoint, and current-only already matches ICRA's own scope above.
# ---------------------------------------------------------------------------

_INFOMERICS_AUTOCOMPLETE_URL = "https://cms.infomerics.com/api/companies/autocomplete"
_INFOMERICS_COMPANY_URL = "https://cms.infomerics.com/api/companies/{slug}"


def _infomerics_search_companies(name: str) -> list:
    resp = requests.get(_INFOMERICS_AUTOCOMPLETE_URL, params={"query": name}, timeout=config.REQUEST_TIMEOUT)
    resp.raise_for_status()
    return resp.json()


def _infomerics_fetch_rating_detail(slug: str) -> dict:
    url = _INFOMERICS_COMPANY_URL.format(slug=slug)
    resp = requests.get(url, timeout=config.REQUEST_TIMEOUT)
    resp.raise_for_status()
    data = resp.json()

    instruments = []
    for inst in data.get("companyInstrument", []):
        if inst.get("isPast"):
            continue
        rating = inst.get("Rating", "")
        outlook = (inst.get("outlook") or {}).get("Title")
        if outlook:
            rating = f"{rating} (Outlook: {outlook})"
        instruments.append({
            "instrument": f"{inst.get('InstrumentTitle', '')} ({inst.get('InstrumentAmount', '')})".strip(),
            "rating": rating,
        })

    return {"instruments": instruments, "url": f"https://www.infomerics.com/pressrelease/{slug}"}


def _lookup_infomerics_rating(company_name: str) -> dict:
    try:
        matches = _infomerics_search_companies(company_name)
    except (requests.RequestException, ValueError) as e:
        return {"found": False, "note": f"Infomerics lookup could not run this pass: {e}"}

    exact = next(
        (m for m in matches if str(m.get("CompanyName", "")).strip().lower() == company_name.strip().lower()),
        None,
    )
    if not exact:
        return {"found": False, "note": "No public Infomerics rating found for this exact legal entity name."}

    try:
        detail = _infomerics_fetch_rating_detail(exact["slug"])
    except requests.RequestException as e:
        return {"found": False, "note": f"Found a matching Infomerics entity ({exact['CompanyName']}) but could not fetch its rating detail this pass: {e}"}

    return {
        "found": True,
        "agency": "Infomerics",
        "company_name": exact["CompanyName"],
        "instruments": detail["instruments"],
        "url": detail["url"],
    }


_CREDIT_RATING_AGENCIES = (
    ("ICRA", _lookup_icra_rating),
    ("Infomerics", _lookup_infomerics_rating),
)


def lookup_credit_rating(company_name: str) -> dict:
    """Checks EVERY agency above (currently ICRA and Infomerics) for a
    public rating under an EXACT (case-insensitive) match on
    `company_name` -- never a fuzzy "probably the same company" guess,
    since misattributing a rating to the wrong legal entity would be a
    serious factual error in a due-diligence document. Always checks all
    agencies, even after an earlier one already found something -- so that
    if two agencies rate the same entity, BOTH ratings surface side by
    side for comparison, rather than silently showing only the first
    match. (This costs one extra request per additional agency versus a
    first-match-wins design, which is negligible next to the value of
    catching a real disagreement between agencies.) Returns:
      {"found": True, "ratings": [{"agency": "ICRA" | "Infomerics",
       "company_name": <matched label>, "instruments": [{"instrument":
       ..., "rating": ...}, ...], "url": ...}, ...], "not_found_agencies":
       [<agency names with no match>]}
        -- `ratings` holds one entry per agency that found something (in
        agency-check order); `not_found_agencies` names the rest so a
        reader knows they were checked, not skipped.
      {"found": False, "note": "..."} -- an honest explanation naming
      every agency actually checked, not an error, when nothing matches
      anywhere or every request failed. A promoter/SPV having no public
      rating from any agency is the NORMAL case, not a red flag: these
      agencies only rate developers that sought a public rating (typically
      larger, listed, or NCD-issuing entities); most MahaRERA promoters
      are too small or private to ever be rated. Callers who also want to
      check a distinct parent/group entity (e.g. "Godrej Properties
      Limited" for promoter "Godrej Skyline Developers Limited") must call
      this a second time with that name explicitly and label the result as
      the PARENT's rating, never as if it were the promoter's own -- this
      function does not guess a parent itself."""
    ratings = []
    not_found_agencies = []
    for agency_name, lookup_fn in _CREDIT_RATING_AGENCIES:
        result = lookup_fn(company_name)
        if result.get("found"):
            ratings.append(result)
        else:
            not_found_agencies.append(agency_name)

    if ratings:
        return {"found": True, "ratings": ratings, "not_found_agencies": not_found_agencies}

    agency_list = ", ".join(name for name, _ in _CREDIT_RATING_AGENCIES)
    return {
        "found": False,
        "note": (
            f"No public rating found for this exact legal entity name from any agency checked "
            f"({agency_list}). This is NOT itself a red flag -- these agencies only rate developers "
            f"that sought a public rating (typically larger, listed, or NCD-issuing entities); most "
            f"MahaRERA promoters are too small or private to ever be rated."
        ),
    }


# ---------------------------------------------------------------------------
# IBBI insolvency check -- confirmed live against IBBI's real Corporate
# Debtor Master Data. The site's own search form (/claims/corporate-personals)
# is CSRF-protected, but submitting it (by CIN) redirects to a plain,
# directly-linkable detail URL -- https://ibbi.gov.in/claims/inner-process/
# <CIN> -- which a fresh, cookie-less requests.get() reaches identically
# (confirmed: no CSRF token or session state needed for this specific URL).
# Also confirmed to accept an LLPIN in the same URL slot (e.g. AAI-5299 for
# Trimity Realty LLP) -- LLPs are "corporate persons" under the IBC too, and
# the identifier isn't validated by format, just passed through.
#
# IMPORTANT CORRECTION (found while testing the LLPIN path): the "no
# process" result is WEAKER evidence than earlier phrasing here implied.
# Feeding this endpoint a deliberately fake identifier (e.g. "ZZZ-0000")
# returns the exact same "ASSIGNMENT NOT APPROVED YET" page as a real,
# genuine CIN/LLPIN with no insolvency history -- the page does not
# validate or discriminate its input at all. So "found_process: False"
# confirms only that IBBI has no ACTIVE PROCESS PAGE at that URL; it is NOT
# proof the identifier itself was recognized as a real, existing entity.
# That confirmation has to come from elsewhere (ZaubaCorp/registry mirror),
# same as always. No real company/LLP with an ACTIVE/PAST insolvency
# process was available to test the positive-match case against, so that
# path still surfaces the raw extracted status text for a human to read
# rather than attempting to classify or summarize content this function was
# never validated against.
# ---------------------------------------------------------------------------

_IBBI_INNER_PROCESS_URL = "https://ibbi.gov.in/claims/inner-process/{cin}"
_IBBI_NO_PROCESS_PHRASE = "ASSIGNMENT NOT APPROVED YET"
_CIN_RE = re.compile(r"\b[UL]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b")
_LLPIN_RE = re.compile(r"\b[A-Z]{3}-\d{4}\b")


def extract_cin(text: str) -> str | None:
    """Pulls a standard-format Indian CIN (e.g. U45309MH2016PLC287858) out
    of a free-text field -- corporate_identity.cin_llpin's value is prose
    ("U45309MH2016PLC287858 (incorporated 22 November 2016; ROC Mumbai) --
    ..."), not a bare code, so this is needed before the IBBI/CIN-keyed
    lookups below can use it."""
    match = _CIN_RE.search(text or "")
    return match.group(0) if match else None


def extract_llpin(text: str) -> str | None:
    """Pulls a standard-format Indian LLPIN (e.g. AAI-5299) out of a
    free-text field, the LLP counterpart of extract_cin() above. A
    completely different format from CIN (3 letters, a hyphen, 4 digits,
    vs. CIN's 21-character U/L-prefixed code) -- confirmed live against
    Trimity Realty LLP's real LLPIN (AAI-5299), and confirmed NOT to
    false-positive-match on real CIN text (extract_cin's own format never
    contains a bare "XXX-9999" substring)."""
    match = _LLPIN_RE.search(text or "")
    return match.group(0) if match else None


def lookup_ibbi_insolvency_status(cin: str) -> dict:
    """Checks IBBI's public Corporate Debtor Master Data for an insolvency
    process tied to `cin` (a CIN or an LLPIN -- confirmed live against
    both). Returns:
      {"found_process": False, "status_text": "ASSIGNMENT NOT APPROVED YET", "url": ...}
        -- the ordinary, clean case: no IBC process recorded against this
        identifier. NOTE: this same result is also what IBBI returns for a
        completely fake/nonexistent identifier (confirmed live) -- it is
        NOT proof the identifier itself is real, only that no active
        process page exists at that URL. Confirm the entity's existence
        via a registry mirror (ZaubaCorp) separately, not via this result.
      {"found_process": True, "status_text": <raw extracted text>, "url": ...}
        -- something other than the known "no process" phrase was found;
        the raw text is surfaced rather than classified, since this
        function's positive-match path was never validated against a real
        example (see module note above) -- a human should read it directly.
      {"found_process": None, "note": "..."} -- the lookup itself failed,
        or no identifier was available to check."""
    if not cin or not cin.strip():
        return {"found_process": None, "note": "no CIN provided to check"}
    url = _IBBI_INNER_PROCESS_URL.format(cin=cin.strip())
    try:
        resp = requests.get(url, timeout=config.REQUEST_TIMEOUT)
        resp.raise_for_status()
    except requests.RequestException as e:
        return {"found_process": None, "note": f"IBBI lookup could not run this pass: {e}", "url": url}

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(resp.text, "html.parser")
    text = soup.get_text("\n", strip=True)

    if _IBBI_NO_PROCESS_PHRASE in text:
        return {"found_process": False, "status_text": _IBBI_NO_PROCESS_PHRASE, "url": url}

    lines = [ln for ln in text.split("\n") if ln.strip()]
    cin_idx = next((i for i, ln in enumerate(lines) if "CIN No" in ln), None)
    context = lines[cin_idx:cin_idx + 10] if cin_idx is not None else lines[:10]
    return {"found_process": True, "status_text": " | ".join(context), "url": url}


# ---------------------------------------------------------------------------
# CIN -> company profile / director-group lookup (ZaubaCorp). Confirmed live
# and scriptable with no browser needed: ZaubaCorp's company page URL only
# needs a correct CIN in the second path segment -- the company-name segment
# is cosmetic and can be any placeholder ("X"), and a real CIN 200s and
# redirects to the true canonical URL (e.g. /company/X/U45500MH2016PTC286108
# -> /WATERTIGHT-DEVELOPERS-PRIVATE-LIMITED-U45500MH2016PTC286108). A CIN
# with no matching company instead redirects to /companysearchresults/X,
# which has zero `li.row` detail elements -- a clean, reliable not-found
# signal. Live-tested against two real, unrelated CINs (Watertight
# Developers Private Limited, a private SPV; India Homes Limited, a listed
# public company formerly India Steel Works Limited) and a deliberately
# fake CIN, confirming both the found and not-found paths.
#
# The page's core details sit in `<li class="row"><span>label</span>
# <label>value</label></li>` pairs -- not a <table> -- which is why earlier
# passes over this page's `<table>` elements alone missed them. Director
# and group-affiliation data IS in tables, each preceded by its own heading:
#   - "Current/Past Directors & Key Managerial Personnel of X" -- this
#     company's own director roster (DIN, name, designation, dates).
#   - "Other Directorships of <NAME>" -- one table per director named in
#     the roster above, listing every OTHER company that same person is or
#     was a director of. This is the group/board-overlap signal, and it
#     arrives on the same page fetch -- no per-director follow-up request
#     is needed.
#   - "Companies with Similar Address" -- a second, independent
#     group-affiliation signal (shared registered office, not shared
#     directors); confirmed several Gupta-family-linked entities share
#     Watertight's exact registered address.
#   - "Subsidiaries, Associate Companies & Joint Ventures of X" -- present
#     for India Homes Limited (a listed company that files this), absent
#     for Watertight (a private SPV that doesn't); a third signal, this one
#     naming an actual percentage-of-shares-held figure.
#
# Shareholding is deliberately NOT surfaced as a precise ownership
# breakdown: per-shareholder/promoter percentage holdings are gated behind
# ZaubaCorp's paid report for every CIN tested here (both the private SPV
# and the listed company) -- only aggregate authorised/paid-up capital and
# (for companies that file it) subsidiary/associate stakes are free. Callers
# needing an actual cap table should treat this the same way the ICRA/IBBI
# checks treat their own gaps: report what's genuinely available, flag the
# rest as not publicly obtainable here rather than guessing.
# ---------------------------------------------------------------------------

_ZAUBACORP_HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"}


def _zaubacorp_fetch(cin: str):
    """Returns (soup, canonical_url) for a real CIN, or None if ZaubaCorp
    has no company matching it. Raises requests.RequestException on a
    network/HTTP failure -- callers convert that into an honest gap."""
    from bs4 import BeautifulSoup

    url = f"https://www.zaubacorp.com/company/X/{cin.strip()}"
    resp = requests.get(url, headers=_ZAUBACORP_HEADERS, timeout=config.REQUEST_TIMEOUT)
    resp.raise_for_status()
    if "companysearchresults" in resp.url:
        return None
    soup = BeautifulSoup(resp.text, "html.parser")
    if not soup.find("li", class_="row"):
        return None
    return soup, resp.url


def _zaubacorp_core_fields(soup) -> dict:
    fields = {}
    for li in soup.find_all("li", class_="row"):
        parts = li.find_all(["span", "label"])
        if len(parts) >= 2:
            fields[parts[0].get_text(strip=True)] = parts[1].get_text(strip=True)
    return fields


def _zaubacorp_director_table(table) -> list:
    rows = []
    header_cells = [th.get_text(strip=True) for th in table.find_all("th")]
    for tr in table.find_all("tr")[1:]:
        cells = [td.get_text(strip=True) for td in tr.find_all("td")]
        if len(cells) == len(header_cells) and any(cells):
            rows.append(dict(zip(header_cells, cells)))
    return rows


_ZAUBACORP_GATED_NOTE = "This information is part of the paid company report.Purchase Report"

# Free-tier mirrors gate some cells behind a paid report and put their upsell
# copy in the cell where the VALUE should be, so a naive scrape reports the
# advert as the data. Confirmed live on a real company: the "Percentage of
# Shares Held" column rendered into the Charter as "subsidiary/associate/JV
# (This information is part of the paid company report.Purchase Report shares
# held)", and on that same table the company NAME cell was gated too --
# _ZAUBACORP_GATED_NOTE's exact-match check above never caught either,
# because the placeholder arrives concatenated with surrounding text.
# Matched case-insensitively on substrings, since the exact wording varies by
# site and by page.
_PAYWALL_MARKERS = (
    "paid company report", "purchase report", "part of the paid",
    "subscribe", "upgrade to", "login to view", "sign up to view",
    "buy full report", "available in the paid",
)


def _looks_paywalled(value: str) -> bool:
    """True when a scraped cell holds a site's paywall upsell copy rather
    than a real value -- such a cell must be treated as "not disclosed",
    never rendered as if it were data."""
    lowered = str(value or "").lower()
    return any(marker in lowered for marker in _PAYWALL_MARKERS)


def _zaubacorp_clean(value: str | None) -> str | None:
    """None instead of ZaubaCorp's free-tier paywall placeholder, so
    callers don't mistake "gated" for a genuine reported value. Delegates to
    _looks_paywalled rather than exact-matching _ZAUBACORP_GATED_NOTE, so
    there is exactly ONE definition of "this cell is an advert, not data" in
    this file -- two competing checks would drift apart, and the exact-match
    version already missed the group-companies table where the same
    placeholder appears in the company-NAME cell (see
    find_group_companies_by_cin)."""
    if value is None or _looks_paywalled(value):
        return None
    return value


def lookup_company_by_cin(cin: str) -> dict:
    """Looks up a company's own registration profile and director roster
    from ZaubaCorp by CIN. Returns:
      {"found": True, "name": ..., "status": ..., "class_of_company": ...,
       "incorporation_date": ..., "roc": ..., "registered_address": ...,
       "authorised_capital": ..., "paid_up_capital": ...,
       "current_directors": [{"din", "name", "designation", "appointment_date"}, ...],
       "past_directors": [...], "shareholding_note": "...", "url": ...}
    or {"found": False, "note": "..."} -- either a genuine no-match or a
    request failure, never guessed."""
    if not cin or not cin.strip():
        return {"found": False, "note": "no CIN provided to look up"}
    try:
        result = _zaubacorp_fetch(cin)
    except requests.RequestException as e:
        return {"found": False, "note": f"ZaubaCorp lookup could not run this pass: {e}"}
    if result is None:
        return {"found": False, "note": f"no ZaubaCorp record found for CIN {cin.strip()}"}
    soup, url = result

    fields = _zaubacorp_core_fields(soup)
    current_directors, past_directors = [], []
    for table in soup.find_all("table"):
        heading = table.find_previous(["h1", "h2", "h3", "h4", "h5"])
        heading_text = heading.get_text(strip=True) if heading else ""
        if heading_text.startswith("Current Directors"):
            current_directors = _zaubacorp_director_table(table)
        elif heading_text.startswith("Past Directors"):
            past_directors = _zaubacorp_director_table(table)

    return {
        "found": True,
        "name": fields.get("Name"),
        "cin": fields.get("CIN", cin.strip()),
        "status": fields.get("Company Status"),
        "class_of_company": fields.get("Class of Company"),
        "company_category": fields.get("Company Category"),
        "roc": fields.get("ROC"),
        "incorporation_date": fields.get("Date of Incorporation"),
        "registered_address": fields.get("Address"),
        "authorised_capital": _zaubacorp_clean(fields.get("Authorised Share Capital")),
        "paid_up_capital": _zaubacorp_clean(fields.get("Paid-up Share Capital")),
        "current_directors": current_directors,
        "past_directors": past_directors,
        "shareholding_note": (
            "Per-shareholder/promoter shareholding percentages are gated behind ZaubaCorp's paid "
            "report and were not available on this pass; only aggregate authorised/paid-up capital "
            "above is free. This is not itself a red flag for a private company -- detailed cap "
            "tables are rarely public for unlisted entities."
        ),
        "url": url,
    }


def find_group_companies_by_cin(cin: str) -> dict:
    """Cross-references directors and registered address of the company at
    `cin` against every OTHER company ZaubaCorp lists them against, to
    surface a likely corporate group -- without guessing: every entry
    names the concrete signal (a specific shared director, or a shared
    registered office) that ties it to the target, so a reader can judge
    the strength of the link themselves. Returns:
      {"found": True, "companies": [{"cin", "name", "basis": [...]}, ...],
       "url": ..., "sources_used": [...], "corroboration": {...}}
    or {"found": False, "note": "..."}.

    Single-sourced to ZaubaCorp by necessity, NOT by preference -- unlike
    the company profile (see _MCA_PROFILE_CHAIN), which queries three
    mirrors and merges them. Checked live against both alternatives for a
    real CIN: InstaFinancials publishes no subsidiary/associate/JV or
    shared-address structure at all (its only "associate" text is prose
    about directors), and Tofler's page carries a long company-name list
    that is a generic industry/directory listing, not verified related
    parties (it contained companies with no plausible connection to the
    subject at all) -- scraping it as a related-party set would invent
    links rather than find them. So the chain here corroborates rather
    than merges: see _instafinancials_directorship_count, which
    independently cross-checks the director-overlap COUNT."""
    if not cin or not cin.strip():
        return {"found": False, "note": "no CIN provided to look up"}
    try:
        result = _zaubacorp_fetch(cin)
    except requests.RequestException as e:
        return {"found": False, "note": f"ZaubaCorp lookup could not run this pass: {e}"}
    if result is None:
        return {"found": False, "note": f"no ZaubaCorp record found for CIN {cin.strip()}"}
    soup, url = result

    by_cin: dict[str, dict] = {}
    # Rows whose IDENTITY (not merely some detail) is paywalled -- see
    # _looks_paywalled. Counted, never listed: the relationship's existence
    # is a real finding, but we do not know who the counterparty is, so
    # inventing a row for it would be worse than reporting the count.
    undisclosed_counts: dict[str, int] = {}

    def _add(other_cin: str, name: str, basis: str, relationship: str = "related"):
        other_cin = (other_cin or "").strip()
        if _looks_paywalled(name) or _looks_paywalled(other_cin):
            undisclosed_counts[relationship] = undisclosed_counts.get(relationship, 0) + 1
            return
        name = _normalise_entity_name(name)
        if not other_cin or not name or other_cin == cin.strip():
            return
        entry = by_cin.setdefault(other_cin, {"cin": other_cin, "name": name, "basis": []})
        entry["basis"].append(basis)

    # Raw (not-yet-formatted) per-(entity, director) state, so a director who
    # appears under TWO of ZaubaCorp's own "Other Directorships of X" rows for
    # the same entity (confirmed live for Mahir Haresh Wadhwani / Suchi
    # Lifespaces LLP -- listed once as "Designated Partner" ceased 2023-04-01
    # and again as "Partner" ceased 2023-06-14) merges into ONE basis entry
    # instead of inflating the "independent links" count or silently dropping
    # the newer role/date. Keyed by (other_cin, director name).
    director_state: dict[tuple[str, str], dict] = {}

    def _format_director_status(designation: str, cessation: str) -> str:
        is_ongoing = not cessation or cessation.strip().lower() == "ongoing"
        status = "ongoing" if is_ongoing else f"resigned {cessation.strip()}"
        return f"{designation}, {status}" if designation else status

    def _add_shared_director(other_cin: str, name: str, director: str, designation: str, cessation: str):
        other_cin = (other_cin or "").strip()
        if _looks_paywalled(name) or _looks_paywalled(other_cin):
            undisclosed_counts["related"] = undisclosed_counts.get("related", 0) + 1
            return
        name = _normalise_entity_name(name)
        if not other_cin or not name or other_cin == cin.strip():
            return

        key = (other_cin, director)
        is_ongoing = not cessation or cessation.strip().lower() == "ongoing"
        existing = director_state.get(key)
        if existing is not None:
            # Prefer whichever occurrence shows the relationship is STILL
            # ONGOING (a person can hold two role-labels for the same entity
            # over time); if both have resigned, keep the LATER cessation
            # date -- that is when they actually finished leaving.
            keep_new = is_ongoing or (not existing["is_ongoing"] and cessation > existing["cessation"])
            if not keep_new:
                return
        director_state[key] = {"cessation": cessation, "is_ongoing": is_ongoing}

        status_text = _format_director_status(designation, cessation)
        basis = f"shared director: {director} ({status_text})"
        entry = by_cin.setdefault(other_cin, {"cin": other_cin, "name": name, "basis": []})
        # Replace rather than append when this key already produced a basis
        # string in a prior call (the merge case above).
        entry["basis"] = [
            b for b in entry["basis"]
            if not (b.startswith("shared director:") and b.split(":", 1)[1].split("(")[0].strip() == director)
        ]
        entry["basis"].append(basis)

    for table in soup.find_all("table"):
        heading = table.find_previous(["h1", "h2", "h3", "h4", "h5"])
        heading_text = heading.get_text(strip=True) if heading else ""

        if heading_text.startswith("Other Directorships of "):
            director = _normalise_entity_name(heading_text[len("Other Directorships of "):])
            for row in _zaubacorp_director_table(table):
                company_name = (row.get("Company Name") or "").strip()
                if not company_name:
                    # A blank row with no company name or CIN -- confirmed
                    # live on ZaubaCorp's own page for a director with a
                    # large directorship count; noise, not a real entity.
                    continue
                designation = row.get("Designation", "").strip()
                cessation = row.get("Cessation", "").strip()
                _add_shared_director(row.get("CIN"), company_name, director, designation, cessation)

        elif heading_text == "Companies with Similar Address":
            for row in _zaubacorp_director_table(table):
                _add(row.get("CIN"), row.get("Company Name"), "shared registered office")

        elif heading_text.startswith("Subsidiaries, Associate Companies"):
            for row in _zaubacorp_director_table(table):
                pct = row.get("Percentage of Shares Held", "").strip()
                if _looks_paywalled(pct):
                    pct = ""  # the site's upsell copy, not a shareholding figure
                basis = "subsidiary/associate/JV" + (
                    f" ({pct} shares held)" if pct else " (shareholding not disclosed on the free tier)"
                )
                _add(row.get("Company Identifier"), row.get("Name"), basis, relationship="subsidiary/associate/JV")

    companies = list(by_cin.values())
    return {
        "found": True,
        "companies": companies,
        "url": url,
        "sources_used": ["zaubacorp.com"],
        # Confirmed live: ZaubaCorp paywalls the Subsidiaries/Associates table
        # in full -- the company NAME cell, not just the shareholding
        # percentage, carries its upsell copy. Those rows are therefore
        # counted here rather than listed as named entities; reporting "5
        # subsidiary/associate/JV relationships exist but their identities are
        # paywalled" is honest, whereas listing 5 rows named after an advert
        # is not.
        "undisclosed_relationship_counts": undisclosed_counts,
        "corroboration": _corroborate_group_companies(cin, companies),
    }


def _instafinancials_directorship_count(cin: str) -> int | None:
    """InstaFinancials states, in prose on its directors page, how many
    directorships this company's directors hold elsewhere ("The directors in
    this company holds a count of 26 directorships in companies, other than
    X") -- confirmed live. That's an INDEPENDENT count of the same
    director-overlap ZaubaCorp's "Other Directorships" tables enumerate, so
    it's usable as a cross-check even though InstaFinancials never names the
    companies. Returns None if the sentence isn't present or the lookup
    fails; never raises."""
    try:
        result = _instafinancials_fetch(cin)
    except Exception:
        return None
    if result is None:
        return None
    _, directors_soup, _, _ = result
    match = re.search(
        r"count of\s+(\d+)\s+directorships",
        directors_soup.get_text(" ", strip=True),
        re.I,
    )
    return int(match.group(1)) if match else None


def _corroborate_group_companies(cin: str, companies: list) -> dict:
    """Cross-checks ZaubaCorp's named director-overlap set against
    InstaFinancials' own independent directorship count (see
    _instafinancials_directorship_count). Returns {"independent_directorship
    _count", "zaubacorp_shared_director_entities", "agrees", "note"} --
    "agrees" is None (not False) when the second source couldn't be reached,
    since "no cross-check ran" is a different finding from "the two sources
    disagree", the same distinction _verify_one_field already draws."""
    shared_director_entities = sum(
        1 for c in companies if any(str(b).startswith("shared director") for b in c.get("basis", []))
    )
    independent_count = _instafinancials_directorship_count(cin)
    if independent_count is None:
        return {
            "independent_directorship_count": None,
            "zaubacorp_shared_director_entities": shared_director_entities,
            "agrees": None,
            "note": "InstaFinancials' independent directorship count could not be read this pass -- "
                    "ZaubaCorp's director-overlap set is uncorroborated, not contradicted.",
        }

    # These count related-but-different things: ZaubaCorp's figure is
    # distinct ENTITIES sharing at least one director, InstaFinancials' is
    # total DIRECTORSHIPS held elsewhere (one director on three boards is 3
    # directorships but up to 3 entities). They should be the same order of
    # magnitude; a large gap means one source is materially incomplete.
    larger = max(independent_count, shared_director_entities)
    agrees = larger == 0 or abs(independent_count - shared_director_entities) <= max(2, round(0.5 * larger))
    return {
        "independent_directorship_count": independent_count,
        "zaubacorp_shared_director_entities": shared_director_entities,
        "agrees": agrees,
        "note": (
            f"InstaFinancials independently reports {independent_count} directorship(s) held elsewhere by this "
            f"company's directors; ZaubaCorp names {shared_director_entities} distinct entit(y/ies) sharing at "
            f"least one director. These count related-but-different things (directorships vs entities), so they "
            f"are expected to be the same order of magnitude rather than identical."
            + ("" if agrees else " The gap is large enough that one source is likely materially incomplete.")
        ),
    }


# ---------------------------------------------------------------------------
# CIN -> company profile, MCA-mirror redundancy chain (ZaubaCorp -> Tofler ->
# InstaFinancials). All three are free-tier scriptable with no login/CAPTCHA
# for the fields below (live-tested against Godrej Skyline Developers
# Limited's real CIN, both its current PLC and superseded PTC variants):
#   - Tofler: CIN resolved via POST /cnamesearch (mode="SCBC", the same web
#     service the site's own JS search bar calls), returning
#     "{internal_id};{NAME};{CIN};{is_previous_name};{state};{city};{status};
#     {type_flag}" plus a ready-made profile "url" field. Unlike ZaubaCorp,
#     Tofler's URL slug is NOT cosmetic -- a placeholder slug 404s -- so it
#     must come from this search response, never guessed. Registered Details
#     (CIN/incorporation/capital) sit in clean <h3>label</h3><span>value
#     </span> pairs inside #registered-details-module; the People table
#     (Designation/Name/DIN-PAN/Tenure) is a normal server-rendered <table>
#     under #people-module; the registered address is populated
#     client-side from an inline `const locationsTableData = [...]` JSON
#     array (the server-rendered table body is empty), extracted here by
#     regex instead of DOM parsing. Tofler's "Network" tab (other-
#     directorships/shared-address crosswalk) is an interactive JS graph
#     widget, not scrapable HTML -- so Tofler contributes to the company-
#     PROFILE chain only, not the group-companies crosswalk below.
#   - InstaFinancials: CIN resolved via POST
#     /ajax-caller.aspx/GetCompanyNames (mode="SCBC"), returning the same
#     kind of semicolon-delimited record (no ready-made URL this time --
#     built from the slugified name + CIN, mirroring the exact JS in
#     setCompanyName()/new-index.js). The profile page's
#     #companyHighlightsDataContainer table and .highlight-card /
#     .capital-bar-item elements are fully server-rendered (CIN, ROC,
#     capital, status, address); its "/company-directors" sub-page has a
#     clean #currentDirectorsArticle table (exact appointment dates, DINs)
#     and a #signatoriesArticle table for KMPs (Company Secretary/CFO/
#     Manager). InstaFinancials' financials are paywalled like ZaubaCorp's,
#     and its own per-director "other directorships" list is login-gated
#     (confirmed: the director page's DirectorshipTable div is empty
#     server-side, behind a "Login to View Full Profile" link) -- so it,
#     too, only contributes to the company-profile chain, not
#     group-companies.
#
# Real-world value of the redundancy, confirmed live against this same CIN:
# the three sources do NOT agree on the current director roster --
# ZaubaCorp lists Aspy Dady Cooper as the third director, Tofler lists
# Abhishek Sahaya only as a KMP (not a director), and InstaFinancials lists
# Abhishek Sahaya as a director appointed 07-10-2025. _run_mca_profile_chain
# below surfaces exactly this kind of cross-source disagreement as an
# unresolved gap rather than silently picking one source's answer.
# ---------------------------------------------------------------------------

_TOFLER_HEADERS = _ZAUBACORP_HEADERS
_INSTAFINANCIALS_HEADERS = _ZAUBACORP_HEADERS


def _tofler_resolve(cin: str, company_name: str):
    """Resolves `cin` to a Tofler profile URL. Confirmed live that Tofler's
    /cnamesearch web service rejects a direct POST with HTTP 401
    "Unauthorized access" -- reproduced identically via plain requests
    (with a real session cookie, matching Referer/Origin/X-Requested-With
    headers, and even the exact non-standard unquoted-key body its own JS
    sends) AND via a raw Playwright page.evaluate(fetch(...)) call, so this
    is not a simple missing-header fix. What DOES work, confirmed live: the
    same endpoint hit via the page's own real keyup-triggered jQuery
    request (i.e. actually typing into the visible #searchbox in a loaded
    browser session) -- so this resolves by NAME through a real headless
    Chromium session (Playwright, already a project dependency -- see
    session_auth.py/mahabhumi.py) and verifies the returned CIN matches
    `cin` exactly before accepting the match, the same "never fuzzy-match,
    always confirm the identifier" principle _zaubacorp_fetch's CIN-based
    URL trick already relies on. Returns {"name", "cin", "status", "url"},
    or None if no company-name search result's CIN matches."""
    if not company_name or not company_name.strip():
        return None

    from playwright.sync_api import sync_playwright

    responses = []

    def _capture(resp):
        if "cnamesearch" in resp.url and resp.status == 200:
            try:
                responses.append(resp.text())
            except Exception:
                pass

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.on("response", _capture)
        page.goto("https://www.tofler.in/", timeout=config.REQUEST_TIMEOUT * 1000)
        page.wait_for_timeout(800)
        page.click("#searchbox")
        page.type("#searchbox", company_name.strip(), delay=40)
        page.wait_for_timeout(1200)
        browser.close()

    for text in reversed(responses):
        try:
            results = json.loads(text)
        except (json.JSONDecodeError, TypeError):
            continue
        for r in results:
            if r.get("value", "").strip().upper() == cin.strip().upper():
                return {
                    "name": r["label"], "cin": r["value"],
                    "status": r.get("status", ""), "url": f"https://www.tofler.in{r['url']}",
                }
    return None


def _tofler_fetch(cin: str, company_name: str):
    """Returns (soup, html_text, canonical_url, resolved_name) for a real
    CIN, or None if Tofler has no matching company. html_text is kept
    alongside soup because the registered address lives in an inline JS
    array, not the DOM (see _tofler_registered_address); resolved_name
    comes from the search match itself (the exact record Tofler filed it
    under), not a page scrape."""
    from bs4 import BeautifulSoup

    match = _tofler_resolve(cin, company_name)
    if match is None:
        return None
    resp = requests.get(match["url"], headers=_TOFLER_HEADERS, timeout=config.REQUEST_TIMEOUT)
    resp.raise_for_status()
    return BeautifulSoup(resp.text, "html.parser"), resp.text, match["url"], match["name"]


def _tofler_core_fields(soup) -> dict:
    fields = {}
    module = soup.find(id="registered-details-module")
    if not module:
        return fields
    for box in module.select(".flex-col.gap-2"):
        label, value = box.find("h3"), box.find("span")
        if label and value:
            fields[label.get_text(strip=True)] = " ".join(value.get_text(strip=True).split())
    type_label = module.find("h3", string="Type")
    if type_label:
        badges_container = type_label.find_next_sibling("div")
        if badges_container:
            badges = [b.get_text(strip=True) for b in badges_container.select(".badge")]
            if badges:
                fields["Type"] = ", ".join(badges)
    return fields


def _tofler_registered_address(html_text: str) -> str | None:
    """Tofler's registered-office address is populated client-side from an
    inline `const locationsTableData = [...]` JSON array -- the static
    server-rendered <table> body is empty, so there's nothing to parse in
    the DOM itself."""
    m = re.search(r"const locationsTableData = (\[.*?\]);", html_text)
    if not m:
        return None
    try:
        locations = json.loads(m.group(1))
    except json.JSONDecodeError:
        return None
    if not locations:
        return None
    for loc in locations:
        if loc.get("type") == "Registered Office" and loc.get("address"):
            return " ".join(loc["address"].split())
    return " ".join((locations[0].get("address") or "").split()) or None


def _tofler_director_table(soup) -> tuple[list, list]:
    """Returns (directors, kmp) from Tofler's People table, normalized to
    the same dict keys ZaubaCorp's own director tables use ("Director
    Name", "Designation", "DIN", "Appointment Date", "Cessation") so
    existing renderers/cross-references written against
    lookup_company_by_cin's shape work unchanged. Tofler only gives a
    relative "Tenure" string (e.g. "3 years"), never an exact date, so
    "Appointment Date"/"Cessation" are left as "-" -- never guessed from
    the tenure string. "Director"-designation rows go to `directors`,
    everything else (Kmp, etc.) to `kmp`."""
    module = soup.find(id="people-module")
    if not module:
        return [], []
    table = module.find("table")
    if not table or not table.find("tbody"):
        return [], []
    directors, kmp = [], []
    for tr in table.find("tbody").find_all("tr"):
        cells = [td.get_text(strip=True) for td in tr.find_all("td")]
        if len(cells) < 4:
            continue
        designation, name, din, tenure = cells[0], cells[1], cells[2], cells[3]
        row = {
            "Director Name": name, "Designation": designation,
            "DIN": din if din and din != "<HIDDEN>" else "",
            "Appointment Date": "-", "Cessation": "-", "Tenure": tenure,
        }
        (directors if designation.lower() == "director" else kmp).append(row)
    return directors, kmp


def lookup_company_tofler(cin: str, company_name: str = "") -> dict:
    """Same return shape as lookup_company_by_cin (plus an extra "kmp" key
    for Tofler's non-director People rows), sourced from Tofler instead of
    ZaubaCorp -- see the module note above this section. "name"/"status"/
    "company_category"/"roc" are left None here (not exposed in the fields
    this scrape covers) rather than guessed; _run_mca_profile_chain fills
    them from whichever other source in the chain has them. Unlike
    ZaubaCorp/InstaFinancials, Tofler can only be resolved by NAME (see
    _tofler_resolve's own note on why its CIN-search endpoint can't be
    called directly) -- `company_name` is required; without it this is
    reported as an honest gap rather than attempted blind."""
    if not cin or not cin.strip():
        return {"found": False, "note": "no CIN provided to look up"}
    if not company_name or not company_name.strip():
        return {"found": False, "note": "no company name provided -- Tofler can only be resolved by name, see _tofler_resolve"}
    try:
        result = _tofler_fetch(cin, company_name)
    except requests.RequestException as e:
        return {"found": False, "note": f"Tofler lookup could not run this pass: {e}"}
    if result is None:
        return {"found": False, "note": f"no Tofler record found for CIN {cin.strip()} under name \"{company_name.strip()}\""}
    soup, html_text, url, resolved_name = result
    fields = _tofler_core_fields(soup)
    current_directors, kmp = _tofler_director_table(soup)

    incorporation = fields.get("Incorporation", "")
    incorporation_year = incorporation.split(",")[0].strip() or None if incorporation else None

    return {
        "found": True,
        "name": resolved_name,
        "cin": fields.get("CIN", cin.strip()),
        "status": None,
        "class_of_company": fields.get("Type"),
        "company_category": None,
        "roc": None,
        "incorporation_date": incorporation_year,
        "registered_address": _tofler_registered_address(html_text),
        "authorised_capital": fields.get("Authorised Capital"),
        "paid_up_capital": fields.get("Paid up Capital"),
        "current_directors": current_directors,
        "past_directors": [],
        "kmp": kmp,
        "shareholding_note": None,
        "url": url,
    }


def _instafinancials_resolve(cin: str):
    """POSTs InstaFinancials' own ASP.NET web service for `cin` (mode
    "SCBC"). Returns {"name", "cin", "status", "url"} for a real CIN, or
    None if InstaFinancials has no match. The profile URL isn't returned by
    the API (unlike Tofler) -- it's built here from the slugified name +
    CIN, replicating setCompanyName()'s exact JS in new-index.js."""
    resp = requests.post(
        "https://www.instafinancials.com/ajax-caller.aspx/GetCompanyNames",
        json={"strSearch": cin.strip(), "mode": "SCBC"},
        headers=_INSTAFINANCIALS_HEADERS, timeout=config.REQUEST_TIMEOUT,
    )
    resp.raise_for_status()
    records = resp.json().get("d", [])
    if not records:
        return None
    parts = records[0].split(";")
    if len(parts) < 7:
        return None
    name, found_cin, status = parts[1], parts[2], parts[6]
    slug = re.sub(r"[^0-9a-zA-Z-]", "-", name.lower().replace("'", "")).replace("--", "-", 1)
    return {"name": name, "cin": found_cin, "status": status, "url": f"https://www.instafinancials.com/company/{slug}-{found_cin}"}


def _instafinancials_fetch(cin: str):
    """Returns (profile_soup, directors_soup, canonical_url, resolved_name)
    for a real CIN, or None if InstaFinancials has no matching company."""
    from bs4 import BeautifulSoup

    match = _instafinancials_resolve(cin)
    if match is None:
        return None
    profile_resp = requests.get(match["url"], headers=_INSTAFINANCIALS_HEADERS, timeout=config.REQUEST_TIMEOUT)
    profile_resp.raise_for_status()
    directors_resp = requests.get(f"{match['url']}/company-directors", headers=_INSTAFINANCIALS_HEADERS, timeout=config.REQUEST_TIMEOUT)
    directors_resp.raise_for_status()
    return (
        BeautifulSoup(profile_resp.text, "html.parser"),
        BeautifulSoup(directors_resp.text, "html.parser"),
        match["url"], match["name"],
    )


def _instafinancials_core_fields(soup) -> dict:
    fields = {}
    container = soup.find(id="companyHighlightsDataContainer")
    table = container.find("table") if container else None
    if table:
        for tr in table.find_all("tr"):
            cells = tr.find_all("td")
            for i in range(0, len(cells) - 1, 2):
                label = cells[i].get_text(strip=True)
                if label:
                    fields[label] = cells[i + 1].get_text(strip=True)

    for card in soup.select(".highlight-card"):
        h3 = card.find("h3")
        if not h3:
            continue
        value_el = card.select_one(".status") or card.select_one(".value")
        if value_el:
            fields[h3.get_text(strip=True)] = " ".join(value_el.get_text(strip=True).split())

    for item in soup.select(".capital-bar-item"):
        title = item.select_one(".bar-title")
        fill = item.select_one(".bar-fill")
        if not title or not fill:
            continue
        m = re.match(r"([A-Za-z ]+)", title.get_text(strip=True))
        if m:
            fields[m.group(1).strip()] = fill.get_text(strip=True)
    return fields


def _instafinancials_table_rows(soup, article_id: str) -> list:
    """Generic parser for InstaFinancials' Current Directors / Signatory
    Details tables -- both are plain <table>s under a named <article>, with
    one hidden (style='display:none') trailing <th>/<td> pair per row that
    must be dropped from both sides before zipping, or headers and cells
    misalign by one column."""
    article = soup.find(id=article_id)
    table = article.find("table") if article else None
    if not table or not table.find("thead") or not table.find("tbody"):
        return []

    def _visible(cells):
        return [c for c in cells if "display:none" not in (c.get("style") or "").replace(" ", "")]

    headers = [th.get_text(strip=True) for th in _visible(table.find("thead").find_all("th"))]
    rows = []
    for tr in table.find("tbody").find_all("tr"):
        cells = [td.get_text(strip=True) for td in _visible(tr.find_all("td"))]
        if len(cells) >= len(headers):
            rows.append(dict(zip(headers, cells)))
    return rows


def _instafinancials_directors(soup) -> tuple[list, list]:
    """Returns (directors, kmp) from InstaFinancials' Current Directors +
    Signatory Details tables, normalized to lookup_company_by_cin's dict
    keys. Only currently-serving people are exposed by these two tables (no
    separate past/resigned-directors table was found on this profile) --
    past_directors is left for the caller to report as [] honestly rather
    than guessed."""
    directors = [
        {
            "Director Name": r.get("Director Name", ""), "Designation": r.get("Designation", "Director"),
            "DIN": r.get("Director DIN", ""), "Appointment Date": r.get("Appointment Date", "-"), "Cessation": "-",
        }
        for r in _instafinancials_table_rows(soup, "currentDirectorsArticle")
    ]
    kmp = [
        {
            "Director Name": r.get("Signatory Name", ""), "Designation": r.get("Designation", ""),
            "DIN": "", "Appointment Date": r.get("Appointment Date", "-"), "Cessation": "-",
        }
        for r in _instafinancials_table_rows(soup, "signatoriesArticle")
    ]
    return directors, kmp


def lookup_company_instafinancials(cin: str) -> dict:
    """Same return shape as lookup_company_by_cin (plus an extra "kmp" key
    for the Signatory Details rows), sourced from InstaFinancials instead
    of ZaubaCorp -- see the module note above this section."""
    if not cin or not cin.strip():
        return {"found": False, "note": "no CIN provided to look up"}
    try:
        result = _instafinancials_fetch(cin)
    except requests.RequestException as e:
        return {"found": False, "note": f"InstaFinancials lookup could not run this pass: {e}"}
    if result is None:
        return {"found": False, "note": f"no InstaFinancials record found for CIN {cin.strip()}"}
    profile_soup, directors_soup, url, resolved_name = result
    fields = _instafinancials_core_fields(profile_soup)
    current_directors, kmp = _instafinancials_directors(directors_soup)

    return {
        "found": True,
        "name": resolved_name,
        "cin": fields.get("Company CIN", cin.strip()),
        "status": fields.get("Company Status"),
        "class_of_company": fields.get("Company Class"),
        "company_category": fields.get("Company Category"),
        "roc": None,
        "incorporation_date": fields.get("Incorp. Date"),
        "registered_address": fields.get("Address"),
        "authorised_capital": fields.get("Authorised Capital"),
        "paid_up_capital": fields.get("Paid up Capital"),
        "current_directors": current_directors,
        "past_directors": [],
        "kmp": kmp,
        "shareholding_note": None,
        "url": url,
    }


# Ordered by confidence/coverage: ZaubaCorp is confirmed most complete
# (full current+past director history with exact dates AND shareholding
# note), Tofler second (exact dates for directors it names, but conflates
# some KMPs, and its resolve step costs a headless-browser launch -- see
# _tofler_resolve), InstaFinancials third (exact dates and DINs but no
# past-director history exposed on this profile).
_MCA_PROFILE_CHAIN = [
    ("zaubacorp.com", lambda cin, company_name: lookup_company_by_cin(cin)),
    ("tofler.in", lambda cin, company_name: lookup_company_tofler(cin, company_name)),
    ("instafinancials.com", lambda cin, company_name: lookup_company_instafinancials(cin)),
]

_SCALAR_PROFILE_FIELDS = (
    "name", "status", "class_of_company", "company_category", "roc",
    "incorporation_date", "registered_address", "authorised_capital", "paid_up_capital",
)


def _normalize_for_compare(value: str) -> str:
    return re.sub(r"[^0-9a-z]", "", value.lower()) if value else ""


_MCA_SOURCE_DISPLAY_NAMES = {
    "zaubacorp.com": "ZaubaCorp", "tofler.in": "Tofler", "instafinancials.com": "InstaFinancials",
}


def _merge_director_rosters(per_source_directors: list[tuple[str, list]]) -> tuple[list, list]:
    """Merges each source's `current_directors` list by DIN (falling back
    to normalized name when a source doesn't expose a DIN, e.g. Tofler's
    <HIDDEN> DINs) into one deduped roster, and returns (merged_directors,
    conflict_notes) -- conflict_notes names every director where sources
    disagree on whether they currently hold the role at all (present in
    some sources' CURRENT list, absent from others'), since that's exactly
    the kind of discrepancy live-tested between ZaubaCorp/Tofler/
    InstaFinancials for this promoter (see the module note above).
    conflict_notes use each source's human-readable display name, not its
    raw domain, since these notes render in gaps -- and thus in the
    External document too (see the doc-variant note in _fill_template) --
    where a bare "zaubacorp.com" would be exactly the kind of internal-
    artifact leakage the External variant is meant to avoid."""
    by_key: dict[str, dict] = {}
    order: list[str] = []
    for domain, directors in per_source_directors:
        for d in directors:
            name = " ".join((d.get("Director Name") or "").split())
            if not name:
                continue
            key = d.get("DIN") or _normalize_for_compare(name)
            if key not in by_key:
                stored = dict(d)
                # Normalise only the DISPLAY name (whitespace + ALL-CAPS), so a
                # roster table doesn't render "SUNDEEP  PODDAR" beside
                # title-cased names from another source. `key` above is
                # deliberately computed from the pre-normalisation name, so this
                # can never change which directors are treated as the same
                # person.
                stored["Director Name"] = _normalise_entity_name(name)
                by_key[key] = {"director": stored, "confirmed_by": [domain]}
                order.append(key)
            else:
                entry = by_key[key]
                entry["confirmed_by"].append(domain)
                if not entry["director"].get("DIN") and d.get("DIN"):
                    entry["director"]["DIN"] = d["DIN"]
                if entry["director"].get("Appointment Date") in (None, "-", "") and d.get("Appointment Date") not in (None, "-", ""):
                    entry["director"]["Appointment Date"] = d["Appointment Date"]

    all_domains = {domain for domain, _ in per_source_directors if _}
    merged, conflicts = [], []
    for key in order:
        entry = by_key[key]
        merged.append(entry["director"])
        missing_from = sorted(all_domains - set(entry["confirmed_by"]))
        if missing_from and len(all_domains) > 1:
            confirmed_names = [_MCA_SOURCE_DISPLAY_NAMES.get(d, d) for d in sorted(entry["confirmed_by"])]
            missing_names = [_MCA_SOURCE_DISPLAY_NAMES.get(d, d) for d in missing_from]
            conflicts.append(
                f"Director roster disagreement for \"{entry['director']['Director Name']}\": listed as a "
                f"current director by {', '.join(confirmed_names)}, but not by "
                f"{', '.join(missing_names)}; not reconciled, flagged for manual verification "
                f"rather than assumed correct or incorrect."
            )
    return merged, conflicts


def _run_mca_profile_chain(cin: str, company_name: str = "") -> dict:
    """Walks _MCA_PROFILE_CHAIN, querying every configured source (not
    stopping at the first success) -- conflict detection on the director
    roster is only possible by comparing sources, and ZaubaCorp/
    InstaFinancials cost only one cheap HTTP request each (Tofler costs one
    headless-browser launch -- see _tofler_resolve -- still a single Charter
    generation's worth of overhead, not a paid lookup). Scalar fields
    (status/capital/address/etc.) take the first source in chain order that
    has a non-empty value for that field, since those have never been
    observed to disagree between sources; the director roster is merged
    across every responding source via _merge_director_rosters, which DOES
    surface disagreements (see its own docstring) rather than silently
    trusting whichever source happened to answer first. `company_name` is
    required for Tofler's resolve step (see lookup_company_tofler) --
    without it Tofler is skipped as a gap, ZaubaCorp/InstaFinancials still
    run normally. Returns the same dict shape as lookup_company_by_cin,
    plus "sources_used" and "roster_conflicts"."""
    per_source_results = []
    for domain, fetch_fn in _MCA_PROFILE_CHAIN:
        try:
            result = fetch_fn(cin, company_name)
        except Exception as e:
            result = {"found": False, "note": f"{domain} lookup could not run this pass: {e}"}
        per_source_results.append((domain, result))

    found_results = [(domain, r) for domain, r in per_source_results if r.get("found")]
    if not found_results:
        notes = "; ".join(f"{domain}: {r.get('note', 'not found')}" for domain, r in per_source_results)
        return {"found": False, "note": f"no MCA-mirror record found for CIN {cin.strip()} in any configured source ({notes})"}

    merged: dict = {"found": True, "cin": cin.strip(), "sources_used": [d for d, _ in found_results]}
    for field in _SCALAR_PROFILE_FIELDS:
        for domain, r in found_results:
            if r.get(field):
                merged[field] = r[field]
                break
        merged.setdefault(field, None)

    merged["current_directors"], merged["roster_conflicts"] = _merge_director_rosters(
        [(domain, r.get("current_directors") or []) for domain, r in found_results]
    )
    # Past-director history and the shareholding-gating note are only ever
    # exposed by ZaubaCorp among the three sources (see each fetcher's own
    # module note) -- taken from there directly rather than merged.
    zaubacorp_result = dict(found_results).get("zaubacorp.com", {})
    merged["past_directors"] = zaubacorp_result.get("past_directors") or []
    merged["shareholding_note"] = zaubacorp_result.get("shareholding_note")
    # "url" always favors ZaubaCorp for backward-compatible citation
    # behavior (existing Charter runs already cite this URL) when present,
    # otherwise whichever source actually responded.
    merged["url"] = zaubacorp_result.get("url") or found_results[0][1].get("url")
    return merged


# ---------------------------------------------------------------------------
# CTS -> land-record lookup (Maha Bhulekh Property Card, see mahabhumi.py).
# The FINAL fetch (mahabhumi.fetch_property_card) is deliberately NOT wired
# to run unconditionally like the CIN checks above: it opens a visible
# browser and blocks waiting for a human to solve a fresh CAPTCHA on every
# single call (that site grants no reusable session -- see mahabhumi.py's
# own module note), so running it unconditionally would silently stall
# every automated Charter pass for up to CAPTCHA_TIMEOUT_SECONDS.
#
# The steps BEFORE that final fetch, however, ARE wired to run
# automatically -- confirmed live that office/village labels on this site
# are Marathi-only (e.g. "उप अधीक्षक भूमि अभिलेख, कल्याण") with no reliable
# way to match them against RERA's own English district/taluka/village
# text, so a human still has to pick the right office and village from a
# list; but FETCHING those lists (list_offices/list_villages -- both
# headless, no CAPTCHA) no longer needs a human to run mahabhumi.py by
# hand first. discover_cts_office_candidates below runs automatically
# inside run_cts_land_lookup whenever a district can be resolved and no
# lookup has been started yet, writing candidates for a human to read and
# choose from (see cts_resolve.py for the rest of the chain: villages,
# then CTS-number candidates, then the final cts_lookup_input.json write
# that this function already knew how to consume).
# ---------------------------------------------------------------------------

_DISTRICT_HINT_RE_CACHE = None


def _extract_district_hint(facts: dict) -> str | None:
    """Best-effort district name extracted from this Charter's own
    land_identification.mandal_taluka_district text (e.g. "Taluka Kalyan,
    District Thane, Maharashtra") -- matched against mahabhumi's own
    verified 34-district English/Marathi map, NEVER against a fuzzy
    guess. Prefers the LONGEST matching district name so "Mumbai
    Suburban" wins over the "Mumbai" substring inside it. Returns None
    (not a guess) if no known district name appears in the text at all."""
    import mahabhumi

    text = ((facts.get("land_identification") or {}).get("mandal_taluka_district") or {}).get("value") or ""
    text_lower = text.lower()
    matches = [name for name in mahabhumi._DISTRICT_NAME_MAP if name in text_lower]
    if not matches:
        return None
    return max(matches, key=len)


def discover_cts_office_candidates(facts: dict, reg_no: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Automatically fetches the real list of Maha Bhulekh offices
    (headless, no CAPTCHA) for whichever district this Charter's own land
    data resolves to, and writes output/<reg_no>/cts_office_candidates
    .json for a human to read and pick from -- see cts_resolve.py.
    Never auto-selects an office itself (see module note above: office
    labels are Marathi-only, a district's offices don't correspond 1:1 to
    RERA's own taluka names, confirmed live that one real project's
    office ["...,Andheri"] named a different place than its own village
    ["Aambivali"]).

    Returns {"found": bool, "district": ..., "offices": [...], "note":
    ...} -- {"found": False, "note": "..."} when no district could be
    resolved from the Charter's own data, or the live fetch itself
    failed; never raises."""
    district_hint = _extract_district_hint(facts)
    if not district_hint:
        return {"found": False, "note": "No known Maharashtra district name could be matched in this Charter's own land_identification.mandal_taluka_district text."}

    import mahabhumi

    result = mahabhumi.list_offices(district_hint)
    if not result.get("found"):
        return {"found": False, "district": district_hint, "note": result.get("note", "Office lookup failed")}

    record = {
        "district": district_hint,
        "district_label": result.get("district_label"),
        "offices": result.get("offices"),
        "generated_at": datetime.now().isoformat(),
        "note": (
            "Pick the office covering this Charter's own recorded taluka/village (see "
            "land_identification in facts.json), then run: python cts_resolve.py villages "
            f"{reg_no} \"<office label from this list, exact>\""
        ),
    }
    project_dir = os.path.join(output_dir, reg_no)
    os.makedirs(project_dir, exist_ok=True)
    with open(os.path.join(project_dir, "cts_office_candidates.json"), "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return {"found": True, "district": district_hint, "offices": result.get("offices"), "note": record["note"]}


def discover_cts_village_candidates(reg_no: str, office_label: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Step 2 of the human-in-the-loop chain: given an office label a
    human picked from cts_office_candidates.json, fetches the real
    village list for it (headless, no CAPTCHA) and writes output/<reg_no>
    /cts_village_candidates.json. Requires discover_cts_office_candidates
    to have already run for this reg_no (reads its district back out).
    Returns {"found": bool, "villages": [...], "note": ...}."""
    office_candidates_path = os.path.join(output_dir, reg_no, "cts_office_candidates.json")
    if not os.path.exists(office_candidates_path):
        return {"found": False, "note": f"{office_candidates_path} not found -- run discover_cts_office_candidates (or company_charter.py {reg_no}) first."}
    with open(office_candidates_path, "r", encoding="utf-8") as f:
        district = json.load(f)["district"]

    import mahabhumi

    try:
        result = mahabhumi.list_villages(district, office_label)
    except mahabhumi.AmbiguousSelectionError as e:
        return {"found": False, "note": str(e), "options": e.options}

    if not result.get("found"):
        return {"found": False, "note": result.get("note", "Village lookup failed")}

    record = {
        "district": district, "office": office_label,
        "villages": result.get("villages"),
        "generated_at": datetime.now().isoformat(),
        "note": (
            "Pick the village covering this Charter's own recorded plot (see land_identification "
            f"in facts.json), then run: python cts_resolve.py candidates {reg_no} \"{office_label}\" "
            "\"<village label from this list, exact>\" <your CTS number>"
        ),
    }
    project_dir = os.path.join(output_dir, reg_no)
    with open(os.path.join(project_dir, "cts_village_candidates.json"), "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return {"found": True, "villages": result.get("villages"), "note": record["note"]}


def discover_cts_number_candidates(reg_no: str, office_label: str, village_label: str, cts_query: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Step 3: given an office+village a human confirmed from the two
    candidate files above, resolves the typed CTS number against the
    site's own valid list for that exact village (headless, no CAPTCHA --
    a CTS number can have multiple sub-divisions, e.g. "100/1", "100/2").
    Writes output/<reg_no>/cts_number_candidates.json. Returns {"found":
    bool, "candidates": [...], "note": ...}."""
    office_candidates_path = os.path.join(output_dir, reg_no, "cts_office_candidates.json")
    if not os.path.exists(office_candidates_path):
        return {"found": False, "note": f"{office_candidates_path} not found -- run discover_cts_office_candidates (or company_charter.py {reg_no}) first."}
    with open(office_candidates_path, "r", encoding="utf-8") as f:
        district = json.load(f)["district"]

    import mahabhumi

    result = mahabhumi.search_cts_candidates(district, office_label, village_label, cts_query)
    if not result.get("found"):
        return {"found": False, "note": result.get("note", "CTS candidate search failed"), "options": result.get("options")}

    record = {
        "district": district, "office": office_label, "village": village_label, "cts_query": cts_query,
        "candidates": result.get("candidates"),
        "generated_at": datetime.now().isoformat(),
        "note": (
            f"Confirm the exact candidate matching your CTS number, then finalize with: python "
            f"cts_resolve.py finalize {reg_no} \"{office_label}\" \"{village_label}\" <exact CTS number "
            f"from candidates> <mobile number>"
        ),
    }
    project_dir = os.path.join(output_dir, reg_no)
    with open(os.path.join(project_dir, "cts_number_candidates.json"), "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return {"found": True, "candidates": result.get("candidates"), "note": record["note"]}


def run_cts_land_lookup(facts: dict, reg_no: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Runs the CTS -> Property Card lookup only if output/<reg_no>/
    cts_lookup_input.json exists, containing:
      {"district": "Pune", "office": "<exact Marathi label from
       mahabhumi.list_offices>", "village": "<exact Marathi label from
       mahabhumi.list_villages>", "cts_number": "100", "mobile": "..."}
    Silently returns facts unchanged if that file is absent -- the ordinary
    case for every automated run. When present, opens a visible browser and
    blocks for up to CAPTCHA_TIMEOUT_SECONDS waiting for a human to solve
    the CAPTCHA (see mahabhumi.fetch_property_card).

    Checked FIRST, before any of that: output/<reg_no>/land_record_carryover
    .json -- written by attach_rera_number when a CTS number reached this
    pipeline before this reg_no did (see run_cts_lookup_standalone). If
    present, its already-fetched cts_land_record_check is reused directly
    (no new browser/CAPTCHA needed), and its CTS number is cross-checked
    against this Charter's own RERA-sourced survey_cts_plot_numbers -- a
    genuine mismatch is recorded as facts["cts_mismatch_note"], which
    _classify_flags promotes to an IMMINENT flag (not a buried gap): a
    developer-supplied CTS not matching RERA's own official record is
    exactly the kind of discrepancy the investment team should be able to
    question the developer about directly, not have it read past unnoticed
    in a gaps list."""
    carryover_path = os.path.join(output_dir, reg_no, "land_record_carryover.json")
    if os.path.exists(carryover_path):
        with open(carryover_path, "r", encoding="utf-8") as f:
            carryover = json.load(f)

        recorded_cts = ((facts.get("land_identification", {}).get("survey_cts_plot_numbers") or {}).get("value") or "")
        carryover_cts = carryover.get("cts_number", "")
        if recorded_cts and carryover_cts and carryover_cts not in recorded_cts:
            facts["cts_mismatch_note"] = (
                f"CTS number mismatch: the CTS/plot number supplied before this project's RERA number was "
                f"known ({carryover_cts!r}, from {carryover.get('village', 'an unspecified village')}) does not "
                f"appear in this Charter's own RERA-sourced survey_cts_plot_numbers ({recorded_cts!r}) -- "
                f"confirm with the developer directly whether these refer to the same plot before trusting "
                f"either record."
            )

        facts["cts_land_record_check"] = carryover.get("cts_land_record_check", {"found": False, "note": "carryover record had no cts_land_record_check"})
        facts.setdefault("sources", []).extend(carryover.get("sources", []))
        facts.setdefault("gaps", []).extend(carryover.get("gaps", []))
        return facts

    input_path = os.path.join(output_dir, reg_no, "cts_lookup_input.json")
    if not os.path.exists(input_path):
        # No manual input yet -- auto-fetch office candidates (headless, no
        # CAPTCHA) so a human's next step is picking from a real list
        # rather than first having to run mahabhumi.py by hand. The live
        # fetch itself only ever runs ONCE per reg_no (checked via the
        # candidates file already existing, so a normal automated re-run
        # doesn't keep re-fetching an unchanging office list) -- but the
        # reminder gap is added on EVERY run until cts_lookup_input.json
        # exists, so this doesn't go quiet and get forgotten after the
        # first mention.
        office_candidates_path = os.path.join(output_dir, reg_no, "cts_office_candidates.json")
        district_for_gap = None
        if not os.path.exists(office_candidates_path):
            discovery = discover_cts_office_candidates(facts, reg_no, output_dir)
            if discovery.get("found"):
                district_for_gap = discovery["district"]
        else:
            with open(office_candidates_path, "r", encoding="utf-8") as f:
                district_for_gap = json.load(f).get("district")

        if district_for_gap:
            facts.setdefault("gaps", []).append(
                f"CTS land-record lookup: office candidates for {district_for_gap} are in "
                f"output/{reg_no}/cts_office_candidates.json. Pick the office covering this project's own "
                f"recorded taluka/village, then run cts_resolve.py to continue (villages, then CTS-number "
                f"candidates, then the final Property Card fetch)."
            )
        return facts

    import mahabhumi

    with open(input_path, "r", encoding="utf-8") as f:
        cts_input = json.load(f)

    required = ("district", "office", "village", "cts_number", "mobile")
    missing = [k for k in required if not cts_input.get(k)]
    if missing:
        facts["cts_land_record_check"] = {"found": False, "note": f"{input_path} is missing required field(s): {', '.join(missing)}"}
        return facts

    recorded_cts = ((facts.get("land_identification", {}).get("survey_cts_plot_numbers") or {}).get("value") or "")
    if recorded_cts and cts_input["cts_number"] not in recorded_cts:
        facts.setdefault("gaps", []).append(
            f"CTS land-record lookup: cts_lookup_input.json's CTS number ({cts_input['cts_number']}) does not "
            f"appear in this Charter's own recorded survey_cts_plot_numbers ({recorded_cts!r}) -- confirm "
            f"these refer to the same plot before trusting the Property Card result below."
        )

    print(f"\n[INFO] {input_path} found -- resolving CTS {cts_input['cts_number']} candidates...")
    try:
        candidates_result = mahabhumi.search_cts_candidates(
            cts_input["district"], cts_input["office"], cts_input["village"], cts_input["cts_number"]
        )
    except Exception as e:
        facts["cts_land_record_check"] = {"found": False, "note": f"CTS candidate search could not run this pass: {e}"}
        return facts

    if not candidates_result.get("found"):
        facts["cts_land_record_check"] = {"found": False, "note": candidates_result.get("note", "CTS candidate search failed")}
        return facts

    candidates = candidates_result["candidates"]
    if cts_input["cts_number"] not in candidates:
        facts["cts_land_record_check"] = {
            "found": False,
            "note": (
                f"CTS number {cts_input['cts_number']!r} is not an exact match against the site's own "
                f"candidates for this village ({candidates}) -- update cts_lookup_input.json with the exact "
                f"value and re-run rather than guessing."
            ),
        }
        return facts

    print(f"[INFO] Opening a browser to fetch the Property Card -- please solve the CAPTCHA when it appears.")
    try:
        result = mahabhumi.fetch_property_card(
            cts_input["district"], cts_input["office"], cts_input["village"],
            cts_input["cts_number"], cts_input["mobile"],
        )
    except (mahabhumi.CaptchaTimeoutError, mahabhumi.BrowserClosedError, mahabhumi.AmbiguousSelectionError) as e:
        result = {"found": False, "note": f"CTS Property Card lookup did not complete: {e}"}
    except Exception as e:
        result = {"found": False, "note": f"CTS Property Card lookup could not run this pass: {e}"}

    facts["cts_land_record_check"] = result
    if result.get("found"):
        facts.setdefault("sources", []).append({
            "label": "Maha Bhulekh Property Card",
            "ref": f"CTS {cts_input['cts_number']}, {cts_input['village']} -- {result.get('url', '')}",
            "topic": "land_record",
            "published_date": "unknown",
            "accessed_date": datetime.now().strftime("%Y-%m-%d"),
        })
    return facts


# ---------------------------------------------------------------------------
# GST filing-compliance intake -- same opt-in convention as
# run_cts_land_lookup/reviews.json just above: does nothing unless a human
# has dropped output/<reg_no>/gst_filing_input.json, since there is no
# automated GST portal scrape here at all (see gst_compliance.py's own
# module docstring: the portal's "Search Taxpayer" filing table sits behind
# a CAPTCHA solved fresh per lookup, the same hard constraint already
# documented for MahaRERA/Maha Bhulekh). The GSTIN is requested directly
# from the developer; the filing dates in that JSON come from either the
# developer's own reply or a human manually reading the portal's filing
# table themselves.
# ---------------------------------------------------------------------------

def run_gst_compliance_check(facts: dict, reg_no: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Runs the GST filing-pattern analysis only if output/<reg_no>/
    gst_filing_input.json exists, containing:
      {"gstin": "27AANCM5273D1ZA", "records": [{"form": "GSTR-3B"|"GSTR-1",
       "period_start": "YYYY-MM-DD", "period_end": "YYYY-MM-DD",
       "filing_date": "YYYY-MM-DD"|null}, ...]}
    Silently returns facts unchanged if that file is absent -- the ordinary
    case for every automated run. Sets facts["gst_compliance_check"] to
    {"found": False, "note": ...} (never raises) on a malformed GSTIN or
    record, so a typo'd input file surfaces as an honest gap rather than
    crashing the whole Charter pass."""
    input_path = os.path.join(output_dir, reg_no, "gst_filing_input.json")
    if not os.path.exists(input_path):
        return facts

    import gst_compliance

    with open(input_path, "r", encoding="utf-8") as f:
        gst_input = json.load(f)

    gstin = (gst_input.get("gstin") or "").strip()
    if not gst_compliance.validate_gstin(gstin):
        facts["gst_compliance_check"] = {"found": False, "note": f"{input_path}'s gstin ({gstin!r}) is not a validly-formatted GSTIN."}
        return facts

    records = []
    for i, rec in enumerate(gst_input.get("records", [])):
        try:
            records.append({
                "form": rec["form"],
                "period_start": datetime.strptime(rec["period_start"], "%Y-%m-%d").date(),
                "period_end": datetime.strptime(rec["period_end"], "%Y-%m-%d").date(),
                "filing_date": datetime.strptime(rec["filing_date"], "%Y-%m-%d").date() if rec.get("filing_date") else None,
            })
        except (KeyError, ValueError) as e:
            facts["gst_compliance_check"] = {"found": False, "note": f"{input_path} record[{i}] is malformed: {e}"}
            return facts

    summary = gst_compliance.summarize_filing_pattern(gstin, records, as_of=datetime.now().date())
    facts["gst_compliance_check"] = {"found": True, "gstin": gstin, "summary": summary}
    facts.setdefault("sources", []).append({
        "label": "GST filing history (developer-supplied)",
        "ref": f"GSTIN {gstin}, {len(records)} filing period(s) on record",
        "topic": "gst_compliance",
        "published_date": "unknown",
        "accessed_date": datetime.now().strftime("%Y-%m-%d"),
    })
    return facts


def _slugify_for_pending_key(*parts: str) -> str:
    """Joins the given parts with underscores into a filesystem-safe
    directory name -- strips characters Windows/NTFS rejects in a path
    segment and collapses whitespace, but otherwise leaves the exact
    Marathi office/village labels intact (NTFS handles Unicode names
    natively, no transliteration needed)."""
    joined = "_".join(str(p).strip() for p in parts if p)
    joined = re.sub(r'[<>:"/\\|?*]', "_", joined)
    joined = re.sub(r"\s+", "_", joined)
    return joined.strip("_") or "unknown"


# ---------------------------------------------------------------------------
# Standalone CTS/land intake -- runs the same CTS -> Property Card lookup
# run_cts_land_lookup runs internally, but from bare district/office/
# village/cts_number/mobile handed to us directly, with no RERA project/
# reg_no/facts dict required. Exists for the case where a CTS number reaches
# this pipeline before any RERA number does (land is acquired, and so has a
# CTS number, well before a project ever registers with RERA).
#
# Deliberately duplicates run_cts_land_lookup's candidate-search/fetch logic
# (rather than refactoring that function to share this code) for the same
# reason run_promoter_intake duplicates run_company_charter's source-shaping
# logic: it keeps this addition from touching a single existing line of
# run_cts_land_lookup, so it stays trivially revertable (delete this
# function and cts_intake.py, nothing else changes) until this standalone
# path has actually proven useful.
#
# Writes its result under facts.json's own key for this check
# ("cts_land_record_check") so a later run_cts_land_lookup pass for the same
# plot can absorb this file with a plain dict assignment, never a reshape.
# ---------------------------------------------------------------------------

def run_cts_lookup_standalone(
    district: str, office: str, village: str, cts_number: str, mobile: str,
    output_dir: str = config.OUTPUT_ROOT,
) -> dict:
    """Runs the CTS -> Property Card lookup from a bare identifier set.
    Candidate search (confirming cts_number is an exact match against the
    site's own valid list for this village) is headless, no CAPTCHA; the
    actual Property Card fetch opens a visible browser and blocks waiting
    for a human to solve a fresh CAPTCHA every call (see mahabhumi.py --
    that site grants no reusable session).

    Persists to output/_pending/<district>_<village>_<cts_number>/
    land_record.json (slugified) -- a bare CTS number isn't globally unique
    (the same number recurs across different villages), so district+
    village+cts_number together form the key. Returns {"cts_land_record_
    check": {...}, "sources": [...], "gaps": []} -- the same shape
    promoter_intake.py's record uses, and the same "cts_land_record_check"
    key facts.json itself uses for this check.

    Confirmed live (a real CAPTCHA solve completed against Pranami Bliss's
    own recorded CTS 183 in Aambivali, Mumbai Suburban): the Property Card
    result page renders as a document/image, not structured HTML text --
    mahabhumi.fetch_property_card's screenshot_path/OCR support (see its
    own docstring) is what actually recovers its content, so the screenshot
    path is computed here BEFORE the CAPTCHA-gated fetch, not after."""
    import mahabhumi

    out_dir = os.path.join(output_dir, "_pending", _slugify_for_pending_key(district, village, cts_number))
    os.makedirs(out_dir, exist_ok=True)
    screenshot_path = os.path.join(out_dir, "property_card_screenshot.png")

    print(f"\n[INFO] Resolving CTS {cts_number} candidates...")
    try:
        candidates_result = mahabhumi.search_cts_candidates(district, office, village, cts_number)
    except Exception as e:
        result = {"found": False, "note": f"CTS candidate search could not run this pass: {e}"}
    else:
        if not candidates_result.get("found"):
            result = {"found": False, "note": candidates_result.get("note", "CTS candidate search failed")}
        elif cts_number not in candidates_result["candidates"]:
            result = {
                "found": False,
                "note": (
                    f"CTS number {cts_number!r} is not an exact match against the site's own "
                    f"candidates for this village ({candidates_result['candidates']}) -- confirm the "
                    f"exact value and re-run rather than guessing."
                ),
            }
        else:
            print("[INFO] Opening a browser to fetch the Property Card -- please solve the CAPTCHA when it appears.")
            try:
                result = mahabhumi.fetch_property_card(
                    district, office, village, cts_number, mobile, screenshot_path=screenshot_path,
                )
            except (mahabhumi.CaptchaTimeoutError, mahabhumi.BrowserClosedError, mahabhumi.AmbiguousSelectionError) as e:
                result = {"found": False, "note": f"CTS Property Card lookup did not complete: {e}"}
            except Exception as e:
                result = {"found": False, "note": f"CTS Property Card lookup could not run this pass: {e}"}

    record = {
        "district": district, "office": office, "village": village,
        "cts_number": cts_number, "generated_at": datetime.now().isoformat(),
        "cts_land_record_check": result,
        "sources": [],
        "gaps": [],
    }
    if result.get("found"):
        record["sources"].append({
            "label": "Maha Bhulekh Property Card",
            "ref": f"CTS {cts_number}, {village} -- {result.get('url', '')}",
            "topic": "land_record",
            "published_date": "unknown",
            "accessed_date": datetime.now().strftime("%Y-%m-%d"),
        })

    with open(os.path.join(out_dir, "land_record.json"), "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return record


# ---------------------------------------------------------------------------
# MahaRERA Orders/Judgments search -- confirmed live and scriptable with no
# browser needed. maharera.maharashtra.gov.in/orders-judgements is a Drupal
# form (POST, requires a fresh form_build_id read from a prior GET -- no
# CSRF/form_token beyond that). Live-tested: searching "Alta Monte" returned
# 11 real results; a search for a project with no published judgment yet
# (Godrej Park Greens) correctly returns zero, not an error.
#
# This complements, not replaces, api_client.download_complaint_orders --
# that mechanism only covers complaints (whose own record already carries
# an order reference); appeals.json carries no such reference at all
# (lastOrder/lastRoznama are always "NA"), which is exactly why this
# external search is needed for appeal-level judgments.
#
# Two real quirks worth documenting:
#   1. Each result's "View Judgement" link is not a URL at all -- the
#      judgment PDF's full bytes are embedded inline as base64 in an
#      `oj-data` attribute. No second request is needed to fetch it.
#   2. The search endpoint is genuinely flaky: because a result page with
#      several embedded PDFs can be 10+ MB, the server sometimes returns a
#      truncated ~60KB placeholder shell (a Drupal BigPipe artifact)
#      instead of the fully-rendered results, non-deterministically --
#      confirmed by requesting the identical search 4 times in a row and
#      getting the full response only 1-3 times. A real browser User-Agent
#      measurably improves (but does not eliminate) this, so both are used
#      together with retries here, rather than trusting a single request.
# ---------------------------------------------------------------------------

_MAHARERA_ORDERS_URL = "https://maharera.maharashtra.gov.in/orders-judgements"
_MAHARERA_COMPLAINT_TYPES = ("rulings_of_MahaRERA", "judgements_by_adjudicating_officers")
_BROWSER_UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"


def _maharera_orders_search_once(project_name: str, complaint_type: str) -> list:
    """Single attempt -- may return [] either because there are genuinely
    no matches, or because of the BigPipe flakiness noted above; the caller
    (search_maharera_judgments) is responsible for retrying."""
    headers = {"User-Agent": _BROWSER_UA}
    session = requests.Session()
    resp1 = session.get(_MAHARERA_ORDERS_URL, headers=headers, timeout=config.REQUEST_TIMEOUT)
    resp1.raise_for_status()
    build_id_match = re.search(r'name="form_build_id" value="([^"]+)"', resp1.text)
    if not build_id_match:
        return []

    data = {
        "order_complaint_type": complaint_type,
        "order_project_name": project_name,
        "form_build_id": build_id_match.group(1),
        "form_id": "orders_judgements_form",
        "op": "Search",
    }
    resp2 = session.post(_MAHARERA_ORDERS_URL, data=data, headers=headers, timeout=config.REQUEST_TIMEOUT)
    resp2.raise_for_status()
    if "bg-body rounded" not in resp2.text:
        return []  # the truncated-shell case -- treat as "try again", not "no results"

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(resp2.text, "lxml")
    cards = [d for d in soup.find_all("div") if d.get("class") and "shadow" in d.get("class")]

    def _label_value(card, label_text):
        label = card.find("label", string=re.compile(re.escape(label_text)))
        if not label:
            return None
        p = label.find_next("p")
        return p.get_text(strip=True) if p else None

    results = []
    for card in cards:
        reg_no_tag = card.find("p", class_="p-0")
        title_tag = card.find("h4")
        pdf_link = card.find("a", attrs={"oj-data": True})
        results.append({
            "reg_no": reg_no_tag.get_text(strip=True).lstrip("#") if reg_no_tag else None,
            "project_name": title_tag.get_text(strip=True) if title_tag else None,
            "complainant_name": _label_value(card, "Complainant Name"),
            "complainant_no": _label_value(card, "Complainant No"),
            "respondent_name": _label_value(card, "Respondent Name"),
            "uploaded_date": _label_value(card, "Uploaded Date"),
            "judgment_pdf_base64": pdf_link.get("oj-data") if pdf_link else None,
            "complaint_type": complaint_type,
        })
    return results


def _maharera_orders_search_with_retry(project_name: str, complaint_type: str, max_attempts: int) -> list:
    """One complaint type's search, retried against the documented BigPipe
    flakiness -- structured to keep the common (non-flaky) case cheap
    while still speeding up genuine retries: tries once, and only if THAT
    comes back empty does it fire the remaining (max_attempts - 1) tries
    CONCURRENTLY, returning as soon as any one of them succeeds rather
    than waiting for all of them in sequence. This was the dominant cost
    of the whole Phase 2 check batch (confirmed live: ~41s of a ~45s
    total, vs. under 2.2s for each of the other four checks combined --
    see the concurrent Phase 2 checks above _safe_credit_rating) precisely
    because it could burn up to `max_attempts` sequential round-trips per
    complaint type. A clean project (no flakiness this pass) still costs
    exactly 1 request, identical to the old fully-sequential version; a
    flaky one now costs at most 2 sequential rounds (1 + 1 concurrent
    batch) instead of up to `max_attempts`."""
    try:
        results = _maharera_orders_search_once(project_name, complaint_type)
    except requests.RequestException:
        results = []
    if results or max_attempts <= 1:
        return results

    executor = concurrent.futures.ThreadPoolExecutor(max_workers=max_attempts - 1)
    try:
        futures = [executor.submit(_maharera_orders_search_once, project_name, complaint_type) for _ in range(max_attempts - 1)]
        for future in concurrent.futures.as_completed(futures):
            try:
                retry_results = future.result()
            except requests.RequestException:
                retry_results = []
            if retry_results:
                return retry_results
        return []
    finally:
        # Deliberately wait=False: once a good result is found (or every
        # retry is exhausted empty), there's no reason to block returning
        # to the caller on some other still-in-flight retry finishing --
        # it'll just complete and be discarded in the background.
        executor.shutdown(wait=False)


def search_maharera_judgments(project_name: str, max_attempts: int = 3) -> list:
    """Searches MahaRERA's public Orders/Judgments page for `project_name`
    under both categories that can carry a genuine adjudicated outcome
    ("Rulings of MahaRERA" and "Judgements by Adjudicating Officers" --
    "Non-Registration Rulings" is a different, irrelevant category). Both
    categories are searched CONCURRENTLY (they're independent of each
    other), and each one's own retries against the BigPipe flakiness
    (see _maharera_orders_search_with_retry) are themselves partly
    concurrent -- so the worst case here is roughly 2 sequential
    round-trips, not up to `max_attempts` x 2 sequential round-trips as in
    the original fully-sequential version. A project with no published
    judgment yet returns [] after exhausting retries -- this is the
    expected, common case (most complaints/appeals are still pending), not
    a failure; there is no way from this function alone to distinguish
    "genuinely nothing published" from "every retry hit the flaky
    empty-shell response", so callers should not treat an empty result as
    a confirmed absence for a project with very few realistic search
    attempts left in a budget."""
    with concurrent.futures.ThreadPoolExecutor(max_workers=len(_MAHARERA_COMPLAINT_TYPES)) as executor:
        futures = [
            executor.submit(_maharera_orders_search_with_retry, project_name, complaint_type, max_attempts)
            for complaint_type in _MAHARERA_COMPLAINT_TYPES
        ]
        all_results = []
        for future in futures:
            all_results.extend(future.result())
    return all_results


def cross_reference_appeals(judgments: list, appeals_data: list) -> list:
    """Matches MahaRERA Orders/Judgments search results against this
    project's own appeals.json records, using `complainant_no` -- confirmed
    to be the exact same ID scheme as appeals.json's own
    `complaintRegistrationNo` field (e.g. CC005000000023263), a real join
    key rather than fuzzy name matching. Returns only judgments that
    matched a real appeal record on this project, each annotated with the
    matched appeal's own fields."""
    by_complaint_no = {a.get("complaintRegistrationNo"): a for a in appeals_data if a.get("complaintRegistrationNo")}
    matched = []
    for j in judgments:
        appeal = by_complaint_no.get(j.get("complainant_no"))
        if appeal:
            matched.append({**j, "matched_appeal": appeal})
    return matched


def _save_appeal_judgment_pdfs(reg_no: str, matched_judgments: list, output_dir: str) -> list:
    """Decodes and saves each matched judgment's embedded base64 PDF to
    output/<reg_no>/appeal_judgments/ -- the same "documents belong on disk,
    referenced by filename" convention as every other document in this
    project, rather than keeping multi-megabyte base64 blobs inline in
    facts.json. Returns the same list with judgment_pdf_base64 replaced by
    saved_filename (or None if that specific save failed)."""
    if not matched_judgments:
        return matched_judgments
    import base64

    out_dir = os.path.join(output_dir, reg_no, "appeal_judgments")
    os.makedirs(out_dir, exist_ok=True)
    result = []
    for j in matched_judgments:
        entry = dict(j)
        b64 = entry.pop("judgment_pdf_base64", None)
        entry["saved_filename"] = None
        if b64:
            try:
                pdf_bytes = base64.b64decode(b64)
                filename = f"{entry.get('complainant_no') or 'judgment'}.pdf"
                with open(os.path.join(out_dir, filename), "wb") as f:
                    f.write(pdf_bytes)
                entry["saved_filename"] = filename
            except Exception:
                pass  # keep saved_filename=None -- the metadata match itself is still real and kept
        result.append(entry)
    return result


# ---------------------------------------------------------------------------
# Review-authenticity heuristic triage (Phase 2) -- approved scope: a
# red-flag surfacing aid, explicitly NOT a fabrication verdict. Confirmed by
# research this session that true fabrication-probability scoring isn't
# achievable without platform-internal data (IP/device signals, removed-
# review logs, verified-purchase flags) this project has no access to.
#
# Each function below is a self-contained analysis primitive taking a plain
# list of review dicts as input ({rating, date (YYYY-MM-DD), text,
# reviewer_name, reviewer_review_count (optional -- the total number of
# reviews that reviewer has ever posted, on any listing, if known)) --
# deliberately decoupled from any specific review source (Google Maps,
# 99acres, MouthShut, ...), since fetching reviews live is a separate,
# harder, source-specific scraping problem (Google Maps in particular is a
# heavy JS SPA, not a plain-HTTP target like the other Phase 2 sources
# above) and was not built as part of this pass -- these functions are
# ready to consume whatever review data is gathered, by hand or by a future
# scraper, without needing to know where it came from.
#
# cross_reference_review_claims is the single highest-value check (per
# research): it costs nothing to fetch (uses only data this project already
# has) and catches the most concrete, checkable kind of dishonesty --
# reviews making a specific factual claim that contradicts this project's
# own confirmed RERA record.
# ---------------------------------------------------------------------------


def analyze_rating_polarization(reviews: list) -> dict:
    """Rating-distribution histogram (1-5 stars) plus the % that sit at the
    extremes (1 or 5 stars) versus the middle (2-4) -- fake-review research
    (Luca & Zervas) finds fabricated reviews skew disproportionately to the
    extremes, so a heavily bimodal distribution is a directional red flag,
    not proof of anything on its own."""
    counts = {i: 0 for i in range(1, 6)}
    for r in reviews:
        rating = r.get("rating")
        if isinstance(rating, (int, float)) and 1 <= rating <= 5:
            counts[int(round(rating))] += 1
    total = sum(counts.values())
    extreme = counts[1] + counts[5]
    return {
        "counts": counts,
        "total": total,
        "pct_extreme": round(100 * extreme / total, 1) if total else None,
    }


def detect_review_bursts(reviews: list, window_days: int = 7, burst_threshold: int = 5) -> list:
    """Flags clusters of `burst_threshold` or more reviews all posted within
    a `window_days`-day span -- an abnormally short posting window is one of
    the standard signals used in spammer-group detection research (a
    kernel-density/burst-clustering approach, simplified here to a sliding
    window over sorted dates, since this project doesn't need the full
    statistical machinery to surface an obvious cluster for a human to look
    at). Requires each review dict to have a parseable "date" (YYYY-MM-DD);
    reviews without one are ignored, not treated as evidence either way."""
    dated = []
    for r in reviews:
        try:
            d = datetime.strptime(str(r.get("date", ""))[:10], "%Y-%m-%d")
            dated.append((d, r))
        except ValueError:
            continue
    dated.sort(key=lambda pair: pair[0])

    bursts = []
    i = 0
    while i < len(dated):
        j = i
        while j < len(dated) and (dated[j][0] - dated[i][0]).days <= window_days:
            j += 1
        cluster_size = j - i
        if cluster_size >= burst_threshold:
            bursts.append({
                "start_date": dated[i][0].strftime("%Y-%m-%d"),
                "end_date": dated[j - 1][0].strftime("%Y-%m-%d"),
                "count": cluster_size,
                "reviews": [r for _, r in dated[i:j]],
            })
            i = j
        else:
            i += 1
    return bursts


def detect_near_duplicate_reviews(reviews: list, similarity_threshold: float = 0.85) -> list:
    """Flags pairs of reviews whose text is near-identical (templated/
    copy-pasted language) -- a documented pattern in fabricated-review
    campaigns, and one of the few fabrication signals checkable purely from
    review text with no platform-internal data. Uses difflib's
    SequenceMatcher (stdlib, no new dependency) rather than a heavier NLP
    similarity model, which is a reasonable trade-off for a triage tool
    flagging pairs for a human to actually read, not an automated verdict."""
    from difflib import SequenceMatcher

    texts = [(i, (r.get("text") or "").strip()) for i, r in enumerate(reviews)]
    texts = [(i, t) for i, t in texts if len(t) > 20]  # too short to meaningfully compare
    pairs = []
    for a in range(len(texts)):
        for b in range(a + 1, len(texts)):
            i1, t1 = texts[a]
            i2, t2 = texts[b]
            ratio = SequenceMatcher(None, t1, t2).ratio()
            if ratio >= similarity_threshold:
                pairs.append({"review_a": reviews[i1], "review_b": reviews[i2], "similarity": round(ratio, 3)})
    return pairs


def flag_one_hit_wonder_reviewers(reviews: list) -> list:
    """Flags reviewers whose total review count (across every listing
    they've ever reviewed, not just this one) is exactly 1 -- a documented
    heuristic from consumer review-integrity tools (ReviewMeta, Fakespot)
    for accounts that exist only to post a single review. Requires the
    caller to supply `reviewer_review_count` per review (obtainable by
    visiting each reviewer's own public profile page -- e.g. clicking
    through on Google Maps -- which this function does not do itself);
    reviews missing that field are skipped, not assumed to be one-hit
    wonders."""
    return [r for r in reviews if r.get("reviewer_review_count") == 1]


_POSSESSION_DELAY_CLAIM_RE = re.compile(
    r"(?:possession|handover|deliver\w*)\D{0,40}?(?:delay\w*|postpone\w*)\D{0,20}?(?:to\s+)?(20\d{2})",
    re.I,
)
_POSSESSION_YEAR_MENTION_RE = re.compile(r"\b(20\d{2})\b")


def cross_reference_review_claims(reviews: list, facts: dict) -> list:
    """The single highest-value check in this module (per Phase 2
    research): extracts specific, checkable factual claims from review text
    -- currently possession-delay year mentions, the most common concrete
    claim in real-estate reviews -- and compares them against this
    project's own confirmed RERA data (rera_core_fields.proposed_completion_date).
    Flags each review as "consistent", "inconsistent" (cites a materially
    different year than the project's own record), or leaves it unflagged
    if it makes no checkable claim at all -- deliberately narrow rather than
    trying to parse every possible claim type, since a wrong classification
    here would be worse than simply not flagging an ambiguous case."""
    completion_date = (facts.get("rera_core_fields", {}) or {}).get("proposed_completion_date", "")
    completion_year_match = re.search(r"\b(20\d{2})\b", completion_date)
    completion_year = int(completion_year_match.group(1)) if completion_year_match else None

    flagged = []
    for r in reviews:
        text = r.get("text") or ""
        delay_match = _POSSESSION_DELAY_CLAIM_RE.search(text)
        if not delay_match:
            continue
        claimed_year = int(delay_match.group(1))
        if completion_year is None:
            flagged.append({"review": r, "claimed_year": claimed_year, "verdict": "unverifiable -- no confirmed completion date on record to check against"})
        elif abs(claimed_year - completion_year) <= 1:
            flagged.append({"review": r, "claimed_year": claimed_year, "project_year": completion_year, "verdict": "consistent"})
        else:
            flagged.append({"review": r, "claimed_year": claimed_year, "project_year": completion_year, "verdict": "inconsistent -- materially different from this project's own confirmed record"})
    return flagged


def run_review_authenticity_triage(reviews: list, facts: dict) -> dict:
    """Runs all four checks above and returns a combined result -- the
    single entry point Company Charter generation (or any other caller)
    should use rather than calling each analysis function separately."""
    return {
        "total_reviews_analyzed": len(reviews),
        "rating_polarization": analyze_rating_polarization(reviews),
        "burst_clusters": detect_review_bursts(reviews),
        "near_duplicate_pairs": detect_near_duplicate_reviews(reviews),
        "one_hit_wonder_reviewers": flag_one_hit_wonder_reviewers(reviews),
        "claim_cross_reference": cross_reference_review_claims(reviews, facts),
    }


# ---------------------------------------------------------------------------
# Template filling
# ---------------------------------------------------------------------------

# Set only for the duration of an "external"-variant _fill_template call
# (via a try/finally there -- never left set across calls). _set_paragraph_text
# is the one low-level primitive nearly everything in this file funnels
# through to write text into the document (directly, or via _set_cell/
# _set_row_cell) -- checking it here, rather than threading a `facts`
# parameter through this function's ~70 call sites individually, is what
# makes External's prose cleanup reach text built from hardcoded Python
# f-strings (flag/reason text assembled at render time, not stored in
# facts.json) as well as facts.json's own free-text fields.
_ACTIVE_EXTERNAL_FACTS: dict | None = None


_CIN_VALUE_PATTERN = r"[UL]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}"
_CIN_LABELED_RE = re.compile(r"\bCIN\s*[:\-]?\s*" + _CIN_VALUE_PATTERN, re.IGNORECASE)
_DIN_LABELED_RE = re.compile(r"\bDIN\s*[:\-]?\s*\d{7,8}\b", re.IGNORECASE)


def _strip_inline_cin_din(text: str) -> str:
    """Removes CIN/DIN VALUES (e.g. "CIN U70102MH2010PTC210775", "DIN
    00448678") from running prose -- CLAUDE.md Section B: CIN/DIN stays
    only in the stakeholder/director identity tables, never inline in
    prose anywhere else. Strips just the "CIN <value>"/"DIN <value>"
    token itself (plus tidying the punctuation left behind), not the
    whole surrounding parenthetical -- "(DIN 00448678, appointed
    2022-09-07)" keeps its other, non-CIN/DIN content: "(appointed
    2022-09-07)"."""
    if not text:
        return text
    text = _CIN_LABELED_RE.sub("", text)
    text = _DIN_LABELED_RE.sub("", text)
    text = re.sub(r"\(\s*,\s*", "(", text)   # "(, appointed" -> "(appointed"
    text = re.sub(r",\s*\)", ")", text)      # "appointed ..., )" -> "appointed ...)"
    text = re.sub(r"\(\s*\)", "", text)      # a parenthetical left fully empty
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([,.])", r"\1", text)
    return text.strip()


def _set_paragraph_text_raw(paragraph, text: str) -> None:
    """The actual paragraph-text setter, with NO CIN/DIN stripping --
    _set_cell/_set_row_cell call this directly so table cells (including
    the stakeholder/director identity tables, where CIN/DIN is required,
    not forbidden) are never touched by _set_paragraph_text's stripping.
    Every other caller should use _set_paragraph_text instead."""
    if _ACTIVE_EXTERNAL_FACTS is not None:
        text = _externalize_prose(_ACTIVE_EXTERNAL_FACTS, str(text))
    for extra_run in paragraph.runs[1:]:
        extra_run.text = ""
        if _ACTIVE_EXTERNAL_FACTS is not None:
            from docx.shared import RGBColor
            extra_run.font.color.rgb = RGBColor.from_string("000000")
            extra_run.italic = False
    if paragraph.runs:
        paragraph.runs[0].text = text
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(text)
    if _ACTIVE_EXTERNAL_FACTS is not None:
        # The template's own placeholder runs (this function's normal path
        # for nearly every body paragraph) carry a grey/italic "note" style
        # baked in at template-authoring time -- fine for Internal's
        # analyst working document, but it rendered almost the ENTIRE
        # External document in low-contrast grey italics, not just the
        # Gaps & Sources list. External gets plain, fully-readable body
        # text instead; severity coloring (flags, KPI cells, the Gaps
        # list) is applied separately, AFTER this runs, so it always wins.
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string("000000")
        run.italic = False


# Proof-of-concept reusable glossary layer for CLAUDE.md Section B's
# "no jargon, plain language, keep key terms" rule -- deliberately a
# SEPARATE, small, general-purpose mechanism from _EXTERNAL_DASH_REWRITES
# (the large hardcoded exact-full-sentence internal->external translation
# dict a few hundred lines below this one). That dict is a one-off lookup
# table keyed on exact sentences from one specific promoter/project (it
# bakes in real DINs/CINs/PANs and falls through as a no-op for any other
# project's sentences -- see its own comment). Whether it should eventually
# be migrated onto a mechanism like this one is flagged back to the user,
# not decided here -- this pass only adds the new jargon-expansion need,
# it does not touch or refactor that existing dict.
_JARGON_GLOSSARY = {
    "CIRP": "insolvency proceedings (CIRP)",
    "NCLT": "the insolvency tribunal (NCLT)",
    "NCLAT": "the insolvency appellate tribunal (NCLAT)",
    "IBC": "the Insolvency and Bankruptcy Code (IBC)",
    "DSRA": "the debt service reserve account (DSRA)",
    "ECLGS": "the Emergency Credit Line Guarantee Scheme (ECLGS)",
    "IGR": "the Inspector General of Registration (IGR)",
    "e-ASR": "the electronic Annual Statement of Rates (e-ASR)",
    "KMP": "Key Managerial Personnel (KMP)",
    "ROC": "the Registrar of Companies (ROC)",
}


def _expand_jargon_first_use(text: str) -> str:
    """CLAUDE.md Section B: on first use per text, expand a known jargon
    term once -- keeping the term itself (still searchable/cross-
    referenceable), just not left unexplained (e.g. "CIRP" -> "insolvency
    proceedings (CIRP)"). Subsequent occurrences of the SAME term within
    the same text are left bare, matching ordinary first-use-expands-then-
    abbreviates writing convention. A plain-language pass, not a
    dumbing-down pass -- the term is kept, not replaced."""
    if not text:
        return text
    for term, expansion in _JARGON_GLOSSARY.items():
        pattern = re.compile(rf"\b{re.escape(term)}\b")
        first = pattern.search(text)
        if not first:
            continue
        start, end = first.span()
        text = text[:start] + expansion + text[end:]
    return text


def _set_paragraph_text(paragraph, text: str) -> None:
    """Body-paragraph text setter for narrative slots -- strips inline
    CIN/DIN values and expands first-use jargon (CLAUDE.md Section B)
    before delegating to _set_paragraph_text_raw. Table cells go through
    _set_cell/_set_row_cell instead, which call the raw setter directly
    and deliberately skip both passes -- see that function's docstring."""
    text = _expand_jargon_first_use(_strip_inline_cin_din(str(text)))
    _set_paragraph_text_raw(paragraph, text)


def _format_elapsed(seconds: float) -> str:
    """Human-readable duration for the version-log line -- "Xh Ym Zs",
    dropping leading zero units (e.g. "4m 12s", not "0h 4m 12s")."""
    total = max(int(seconds), 0)
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours}h {minutes}m {secs}s"
    if minutes:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


def _insert_version_log(doc, p, elapsed_seconds: float | None, cost_usd: float | None, api_calls: int | None) -> None:
    """Inserts a small metadata line on the cover page (run time + Anthropic
    API cost for this specific run), right before the "Overview & Flags"
    heading. Caller only invokes this for the Internal variant -- cost/
    timing is operational detail, not something to hand an external
    counterparty. Inserted as a brand-new paragraph via python-docx rather
    than a template slot, since the .docx template predates this field --
    this only touches the in-memory Document built at render time, never
    output/company_charters/Company_Charter_TEMPLATE_WebSourced.docx itself,
    so no template backup is needed (see CLAUDE.md's template-safety note,
    which is about edits to the template FILE).

    Deliberately does NOT anchor on the `p` list captured near the top of
    _fill_template: "Overview & Flags" is assembled/repositioned to the
    front of the document by later code in this function, using its own
    fixed-index lookups against the ORIGINAL template layout (where that
    position instead holds "Methodology Note", which gets deleted
    entirely) -- p[4] is stale by the time this runs. Querying doc.paragraphs
    fresh here reflects the document's true, final, already-reordered shape."""
    from docx.shared import Pt

    if elapsed_seconds is None and cost_usd is None:
        return
    if api_calls:
        cost_text = f"API cost: ${cost_usd:,.2f} ({api_calls} Anthropic API call{'s' if api_calls != 1 else ''})"
    else:
        cost_text = "API cost: no Anthropic API calls recorded this run (content supplied via a pre-built facts file)"
    time_text = f"Generation time: {_format_elapsed(elapsed_seconds)}" if elapsed_seconds is not None else ""
    line = "  |  ".join(t for t in (time_text, cost_text) if t)

    anchor = next((para for para in doc.paragraphs if para.text.strip() == "Overview & Flags"), None)
    if anchor is None:
        # Heading text not found (template changed?) -- fall back to right
        # after the title block rather than silently dropping the line.
        anchor = p[4]
    version_log_p = anchor.insert_paragraph_before("")
    run = version_log_p.add_run(line)
    run.font.size = Pt(9)
    run.italic = True
    _color_run(run, "808080")


def _remove_paragraph(paragraph) -> None:
    """Deletes `paragraph` from the document entirely (not just its text) --
    a suppressed field that still occupies a bulleted-list template slot
    would otherwise render as a dangling empty bullet or a blank line, which
    is exactly the clutter External is trying to remove."""
    p_element = paragraph._p
    p_element.getparent().remove(p_element)


_HEADING_LEVEL_BY_STYLE = {"Heading 1": 1, "Heading 2": 2}


def _remove_empty_section_headings(doc) -> None:
    """Removes a Heading 1/2 paragraph if nothing but blank paragraphs
    and other headings survive between it and the next heading of the
    same or shallower level -- called after _apply_deferred_bullets, so
    gap-only content has already been stripped out by then (e.g. an FSI
    sub-section whose only two fields both turned out to be unresolved
    gaps). Checks for a TABLE as well as paragraph text -- some sections'
    real content lives entirely in a table (Land Identification,
    Directors, Group Companies, ...), which doc.paragraphs alone would
    never see, and wrongly deleting a heading that has real tabular
    content under it would be a serious regression, not a cleanup.

    Runs to a fixed point (repeats until a pass removes nothing): deleting
    an empty child heading (e.g. "Unit / Building Summary" with nothing
    left under it) can itself leave a PARENT heading ("RERA Core Data")
    with nothing left either, only visible on the next pass once the
    child is actually gone."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    while True:
        removed_any = False
        body = doc.element.body
        children = list(body)
        heading_positions = []
        for i, el in enumerate(children):
            if el.tag != qn("w:p"):
                continue
            style = Paragraph(el, doc).style
            if style is not None and style.name in _HEADING_LEVEL_BY_STYLE:
                heading_positions.append((i, _HEADING_LEVEL_BY_STYLE[style.name]))
        for idx, (pos, level) in enumerate(heading_positions):
            end = len(children)
            for later_pos, later_level in heading_positions[idx + 1 :]:
                if later_level <= level:
                    end = later_pos
                    break
            has_content = False
            empty_tables = []
            for el in children[pos + 1 : end]:
                if el.tag == qn("w:tbl"):
                    # A table with only its header row (no data rows --
                    # e.g. facts["blocks"] == [], or every data row was
                    # gap-filtered by _remove_gap_rows) is exactly as
                    # empty as an unresolved-gap paragraph; a table with
                    # 2+ rows has at least one real data row and counts
                    # as content. Deliberately no early `break` here --
                    # a section can have real paragraph content (a note
                    # explaining why the table is empty) sitting right
                    # next to a header-only table shell, and both need
                    # to be seen in the same pass or the table survives
                    # as an orphan just because the paragraph came first.
                    if len(Table(el, doc).rows) > 1:
                        has_content = True
                    else:
                        empty_tables.append(el)
                    continue
                if el.tag == qn("w:p") and Paragraph(el, doc).text.strip():
                    has_content = True
            if not has_content:
                # The heading AND any now-orphaned header-only table in
                # its range both go -- removing just the heading and
                # leaving a table shell behind (a header row floating
                # with no heading above it and no data below it) is
                # exactly the kind of dangling clutter this function
                # exists to prevent, not a lesser version of it.
                for empty_table in empty_tables:
                    body.remove(empty_table)
                body.remove(children[pos])
                removed_any = True
                break  # positions are now stale; restart the scan
            elif empty_tables:
                # The heading survives (real content exists elsewhere in
                # the section), but a header-only table shell alongside
                # that content is still pure clutter -- strip just the
                # table(s), leave the heading and real content in place.
                for empty_table in empty_tables:
                    body.remove(empty_table)
                removed_any = True
                break  # positions are now stale; restart the scan
        if not removed_any:
            break


def _split_into_bullet_clauses(text: str) -> list[str]:
    """Splits one long block of prose into short, standalone bullet-point
    clauses at SENTENCE boundaries only -- a period followed by
    whitespace and a capital letter/digit/quote next (so a decimal or a
    domain like "propertyok.com" is never mistaken for a sentence end).
    Depth-aware like _split_outside_parens, so a period INSIDE a
    parenthetical aside never triggers a false split.

    Deliberately does NOT split on semicolons (an earlier version did,
    and it produced messy, fragmented, hard-to-read bullets -- confirmed
    live: a two-sentence passage like "Bombay HC's search was
    CAPTCHA-gated; e-Daakhil's search was unreachable; these two checks
    remain unknowns." became THREE separate bullets, one of them a bare
    lowercase continuation ("a genuine negative result, not an inability
    to check.") that reads as a disconnected fragment rather than a
    standalone point. A semicolon in this codebase's prose usually joins
    two closely related clauses of the SAME thought (or is a leftover
    dash-replacement from an earlier sanitization pass -- see
    charter_citation_system_fixes memory), not two separate findings, so
    splitting there produces MORE, WORSE bullets, not briefer ones. A
    period is a real, unambiguous sentence boundary; that's the right and
    only place to split.

    This is the "informational pointers, not narrated reasoning" rule
    from the 2026-07-30 Charter-restructure design session -- that
    session designed it for charter_document.py's restructure, which was
    never actually wired into production (see charter_two_builders_gotcha
    memory); this is that same principle finally applied to the real
    production template (_fill_template) instead."""
    pieces = []
    depth = 0
    start = 0
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        if depth == 0 and ch == ".":
            m = re.match(r"\s+(?=[A-Z0-9'\"])", text[i + 1 :])
            if m and not _period_is_abbreviation(text, i):
                pieces.append(text[start : i + 1].strip())
                start = i + 1 + m.end()
                i = start
                continue
        i += 1
    pieces.append(text[start:].strip())
    return [p for p in pieces if p]


# Tokens that end in a period WITHOUT ending a sentence. Without this guard the
# splitter fragments real bullets mid-phrase, because the character after the
# period is a capital or digit and looks exactly like a sentence start:
#   "Dwarkadas J. Sanghvi College"          -> "Dwarkadas J." + "Sanghvi College"
#   "...up to 25 minutes via S.V. Road"     -> "...via S.V."  + "Road/Western..."
#   "Government Regulation No. TPB4315/..." -> "...No."       + "TPB4315/..."
#   "62.24 sq. m. MHADA tit-bit area"       -> "...sq. m."    + "MHADA tit-bit..."
# All four shipped in the 2026-08 Pranami Bliss charter as broken bullets.
_NON_TERMINAL_ABBREVIATIONS = {
    # reference / numbering
    "no", "nos", "sr", "jr", "art", "cl", "para", "pp", "vol", "ed", "fig",
    "sec", "regn", "reg", "ch", "cf", "al", "ors", "anr", "vs", "viz", "etc",
    # units
    "sq", "m", "ft", "mtr", "mtrs", "km", "cm", "kg", "hrs", "approx", "wt",
    # entity / honorific
    "pvt", "ltd", "co", "inc", "corp", "llp", "mr", "mrs", "ms", "dr", "adv",
    "prof", "hon", "smt", "shri", "st", "rd", "rs",
    # organisational
    "govt", "dept", "dist", "opp",
}


def _period_is_abbreviation(text: str, dot_index: int) -> bool:
    """True when the period at `dot_index` closes an abbreviation or an
    initial rather than a sentence.

    Three cases, in order of how often they bite here:
      * a dotted initialism -- "S.V.", "N.A." -- detected by the captured
        token still containing a period;
      * a single-letter initial in a name -- "Dwarkadas J. Sanghvi";
      * a known abbreviation from _NON_TERMINAL_ABBREVIATIONS.

    Deliberately biased toward NOT splitting. A missed split merges two
    sentences into one slightly long bullet; a false split severs a name or a
    regulation number across two bullets and leaves a dangling fragment, which
    CLAUDE.md's "Bullets and grammar" rule calls out specifically. The cost of
    the two failures is not symmetric, so ambiguous tokens ("Ltd.", "Co.")
    stay in the set."""
    m = re.search(r"([A-Za-z]+(?:\.[A-Za-z]+)*)\.$", text[: dot_index + 1])
    if not m:
        return False
    token = m.group(1)
    if "." in token:          # S.V, N.A, i.e -- a dotted initialism
        return True
    if len(token) == 1:       # a personal initial: "J.", "A."
        return True
    return token.lower() in _NON_TERMINAL_ABBREVIATIONS


_MIN_STANDALONE_BULLET_WORDS = 5


def _consolidate_bullet_clauses(clauses: list[str]) -> list[str]:
    """Merges a clause into the previous one when it's too short to read
    as a standalone point (CLAUDE.md Section B: consolidate split/
    fragmented bullets into one well-formed bullet rather than a fragment
    trail) -- _split_into_bullet_clauses correctly splits on every real
    sentence boundary, but a short trailing sentence (e.g. "Confirmed as
    of 2026.") reads as an orphaned fragment, not its own point, when left
    as a separate bullet. The first clause never merges backward (nothing
    to merge into)."""
    if len(clauses) <= 1:
        return clauses
    merged = [clauses[0]]
    for clause in clauses[1:]:
        if len(clause.split()) < _MIN_STANDALONE_BULLET_WORDS:
            merged[-1] = f"{merged[-1]} {clause}"
        else:
            merged.append(clause)
    return merged


def _fix_bullet_capitalization(clause: str) -> str:
    """Deterministic capitalization/punctuation cleanup (CLAUDE.md Section
    B) -- capitalizes the clause's first letter if it isn't already, and
    ensures it ends with terminal punctuation, so a bullet reads as a
    complete sentence rather than a raw fragment. Deliberately NOT full
    grammar correction (an LLM/NLP job, with real risk of silently
    altering a factual claim in a due-diligence document) -- this is the
    safe, deterministic slice of "grammar and capitalization" code can do
    without a human or model re-checking every rewrite. Scoped to bullet
    clauses (narrative prose), never applied to table cells -- a blanket
    pass there would corrupt structured values (dates, CIN/DIN, currency
    figures) that should never gain a forced capital or trailing period."""
    clause = clause.strip()
    if not clause:
        return clause
    if clause[0].islower():
        clause = clause[0].upper() + clause[1:]
    if clause[-1] not in ".!?\"')":
        clause += "."
    return clause


_GAP_PREFIXES = (
    "not confirmed", "not disclosed", "not established", "not applicable",
    "not yet applicable", "not computed", "not stated", "not individually tabulated",
    "not separately disclosed", "not conclusively confirmed", "not independently confirmed",
    "gap", "unknown",
)


def _looks_like_unresolved_gap(text: str) -> bool:
    """True for text reporting an ABSENCE OF INFORMATION (a research
    limitation -- "we don't know"); false for text reporting a confirmed,
    POSITIVE absence of something bad ("we checked and there's no
    litigation/discrepancy/complaint" -- a real, valuable finding, not a
    gap, and never dropped). The distinguishing signal, confirmed against
    every real gap-vs-finding sentence already in this pipeline's own
    facts.json files (2026-08-03 survey of Pranami's and Lavina Estates'):
    a genuine gap starts with "Not ..." ("Not confirmed", "Not disclosed",
    "Not established", "Not applicable", ...), "gap", or "unknown"; a
    positive clean finding is instead phrased "No ..." / "None ..." ("No
    litigation was found", "No discrepancy found", "No address
    discrepancy", "NO MahaRERA registration found" -- this pipeline's own
    single most important finding for an unregistered project) -- which
    always stays. Never used to silently drop a finding; only to skip
    repeating an already-tracked gap (see facts["gaps"]/_classify_flags)
    as a sentence in the main narrative. A literally blank value is
    treated the same as a gap prefix -- an empty table cell is exactly
    the "bare, un-informative row" this check exists to catch, not a
    third category of its own."""
    lowered = text.strip().lower()
    if not lowered:
        return True
    return any(lowered.startswith(prefix) for prefix in _GAP_PREFIXES)


def _remove_gap_rows(table, value_col: int, header_rows: int = 1) -> None:
    """Removes any DATA row (below `header_rows`) whose `value_col` cell
    is an unresolved gap -- same principle as narrative-bullet removal
    (_looks_like_unresolved_gap), applied to a fixed-template Field/Value
    -style table (Land Identification, Corporate Identity, Neighbourhood,
    FSI Metrics) instead of a paragraph. The gap stays tracked once in
    facts["gaps"]; this only stops it from ALSO being repeated as a bare
    "Not disclosed" table row a reader has to scan past.

    Deliberately NOT applied to every table in the document -- a
    Director or Group-Companies table has a DIFFERENT column semantics
    where one column reading "unknown" (e.g. a CIN not on public record)
    does not mean the whole row (a real, named company) is worthless;
    only call this on a table where the checked column genuinely IS the
    row's entire payload."""
    for row in list(table.rows[header_rows:]):
        cell_text = row.cells[value_col].text
        if _looks_like_unresolved_gap(cell_text):
            table._tbl.remove(row._tr)


def _remove_fully_empty_rows(table, header_rows: int = 1) -> None:
    """Removes a DATA row only if EVERY column in it is empty/an
    unresolved gap (CLAUDE.md Section B) -- the multi-column counterpart
    to _remove_gap_rows above, for tables like Group/Affiliated Companies
    where one gap column (e.g. an unconfirmed CIN) does NOT mean the rest
    of the row is worthless. A row survives if even one cell holds a
    real, confirmed value -- e.g. a real company name and a real "shared
    director" basis survive even if the CIN column reads "unknown"."""
    for row in list(table.rows[header_rows:]):
        if all(not cell.text.strip() or _looks_like_unresolved_gap(cell.text) for cell in row.cells):
            table._tbl.remove(row._tr)


# Maps what a clause is ABOUT to the facts["sources"] topic that can support
# it. Ordered: the first pattern that matches wins, so the more specific
# subjects come first. Used only when no field-level source exists -- see
# _set_paragraph_as_bullets. Anything not matched here stays uncited on
# purpose; Section C would rather have a missing marker than a wrong one.
_CLAUSE_TOPIC_PATTERNS = (
    (r"\b(?:CRISIL|ICRA|credit rating|rating rationale|downgrade[ds]?)\b", ("credit_rating",)),
    (r"\b(?:insolvency|IBBI|NCLT|CIRP)\b", ("insolvency_status",)),
    (r"\b(?:title|encumbrance|sale deed|conveyance|lis pendens|MHADA|land|plot|CTS)\b", ("land_title", "legal_documents")),
    (r"\b(?:FSI|floor space|built-up|BUA|sanctioned plan|IOD|DCPR|setback|fungible)\b", ("legal_documents",)),
    (r"\b(?:sold|unsold|units?|carpet area|booking|inventory|completion|quarterly)\b", ("project_registration",)),
    (r"\b(?:director|CIN|incorporat|registered office|shareholding|paid-up|authorized capital)\b", ("company_profile",)),
    (r"\b(?:price|pricing|rate|psf|per sq)\b", ("pricing", "market_trend")),
    (r"\b(?:school|college|hospital|mall|station|metro|connectivity|locality|neighbourhood|located|situated|complex)\b", ("distance", "market_trend")),
    (r"\b(?:complaint|appeal|warrant|litigation)\b", ("project_registration",)),
)


def _clause_topic_citation(facts: dict, clause: str) -> str | None:
    """Resolves ONE clause to the "[N]" of a source that actually supports it,
    by what the clause is about (see _CLAUSE_TOPIC_PATTERNS).

    Returns None when nothing matches, or when no source carries the matched
    topic -- never a fabricated or merely adjacent citation. Section C's own
    worked failure is citing the MahaRERA complaints record for a sentence
    beginning "independent web research found...": both real sources, wrong
    pairing. Leaving that clause uncited is the better failure."""
    if facts.get("_doc_variant") != "external":
        return None

    # A model match, if run_editorial_passes cached one, cites the source that
    # actually establishes THIS clause rather than the first source carrying a
    # keyword-guessed topic. A miss falls through to the keyword table, so no
    # match means today's behaviour exactly.
    matched = facts.get("_claim_source_matches", {}).get((clause or "").strip())
    if matched is not None:
        sources = facts.get("sources") or []
        if 0 <= matched < len(sources):
            label = _clean_source_label(sources[matched].get("label") or "") or (sources[matched].get("ref") or "")
            citation = _citation_text(facts, label)
            if citation:
                return citation

    for pattern, topics in _CLAUSE_TOPIC_PATTERNS:
        if re.search(pattern, clause, re.IGNORECASE):
            return _cite_marker(*topics, facts=facts)
    return None


def _set_paragraph_as_bullets(facts: dict, paragraph, text, *, gap_check_text: str | None = None, citation: str | None = None) -> None:
    """Renders `text` as brief bullet-point pointers instead of one
    flowing paragraph, when it's actually long enough to benefit (splits
    to 2+ clauses via _split_into_bullet_clauses) -- a short value
    (splits to just one clause) is left as an ordinary paragraph, pixel-
    identical to what _set_paragraph_text alone would have produced, and
    updated immediately (no paragraph is inserted, so nothing shifts).

    A genuinely long value is NOT split immediately, though -- it's
    QUEUED on facts["_deferred_bullets"] and only actually inserted by
    _apply_deferred_bullets, called once at the very end of
    _fill_template. Confirmed live why this matters: several things later
    in this same function (the Sources-list fill, section-consolidation,
    the Methodology Note removal) look up doc.paragraphs by a fixed LIVE
    index, not the stable p[N] object list this function itself uses --
    inserting new paragraphs immediately, at an early field like
    Executive Summary, silently shifted every later fixed-index lookup
    and corrupted the Sources list with stale, unrelated template
    placeholder text. Deferring every insertion until after all of that
    has already run avoids the shift entirely.

    Checked WHOLE, before any splitting -- not per-clause. A gap-only
    field is dropped from the narrative entirely (the paragraph itself is
    queued for removal, not left empty); the gap is still tracked once in
    facts["gaps"]/the Monitor flag list, just never repeated as a
    sentence here too. Confirmed live why per-clause filtering doesn't
    work: a real field is often phrased "Not confirmed; no X was found in
    open sources. Y is not established either." -- splitting on the
    semicolon/period separates the "Not confirmed" lead-in from its own
    continuation clauses, and per-clause checking then only recognizes
    the FIRST fragment as a gap, leaving the rest as a dangling,
    context-free bullet ("no X was found in open sources.") that reads as
    broken. Checking gap-ness once, on the whole value, avoids this: if
    the field opens with a gap phrase, the entire thing goes, continuation
    clauses included; if it doesn't, every clause stays (a real finding's
    OWN methodology caveats, e.g. "Bombay HC's search was CAPTCHA-gated,"
    are legitimate supporting context, not something to strip out
    individually).

    `gap_check_text`, if given, is what gets checked for gap-ness INSTEAD
    of `text` -- needed at a call site that prepends a fixed label
    ("Collection Account of the Project (100%): {value}"), since the
    label itself never looks like a gap even when the value is one.

    `citation`, if given, is a bare marker string (see _cite_marker) that
    gets appended to EVERY clause AFTER splitting -- never build the
    citation into `text` yourself and rely on this function to split it.
    Doing that used to be the actual bug here: the marker was glued once
    onto the whole multi-sentence blob before _split_into_bullet_clauses
    ran, so it landed on whichever clause happened to be last, and every
    earlier clause rendered with no citation at all."""
    text = str(text)
    check_text = str(gap_check_text) if gap_check_text is not None else text
    if _looks_like_unresolved_gap(check_text):
        facts.setdefault("_deferred_bullets", []).append((paragraph, []))
        return

    clauses = _split_into_bullet_clauses(text)
    clauses = _consolidate_bullet_clauses(clauses)
    clauses = [_fix_bullet_capitalization(c) for c in clauses]
    if citation and clauses:
        clauses = [f"{clause} {citation}" for clause in clauses]
    elif clauses and facts.get("_doc_variant") == "external":
        # No single source covers this field, so resolve per clause instead of
        # blanket-citing. CLAUDE.md Section C is explicit that a marker "must
        # actually support the claim" -- attaching one field-level source to
        # eight Executive Summary bullets making eight different claims is the
        # exact mis-citation the rule names, and a wrong marker is worse than
        # none. _clause_topic_citation returns None when it cannot tell, which
        # leaves the clause uncited rather than mis-cited.
        clauses = [
            f"{clause} {marker}" if (marker := _clause_topic_citation(facts, clause)) else clause
            for clause in clauses
        ]
    if len(clauses) <= 1:
        _set_paragraph_text(paragraph, clauses[0] if clauses else text)
        return
    facts.setdefault("_deferred_bullets", []).append((paragraph, clauses))


def _apply_bullet_hanging_indent(paragraph) -> None:
    """A "•" text character glued to the front of a paragraph is NOT a
    bulleted list on its own -- with no hanging indent, a wrapped long
    sentence flows all the way back to the left margin under the bullet,
    which reads as an ordinary paragraph that happens to start with a
    dot, not a genuine short pointer. Confirmed live: this is exactly
    what the user saw and flagged as "no brief pointers," even though the
    text itself had already been split into short clauses. left_indent
    pulls the WHOLE paragraph in; a negative first_line_indent of the
    same magnitude pulls just the first line (the one starting with •)
    back out to the margin -- the standard hanging-indent bullet shape,
    so a wrapped second line aligns under the text, not under the bullet."""
    from docx.shared import Inches
    pf = paragraph.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.25)


def _apply_justify(paragraph) -> None:
    """"Aligned from both sides" -- flush against both margins, wrapped
    lines stretched to fill the full width (Word's own "Justify" rule
    only stretches lines that actually wrap; a short one-line heading or
    label is untouched, so this is safe to apply broadly to narrative
    body text without disturbing single-line content)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _justify_body_paragraphs(doc) -> None:
    """Document-wide pass: every top-level body paragraph (never table
    cells -- those get centered instead, see _center_all_table_cells)
    that has NO explicit alignment already set gets justified. Skipping
    paragraphs with an explicit alignment (e.g. a centered cover title or
    section badge) avoids the one real regression risk: Word's "last
    line" of a justified paragraph always renders flush-left, so
    overwriting an intentionally-centered single-line heading with
    JUSTIFY would visibly shift it left."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        if paragraph.paragraph_format.alignment is not None:
            continue
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# A cell longer than this carries narrative prose rather than a value, and is
# left-aligned instead of centered -- centered prose is ragged on both edges
# and unreadable (CLAUDE.md Section A, "Alignment", amended 2026-08-10). Tuned
# so dates, names, statuses, figures and short phrases still center, while
# cells like Landowner/Investor, the Board Resolution row and the Mortgage row
# (all multi-sentence) go left.
_LONG_CELL_CHARS = 110


def _center_all_table_cells(doc) -> None:
    """"Center in the table" -- every table in the document (Field/Value
    tables, Director/Group-Companies tables, scoring tables, all of it)
    gets its cell content centered, EXCEPT cells carrying a paragraph of
    narrative prose, which are left-aligned (see _LONG_CELL_CHARS).
    Table-cell writes are scattered across dozens of call sites in this file
    (_set_cell/_set_row_cell plus many direct `row.cells[i].text = ...`
    assignments), so this runs as a single blanket pass over the finished
    tables rather than needing every call site touched individually.

    Header rows always center regardless of length -- a long header label is
    still a label, not prose."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                is_prose = (row_idx > 0
                            and len(cell.text.strip()) > _LONG_CELL_CHARS)
                target = (WD_ALIGN_PARAGRAPH.LEFT if is_prose
                          else WD_ALIGN_PARAGRAPH.CENTER)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = target


def _apply_table_pagination(doc) -> None:
    """Stops tables breaking mid-row across a page boundary, and repeats the
    header row at the top of each continuation page (CLAUDE.md Section A,
    "Table pagination"). Runs as a blanket pass alongside
    _center_all_table_cells, for the same reason: table construction is
    spread across too many call sites to set this at each one."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            if not trPr.findall(qn('w:cantSplit')):
                trPr.append(OxmlElement('w:cantSplit'))
            if row_idx == 0 and not trPr.findall(qn('w:tblHeader')):
                hdr = OxmlElement('w:tblHeader')
                hdr.set(qn('w:val'), 'true')
                trPr.append(hdr)


def _apply_deferred_bullets(facts: dict) -> None:
    """Actually inserts the bulleted paragraphs queued by
    _set_paragraph_as_bullets, or removes the paragraph entirely if every
    clause turned out to be an unresolved gap (an empty `clauses` list).
    Must run as the LAST content change in _fill_template, after every
    fixed-index doc.paragraphs[...] lookup elsewhere in that function has
    already happened -- see _set_paragraph_as_bullets' own docstring for
    why.

    Confirmed live: roughly half of this template's target paragraph
    slots (e.g. p[41]/p[42]/.../p[51], the RERA Compliance/Local Planning
    fields) already carry native Word list numbering ("List Paragraph"
    style with a real numPr) baked in at template-authoring time --
    unconditionally prepending a manual "• " text character produced a
    visible DOUBLE bullet ("•  • NO MahaRERA registration found.") on
    those specific slots. Detected here instead of assumed: a paragraph
    that already has numPr gets its clauses split with NO added "• "
    text and NO manual hanging-indent (Word's own list style already
    supplies both, correctly, via the numbering definition); a plain
    paragraph with no numPr keeps this session's original text-based
    "• " + explicit hanging-indent treatment. Same pattern
    _fill_variable_paragraphs already uses for exactly this reason on its
    own overflow rows."""
    for paragraph, clauses in facts.pop("_deferred_bullets", []):
        if not clauses:
            _remove_paragraph(paragraph)
            continue
        num_pr = None
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            num_pr = paragraph._p.pPr.numPr

        bullets = clauses if num_pr is not None else [f"• {c}" for c in clauses]
        _set_paragraph_text(paragraph, bullets[0])
        if num_pr is None:
            _apply_bullet_hanging_indent(paragraph)
        last = paragraph
        for extra in bullets[1:]:
            new_para = last.insert_paragraph_before("", style=paragraph.style)
            last._p.addnext(new_para._p)
            if num_pr is not None:
                new_para._p.get_or_add_pPr().append(copy.deepcopy(num_pr))
            else:
                _apply_bullet_hanging_indent(new_para)
            _set_paragraph_text(new_para, extra)
            last = new_para


# --- doc_variant consolidation primitives --------------------------------
# A whole session's worth of "if doc_variant == external: SHORT else: LONG"
# blocks accumulated across _fill_template and the _append_*_section
# functions -- each written by hand, each a fresh chance to bypass
# _externalize_prose or the paragraph-removal convention by accident (two
# real bugs this session came from exactly that: a KPI-card cell and a
# Documentation-Confidence score run that skipped the usual helper and
# wrote text some other way). These three primitives are the ONE path every
# variant-aware call site should now go through, so a future field is
# structurally unable to bypass them the way those two did.

def _variant(facts: dict, internal_value, external_value):
    """Returns `external_value` for the External variant, `internal_value`
    otherwise. The one place doc_variant branching happens for a plain
    value -- every other helper below is built on this."""
    return external_value if facts.get("_doc_variant") == "external" else internal_value


def _variant_paragraph(doc, facts: dict, internal_text: str, external_text: str):
    """Appends one paragraph with different text per variant. Replaces the
    ~8 scattered "if variant == external: doc.add_paragraph(SHORT) else:
    doc.add_paragraph(LONG)" blocks this session added to the Credit
    Rating/Insolvency/Company Registration/Group Companies/Developer
    Score/Documentation Confidence sections -- always routes through
    doc.add_paragraph, which is itself monkey-patched to run
    _externalize_prose for External (see _fill_template), so this can't
    silently skip that step the way a stray direct .add_run() could."""
    return doc.add_paragraph(_variant(facts, internal_text, external_text))


def _variant_sep(facts: dict) -> str:
    """The ":" vs " -- " choice repeated at every dynamic-value call site
    (a score, a grade, a CIN) where the surrounding dash can't be reached
    by the literal-text _EXTERNAL_DASH_REWRITES dict because a variable
    sits right where the dash does."""
    return _variant(facts, " --", ":")


def _fix_bullet_hanging_indent(doc) -> None:
    """The template's own numbering.xml defines the bullet list every
    "List Paragraph" in this document uses (abstractNum id 2, referenced by
    numId 2) with NO indent at all on its level-0 definition -- unlike a
    second, unused abstractNum (id 1) that correctly sets a 720/360
    hanging indent. Word's fallback for a bullet with no indent puts the
    bullet glyph and first line at the left margin with only a tab before
    the text, but a wrapped second line has nothing to hang from and falls
    all the way back to the left margin -- the "text runs off alignment"
    look throughout every bulleted sub-section. Both doc variants get this
    fix (it's a genuine template defect, not a content/scope difference)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    numbering_root = doc.part.numbering_part.element
    for abstract_num in numbering_root.findall(qn("w:abstractNum")):
        if abstract_num.get(qn("w:abstractNumId")) != "2":
            continue
        lvl0 = abstract_num.find(qn("w:lvl"))
        if lvl0 is None or lvl0.find(qn("w:pPr")) is not None:
            continue
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl_jc = lvl0.find(qn("w:lvlJc"))
        if lvl_jc is not None:
            lvl_jc.addnext(p_pr)
        else:
            lvl0.append(p_pr)


# Mirrors _TEXT_RED/_TEXT_AMBER/_TEXT_GREEN (defined later in this file,
# alongside the other document-styling constants) -- duplicated here as
# literals rather than referenced, since this set is built at module-import
# time, before those later assignments have run. "1F3864" is the template's
# own Heading style color (navy), applied to every section heading in both
# variants. "375623" (_TEXT_GREEN) completes the same red/amber/green
# traffic-light convention as a positive-confirmation color in
# charter_document.py's "What Checks Out" section, which didn't exist when
# this gate was first written -- none of these are the grey-placeholder bug
# this gate exists to catch.
_EXTERNAL_ALLOWED_RUN_COLORS = {None, "000000", "C00000", "BF8F00", "1F3864", "375623"}


def _iter_all_paragraphs(doc):
    """Yields every paragraph in the document body plus every table cell
    (recursively, in case a cell itself contains a nested table) -- the
    verification gate below needs to see ALL rendered text, not just
    doc.paragraphs (which skips table contents entirely)."""
    def _from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from _from_tables(cell.tables)

    yield from doc.paragraphs
    yield from _from_tables(doc.tables)


_FACTUAL_CLAIM_MARKER_RE = re.compile(
    r"\b(19|20)\d{2}\b"                                    # a year
    r"|\bDIN\s*\d{7,8}\b"                                   # a DIN
    r"|\b[UL]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b"             # a CIN
    r"|\bP\d{11}\b"                                         # a MahaRERA registration number
    r"|\b(?:Rs\.?|INR|₹)\s?[\d,]+"                     # a rupee figure
    r"|\b[\d,]+\s*(?:crore|lakh|sq\.?\s?ft|units?)\b"       # a scale figure
    r"|\b(?:High Court|Supreme Court|NCLT|NCLAT|Sessions Court|Tribunal)\b",
    re.IGNORECASE,
)
# Lines that legitimately state a fact with no [N] marker because they're
# not a sourced claim at all -- a plain absence statement (a genuine
# finding, not a gap -- see _looks_like_unresolved_gap's own "No X was
# found" vs "Not confirmed" distinction), a Developer Score/Documentation
# Confidence line (code-computed from facts already cited elsewhere, not
# an independent claim needing its own source), or a standing/gap notice.
# Matched as a lowercase substring.
_CITATION_EXEMPT_SUBSTRINGS = (
    "nothing found", "none identified", "no additional material gaps",
    "no additional gaps identified", "developer score", "documentation confidence",
    "not independently confirmed", "not disclosed", "not confirmed", "not established",
    "not applicable", "deep market research", "prepared ",
)
_ABSENCE_FINDING_RE = re.compile(r"^\W*no\b", re.IGNORECASE)
# Short lines (cover-page titles/subtitles, table-adjacent labels) aren't
# claims a reader parses as prose needing a citation -- only worth
# checking once there's enough text to actually read as a sentence.
_MIN_CLAIM_WORDS = 6


def _looks_like_uncited_factual_claim(text: str) -> bool:
    """True when `text` states something specific enough (a year, a CIN/
    DIN, a registration number, a rupee figure, a scale figure, or a named
    court/tribunal) to be a factual claim a reader should be able to trace
    to a source, but carries no "[N]" citation marker anywhere in it.
    Deliberately pattern-based (fast, deterministic, no API call) --
    see _llm_verify_citation_completeness for the judgment-based
    companion pass that catches claims with no fixed pattern to match."""
    stripped = text.strip()
    if not stripped or "[" in stripped:
        return False
    if len(stripped.split()) < _MIN_CLAIM_WORDS:
        return False
    if _ABSENCE_FINDING_RE.match(stripped.lstrip("•").strip()):
        return False
    lowered = stripped.lower()
    if any(s in lowered for s in _CITATION_EXEMPT_SUBSTRINGS):
        return False
    return bool(_FACTUAL_CLAIM_MARKER_RE.search(stripped))


def _check_citation_completeness(docx_path: str) -> list[str]:
    """Re-opens a just-saved External Charter and flags any body paragraph
    that reads as an uncited factual claim (see
    _looks_like_uncited_factual_claim) -- the CLAUDE.md Section C rule
    that every inline factual claim must resolve to a numbered "[N]"
    marker. Deliberately kept SEPARATE from _verify_external_document_
    quality's hard-fail violations: a missing citation marker can reflect
    a genuine, pre-existing gap in the underlying facts' source tagging
    (real project data, not a code bug), and that shouldn't be able to
    sink an otherwise-good document out of existing entirely the way a
    deterministic rendering regression (a stray dash, a lost bullet)
    should. Callers should log this, not raise on it -- same advisory
    treatment as _llm_verify_citation_completeness, its judgment-based
    companion check."""
    import docx as _docx

    doc = _docx.Document(docx_path)
    flags = []
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if re.match(r"^\[(\d+)\]\s+\S", stripped):
            continue  # a Sources-list entry itself, not a claim needing one
        style_name = para.style.name if para.style is not None else ""
        if style_name in _HEADING_LEVEL_BY_STYLE:
            continue
        if _looks_like_uncited_factual_claim(stripped):
            flags.append(f"factual-looking claim with no citation marker: {stripped[:100]!r}")
    return flags


# The External document is now allowed its own Document Library section -- an
# opened-and-extracted-only table, agreed 2026-08-10 -- so the old blanket ban
# on the string "Document Library" would hard-fail every External save. What
# must still never appear is the INTERNAL table's per-document status column,
# which narrates what was NOT opened ("Not opened this pass -- not a
# high-priority legal/regulatory document type") and is exactly the
# absence-reporting CLAUDE.md Section B deletes. This targets that text, not
# the heading.
_EXTERNAL_DOC_LIBRARY_LEFTOVER_RE = re.compile(
    r"Not opened this pass|another document already opened under this same shared label",
    re.IGNORECASE,
)


def _verify_external_document_quality(docx_path: str) -> list[str]:
    """Re-opens a just-saved External Charter and checks for the exact
    regressions this session found and fixed by hand (grey/italic body
    text, hyphen-pair "dashes", the Weight column, the Document Library
    section, missing bullet numbering, and dangling-paren citations) --
    not a general style linter, just a guard against these specific bugs
    coming back the next time someone edits _fill_template. Returns a
    list of violation strings; empty means clean. Raises nothing itself --
    the caller decides how loudly to fail. See _check_citation_completeness
    for the separate, advisory-only citation-marker check."""
    import re
    import docx as _docx

    doc = _docx.Document(docx_path)
    violations = []

    for para in _iter_all_paragraphs(doc):
        text = para.text
        if " -- " in text:
            violations.append(f"hyphen-pair dash found in paragraph: {text[:80]!r}")
        if "—" in text:  # em dash
            violations.append(f"em dash found in paragraph: {text[:80]!r}")
        if _EXTERNAL_DOC_LIBRARY_LEFTOVER_RE.search(text):
            violations.append(f"leftover Internal Document Library status text survived in External: {text[:80]!r}")
        if _is_cited_absence(text):
            # CLAUDE.md Section B: "Never attach a citation marker to an
            # absence." A citation spent establishing that there is nothing to
            # report is the exact defect the clean-check rule exists to
            # prevent, and unlike the rule's softer cases it is unambiguous --
            # so it hard-fails the save rather than merely being scrubbed.
            violations.append(f"citation marker attached to an absence: {text[:80]!r}")

        for run in para.runs:
            color = run.font.color
            rgb = str(color.rgb) if color is not None and color.rgb is not None else None
            if run.italic and rgb != "C00000":
                # The one deliberate exception: the Standing Gap paragraph
                # in the Gaps & Sources section is colored red AND italic on
                # purpose -- red pairs only with that intentional styling,
                # never with the grey-placeholder bug this check exists to
                # catch.
                violations.append(f"italic run (leftover template placeholder styling) in: {text[:80]!r}")
            if rgb not in _EXTERNAL_ALLOWED_RUN_COLORS:
                violations.append(f"unexpected run color {rgb} (expected black/red/amber/heading-navy only) in: {text[:80]!r}")

    # Citation-entry checks (dangling parens, missing bullet) apply only to
    # the real Sources list in the document body -- NOT to every table cell
    # that happens to hold a bare "[N]" marker (e.g. a "Source" column),
    # which was never a bulleted list item to begin with.
    for para in doc.paragraphs:
        stripped = para.text.strip()
        m = re.match(r"^\[(\d+)\]\s+\S", stripped)
        if m:
            if stripped.count("(") != stripped.count(")"):
                violations.append(f"unbalanced parentheses in citation entry: {stripped[:80]!r}")
            p_pr = para._p.pPr
            if p_pr is None or p_pr.numPr is None:
                violations.append(f"citation entry lost its bullet numbering: {stripped[:80]!r}")

    for table in doc.tables:
        if not table.rows:
            continue
        header_texts = [c.text.strip() for c in table.rows[0].cells]
        if header_texts[:2] == ["Bucket", "Sub-metric"] and "Weight" in header_texts:
            violations.append("Developer Score table still has a Weight column in External")
        if header_texts[:2] == ["Bucket", "Criterion"] and "Weight" in header_texts:
            violations.append("Documentation Confidence table still has a Weight column in External")

    from docx.oxml.ns import qn
    numbering_root = doc.part.numbering_part.element
    for abstract_num in numbering_root.findall(qn("w:abstractNum")):
        if abstract_num.get(qn("w:abstractNumId")) != "2":
            continue
        lvl0 = abstract_num.find(qn("w:lvl"))
        if lvl0 is None or lvl0.find(qn("w:pPr")) is None:
            violations.append("bullet numbering (abstractNum id=2) is missing its hanging-indent fix")

    return violations


_CITATION_JUDGE_SCHEMA = {
    "type": "object",
    "properties": {
        "uncited_claims": {
            "type": "array",
            "items": {"type": "string"},
        },
    },
    "required": ["uncited_claims"],
}

_CITATION_JUDGE_INSTRUCTIONS = f"""You are reviewing the rendered body text of an External \
Company Charter for citation completeness, per the numbered-citations rule above. You will be \
given the document's paragraphs/bullets, one per line. For each line that states a specific \
factual claim (a name, a date, a number, a registry/registration status, or a court/tribunal \
finding) with NO "[N]"-style citation marker anywhere in it, add its exact verbatim text to \
`uncited_claims`. Do NOT flag: headings, the Sources list itself (lines already starting with \
"[N]"), Developer Score / Documentation Confidence Summary lines (these are code-computed from \
facts already cited elsewhere in the document, not independent claims needing their own \
source), or a plain "Nothing found" / gap / not-disclosed statement (an absence, not a claim).

Your FINAL reply must be ONLY a single raw JSON object -- no prose, no markdown code fences -- \
matching exactly this JSON Schema: {json.dumps(_CITATION_JUDGE_SCHEMA)}"""


def _llm_verify_citation_completeness(docx_path: str) -> list[str]:
    """The Task-8 gate-check: a genuine LLM judgment pass over the rendered
    External Charter's body text, run IN ADDITION to (never instead of)
    the mechanical regex heuristic in _verify_external_document_quality.
    The regex only catches claims matching a known fixed pattern (a year,
    a CIN, a rupee figure...); this catches claims a human editor would
    recognize as needing a source but that match no fixed pattern (e.g. a
    named individual's role, a specific allegation, a dated event
    described in prose without a numeral). Gets the CLAUDE.md Section B +
    Section C system-prompt blocks (external doc_variant only) via
    _charter_system_blocks -- this is the one call in this file that
    exists specifically because doc_variant == "external".

    Degrades to an empty list on ANY failure (missing ANTHROPIC_API_KEY,
    network error, bad JSON reply) rather than raising -- same philosophy
    as deep_research._verify_claim's verification_error handling; an LLM
    judgment pass is advisory, not something a flaky network call should
    be able to discard an otherwise-good run over. Callers should log
    (not silently drop) whatever this returns."""
    import docx as _docx

    doc = _docx.Document(docx_path)
    body_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
    if not body_text:
        return []

    system = _charter_system_blocks(external=True, extra=_CITATION_JUDGE_INSTRUCTIONS)
    prompt = f"External Charter body text (one paragraph/bullet per line):\n\n{body_text}"
    try:
        result = deep_research._run_agentic_pass(prompt, system, label="citation_completeness_judge")
    except Exception as e:
        return [f"(citation-completeness judge could not run this pass -- {e})"]
    claims = result.get("uncited_claims", []) if isinstance(result, dict) else []
    return [str(c) for c in claims if str(c).strip()]


def _set_cell(table, row: int, col: int, text: str) -> None:
    # _set_paragraph_text_raw, not _set_paragraph_text -- table cells (the
    # Corporate/Promoter Identity, Directors, Company Registration Profile
    # tables among them) are exactly where CIN/DIN is REQUIRED, so cell
    # writes must never go through the prose CIN/DIN stripper.
    _set_paragraph_text_raw(table.rows[row].cells[col].paragraphs[0], str(text))


def _set_row_cell(row, col: int, text: str) -> None:
    _set_paragraph_text_raw(row.cells[col].paragraphs[0], str(text))


def _fill_variable_rows(table, header_rows: int, items: list, fill_row) -> None:
    """Grows/shrinks table to exactly len(items) data rows below
    header_rows, calling fill_row(row, item) for each new/existing row.
    Extra template rows are deleted from the end; missing ones are added
    with table.add_row() (already in correct trailing order)."""
    while len(table.rows) - header_rows < len(items):
        table.add_row()
    while len(table.rows) - header_rows > len(items):
        table._tbl.remove(table.rows[-1]._tr)

    for i, item in enumerate(items):
        fill_row(table.rows[header_rows + i], item)


def _fill_variable_paragraphs(doc, start_index: int, slot_count: int, texts: list) -> None:
    """Fills the slot_count existing template paragraphs starting at
    start_index with the first items of `texts`; deletes unused slots if
    there are fewer, or appends new paragraphs (cloning the last slot's
    style) if there are more."""
    slots = doc.paragraphs[start_index:start_index + slot_count]

    for i in range(min(slot_count, len(texts))):
        _set_paragraph_text(slots[i], texts[i])
    for i in range(len(texts), slot_count):
        slots[i]._element.getparent().remove(slots[i]._element)

    # `style=` only sets pStyle -- it does NOT carry over the bullet's own
    # <w:numPr> (ilvl/numId), which lives separately in pPr. Past slot_count
    # items were rendering as plain unbulleted, unindented paragraphs for
    # exactly that reason. Captured once, before the loop reassigns `last`,
    # so every overflow paragraph gets the SAME numbering as the real last
    # template slot, not whatever the immediately-preceding new paragraph
    # ended up with.
    last = slots[-1] if slots else None
    num_pr = None
    if last is not None and last._p.pPr is not None and last._p.pPr.numPr is not None:
        num_pr = last._p.pPr.numPr

    for text in texts[slot_count:]:
        if last is None:
            break
        new_para = last.insert_paragraph_before(text, style=last.style)
        if num_pr is not None:
            new_para._p.get_or_add_pPr().append(copy.deepcopy(num_pr))
        last._p.addnext(new_para._p)
        last = new_para


def _convert_docx_to_pdf(docx_path: str) -> str | None:
    """Converts a just-saved .docx to a same-named .pdf via docx2pdf (Word
    COM automation on Windows/macOS) -- PDF is the actual final deliverable
    format (2026-08-04): a due-diligence document should hand off as
    something any recipient can open faithfully without Word installed,
    not a .docx whose fonts/template styles can render differently
    elsewhere. Returns the PDF path, or None on any failure (no Word
    installed, a COM error, an unsupported platform) -- degrades to
    leaving the .docx as the only output for that file rather than
    crashing the whole generation run; the caller logs this, never
    treats a missing PDF as fatal."""
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except Exception as e:
        print(f"[!] Could not convert {docx_path} to PDF ({e}) -- .docx is the only output for this file.")
        return None


# ---------------------------------------------------------------------------
# Clean-check scrubber -- CLAUDE.md Section B, "A clean check produces no
# sentence": when a check ran and came back clear, the document says nothing
# at all about it. The prompt asks for this on every generating call, but
# model compliance is probabilistic, so this is the deterministic layer.
#
# Two independent guards, because shape alone provably cannot do the job.
# These two sentences are grammatically identical:
#     "No litigation is disclosed against the promoter."       -> delete
#     "Not found among the documents reviewed: an FSI certificate." -> KEEP
# The first reports that a RISK is absent (reassuring, and the rule's whole
# target). The second reports that EVIDENCE is absent -- a gap, which the same
# section says to keep in full and "never compress or delete under the
# clean-check rule". So a clause is only scrubbed when BOTH hold: it sits in a
# field on the reviewed list below, AND the thing it reports missing is a risk.
# A field not on the list is still covered by the prompt layer and by
# _verify_external_document_quality's own absence-with-a-citation check.
# ---------------------------------------------------------------------------

# Narrative fields audited against the real Pranami Bliss / IRA Insignia facts
# and confirmed to carry clean-check prose. Dotted path into `facts`.
_CLEAN_CHECK_FIELDS = (
    "executive_summary",
    "address_discrepancy_note",
    "litigation_status.value",
    "rera_core_fields.litigations_per_record",
    "rera_compliance.litigations_complaints_appeals",
    "fsi_metrics.mortgage_area",
    "fsi_metrics.mortgage_lender.value",
)

# Things whose ABSENCE is a reassuring finding rather than an open unknown.
# The distinction that makes this list necessary is in the block comment above:
# a missing encumbrance is a clean check, a missing certificate is a gap.
_RISK_NOUNS = (
    "litigation", "litigations", "complaint", "complaints", "appeal", "appeals",
    "warrant", "warrants", "encumbrance", "encumbrances", "charge", "charges",
    "mortgage", "mortgages", "lien", "liens", "default", "defaults", "dispute",
    "disputes", "discrepancy", "discrepancies", "proceeding", "proceedings",
    "insolvency", "adverse",
)

# Structural absence shapes. Deliberately not anchored to any topic -- the
# topic test is _RISK_NOUNS, applied separately.
_ABSENCE_SHAPES = (
    r"\bno\s+[\w\s,'()/-]{0,60}?\b(?:found|disclosed|recorded|reported|identified|registered|filed|raised|pending|on record|against)\b",
    r"\b(?:found|returned|revealed|showed|turned up)\s+(?:nothing|no\s+\w+)\b",
    r"\b(?:records?|register|registers|search|searches|filings?|data)\b[\w\s,'()/-]{0,40}\b(?:are|is|was|were)\s+(?:all\s+)?empty\b",
    r"\bnone\b",
    r"\bno\s+(?:%s)\b" % "|".join(_RISK_NOUNS),
    r"\bnot\s+(?:found|disclosed|recorded|reported|identified)\b",
)

# Any one of these means the clause is doing more than reporting a nothing, so
# it survives regardless. "except"/"other than" introduce the actual finding
# ("...found nothing except a Notice of Lis Pendens"); a calendar date means
# the clause names a specific real event, which a pure absence never does.
_ABSENCE_RESCUE_RE = re.compile(
    r"\b(?:except|other than|save for|apart from|aside from|barring|besides|with the exception of)\b"
    r"|\b(?:\d{1,2}\s+)?(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b"
    r"|\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b"
    # ISO form too. Omitting it made the gate reject two real IRA Insignia
    # paragraphs that cite dated MahaRERA orders ("Order No. 35 of 2022 (dated
    # 2022-08-12)") -- a clause naming a specific dated instrument is a
    # finding, which is the whole principle this rescue encodes.
    r"|\b\d{4}-\d{2}-\d{2}\b",
    re.IGNORECASE,
)


def _matches_absence_shape(text: str) -> bool:
    """Cheap prefilter: does this clause assert that something is absent?

    Deliberately over-inclusive. It answers "is this worth judging at all",
    not "should this be deleted", so it costs nothing to let a gap through
    here: the risk-noun test or the model verdict decides that afterwards."""
    t = (text or "").strip()
    if not t:
        return False
    return any(re.search(p, t, re.IGNORECASE) for p in _ABSENCE_SHAPES)


def _is_clean_check_clause(text: str, risk_context: bool = False, facts: dict | None = None) -> bool:
    """True when `text` is a bare report that a risk was looked for and not
    found -- the sentence CLAUDE.md Section B deletes outright.

    `risk_context` lets a caller supply the risk noun from the field name when
    the clause itself is too terse to carry one: fsi_metrics.mortgage_lender's
    entire value is "None disclosed.", which is only interpretable as a clean
    check because of the field it sits in."""
    t = (text or "").strip()
    if not t:
        return False
    if not any(re.search(p, t, re.IGNORECASE) for p in _ABSENCE_SHAPES):
        return False

    # A model verdict, if run_editorial_passes cached one for this exact clause,
    # decides the question the risk-noun list only approximates: is the absent
    # thing a RISK (delete) or EVIDENCE (a gap, keep)? A miss falls through to
    # the keyword path below, so no verdict means today's behaviour exactly.
    verdict = (facts or {}).get("_clean_check_verdicts", {}).get(t)
    if verdict:
        return verdict == "clean_check"

    if not risk_context and not re.search(r"\b(?:%s)\b" % "|".join(_RISK_NOUNS), t, re.IGNORECASE):
        return False
    return not _ABSENCE_RESCUE_RE.search(t)


def _strip_absence_tail(clause: str, risk_context: bool = False, facts: dict | None = None) -> str:
    """Drops semicolon-delimited segments that are themselves pure clean checks,
    from a clause that survives as a whole because it also carries a finding.

    Section B's own worked example of what to delete is exactly this shape:
    "No litigation found...; MahaRERA's complaint, appeal and warrant records
    are empty[18]; the Title Report's 30-year search... returned nothing[1]".
    Bolting that tail onto a real finding is how it survives clause-level
    scrubbing, e.g. the Executive Summary's "...the Title Report (18 April
    2024) opines the land title is clear...; no litigation was found... and
    MahaRERA's records are empty."

    Only semicolons are considered. _split_into_bullet_clauses deliberately
    refuses to split on them (it fragments real bullets), and that reasoning
    holds for rendering -- but a semicolon IS a reliable boundary for deciding
    whether a segment is self-contained enough to test on its own. The leading
    segment is never dropped: it carries the finding the clause was kept for,
    and a bullet that opens mid-thought is worse than one that runs long."""
    if ";" not in clause:
        return clause
    segments = [s.strip() for s in clause.split(";")]
    kept = [segments[0]] + [
        s for s in segments[1:]
        if not _is_clean_check_clause(s.rstrip("."), risk_context=risk_context, facts=facts)
    ]
    if len(kept) == len(segments):
        return clause
    rejoined = "; ".join(s for s in kept if s).rstrip(" ;,")
    return rejoined if rejoined.endswith(".") else rejoined + "."


def _resolve_path(facts: dict, path: str):
    """Returns (container, key) for a dotted path, or (None, None) if any step
    is missing or is not a dict. Never creates anything."""
    node = facts
    parts = path.split(".")
    for part in parts[:-1]:
        if not isinstance(node, dict) or part not in node:
            return None, None
        node = node[part]
    if not isinstance(node, dict) or parts[-1] not in node:
        return None, None
    return node, parts[-1]


def _scrub_clean_checks(facts: dict) -> dict:
    """Deletes clean-check clauses from the reviewed narrative fields, in place.

    Returns {path: original_text} for everything it changed, so the caller can
    put the full text back before persisting. That restoration is not optional:
    CLAUDE.md Section B is explicit that the scope of what was checked "stay[s]
    in the facts file; they do not reach the page" -- this removes them from
    the PAGE only, never from the record.

    Idempotent, so rendering both variants from one dict is safe: a second pass
    finds nothing left to cut and reports no further changes."""
    changed = {}
    for path in _CLEAN_CHECK_FIELDS:
        container, key = _resolve_path(facts, path)
        if container is None:
            continue
        original = container[key]
        if not isinstance(original, str) or not original.strip():
            continue

        risk_context = any(noun in path.lower() for noun in _RISK_NOUNS)
        kept = []
        for clause in _split_into_bullet_clauses(original):
            if _is_clean_check_clause(clause, risk_context=risk_context, facts=facts):
                continue
            kept.append(_strip_absence_tail(clause, risk_context, facts=facts))
        scrubbed = " ".join(c for c in kept if c).strip()
        if scrubbed != original.strip():
            changed[path] = original
            # Emptied entirely: every clause was a clean check. The field's own
            # empty-value handling (_remove_gap_rows / the empty-section rule)
            # takes it from here rather than printing a blank bullet.
            container[key] = scrubbed
    return changed


# ---------------------------------------------------------------------------
# Facts normalization -- three specific misfilings, corrected before render.
#
# Unlike the clean-check scrubber above, these are NOT reversed before the
# facts file is persisted. A scrub removes information from the page and the
# record must keep it; a relocation moves the same text to the field it always
# belonged in, so nothing is lost and the corrected record is the better one.
# All three are idempotent: once moved, the source no longer matches.
# ---------------------------------------------------------------------------

# "The two 6 March 2024 orders sometimes referenced as a 'merge order' are
# administrative Deputy Registrar of Co-operative Societies actions
# (amalgamating ... under the Maharashtra Co-operative Societies Act, 1960),
# not litigation or a tribunal proceeding."
_MERGE_ORDER_SENTENCE_RE = re.compile(
    r"[^.]*\bDeputy Registrar of Co-?operative Societies\b[^.]*\.(?:\s|$)",
    re.IGNORECASE,
)


def _relocate_merge_orders(facts: dict) -> bool:
    """Moves society-amalgamation orders out of litigation_status.

    CLAUDE.md Section B, "Ruled-out items survive, in the right section",
    names this exact case: "Deputy Registrar society-amalgamation orders are
    land assembly, not litigation - they belong under Land Identification".
    Sitting in litigation_status, they existed only to be denied ("...not
    litigation or a tribunal proceeding"); under Land Identification they read
    as what they are, which is how the two plots became one."""
    litigation = facts.get("litigation_status")
    if not isinstance(litigation, dict):
        return False
    value = litigation.get("value") or ""
    match = _MERGE_ORDER_SENTENCE_RE.search(value)
    if not match:
        return False

    # Strip a leading closing quote/bracket: the preceding sentence can end
    # "...on the said Land.'", so the match legitimately begins after that
    # period but before the quote that closes it.
    sentence = match.group(0).strip().lstrip("'\"’”) ").strip()
    litigation["value"] = (value[: match.start()] + value[match.end():]).strip()

    # Reframed on arrival: the trailing "not litigation or a tribunal
    # proceeding" only made sense as a denial in the section it just left.
    sentence = re.sub(r",?\s*not litigation or a tribunal proceeding\.?\s*$", ".", sentence).strip()
    sentence = re.sub(r"^The two", "The two", sentence)

    land = facts.setdefault("land_identification", {})
    existing = land.get("land_assembly")
    if isinstance(existing, dict):
        return False  # already relocated on an earlier pass
    land["land_assembly"] = {
        "value": sentence,
        "source": (litigation.get("source") or ""),
    }
    return True


def _point_rera_landowner_at_identity_table(facts: dict) -> bool:
    """Replaces RERA Core Data's promoter_land_owner_investor with a pointer.

    It restates the Corporate/Promoter Identity table's "Landowner / Investor
    on record" row almost verbatim -- same chain, same conclusion, different
    wording. CLAUDE.md Section B, "Say it once": a fact appears in exactly one
    place, and a second mention is a pointer, never a restatement."""
    core = facts.get("rera_core_fields")
    corp = facts.get("corporate_identity") or {}
    if not isinstance(core, dict):
        return False
    current = (core.get("promoter_land_owner_investor") or "").strip()
    other = ((corp.get("landowner_investor") or {}) if isinstance(corp.get("landowner_investor"), dict) else {}).get("value", "")
    if not current or not other.strip():
        return False
    if current.startswith("See the Corporate"):
        return False  # already a pointer
    core["promoter_land_owner_investor"] = (
        "See the Corporate/Promoter Identity table's Landowner / Investor row, "
        "which sets out the ownership chain in full."
    )
    return True


# Narrative fields that carry confirmed findings worth researching in depth.
# Deliberately NOT every narrative field: CLAUDE.md Section B scopes this to
# findings ("not gaps, not clean checks"), and a gap sent to a research pass
# would come back either padded or rewritten into something that reads more
# certain than the evidence supports.
_FINDING_FIELDS = (
    "litigation_status.value",
    "land_identification.land_assembly.value",
    "fsi_metrics.mortgage_area",
)

# Below this, a clause is a label or a stub rather than a finding with anything
# to research.
_MIN_FINDING_LENGTH = 80

# Narrative fields whose clauses carry factual claims needing a citation in the
# External document. Used only to assemble the claim list for the editorial
# citation-matching pass; the renderer still decides per clause.
_CITED_NARRATIVE_FIELDS = (
    "executive_summary",
    "litigation_status.value",
    "land_identification.land_assembly.value",
    "fsi_metrics.mortgage_area",
    "fsi_interpretation",
    "fsi_governing_framework",
    "social_infrastructure",
    "micro_market_overview",
    "area_intelligence_trend",
    "unit_summary_note",
)


def _collect_findings(facts: dict) -> list:
    """Returns [{"path", "clause"}] for every confirmed finding worth a
    research pass.

    Runs over the scrubbed view, so a clean check is never sent for research
    (there is nothing to research about a nothing) and neither is a gap. What
    survives the scrubber in these fields IS the finding, which is why this
    runs after _normalize_misfiled_facts and _scrub_clean_checks."""
    findings = []
    for path in _FINDING_FIELDS:
        container, key = _resolve_path(facts, path)
        if container is None or not isinstance(container[key], str):
            continue
        for clause in _split_into_bullet_clauses(container[key]):
            if len(clause) >= _MIN_FINDING_LENGTH and not _is_clean_check_clause(clause):
                findings.append({"path": path, "clause": clause})
    return findings


def _finding_research_context(facts: dict) -> str:
    """Project identity for disambiguating a search. A Notice of Lis Pendens
    means nothing without knowing which plot it is being researched against."""
    core = facts.get("rera_core_fields", {}) or {}
    corp = facts.get("corporate_identity", {}) or {}
    land = facts.get("land_identification", {}) or {}
    fld = lambda d, k: (d.get(k) or {}).get("value", "") if isinstance(d.get(k), dict) else ""
    bits = [
        f"Project: {core.get('project_name', '')}",
        f"MahaRERA registration: {core.get('registration_number', '')}",
        f"Promoter: {fld(corp, 'promoter_name')}",
        f"Location: {fld(land, 'village_locality')}, {fld(land, 'mandal_taluka_district')}",
        f"Plot: {fld(land, 'survey_cts_plot_numbers')}",
    ]
    return "\n".join(b for b in bits if b.split(": ", 1)[-1].strip())


def run_finding_research(facts: dict, researcher=None) -> dict:
    """Researches every confirmed finding in depth and writes the resolved text
    back into the field it came from.

    CLAUDE.md Section B, "Deep research on every finding": a found item is
    written out with what it is, who it involves, when it arose, whether it is
    still live, and what it means for this project. This is a per-finding
    research stage, not a formatting pass, and it is budgeted as one: each call
    logs under its own "finding_research" label so its cost shows up separately
    in usage_summary.json, and the fan-out is capped by
    deep_research.MAX_FINDING_RESEARCH_CALLS.

    Degrades safely and unconditionally. A finding whose research call fails,
    or that falls beyond the cap, keeps its original text verbatim -- an auth
    failure or a rate limit must never silently delete a finding, which would
    be worse than never running the stage. `researcher` is injectable so the
    whole stage can be exercised without spending anything.

    Returns a summary dict; also records it on facts["finding_research"] for
    auditability."""
    researcher = researcher or deep_research.research_finding
    findings = _collect_findings(facts)
    cap = deep_research.MAX_FINDING_RESEARCH_CALLS
    context = _finding_research_context(facts)

    results, enriched, failed = [], 0, 0
    for i, finding in enumerate(findings):
        if i >= cap:
            results.append({**finding, "resolved": False, "note": f"beyond the {cap}-finding cap for one run"})
            continue
        try:
            outcome = researcher(finding["clause"], context)
        except Exception as e:
            # Per finding, not per run: deep_research.research_finding already
            # swallows its own failures, but an injected or future researcher
            # might not, and one bad finding must not cost the others their
            # research. This finding keeps its original text.
            outcome = {"resolved": False, "text": finding["clause"],
                       "still_live": "unknown", "note": f"deeper research could not run: {e}"}
        record = {**finding, "resolved": bool(outcome.get("resolved")),
                  "still_live": outcome.get("still_live", "unknown"),
                  "note": outcome.get("note", "")}
        if record["resolved"]:
            new_text = (outcome.get("text") or "").strip()
            container, key = _resolve_path(facts, finding["path"])
            if container is not None and new_text and finding["clause"] in container[key]:
                container[key] = container[key].replace(finding["clause"], new_text)
                record["text"] = new_text
                enriched += 1
            else:
                record["resolved"] = False
                failed += 1
        else:
            failed += 1
        results.append(record)

    summary = {
        "findings_seen": len(findings),
        "enriched": enriched,
        "kept_original": failed + max(0, len(findings) - cap),
        "cap": cap,
        "results": results,
    }
    facts["finding_research"] = summary
    return summary


def run_editorial_passes(facts: dict, judge=None, matcher=None, headline_writer=None) -> dict:
    """Precomputes the three editorial judgements that a model does better than
    a keyword table, and caches them on `facts` for the renderer to consult.

    Each one replaces a heuristic that was written deterministically because
    the reshape spec asked for determinism, and each was visibly limited by it:

      * clean-check classification. Shape cannot separate "no litigation found"
        (delete) from "no FSI certificate found" (keep, it is a gap), so
        _is_clean_check_clause needed a risk-noun list AND a field allow-list,
        and needs a code change per new field.
      * citation matching. A keyword table mapping clause subject to source
        topic, which returns nothing on unanticipated wording. That is the
        right failure but a frequent one, and it is why External coverage sits
        at 83% rather than near-total.
      * flag headlines. "First sentence" compresses nothing when a gap IS one
        sentence, which was 11 of 17 gaps on the Pranami data, leaving the flag
        line identical to the gap entry it points at.

    Runs ONCE per document set, before either variant renders, so it costs
    three calls per run rather than three per clause.

    Every consumer falls back to its existing deterministic behaviour when a
    verdict is missing, so a failed call, a partial reply, or no API key at all
    leaves output exactly as it is today rather than degraded. That is why the
    caches are keyed by clause text: a lookup miss is indistinguishable from
    "no model ran", and both take the old path."""
    judge = judge or deep_research.classify_clean_checks
    matcher = matcher or deep_research.match_claims_to_sources
    headline_writer = headline_writer or deep_research.write_flag_headlines

    summary = {"clean_checks": 0, "citations": 0, "headlines": 0}

    # 1. Clean-check verdicts, over clauses the cheap shape test flags as
    #    absence-shaped. The shape test stays as a prefilter: it costs nothing
    #    and keeps the model's input to the genuinely ambiguous cases.
    candidates = []
    for path in _CLEAN_CHECK_FIELDS:
        container, key = _resolve_path(facts, path)
        if container is None or not isinstance(container[key], str):
            continue
        for clause in _split_into_bullet_clauses(container[key]):
            if _matches_absence_shape(clause) and clause not in candidates:
                candidates.append(clause)
    if candidates:
        verdicts = judge(candidates)
        if verdicts:
            facts["_clean_check_verdicts"] = {candidates[i]: k for i, k in verdicts.items()}
            summary["clean_checks"] = len(verdicts)

    # 2. Claim-to-source matches, for External citation attachment.
    sources = [(s.get("ref") or s.get("label") or "").strip() for s in (facts.get("sources") or [])]
    claims = []
    for path in _CITED_NARRATIVE_FIELDS:
        container, key = _resolve_path(facts, path)
        if container is None or not isinstance(container[key], str):
            continue
        for clause in _split_into_bullet_clauses(container[key]):
            if len(clause) >= 40 and clause not in claims:
                claims.append(clause)
    if claims and any(sources):
        matches = matcher(claims, sources)
        if matches:
            facts["_claim_source_matches"] = {claims[c]: s for c, s in matches.items()}
            summary["citations"] = len(matches)

    # 3. Flag headlines, one per gap.
    gaps = facts.get("gaps") or []
    if gaps:
        headlines = headline_writer(list(gaps))
        if headlines:
            facts["_flag_headlines"] = {i + 1: h for i, h in headlines.items()}
            summary["headlines"] = len(headlines)

    facts["_editorial_passes"] = summary
    return summary


def _preflight_rules(doc_variant: str) -> dict:
    """Confirms rules.md is present and intact BEFORE any document is built.

    Most of rules.md is enforced by code rather than by a model: _fill_template
    renders in pure Python, so a missing or malformed rules file would not
    announce itself, it would just quietly stop constraining output. This is
    the one place that fails loudly instead.

    Three checks:
      * all three sections load and are non-empty (a renamed file, a broken
        `--- Section X: ... ---` marker, or an accidentally emptied section);
      * Section B, and Section C for External, carry no em dash and no double
        hyphen used as a dash. They are injected verbatim into External-facing
        prompts, prompt punctuation bleeds into model output, and
        _verify_external_document_quality hard-fails the save on either. This
        check exists because that trap was actually sprung once, by a rewrite
        of the very sentence forbidding it;
      * Section A is NOT returned to any caller here, so this function cannot
        become a route by which coding-time guidance reaches an API call.

    Raises RuntimeError on any failure: generating a document against rules
    that are missing or self-contradictory is worse than not generating one."""
    sections = {"A": _coding_time_notes(), "B": _common_content_rules()}
    if doc_variant == "external":
        sections["C"] = _external_citation_rule()

    for marker, body in sections.items():
        if not body.strip():
            raise RuntimeError(f"rules.md Section {marker} is empty -- refusing to generate against it")

    for marker in ("B", "C"):
        body = sections.get(marker)
        if body is None:
            continue
        if "—" in body or " -- " in body:
            raise RuntimeError(
                f"rules.md Section {marker} contains an em dash or a double-hyphen dash. It is injected "
                f"verbatim into External prompts and _verify_external_document_quality hard-fails on "
                f"either, so this would corrupt output. Use commas, colons or a single spaced hyphen."
            )

    return {"sections_loaded": sorted(sections), "variant": doc_variant}


def _rendered_document_text(docx_path: str, limit: int = 60000) -> str:
    """Flattens a saved .docx to plain text for review: body paragraphs and
    table cells, in document order, which is what a reader actually sees."""
    import docx as _docx

    doc = _docx.Document(docx_path)
    parts = [p.text for p in _iter_all_paragraphs(doc)]
    text = "\n".join(p for p in parts if p and p.strip())
    return text[:limit]


def run_claude_md_document_review(paths_by_variant: dict, output_dir: str = ".", reg_no: str = "") -> dict:
    """Final pipeline stage: re-reads each rendered Charter and audits it
    against the CLAUDE.md rules it was written under, via the Claude API.

    Section B goes to both variants; Section C is added for External only,
    exactly as it is at generation time. Section A is never sent: it is
    coding-time guidance, and _coding_time_notes' own docstring says so.

    ADVISORY BY DESIGN. It reports and records, and does not block the PDF.
    Two reasons: the deterministic gate (_verify_external_document_quality)
    already hard-fails a genuinely bad save before this point, and a language
    model's opinion should not be able to stop a finished document being
    delivered. Findings are printed and written to a JSON report next to the
    documents so a human can act on them.

    Never raises. No API key, a rate limit or a malformed reply all come back
    as reviewed=False with the reason, and the run continues."""
    section_b = _common_content_rules()
    section_c = _external_citation_rule()

    results = {}
    for variant, path in (paths_by_variant or {}).items():
        if not path or not os.path.exists(path):
            continue
        rules = section_b if variant != "external" else f"{section_b}\n\n{section_c}"
        try:
            text = _rendered_document_text(path)
            outcome = deep_research.review_document_against_rules(text, rules, variant)
        except Exception as e:
            outcome = {"reviewed": False, "compliant": None, "violations": [],
                       "summary": f"CLAUDE.md document review could not run: {e}"}
        outcome["document"] = os.path.basename(path)
        results[variant] = outcome

        if not outcome["reviewed"]:
            print(f"[WARN] CLAUDE.md review ({variant}): {outcome['summary']}")
        elif outcome["compliant"]:
            print(f"[OK] CLAUDE.md review ({variant}): compliant. {outcome['summary']}")
        else:
            print(f"[!] CLAUDE.md review ({variant}): {len(outcome['violations'])} issue(s) to review.")
            for v in outcome["violations"][:10]:
                print(f"    - {v.get('rule', 'rule')}: {str(v.get('quote', ''))[:100]!r}")

    if results and reg_no:
        report_dir = os.path.join(output_dir, "company_charters")
        os.makedirs(report_dir, exist_ok=True)
        report_path = os.path.join(report_dir, f"Company_Charter_{reg_no}_claude_md_review.json")
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"[INFO] CLAUDE.md review report written to {report_path}")
    return results


def _merged_mortgage_value(facts: dict, fsi_metrics: dict) -> tuple:
    """Collapses fsi_metrics.mortgage_area and fsi_metrics.mortgage_lender into
    one "Mortgage / charge on the land" value, returning (text, citation).

    Kept as a render-time merge rather than a facts rewrite because the two
    fields have genuinely different shapes -- mortgage_area is plain text,
    mortgage_lender is {value, source} and feeds _diff_mortgage_lender's
    run-over-run comparison, which needs the field to keep existing.

    Order matters: whatever is left after the clean-check scrubber is a
    finding, and findings lead. Only the lender field carries a source, so
    that is the citation for the merged row."""
    area = (fsi_metrics.get("mortgage_area") or "").strip()
    lender_field = fsi_metrics.get("mortgage_lender")
    lender = ((lender_field or {}).get("value") or "").strip() if isinstance(lender_field, dict) else ""

    parts = [p for p in (area, lender) if p]
    # Both fields habitually open on the same "None disclosed" phrasing; if the
    # scrubber left them saying substantially the same thing, print it once.
    if len(parts) == 2 and (parts[0].startswith(parts[1][:40]) or parts[1].startswith(parts[0][:40])):
        parts = [max(parts, key=len)]
    text = " ".join(parts).strip() or "None disclosed."

    citation = None
    if isinstance(lender_field, dict):
        citation = _citation_text(facts, _clean_source_label(lender_field.get("source") or ""))
    return text, citation


def _normalize_misfiled_facts(facts: dict) -> list:
    """Runs every facts correction and returns the names of those that fired,
    for the caller to log. Safe to call repeatedly."""
    applied = []
    if _relocate_merge_orders(facts):
        applied.append("merge_orders_to_land_identification")
    if _point_rera_landowner_at_identity_table(facts):
        applied.append("rera_landowner_pointer")
    return applied


# A leading "some_field:" or "some.dotted.path:" used as a label. Lowercase and
# underscores only, so real prose labels ("Cross-corroboration:", "CTS
# land-record lookup:") are never touched.
_KEY_LABEL_RE = re.compile(r"^([a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)*)\s*:\s*")
# A raw API exception quoted inside a gap.
_QUOTED_EXCEPTION_RE = re.compile(r"verification could not run:\s*\"[^\"]*\"", re.IGNORECASE)
# A path under the run output directory.
# Must end on a word character, so the full stop closing the sentence is not
# swallowed along with the filename.
_OUTPUT_PATH_RE = re.compile(r"\boutput/[\w./-]*\w")

# Field names whose humanized form would otherwise read badly ("Cin llpin").
_PROCESS_LABEL_OVERRIDES = {
    "cin_llpin": "CIN / LLPIN",
    "din": "DIN",
    "gstin": "GSTIN",
    "pan": "PAN",
    "rera_core_fields": "RERA core fields",
    "cts": "CTS",
}
# A module named as a next step.
_MODULE_NAME_RE = re.compile(r"\b([a-z][a-z0-9_]*)\.py\b")


def _sanitize_process_text(text: str) -> str:
    """Strips internals out of a process-failure item while keeping what it
    tells the reader.

    CLAUDE.md Section B pulls in two directions here, and both have to be
    honoured. "Internal keeps process failures, External does not" means these
    items STAY in the Internal document, so deleting them is wrong. But the
    same section forbids "a file path, module name, function or parameter name,
    JSON key, or raw exception string into EITHER document" -- so what stays
    has to be rewritten, not passed through.

    Four rewrites, each preserving the information:
      * "promoter_name:" becomes "Promoter name:", and a dotted path is
        reduced to its last segment ("land_identification.total_gross_area:"
        becomes "Total gross area:");
      * a quoted API exception becomes a plain statement that the step could
        not run, which is the only part a reader can act on anyway;
      * a run-output path becomes a description of where the data went;
      * a named module becomes a description of the step it performs.
    """
    text = (text or "").strip()
    if not text:
        return text

    match = _KEY_LABEL_RE.match(text)
    if match:
        # Humanized locally rather than via charter_document._field_display_name:
        # that module is the dead builder, and the live pipeline should not grow
        # an import into it just to title-case a word.
        raw_leaf = match.group(1).split(".")[-1]
        leaf = _PROCESS_LABEL_OVERRIDES.get(raw_leaf)
        if not leaf:
            leaf = raw_leaf.replace("_", " ")
            leaf = f"{leaf[:1].upper()}{leaf[1:]}"
        text = f"{leaf}: {text[match.end():]}"

    # No "this pass" here: the sentences carrying this already say it, and
    # doubling it read as "not re-verified this pass, could not run this pass".
    text = _QUOTED_EXCEPTION_RE.sub("the verification step could not run", text)
    text = _OUTPUT_PATH_RE.sub("this project's own run output", text)
    text = _MODULE_NAME_RE.sub(lambda m: f"the {m.group(1).replace('_', ' ')} step", text)
    return text.strip()


def _sanitize_process_gaps(facts: dict):
    """Applies _sanitize_process_text across facts["gaps"], in place.

    Returns the ORIGINAL list when anything changed, so the caller can restore
    it before persisting: the raw exception text and the path are genuine
    diagnostics and belong in the record, exactly like the clean checks the
    scrubber removes. They just do not belong on the page. Returns None when
    nothing changed, which makes this idempotent across two renders."""
    gaps = facts.get("gaps") or []
    if not gaps:
        return None
    cleaned = [_sanitize_process_text(g) for g in gaps]
    if cleaned == gaps:
        return None
    original = list(gaps)
    facts["gaps"] = cleaned
    return original


def _is_cited_absence(text: str) -> bool:
    """True when a rendered paragraph both reports a nothing AND cites a source
    for it -- the one clean-check failure that hard-fails an External save.

    Two carve-outs are honoured here rather than by the caller. A Gaps & Sources
    entry ("Gap 4. ...") is an open unknown, not a clean result, and Section B
    keeps those in full. A score's stated basis is exempt too, but needs no test
    of its own: a sub-metric note reads "0 complaints, 0 appeals", which states
    a figure rather than asserting a non-existence, so no absence shape matches
    it in the first place."""
    stripped = re.sub(r"\[\d+\]", "", text or "").strip()
    if not stripped or stripped.startswith("Gap "):
        return False
    if not re.search(r"\[\d+\]", text or ""):
        return False
    clauses = _split_into_bullet_clauses(stripped.lstrip("• ").strip())
    return bool(clauses) and all(_is_clean_check_clause(c) for c in clauses)


def _restore_clean_checks(facts: dict, changed: dict) -> None:
    """Puts the pre-scrub text back, so what gets persisted to .facts.json is
    the complete record. See _scrub_clean_checks."""
    for path, original in (changed or {}).items():
        container, key = _resolve_path(facts, path)
        if container is not None:
            container[key] = original


def _fill_template(
    reg_no: str, facts: dict, out_path: str, doc_variant: str = "internal",
    elapsed_seconds: float | None = None, cost_usd: float | None = None, api_calls: int | None = None,
) -> None:
    """doc_variant is "internal" (default -- today's behavior: inline
    "(label)" citations, "(Code-Computed)" labels, verbatim prose) or
    "external" (deduped, sequentially-numbered "[N]" citations resolving
    to a generic-language Sources list at the end, no "(Code-Computed)"
    labels, and known internal-process phrases generalized in prose --
    see _citation_text/_generic_one_label/_externalize_prose). Reset
    fresh on every call (not merely defaulted) so two calls against the
    same `facts` dict -- once per variant, the normal calling pattern --
    never leak state from one into the other."""
    import docx

    if doc_variant == "external":
        # A deep copy, not an in-place rewrite: the caller's own `facts`
        # dict (run_company_charter persists it to .facts.json after
        # rendering) must keep the real, un-rewritten internal content
        # regardless of whether/when an External pass runs against it --
        # see _externalized_facts_copy's own docstring.
        facts = _externalized_facts_copy(facts)
    # Rules first, before a single paragraph is written. See _preflight_rules:
    # rendering is pure code, so a missing or self-contradictory rules file
    # would silently stop constraining output rather than fail.
    _preflight_rules(doc_variant)

    facts["_doc_variant"] = doc_variant
    facts["_citation_registry"] = {"order": [], "index": {}}

    # CLAUDE.md Section B, "A clean check produces no sentence". Runs here so
    # EVERY caller gets it, including scripts that hit _fill_template directly.
    # The pre-scrub text is stashed on `facts` rather than returned, because
    # run_company_charter renders twice and must restore from the FIRST pass;
    # it restores and pops this key before persisting .facts.json, so the
    # record keeps what the page drops.
    # Corrections first, scrub second: relocating the merge-order sentence out
    # of litigation_status changes what the scrubber then sees there (what is
    # left is a bare Form B "no litigation" declaration, which is a clean check
    # and goes). Running them the other way round would leave that behind.
    _normalize_misfiled_facts(facts)

    _scrubbed = _scrub_clean_checks(facts)
    if _scrubbed and not facts.get("_pre_scrub_narrative"):
        facts["_pre_scrub_narrative"] = _scrubbed

    # Process failures stay in the Internal document (Section B), but the raw
    # path, JSON key and API exception string inside them may not appear in
    # EITHER document. Same restore-before-persist contract as the scrub above:
    # those internals are real diagnostics and belong in the record.
    _pre_sanitize_gaps = _sanitize_process_gaps(facts)
    if _pre_sanitize_gaps and not facts.get("_pre_sanitize_gaps"):
        facts["_pre_sanitize_gaps"] = _pre_sanitize_gaps

    shutil.copy2(TEMPLATE_PATH, out_path)
    doc = docx.Document(out_path)
    _fix_bullet_hanging_indent(doc)
    p = doc.paragraphs
    # Paragraphs queued for deletion (suppressed External-only content) --
    # removed in one batch right before save, not as each is decided. Several
    # OTHER things below (the Sources list, section-consolidation heading
    # renames) look up doc.paragraphs by a fixed numeric index; deleting
    # mid-function would shift every later index and break those lookups.
    # `p` itself stays safe to keep indexing throughout (it holds resolved
    # Paragraph objects from before any deletion, not positions).
    _paragraphs_to_remove = []

    global _ACTIVE_EXTERNAL_FACTS
    _ACTIVE_EXTERNAL_FACTS = facts if doc_variant == "external" else None
    # doc.add_paragraph(...) is how every _append_*_section function writes
    # NEW content (as opposed to _set_paragraph_text filling a pre-existing
    # template slot) -- patching this one fresh, single-use `doc` instance's
    # bound method reaches all of them uniformly, without touching each call
    # site. CIN/DIN stripping (CLAUDE.md Section B) applies to BOTH variants
    # here, unlike _externalize_prose below (External-only).
    _real_add_paragraph = doc.add_paragraph

    def _add_paragraph_sanitized(text="", *args, **kwargs):
        text = _expand_jargon_first_use(_strip_inline_cin_din(str(text)))
        if doc_variant == "external":
            text = _externalize_prose(facts, text)
        return _real_add_paragraph(text, *args, **kwargs)

    doc.add_paragraph = _add_paragraph_sanitized

    li = facts["land_identification"]
    ci = facts["corporate_identity"]
    conn = facts["connectivity"]
    fsi_m = facts["fsi_metrics"]
    rs = facts["rules_statutory"]
    rc = facts["rera_compliance"]
    lp = facts["local_planning"]
    core = facts["rera_core_fields"]
    nb = facts["neighbourhood"]

    fld = lambda d, k: (d.get(k) or {}).get("value", "") if isinstance(d.get(k), dict) else ""
    src = lambda d, k: (d.get(k) or {}).get("source", "") if isinstance(d.get(k), dict) else ""

    _set_paragraph_text(p[1], f'Project: {core.get("project_name", "[Unknown]")} | Promoter: {fld(ci, "promoter_name") or "[Unknown]"}')
    _set_paragraph_text(p[2], "Public Web-Sourced Edition -- Adapted for Maharashtra (MahaRERA)")
    _month_year = datetime.now().strftime('%B %Y')
    _set_paragraph_text(p[3], f"Deep Market Research (Prepared {_month_year})" if doc_variant == "external" else f"Deep Market Research -- Prepared {_month_year}")
    # p[5] (Methodology Note) is deleted entirely further down regardless
    # of variant ("per request") -- never worth bulleting content that's
    # about to be thrown away, and _apply_deferred_bullets would crash on
    # an already-detached paragraph if it tried.
    _set_paragraph_text(p[5], facts["methodology_note"])
    _set_paragraph_as_bullets(facts, p[7], facts["executive_summary"])
    litigation_citation = _citation_text(facts, _clean_source_label(src(facts, "litigation_status")))
    litigation_text = fld(facts, "litigation_status")
    _set_paragraph_as_bullets(facts, p[15], litigation_text, citation=litigation_citation)
    _set_paragraph_as_bullets(facts, p[24], f"Road: {conn.get('road', '')}", gap_check_text=conn.get("road", ""), citation=_cite_marker("distance", facts=facts))
    _set_paragraph_as_bullets(facts, p[25], f"Rail: {conn.get('rail', '')}", gap_check_text=conn.get("rail", ""), citation=_cite_marker("distance", facts=facts))
    _set_paragraph_as_bullets(facts, p[26], f"Metro: {conn.get('metro', '')}", gap_check_text=conn.get("metro", ""), citation=_cite_marker("distance", facts=facts))
    _set_paragraph_as_bullets(facts, p[27], f"Air: {conn.get('air', '')}", gap_check_text=conn.get("air", ""), citation=_cite_marker("distance", facts=facts))
    _set_paragraph_as_bullets(facts, p[30], facts["social_infrastructure"])
    _set_paragraph_as_bullets(facts, p[32], facts["fsi_governing_framework"])
    _set_paragraph_as_bullets(facts, p[33], facts["fsi_interpretation"])
    _set_paragraph_as_bullets(facts, p[37], f"Planning approval sequence: {rs.get('planning_approval_sequence', '')}", gap_check_text=rs.get("planning_approval_sequence", ""), citation=_cite_marker("project_registration", facts=facts))
    allotment_text = f"Allotment mechanics: {rs.get('allotment_mechanics', '')}"
    _set_paragraph_as_bullets(facts, p[41], rc.get("registration_summary", ""))
    _set_paragraph_as_bullets(facts, p[42], f"Collection Account of the Project (100%): {rc.get('collection_account', '')}", gap_check_text=rc.get("collection_account", ""))
    _set_paragraph_as_bullets(facts, p[43], f"Separate/Transaction RERA escrow sub-accounts: {rc.get('escrow_subaccounts', '')}", gap_check_text=rc.get("escrow_subaccounts", ""))
    _set_paragraph_as_bullets(facts, p[44], f"Litigations/complaints/appeals related to the project: {rc.get('litigations_complaints_appeals', '')}", gap_check_text=rc.get("litigations_complaints_appeals", ""))
    _set_paragraph_as_bullets(facts, p[45], rc.get("statutory_declaration", ""))
    _set_paragraph_as_bullets(facts, p[46], f"Construction progress: {rc.get('construction_progress', '')}", gap_check_text=rc.get("construction_progress", ""))
    _set_paragraph_as_bullets(facts, p[49], f"Authority of record: {lp.get('authority_of_record', '')}", gap_check_text=lp.get("authority_of_record", ""), citation=_cite_marker("project_registration", facts=facts))
    _set_paragraph_as_bullets(facts, p[50], f"Project type: {lp.get('project_type', '')}", gap_check_text=lp.get("project_type", ""), citation=_cite_marker("project_registration", facts=facts))
    _set_paragraph_as_bullets(facts, p[51], f"Professionals of record: {lp.get('professionals_of_record', '')}", gap_check_text=lp.get("professionals_of_record", ""), citation=_cite_marker("project_registration", facts=facts))
    _set_paragraph_as_bullets(facts, p[54], facts["micro_market_overview"], citation=_cite_marker("pricing", "market_trend", facts=facts))
    _set_paragraph_as_bullets(facts, p[56], facts["area_intelligence_trend"], citation=_cite_marker("market_trend", "pricing", facts=facts))
    _set_paragraph_as_bullets(facts, p[61], facts["unit_summary_note"], citation=_cite_marker("project_registration", facts=facts))
    _set_paragraph_as_bullets(facts, p[64], facts["documents_absent_note"])

    # Every entry below is "one template paragraph, either rendered with its
    # Internal text or deleted outright for External" -- previously ~7
    # separate hand-written if/else blocks, each a fresh chance to bypass
    # _set_paragraph_text or the removal convention by accident. text_fn is
    # LAZY (only called when actually used) so a _cite(...) call never
    # registers a citation for a paragraph about to be deleted -- the exact
    # behavior the original if/else blocks had by construction. suppress_fn
    # defaults to "always suppress for External"; only the Allotment
    # mechanics row overrides it, since only its generic "nothing special
    # disclosed" boilerplate should disappear, not a genuinely unusual
    # mechanism that happens to land in the same slot.
    _always_suppress = lambda _facts: True
    _external_suppressed_paragraphs = (
        # (index, text_fn(facts) -> str, suppress_for_external_fn(facts) -> bool, gap_check_fn(facts) -> str | None, citation_fn(facts) -> str | None)
        (12, lambda f: f["address_discrepancy_note"], _always_suppress, None, None),
        (13, lambda f: f["corporate_registry_cross_check"], _always_suppress, None, None),
        (17, lambda f: f["location_coordinates_note"], _always_suppress, None, lambda f: _cite_marker("distance", facts=f)),
        (18, lambda f: "Map screenshot not embedded -- a live map cannot be fetched programmatically, so distances below were sourced from mapping-service queries instead (see Sources).", _always_suppress, None, None),
        (36, lambda f: f"Governing act: {rs.get('governing_act', '')}", _always_suppress, lambda f: rs.get("governing_act", ""), lambda f: _cite_marker("project_registration", facts=f)),
        (38, lambda f: allotment_text, lambda f: _externalize_prose(f, allotment_text) == "", lambda f: rs.get("allotment_mechanics", ""), None),
        (58, lambda f: f.get("rera_scraping_note", f"Extracted directly from the live MahaRERA public project page for registration number {reg_no}."), _always_suppress, None, None),
        (63, lambda f: f["documents_reviewed_note"], _always_suppress, None, None),
    )
    # Why each is suppressed for External (kept here, once, rather than
    # repeated as a comment on every removed if/else block above):
    #   12/13 -- QA cross-check narrations ("we compared N sources and they
    #     agree/disagree"); a real disagreement already surfaces as a flag
    #     or gap elsewhere, and these two otherwise just restate facts
    #     already in the Corporate Identity table (CIN, registered office,
    #     entity type, partner/director names) plus audit-trail detail
    #     ("independently retrieved", "fully corroborated") about this
    #     pipeline's own research process, not the promoter.
    #   17/18 -- pure measurement-methodology caveats (no map was plotted;
    #     distances are estimated from the locality, not an exact pin) --
    #     no decision-relevant content given the actual Distances table
    #     right below.
    #   36 -- every MahaRERA-registered project is governed by the same Act
    #     by definition; carries no signal about THIS project.
    #   38 -- see suppress_fn above.
    #   58 -- a sourcing/methodology note ("where did this data come
    #     from"), not a project fact.
    #   63 -- "which files I personally opened vs. just confirmed present"
    #     is research-scope bookkeeping, not a project fact.
    # Internal keeps every one of these in full -- that process detail IS
    # the point there.
    for _idx, _text_fn, _suppress_fn, _gap_check_fn, _citation_fn in _external_suppressed_paragraphs:
        if doc_variant == "external" and _suppress_fn(facts):
            _paragraphs_to_remove.append(p[_idx])
        else:
            _set_paragraph_as_bullets(
                facts, p[_idx], _text_fn(facts),
                gap_check_text=(_gap_check_fn(facts) if _gap_check_fn else None),
                citation=(_citation_fn(facts) if _citation_fn else None),
            )

    gaps = facts.get("gaps", [])
    if facts.get("_doc_variant") == "external":
        # External: show ONLY gaps that were serious enough to also earn an
        # Imminent/Structural flag above (see _classify_flags), colored to
        # match that same severity -- a long, uniformly grey/italic list of
        # every minor caveat wasn't legible or prioritized for a reader who
        # skips straight to this section. Monitor-only/unflagged gaps are
        # dropped from THIS section entirely, not lost -- they're still
        # visible in Overview & Flags' own Monitor list. _classify_flags is
        # a pure reader of facts with no side effects, safe to call again
        # here rather than threading its result through the call chain.
        gap_severity = {}
        flags_for_gaps = _classify_flags(facts)
        for severity in ("structural", "imminent"):  # imminent checked last so it wins on the (should never happen) double-tag case
            for item in flags_for_gaps.get(severity, []):
                m = re.match(r"gaps\[(\d+)\]$", item.get("field", ""))
                if m:
                    gap_severity[int(m.group(1))] = severity

        for run in list(p[66].runs):
            run.text = ""
            run._r.getparent().remove(run._r)
        material = [(i, g) for i, g in enumerate(gaps) if i in gap_severity]
        if not material:
            p[66].add_run(_externalize_prose(facts, "No additional material gaps identified beyond the standing gap below."))
        else:
            # Numbered to match the "(Gap N)" pointers in Overview & Flags.
            # The number is the gap's stable identity (see _classify_flags),
            # so External's subset reads "Gap 3.", "Gap 7." rather than
            # renumbering -- a pointer means the same thing in both variants.
            for idx, (i, g) in enumerate(material):
                run = p[66].add_run(_externalize_prose(facts, f"Gap {i + 1}. {g}"))
                _color_run(run, _TEXT_RED if gap_severity[i] == "imminent" else _TEXT_AMBER)
                if idx < len(material) - 1:
                    run.add_break()
    else:
        _set_paragraph_text(
            p[66],
            "\n".join(f"Gap {i + 1}. {g}" for i, g in enumerate(gaps)) if gaps
            else "No additional gaps identified beyond the standing gap below.",
        )
    # p[67] (the permanent standing gap) is left untouched deliberately.

    t = doc.tables
    for row, key in zip(range(1, 8), (
        "survey_cts_plot_numbers", "village_locality", "mandal_taluka_district",
        "pincode", "total_gross_area", "area_affected", "net_area",
    )):
        _set_cell(t[0], row, 1, fld(li, key))
        # The template's own "Source" column showed the raw output/<reg_no>/...
        # path verbatim -- cleaned to a short label here, same convention as
        # every other citation added below.
        _set_cell(t[0], row, 2, _citation_text(facts, _clean_source_label(src(li, key))) or "")

    # Grown row: how the plots were assembled (see _relocate_merge_orders).
    # The template predates this field, which only exists because society
    # amalgamation orders were being filed under litigation_status.
    assembly = fld(li, "land_assembly")
    if assembly:
        assembly_row = t[0].add_row()
        _set_row_cell(assembly_row, 0, "Land assembly")
        _set_row_cell(assembly_row, 1, assembly)
        _set_row_cell(assembly_row, 2, _citation_text(facts, _clean_source_label(src(li, "land_assembly"))) or "")

    _remove_gap_rows(t[0], value_col=1)

    for row, key in zip(range(1, 10), (
        "promoter_name", "organization_type", "cin_llpin", "registered_office_main",
        "registered_office_board_resolution", "registered_office_planning_stage",
        "authorized_signatory", "partners_directors", "landowner_investor",
    )):
        # This table has no Source column in the template (unlike Land
        # Identification) -- each fact carries a real source in facts.json
        # that was never shown anywhere, so it's appended inline instead.
        value = fld(ci, key)
        citation = _citation_text(facts, _clean_source_label(src(ci, key)))
        _set_cell(t[1], row, 1, f"{value} {citation}" if value and citation else value)

    if facts.get("_doc_variant") == "external":
        # The "Field" column (col 0) is static text baked directly into the
        # template docx -- unlike col 1 above, nothing ever calls _set_cell
        # on it, so it never passes through _externalize_prose on its own.
        # This one row's static label happens to use " -- ", so it needs an
        # explicit touch here. Internal leaves the template's own text
        # completely alone, as always.
        _set_cell(t[1], 3, 0, _externalize_prose(facts, t[1].rows[3].cells[0].text))
    # Row-index-sensitive fix above must run BEFORE any row removal --
    # removing a row shifts every later fixed index.
    _remove_gap_rows(t[1], value_col=1)

    for row, key in zip(range(1, 5), ("east", "west", "north", "south")):
        _set_cell(t[2], row, 1, _cite(nb.get(key, ""), "project_registration", "legal_documents", facts=facts))
    _remove_gap_rows(t[2], value_col=1)

    def _fill_distance_row(row, item):
        _set_row_cell(row, 0, item["landmark"])
        _set_row_cell(row, 1, item["distance_time"])
        _set_row_cell(row, 2, item["route_note"])

    _fill_variable_rows(t[3], 1, facts["distances"], _fill_distance_row)

    for row, key in zip(range(1, 6), ("net_land_area", "approved_bua", "sanctioned_bua", "mortgage_area", "implied_fsi")):
        _set_cell(t[4], row, 1, fsi_m.get(key, ""))

    # One row, not two. "Mortgage area" and "Mortgage lender" sat adjacent and
    # both reported the same absence -- CLAUDE.md Section B: "Two table rows
    # must not both exist to report the same absence (Mortgage area and
    # Mortgage lender both reading 'none' is one row's worth of information)."
    # The row that survives is the one carrying the finding: development
    # agreements permit the developer to mortgage its free-sale area, and no
    # mortgage has been taken. A live right is a finding even unexercised.
    mortgage_value, mortgage_citation = _merged_mortgage_value(facts, fsi_m)
    _set_cell(t[4], 4, 0, "Mortgage / charge on the land")
    _set_cell(t[4], 4, 1, f"{mortgage_value} {mortgage_citation}" if mortgage_citation else mortgage_value)
    lender_history_note = facts.get("mortgage_lender_history_note")
    if lender_history_note:
        history_row = t[4].add_row()
        _set_row_cell(history_row, 0, "Mortgage lender -- change since prior run")
        _set_row_cell(history_row, 1, lender_history_note)
    _remove_gap_rows(t[4], value_col=1)

    def _fill_comparable_row(row, item):
        distance = item.get("distance_km")
        project_label = f"{item['project']} ({distance} km)" if distance else item["project"]
        _set_row_cell(row, 0, project_label)
        _set_row_cell(row, 1, item["configuration"])
        _set_row_cell(row, 2, item["pricing"])
        # CLAUDE.md Section A, "No URLs in narrow table columns": a full listing
        # URL in this four-column table wraps mid-token and can split across a
        # page break. The cell carries a short label (Internal) or an "[N]"
        # marker (External) and the URL itself lives once, in the Sources list,
        # which is what _citation_text already does everywhere else.
        _set_row_cell(row, 3, _citation_text(facts, _clean_source_label(item.get("source", ""))) or "")

    _fill_variable_rows(t[5], 1, facts.get("comparables", []), _fill_comparable_row)

    for row, key in zip(range(1, 14), (
        "project_name", "registration_number", "promoter_name", "authority",
        "plan_approval_number", "project_status", "approved_date",
        "proposed_completion_date", "project_type", "litigations_per_record",
        "promoter_land_owner_investor", "collection_bank_account", "total_building_units",
    )):
        _set_cell(t[6], row, 1, core.get(key, ""))
    # Same single-value-column Field/Value shape as Land Identification/
    # Corporate Identity/Neighbourhood/FSI Metrics above -- this table was
    # simply never wired into _remove_gap_rows, which is why "RERA Core
    # Data" used to survive fully intact even when nearly every row read
    # "Not disclosed"/"Not applicable"/"Not confirmed" (CLAUDE.md Section
    # B: drop the whole table if every row is a gap -- already-existing
    # logic in _remove_empty_section_headings handles that once this
    # table's data rows are actually stripped down to none).
    _remove_gap_rows(t[6], value_col=1)

    def _fill_block_row(row, item):
        _set_row_cell(row, 0, item["block_wing"])
        _set_row_cell(row, 1, item["floors"])
        _set_row_cell(row, 2, item["config"])
        _set_row_cell(row, 3, item["units_counted"])
        _set_row_cell(row, 4, item["note"])

    _fill_variable_rows(t[7], 1, facts.get("blocks", []), _fill_block_row)

    if facts.get("_doc_variant") == "external":
        # A 50-70-row table that's "Downloaded" repeated over and over is
        # diligence bookkeeping, not a finding -- drops the heading and the
        # whole table for External. Internal keeps the full per-document
        # checklist.
        _paragraphs_to_remove.append(p[62])
        t[8]._tbl.getparent().remove(t[8]._tbl)
    else:
        def _fill_doc_library_row(row, item):
            _set_row_cell(row, 0, item["document_name"])
            _set_row_cell(row, 1, item["status"])

        _fill_variable_rows(t[8], 1, facts.get("document_library", []), _fill_doc_library_row)

    def _format_source_line(s: dict) -> str:
        line = f'{s.get("label", "")} -- {s.get("ref", "")}'
        # published_date/accessed_date are new fields (see _CHARTER_FACTS_SCHEMA) -- guard
        # with .get() so a facts dict built before this change still renders correctly.
        published = s.get("published_date")
        accessed = s.get("accessed_date")
        if published or accessed:
            line += f' [published: {published or "unknown"}, accessed: {accessed or "unknown"}]'
        return line

    if facts.get("_doc_variant") == "external":
        # This Sources list is built from the citation registry BEFORE
        # _append_overview_section (called later, at the end of this
        # function -- see its own _capture_batch call) ever runs, and that
        # section is what registers every Overview & Flags citation (see
        # _annotate_flag_citations). Confirmed live: without this
        # pre-registration pass, a flag line got a real, working "[N]"
        # marker in the body, but that citation never made it into the
        # printed Sources list at all, since the list below was already
        # written by the time it was registered. _register_citation is
        # idempotent (same generic label -> same number, however many
        # times it's called), so re-deriving these same flag texts here
        # purely to populate the registry, then letting
        # _append_overview_section do it again later for the real
        # rendering, can never produce two different numbers for the same
        # source.
        _flags_preview = _classify_flags(facts)
        for _flag_item in (
            _flags_preview.get("imminent", [])
            + _flags_preview.get("structural", [])
            + _flags_preview.get("monitor", [])
        ):
            _annotate_flag_citations(facts, _externalize_prose(facts, f"• {_flag_item['text']}"))

        # Every "[N]" marker in the body was registered against this exact
        # list, in this exact order (see _register_citation) -- replaces
        # facts["sources"] entirely for External, since that list's
        # internal topic/date bookkeeping isn't what the numbered
        # in-body citations actually point to.
        registry = facts["_citation_registry"]
        sources = [f"[{i + 1}] {label}" for i, label in enumerate(registry["order"])]
    else:
        sources = [_format_source_line(s) for s in facts.get("sources", [])]
    _fill_variable_paragraphs(doc, 69, 8, sources)

    # Source-trust-registry promotions from THIS run are internal pipeline
    # bookkeeping (see _record_source_hits_and_promote's own module note on
    # why an algorithmic, frequency-based promotion isn't the same thing as
    # a manually-vetted trusted source) -- Internal-only, never External.
    if facts.get("_doc_variant") != "external" and facts.get("source_promotion_notes"):
        doc.add_paragraph()
        note_heading = doc.add_paragraph("Source Trust Registry Updates (this run)")
        for run in note_heading.runs:
            run.bold = True
        for note in facts["source_promotion_notes"]:
            _apply_bullet_hanging_indent(doc.add_paragraph(f"• {note}"))

    # ---------------------------------------------------------------------
    # Section consolidation -- renames existing template headings to the
    # target vocabulary (docs/Company_Charter_Executive_Design.docx section
    # 3's 9-section table) and regroups content to match, without touching
    # any of the prose-generation logic above or in any _append_* function.
    # Three kinds of change, all via the same two primitives (retext an
    # EXISTING heading paragraph, or capture-then-relocate NEWLY appended
    # content -- the latter already proven safe for the checks bolted onto
    # the end in the previous pass):
    #   1. Renames: "1. Legal Identifiers..." -> "Counterparty", "2. Location
    #      Map" -> "The Asset", "4. Rules" -> "Compliance & Legal Detail",
    #      "5. Area Intelligence" -> "Market & Area Intelligence", "6. RERA
    #      Scraping..." -> "RERA Core Data", "Gaps & Limitations..." ->
    #      "Gaps & Sources". "3. FSI", "Document Library Contents", and
    #      "Sources" are demoted from Heading 1 to Heading 2, becoming
    #      subsections of "The Asset" / a new "Diligence Appendix" / the
    #      merged "Gaps & Sources" respectively, instead of standalone
    #      top-level sections.
    #   2. Land Identification is pulled out from under the old combined
    #      "Legal Identifiers" heading to sit under "The Asset" instead,
    #      alongside Location Map/FSI -- the one case here of relocating
    #      EXISTING template content, not newly-appended content.
    #   3. The 5 check sections (credit rating/IBBI/company profile/group
    #      companies/Developer Score) get a short, one-line-each SUMMARY
    #      under Counterparty (_append_counterparty_summary, new) instead
    #      of their full detail sitting there -- the full tables move into
    #      a new "Diligence Appendix" (alongside Document Library and Data
    #      Authenticity), demoted to Heading 2 subsections of it. Complaint
    #      Order Outcomes and Appeal-Level Judgments move into "Compliance
    #      & Legal Detail" (same anchor Appeal Judgments already used).
    # ---------------------------------------------------------------------

    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as _Paragraph

    _set_paragraph_text(p[8], "Counterparty")
    _set_paragraph_text(p[16], "The Asset")
    _set_paragraph_text(p[34], "Compliance & Legal Detail")
    _set_paragraph_text(p[52], "Market & Area Intelligence")
    _set_paragraph_text(p[57], "RERA Core Data")
    _set_paragraph_text(p[65], "Gaps & Sources")

    h1_style = p[8].style
    h2_style = p[9].style  # "Land Identification" -- a known Heading 2
    _set_paragraph_text(p[31], "FSI (Floor Space Index)")
    p[31].style = h2_style  # was Heading 1 ("3. FSI...") -- now a subsection of "The Asset"
    p[62].style = h2_style  # "Document Library Contents..." -- now a subsection of the new Diligence Appendix
    p[68].style = h2_style  # "Sources" -- now a subsection of the merged "Gaps & Sources"

    # Every _append_* function below appends new content at the document's
    # true end (doc.add_paragraph()/doc.add_table()), same as they always
    # have -- several of them internally look up doc.paragraphs[4] fresh
    # for heading style, so ALL of them must run before anything gets
    # relocated below; relocating first would shift a fresh paragraphs[4]
    # lookup onto whatever new content happened to land there instead of
    # the real "Methodology Note" heading. `p` (captured at the top of this
    # function, before any of this) stays valid as a relocation anchor
    # throughout, since it holds resolved Paragraph objects, not indices
    # that shift as new paragraphs are inserted elsewhere in the tree --
    # true whether that paragraph's own text/style was just edited above
    # too, since neither changes which _p element `p[N]` wraps.
    body = doc.element.body
    sectPr = body[-1]

    def _capture_batch(append_fn) -> list:
        """Runs append_fn() (which appends new paragraphs/tables at the
        document's true end via doc.add_paragraph()/doc.add_table()) and
        returns exactly the elements it just added, in document order --
        found via XML tree navigation (getnext/getprevious), not Python
        object identity, since lxml doesn't guarantee stable proxy objects
        across separate accesses to the same underlying node."""
        marker = sectPr.getprevious()
        append_fn()
        new_elements = []
        el = marker.getnext() if marker is not None else (body[0] if body[0] is not sectPr else None)
        while el is not None and el is not sectPr:
            new_elements.append(el)
            el = el.getnext()
        return new_elements

    def _demote_first_heading(batch: list, new_style) -> None:
        """Finds the first Heading-1-styled paragraph in a captured batch
        (every _append_*_section function's own heading, since each was
        written as a standalone top-level section) and demotes it to
        `new_style` -- turns it into a subsection of whatever umbrella
        heading it's being relocated under, without touching that
        section's own rendering code at all."""
        for el in batch:
            if el.tag == qn("w:p"):
                para = _Paragraph(el, doc._body)
                if para.style and para.style.name == "Heading 1":
                    para.style = new_style
                    return

    # facts["developer_score"] must exist before _append_developer_score_section
    # / _append_counterparty_summary read it -- computed once here (Overview
    # needs the same flags/score too).
    flags = _classify_flags(facts)
    facts["developer_score"] = _compute_developer_score(facts, flags)

    overview_batch = _capture_batch(lambda: _append_overview_section(doc, facts, flags))

    # The material-findings cards (Completion Slippage/Units Sold/
    # Litigation Load/Land Title) land right under the Executive Summary's
    # own prose (p[7]) -- captured separately from Overview & Flags so it
    # can anchor at a different spot (immediately before p[8],
    # "Counterparty") while Overview & Flags (which carries its own,
    # differently-scoped scorecard) anchors at p[6].
    exec_summary_kpi_batch = _capture_batch(lambda: _append_executive_summary_kpis(doc, facts, flags))

    # Short summaries under Counterparty -- full detail lives in the
    # Diligence Appendix batch built right after.
    counterparty_summary_batch = _capture_batch(lambda: _append_counterparty_summary(doc, facts))

    diligence_appendix_batch = []
    for append_fn in (
        lambda: _append_credit_rating_section(doc, facts),
        lambda: _append_ibbi_check_section(doc, facts),
        lambda: _append_company_profile_section(doc, facts),
        lambda: _append_group_companies_section(doc, facts),
        lambda: _append_developer_score_section(doc, facts),
    ):
        batch = _capture_batch(append_fn)
        _demote_first_heading(batch, h2_style)
        diligence_appendix_batch.extend(batch)

    def _add_diligence_appendix_heading() -> None:
        heading_para = doc.add_paragraph("Diligence Appendix")
        heading_para.style = h1_style

    diligence_appendix_heading_batch = _capture_batch(_add_diligence_appendix_heading)

    # Compliance & Legal Detail gains the full litigation/complaint detail
    # -- Complaint Order Outcomes joins Appeal-Level Judgments there (same
    # anchor the latter already used in the previous pass). Both are
    # demoted to Heading 2, same reasoning as the Diligence Appendix
    # checks: each was written as its own standalone top-level section, so
    # becoming a subsection of an umbrella heading needs its own demotion.
    complaint_outcomes_batch = _capture_batch(lambda: _append_complaint_outcomes_section(doc, facts))
    _demote_first_heading(complaint_outcomes_batch, h2_style)
    _append_cts_land_record_section(doc, facts)  # unmoved, same relative position as before
    appeal_judgments_only_batch = _capture_batch(lambda: _append_appeal_judgments_section(doc, facts))
    _demote_first_heading(appeal_judgments_only_batch, h2_style)
    compliance_batch = complaint_outcomes_batch + appeal_judgments_only_batch

    _append_review_authenticity_section(doc, facts)  # unmoved

    authenticity_batch = _capture_batch(lambda: _append_authenticity_page(doc, facts))
    _demote_first_heading(authenticity_batch, h2_style)
    diligence_appendix_batch.extend(authenticity_batch)

    # Relocate each batch now that every _append_* call above is done, so
    # none of them could have been tripped up by an earlier relocation --
    # p[4]/p[16]/p[17]/p[48]/p[62]/p[65] are still exactly the same
    # paragraph objects as when this function started, since nothing has
    # moved yet (only text/style on some of them changed, which doesn't
    # affect their position).
    # Overview & Flags anchors at p[6] ("Executive Summary" heading), not
    # p[4] -- the Methodology Note (p[4]/p[5]) is removed entirely below,
    # so the batch lands directly before Executive Summary instead of
    # before what used to be the Methodology Note.
    for el in overview_batch:
        p[6]._p.addprevious(el)
    for el in exec_summary_kpi_batch:
        p[8]._p.addprevious(el)
    for el in counterparty_summary_batch:
        p[16]._p.addprevious(el)
    for el in diligence_appendix_heading_batch:
        p[62]._p.addprevious(el)
    for el in diligence_appendix_batch:
        p[65]._p.addprevious(el)
    for el in compliance_batch:
        p[48]._p.addprevious(el)

    # Land Identification -- the one relocation of EXISTING template
    # content rather than newly-appended content: pulled from under
    # Counterparty (where it used to anchor the old combined "Legal
    # Identifiers" section) to sit under "The Asset" instead, right before
    # its first piece of content (location_coordinates_note).
    for el in (p[9]._p, t[0]._tbl, p[10]._p):
        p[17]._p.addprevious(el)

    # Methodology Note removed entirely, per request -- safe only now that
    # every _append_*_section call above is done (several of them looked up
    # doc.paragraphs[4].style fresh for their own heading style, so this
    # element had to survive until the very last _append_* call finished).
    p[4]._p.getparent().remove(p[4]._p)
    p[5]._p.getparent().remove(p[5]._p)

    # Deferred from wherever they were decided above (see the comment by
    # `_paragraphs_to_remove`'s declaration) -- last thing before save, once
    # every fixed-index doc.paragraphs[...] lookup (Sources list, section
    # consolidation, Methodology Note removal) is done depending on stable
    # positions.
    for paragraph in _paragraphs_to_remove:
        _remove_paragraph(paragraph)

    # Genuinely the LAST content change before save -- see
    # _set_paragraph_as_bullets' own docstring for why this must run only
    # after every fixed-index doc.paragraphs[...] lookup above (Sources
    # list, section consolidation, Methodology Note removal) is done.
    _apply_deferred_bullets(facts)

    # Only safe AFTER _apply_deferred_bullets -- a heading whose only
    # content was a gap-only field can't be judged empty until that
    # field's paragraph has actually been removed.
    _remove_empty_section_headings(doc)

    # Version log (run time + Anthropic API cost) -- Internal only, on the
    # cover page, right above "Overview & Flags". Before the alignment pass
    # below so it's treated the same as everything else.
    if doc_variant == "internal":
        _insert_version_log(doc, p, elapsed_seconds, cost_usd, api_calls)

    # Alignment pass: flowing body text justified (both margins), table
    # content centered for short values and left-aligned for prose cells.
    # Last thing before save -- runs after every row/paragraph add-or-remove
    # above so it sees the document's final shape.
    _justify_body_paragraphs(doc)
    _center_all_table_cells(doc)
    _apply_table_pagination(doc)

    doc.save(out_path)
    _ACTIVE_EXTERNAL_FACTS = None

    if doc_variant == "external":
        # Re-opens the file we just saved (not the in-memory `doc`) so this
        # checks exactly what a reader will actually open -- catches a
        # future code change silently reintroducing any of the specific
        # bugs this session found and fixed by hand (see the function's own
        # docstring). Fails loudly rather than silently shipping a
        # regressed document.
        violations = _verify_external_document_quality(out_path)
        if violations:
            raise RuntimeError(
                f"External Charter quality gate failed for {out_path} "
                f"({len(violations)} violation(s)):\n" + "\n".join(f"  - {v}" for v in violations)
            )

        # Citation-completeness checks (the mechanical regex heuristic AND
        # the Task-8 LLM judgment pass) are ADVISORY, not fatal -- see both
        # functions' own docstrings on why a missing citation marker (which
        # can reflect a genuine, pre-existing gap in the underlying facts'
        # source tagging, not a code bug) shouldn't be able to discard an
        # otherwise-good run the way the hard style-regression gate above
        # does. Printed for the operator to review, same visibility level
        # as this pipeline's other console progress output.
        citation_flags = _check_citation_completeness(out_path) + _llm_verify_citation_completeness(out_path)
        if citation_flags:
            print(f"[!] Citation-completeness check flagged {len(citation_flags)} claim(s) for review in {out_path}:")
            for flag in citation_flags:
                print(f"    - {flag}")

    # For doc_variant="internal", `facts` above is the caller's own object
    # (only "external" works on a deep copy), so these two rendering-only
    # keys must be stripped before returning -- otherwise they'd leak into
    # whatever the caller persists (run_company_charter dumps this same
    # dict to .facts.json right after calling us).
    facts.pop("_doc_variant", None)
    facts.pop("_citation_registry", None)


# ---------------------------------------------------------------------------
# Authenticity page -- computed from the actual sources/gaps data, not
# self-reported by the model. A narrative claim like "this report is highly
# reliable" would itself just be another unverified claim; a count of what
# tier each cited source actually falls into is auditable against the
# Sources section a reader can already see.
# ---------------------------------------------------------------------------

_SOURCE_TIERS = [
    ("Primary regulatory record (MahaRERA/MCA, or a document opened from it)", (
        "maharerait.maharashtra.gov.in", "maharera.maharashtra.gov.in", "mca.gov.in",
        "output" + os.sep, "output/",
    )),
    # Credit-rating and government legal/insolvency records: not yet produced by any pass
    # in this file as of this session (the credit-rating lookup and IBBI/NCLT check are
    # separate, not-yet-built follow-on work) -- listed here now so the tier taxonomy is
    # ready the moment those checks start citing sources, rather than everything from them
    # silently landing in "Other" until someone remembers to add the tier.
    ("Credit rating agency (CRISIL/ICRA/CARE/India Ratings)", (
        "crisil.com", "crisilratings.com", "icra.in", "careratings.com", "careedge.in", "indiaratings.co.in",
    )),
    ("Government legal/insolvency record (IBBI, NCLT, NCDRC, MahaREAT judgments)", (
        "ibbi.gov.in", "nclt.gov.in", "ncdrc.nic.in", "mahareat.maharashtra.gov.in", "indiankanoon.org",
    )),
    ("Corporate-registry mirror", ("zaubacorp.com", "tofler.in", "instafinancials.com")),
    ("Live Google Maps verification", ("google.com/maps", "Google Maps")),
    # Genuine third-party homebuyer-advocacy/watchdog bodies -- explicitly documented (this
    # session's research) as more reliable than a review site or listing portal, since they
    # aren't commercial review-hosting businesses with their own credibility problems.
    ("Homebuyer-advocacy / watchdog body", ("fpce.in", "indianrealestateforum.com")),
    ("Real-estate aggregator / listing site", (
        "99acres.com", "nobroker.in", "squareyards.com", "magicbricks.com", "housing.com",
        "indextap.com", "propertypistol.com", "navi.com", "regrob.com",
    )),
    # A developer/promoter's own website is self-published and not independently verified --
    # kept as its own, lower tier rather than folded into independent aggregators above (it
    # previously was; that conflated "the subject describing itself" with "an independent
    # third party describing the subject", which are very different trust levels).
    ("Developer/promoter's own website (self-published)", (
        "godrejproperties.com", "godrejvanantara.com", "godrejgroup.org.in",
    )),
    ("News / press coverage", ("sproutsnews.com", "economictimes", "moneycontrol.com", "livemint.com")),
    ("Social media / user-generated content", ("facebook.com", "twitter.com", "x.com", "reddit.com", "mouthshut.com", "quora.com")),
    # Crowd-sourced/user-editable references: not cited anywhere in this project to date,
    # listed for when one ever is, rather than letting it default to "Other / unclassified".
    ("Crowd-sourced / user-editable reference (Wikipedia, etc.)", ("wikipedia.org", "wikidata.org")),
]


def _classify_source_tier(ref: str) -> str:
    ref_lower = (ref or "").lower()
    for tier_name, markers in _SOURCE_TIERS:
        if any(marker.lower() in ref_lower for marker in markers):
            return tier_name
    return "Other / unclassified source"


# ---------------------------------------------------------------------------
# Cross-run source-trust registry -- the "trusted list vs. web list" design:
# a source found via open-web research (i.e. any _SOURCE_TIERS marker below
# the already-trusted MahaRERA/MCA-mirror tiers) that keeps coming up as a
# corroborating source across separate projects gets auto-promoted once it
# crosses _SOURCE_PROMOTION_HIT_THRESHOLD distinct projects. Promotion is
# fully automatic (no human gate), but every promotion is logged with a
# review note here AND surfaced in the Internal Charter only (see
# _fill_template) -- an algorithmic promotion is not the same thing as a
# manually-vetted trusted source, and a reader of the Internal document
# should be able to see which trust-tier upgrades haven't had a human look
# at them yet. This is deliberately a plain JSON file at the repo root, not
# per-project output/ state -- the whole point is that it persists and
# accumulates ACROSS every project this pipeline ever runs against.
# ---------------------------------------------------------------------------

_SOURCE_TRUST_REGISTRY_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "source_trust_registry.json")
_SOURCE_PROMOTION_HIT_THRESHOLD = 5

# Tiers this pipeline deliberately built a dedicated, code-driven lookup
# for (MahaRERA/MCA records, credit-rating agencies, IBBI/NCLT, live Google
# Maps) -- these are the "trusted list" and never need promoting, no matter
# how often they're cited. Only tiers an LLM's own open-web research
# actually DISCOVERED (aggregators, press, social media, watchdog sites,
# a developer's own site, Wikipedia) are eligible to graduate into that
# trusted set via repeated corroboration.
_ALREADY_TRUSTED_TIERS = frozenset({
    "Primary regulatory record (MahaRERA/MCA, or a document opened from it)",
    "Credit rating agency (CRISIL/ICRA/CARE/India Ratings)",
    "Government legal/insolvency record (IBBI, NCLT, NCDRC, MahaREAT judgments)",
    "Corporate-registry mirror",
    "Live Google Maps verification",
})


def _load_source_trust_registry() -> dict:
    if not os.path.exists(_SOURCE_TRUST_REGISTRY_PATH):
        return {"domains": {}}
    try:
        with open(_SOURCE_TRUST_REGISTRY_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return {"domains": {}}


def _save_source_trust_registry(registry: dict) -> None:
    with open(_SOURCE_TRUST_REGISTRY_PATH, "w", encoding="utf-8") as f:
        json.dump(registry, f, indent=2, ensure_ascii=False)


def _extract_open_web_domain(ref: str) -> str | None:
    """Returns the specific _SOURCE_TIERS marker (e.g. "99acres.com") that
    `ref` matches, restricted to tiers NOT already considered trusted --
    i.e. only a genuine open-web source is eligible for promotion, never a
    MahaRERA/MCA-mirror source that's trusted by design already."""
    ref_lower = (ref or "").lower()
    for tier_name, markers in _SOURCE_TIERS:
        if tier_name in _ALREADY_TRUSTED_TIERS:
            continue
        for marker in markers:
            if marker.lower() in ref_lower:
                return marker
    return None


def _record_source_hits_and_promote(facts: dict, reg_no: str) -> list[str]:
    """Tallies each of this run's cited open-web sources into the
    persistent, cross-run source_trust_registry.json (one hit per distinct
    project a domain has corroborated, not one per Charter re-run of the
    same project), auto-promoting any domain that reaches
    _SOURCE_PROMOTION_HIT_THRESHOLD. Returns review-note strings for
    promotions that happened on THIS run specifically, for Internal-only
    rendering. The registry write itself is unconditional regardless of
    doc_variant -- it's cross-run pipeline state, not document content, so
    it happens once per run_company_charter call, not once per variant."""
    registry = _load_source_trust_registry()
    domains = registry.setdefault("domains", {})
    promotion_notes = []
    seen_this_run = set()

    for s in facts.get("sources", []) or []:
        domain = _extract_open_web_domain(s.get("ref") or s.get("label") or "")
        if not domain or domain in seen_this_run:
            continue
        seen_this_run.add(domain)
        entry = domains.setdefault(domain, {
            "hit_count": 0, "projects": [], "promoted": False,
            "promoted_on_project": None, "review_note": None,
        })
        if reg_no not in entry["projects"]:
            entry["projects"].append(reg_no)
            entry["hit_count"] = len(entry["projects"])
        if not entry["promoted"] and entry["hit_count"] >= _SOURCE_PROMOTION_HIT_THRESHOLD:
            entry["promoted"] = True
            entry["promoted_on_project"] = reg_no
            note = (
                f"Source \"{domain}\" auto-promoted to trusted status after {entry['hit_count']} "
                f"corroborating uses across separate projects (most recently {reg_no}) -- this was an "
                "automatic, frequency-based promotion, not a manual accuracy review; recommend a "
                "one-time spot-check of this source's reliability."
            )
            entry["review_note"] = note
            promotion_notes.append(note)

    _save_source_trust_registry(registry)
    return promotion_notes


def _is_identifier_annotation(annotation: str) -> bool:
    """True when a "(...)" annotation on a source label is a raw code
    identifier -- an API/JSON field name, a dotted path, a snake_case key --
    rather than something written for a reader.

    A human annotation contains a space ("First Schedule", "Performa A-1"); an
    identifier is a single token, and is additionally either dotted or
    camelCase. Both conditions are required so an ordinary one-word annotation
    ("(notarized)") is not mistaken for code."""
    inner = annotation.strip().strip("()").strip()
    if not inner or " " in inner:
        return False
    return "." in inner or bool(re.search(r"[a-z][A-Z]", inner)) or "_" in inner


def _clean_source_label(raw_source: str) -> str | None:
    """Turns a _FIELD_WITH_SOURCE value's raw `source` string -- a
    filesystem path under output/<reg_no>/, a URL, or several of either
    joined with "; " -- into a short, plain-text inline citation. Strips
    the output/<reg_no>/documents|raw/ path prefix down to just the
    filename (prefixing a bare raw/*.json file with "MahaRERA", since
    that's literally what those files are), reduces a URL to its domain,
    and preserves any trailing parenthetical annotation the source itself
    already carries (e.g. "(First Schedule)"). Returns None for a source
    that explicitly says it's an unconfirmed gap ("gap -- see Gaps
    section") -- there is nothing to cite there, not a citation to
    invent.

    One kind of annotation is dropped rather than preserved: a raw API/JSON
    key, e.g. "MahaRERA project registration data
    (projectLegalLandAddressDetails.boundariesSouth)". CLAUDE.md Section B
    forbids a JSON key in EITHER document, and this one reached the page --
    it is also what made charter_report.py's own gate report a "bare-domain
    mention", since "...Details.boundariesSouth" pattern-matches a hostname.
    Human annotations ("First Schedule", "Performa A-1") contain a space and
    are kept; an identifier token never does. See _is_identifier_annotation."""
    if not raw_source or not raw_source.strip():
        return None
    if raw_source.strip().lower().startswith("gap"):
        return None

    def _clean_one(piece: str) -> str:
        piece = piece.strip()
        annotation = ""
        m = re.search(r"\s*(\([^)]*\))\s*$", piece)
        if m:
            if not _is_identifier_annotation(m.group(1)):
                annotation = " " + m.group(1)
            piece = piece[: m.start()].strip()
        if piece.lower().startswith(("http://", "https://")):
            domain = re.sub(r"^https?://(www\.)?", "", piece).split("/")[0]
            return domain + annotation
        piece = piece.replace("\\", "/")
        name = piece.split("/")[-1]
        folder = piece.rsplit("/", 2)[-2] if "/" in piece else ""
        if folder == "raw" and name.endswith(".json"):
            name = f"MahaRERA {name}"
        return name + annotation

    pieces = [p for p in _split_outside_parens(raw_source, ";") if p.strip()]
    return "; ".join(_clean_one(p) for p in pieces)


def _split_outside_parens(text: str, sep: str) -> list[str]:
    """Splits `text` on `sep`, but only where `sep` sits outside any
    "(...)" span -- a real bug this caught: a single source string like
    "Title Report.pdf (First Schedule; land-title chain; 30-year
    litigation search)" has semicolons INSIDE its own trailing
    parenthetical (several related concepts, not several sources), and a
    plain str.split(";") tore it into 3 fake separate citations, one of
    them left with a dangling, unmatched ")"."""
    pieces = []
    depth = 0
    current = []
    for ch in text:
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        if ch == sep and depth == 0:
            pieces.append("".join(current))
            current = []
        else:
            current.append(ch)
    pieces.append("".join(current))
    return pieces


# Keyword rules for _generic_one_label, tried in order (most specific
# first) -- map an Internal-doc cleaned label (an internal artifact
# filename or a bare domain) to a generic, client-facing description for
# the External doc. Matched against the lowercased label with any trailing
# "(annotation)" stripped off entirely (never reattached); see
# _generic_one_label.
_MAHARERA_JSON_GENERIC = (
    ("partners.json", "MahaRERA promoter/partner filing"),
    ("complaints.json", "MahaRERA complaint record"),
    ("appeals.json", "MahaRERA appeal record"),
    ("past_experiences.json", "MahaRERA past-experience filing"),
    ("projects.json", "MahaRERA project filing"),
    # Hand-written source strings, not filenames: a facts.json `source` can say
    # "MahaRERA complaints/appeals data" directly. Left unmapped, the generic
    # fallback reduced that to the bare category label "appeals data", which
    # Section C names as exactly what an External citation must never be.
    ("complaints/appeals", "MahaRERA complaint and appeal records for this project registration"),
    ("complaint/appeal", "MahaRERA complaint and appeal records for this project registration"),
    ("appeals data", "MahaRERA appeal records for this project registration"),
    ("complaints data", "MahaRERA complaint records for this project registration"),
)
_DOMAIN_GENERIC = (
    ("zaubacorp.com", "ZaubaCorp (Corporate Registry)"),
    ("tofler.in", "Tofler (Corporate Registry)"),
    ("instafinancials.com", "InstaFinancials (Corporate Registry)"),
    ("thecompanycheck.com", "TheCompanyCheck (Corporate Registry)"),
    ("ibbi.gov.in", "IBBI (Insolvency and Bankruptcy Board of India)"),
    ("nclt.gov.in", "NCLT (National Company Law Tribunal)"),
    ("efiling.nclt.gov.in", "NCLT (National Company Law Tribunal)"),
    ("indiankanoon.org", "Indian Kanoon (Case Law Record)"),
    ("insolvencytracker.in", "Insolvency Tracker"),
    ("linkedin.com", "LinkedIn (Professional Profile)"),
    ("rocketreach.co", "RocketReach (Organisation Directory)"),
    ("igrmaharashtra.gov.in", "IGR Maharashtra (Dept. of Registration & Stamps)"),
    ("maharerait.maharashtra.gov.in", "MahaRERA"),
    ("maharera.maharashtra.gov.in", "MahaRERA"),
)

# Keyword -> canonical domain, for matching an ORGANISATION NAME as it might
# appear in free-form flag/gap prose (e.g. "TheCompanyCheck lists three
# charges...") back to the SAME generic label a structured {"value",
# "source"} citation elsewhere in this same document would produce for that
# source -- e.g. company_profile_check.url citing thecompanycheck.com. This
# is what lets a flag item and an Entity Profile table row citing the same
# real-world source share the exact same "[N]", rather than minting a
# second, inconsistent citation number for a name-only mention. Checked in
# order, longest/most-specific phrase first, against the lowercased flag
# text -- never against a domain (that path already goes through
# _clean_source_label/_generic_one_label directly).
_FLAG_TEXT_ORG_ALIASES = (
    ("thecompanycheck", "thecompanycheck.com"),
    ("instafinancials", "instafinancials.com"),
    ("zaubacorp", "zaubacorp.com"),
    ("tofler.in", "tofler.in"),
    ("tofler", "tofler.in"),
    ("ibbi", "ibbi.gov.in"),
    ("nclt", "nclt.gov.in"),
    ("indian kanoon", "indiankanoon.org"),
    ("indiankanoon", "indiankanoon.org"),
    ("insolvency tracker", "insolvencytracker.in"),
    ("linkedin", "linkedin.com"),
    ("rocketreach", "rocketreach.co"),
    ("igr maharashtra", "igrmaharashtra.gov.in"),
    ("maharera", "maharera.maharashtra.gov.in"),
)
_DOCUMENT_KEYWORD_GENERIC = (
    ("form b", "Form B Declaration"),
    ("title report", "Title Report"),
    ("supplemental title", "Title Report"),
    ("non encubrance", "Non-Encumbrance Declaration"),
    ("non-encumbrance", "Non-Encumbrance Declaration"),
    ("layout", "Sanctioned Layout Drawing"),
    ("iod", "IOD (Intimation of Disapproval)"),
    ("encumbrance", "Encumbrance Certificate"),
)


def _generic_one_label(label: str) -> str:
    """Maps ONE source label (see _clean_source_label/_topic_citation) to
    a generic, client-facing description for the External doc -- never a
    literal internal filename/path. Any trailing "(annotation)" (e.g.
    "(First Schedule -- exact phrase found on re-extraction after fixing
    this session's OCR pipeline)") is dropped entirely, not reworded --
    that kind of internal-QA precision note isn't something a numbered
    external citation needs to carry, and some of it is exactly the
    internal-process language this variant exists to keep out. A label
    that's already generic (an organization/product name from a
    facts['sources'] entry, e.g. "99acres, accessed 2026-07-17", "Google
    Maps, accessed ...") is kept as-is. Falls back to "Project record"
    for an unrecognized document filename rather than showing it
    verbatim."""
    m = re.search(r"\s*\([^)]*\)\s*$", label)
    core = label[: m.start()].strip() if m else label
    lowered = core.lower()

    for keyword, generic in _MAHARERA_JSON_GENERIC:
        if keyword in lowered:
            return generic
    for domain, generic in _DOMAIN_GENERIC:
        if domain in lowered:
            return generic
    for keyword, generic in _DOCUMENT_KEYWORD_GENERIC:
        if keyword in lowered:
            return generic
    if lowered.endswith(".pdf") or lowered.endswith(".json"):
        # Last resort, for a document with no facts["sources"] entry to describe
        # it (see _external_source_label, which is tried first). "Project
        # record" was worse than useless -- a bare category label naming
        # nothing, which Section C forbids -- and it also collapsed several
        # genuinely different documents onto one citation number. This at least
        # tells the reader what kind of thing it is and where it came from.
        # A source landing here is a facts-quality problem: the document was
        # cited without ever being catalogued in facts["sources"].
        return "Document supplied by the promoter for this project"
    # A bare organisation NAME (e.g. "TheCompanyCheck", written by hand in a
    # source field) rather than a domain/URL -- resolve it to the SAME
    # generic label its domain form would produce (see _DOMAIN_GENERIC),
    # so "TheCompanyCheck" and "thecompanycheck.com" cited for the same
    # real source in two different facts.json fields converge on one
    # citation number instead of minting two. Checked last, after every
    # more specific table above, so it can never shadow a real match.
    for keyword, domain in _FLAG_TEXT_ORG_ALIASES:
        if keyword in lowered:
            for d, generic in _DOMAIN_GENERIC:
                if d == domain:
                    return generic
            return domain
    return core


def _external_source_label(facts: dict, label: str) -> str:
    """Resolves one source label to the descriptive citation the External
    Sources list prints -- issuer, what the document is, and its date.

    CLAUDE.md Section C: an External Sources entry is "never a raw internal
    filename, never a bare category label like 'Project record' or 'appeals
    data' that a reader cannot check". _generic_one_label alone produced both:
    it falls back to the literal string "Project record" for any unrecognised
    .pdf/.json, and it shipped 'Declaration for one registration 18 March
    24.pdf' verbatim to a client.

    Every facts["sources"] entry already carries exactly the right text in its
    `ref` ("Legal Title Report and Title Certificate, Adv. Preet J. Chheda, 18
    April 2024"), which nothing previously used. This prefers that, and falls
    back to _generic_one_label only for labels that are not documents at all
    (a bare domain from _annotate_flag_citations, say).

    Deduplication still works on the returned string, so two fields citing the
    same document converge on one number -- and, unlike the old generic
    labels, two DIFFERENT documents no longer collapse into a single
    "Project record" entry."""
    core = re.sub(r"\s*\([^)]*\)\s*$", "", label or "").strip()
    # _topic_citation appends ", accessed YYYY-MM-DD"; strip it before matching.
    core = re.sub(r",\s*accessed\s+\d{4}-\d{2}-\d{2}$", "", core, flags=re.IGNORECASE).strip()
    if core:
        for s in facts.get("sources", []) or []:
            raw = (s.get("label") or "").strip()
            if raw and raw.lower() == core.lower():
                ref = (s.get("ref") or "").strip()
                if ref:
                    return ref
                break
    return _generic_one_label(label)


def _register_citation(facts: dict, generic_label: str) -> str:
    """External-variant only: returns "[N]" for `generic_label`, assigning
    the next sequential number the first time this exact generic label is
    seen this document and reusing that same number every later time the
    same source is cited -- a deduped bibliography, never two different
    numbers for one source."""
    registry = facts["_citation_registry"]
    if generic_label not in registry["index"]:
        registry["index"][generic_label] = len(registry["order"]) + 1
        registry["order"].append(generic_label)
    return f"[{registry['index'][generic_label]}]"


def _annotate_flag_citations(facts: dict, text: str) -> str:
    """External-variant only (returns `text` unchanged for Internal):
    Overview & Flags' Imminent/Structural/Monitor items come out of
    _classify_flags as {"text", "field"} -- a facts.json path, never a
    source -- so unlike every other claim in this Charter, a flag line
    naming a real source in plain prose (e.g. "Tofler.in's own charges
    section states...") never resolved to a numbered "[N]", even when
    that exact same source is cited elsewhere in this same document.
    Confirmed live: this is why the External Charter's flag lists showed
    zero citations while the rest of the document did not.

    This scans the already-externalized flag text for a known set of
    organisation-name mentions (_FLAG_TEXT_ORG_ALIASES) and, for the
    FIRST mention of each distinct source, inserts the SAME "[N]" that
    source's own structured citation elsewhere in this document already
    uses (via the same _generic_one_label/_register_citation this whole
    Charter runs on) -- never a fabricated citation for a source that
    isn't actually named in the text. A source mentioned twice in one
    flag (or under two aliases, e.g. "Tofler.in" then "Tofler") is only
    cited once, matching how a citation attaches once per claim
    elsewhere in this Charter rather than once per word."""
    if facts.get("_doc_variant") != "external":
        return text

    cited_domains = set()
    for keyword, domain in _FLAG_TEXT_ORG_ALIASES:
        if domain in cited_domains:
            continue
        match = re.search(re.escape(keyword), text, re.IGNORECASE)
        if not match:
            continue
        marker = _register_citation(facts, _external_source_label(facts, domain))
        text = _insert_marker_at_clause_end(text, match.end(), marker)
        cited_domains.add(domain)
    return text


def _insert_marker_at_clause_end(text: str, from_index: int, marker: str) -> str:
    """Places `marker` at the end of the clause containing `from_index`.

    CLAUDE.md Section C, both halves of the placement rule: a marker "sits at
    the end of the clause it supports, not at the end of the paragraph", and it
    "never lands mid-word or mid-token". Inserting immediately after the matched
    keyword produced literally the example the rule cites against --
    "MahaRERA[10]-registered" -- because the mention sits inside a compound.
    Scanning to the clause boundary satisfies both halves at once."""
    end = len(text)
    depth = 0
    for i in range(from_index, len(text)):
        ch = text[i]
        if ch == "(":
            depth += 1
        elif ch == ")":
            depth = max(0, depth - 1)
        elif depth == 0 and ch in ".;":
            end = i
            break
    return text[:end] + marker + text[end:]


def _citation_text(facts: dict, cleaned_label: str | None) -> str | None:
    """Renders a cleaned source label (see _clean_source_label/
    _topic_citation -- possibly several joined with "; ", when one fact
    cites more than one source) into the citation representation for this
    Charter's doc_variant. "internal": the literal label in parentheses
    (today's behavior, unchanged). "external": each individual source is
    generalized and registered separately -- so a source already cited
    elsewhere keeps the SAME number even when it's now combined with a
    different second source -- and the resulting numbers are joined with
    no separator ("[6][11]"), standard multi-citation style; a citation
    combining two sources that generalize to the SAME description (e.g.
    two different MahaRERA filings of the same category) shows that
    number only once. Returns None (never a fabricated citation) if
    there's no label to cite in the first place."""
    if not cleaned_label:
        return None
    if facts.get("_doc_variant") != "external":
        return f"({cleaned_label})"
    markers = []
    for piece in (p.strip() for p in _split_outside_parens(cleaned_label, ";")):
        if not piece:
            continue
        marker = _register_citation(facts, _external_source_label(facts, piece))
        if marker not in markers:
            markers.append(marker)
    return "".join(markers) if markers else None


# Exact-text rewrites of known facts.json prose, applied via plain
# substring replacement (not regex -- these are long, literal, already-
# written sentences, not patterns) before _EXTERNAL_PROSE_SUBSTITUTIONS
# runs. Purpose: External wants no " -- " (em-dash-style) punctuation
# anywhere, per explicit request -- but a blind regex swap of every " -- "
# for a comma would produce comma splices and broken grammar in a good
# fraction of cases (a dash sometimes joins two independent clauses, which
# needs a period/semicolon, not a comma). Each entry here is a hand-checked,
# grammatically-correct rewrite of one specific known sentence, keyed by
# its exact original text -- Internal is untouched (this dict is only ever
# consulted when doc_variant == "external", same as every other transform
# in _externalize_prose). Not exhaustive against future/unseen prose (see
# _EXTERNAL_PROSE_SUBSTITUTIONS's own note below on that same limitation);
# new fields written with " -- " will need a new entry added here.
_EXTERNAL_DASH_REWRITES = {
    # --- static template labels/headings, hardcoded in this file (not
    # facts.json), shared by every Charter this pipeline generates ---
    "Public Web-Sourced Edition -- Adapted for Maharashtra (MahaRERA)":
        "Public Web-Sourced Edition: Adapted for Maharashtra (MahaRERA)",
    "Every figure above is drawn directly from the underlying facts already researched for this project -- see the flag lists below for detail and the Diligence Appendix for the per-pillar/per-check breakdown behind the Developer Score and Data Authenticity figures.":
        "Every figure above is drawn directly from the underlying facts already researched for this project; see the flag lists below for detail and the Diligence Appendix for the per-pillar/per-check breakdown behind the Developer Score and Data Authenticity figures.",
    "Imminent Red Flags -- act on these before proceeding": "Imminent Red Flags: act on these before proceeding",
    "Structural Flags -- standing characteristics, raise directly with the developer": "Structural Flags: standing characteristics, raise directly with the developer",
    "Monitor Flags -- re-check on a future pass": "Monitor Flags: re-check on a future review",
    "The cards above are the material findings from this Charter's own research -- see Overview & Flags below for the full flag detail behind each one.":
        "The cards above are the material findings from this Charter's own research; see Overview & Flags below for the full flag detail behind each one.",
    "Insolvency Check -- IBBI Corporate Debtor Master Data": "Insolvency Check: IBBI Corporate Debtor Master Data",
    "Checked directly against the Insolvency and Bankruptcy Board of India's public Corporate Debtor Master Data, by the promoter's own CIN -- an exact-identifier lookup, not a name-based guess.":
        "Checked directly against the Insolvency and Bankruptcy Board of India's public Corporate Debtor Master Data, by the promoter's own CIN: an exact-identifier lookup, not a name-based guess.",
    "Checked directly against every rating agency's public database (currently ICRA and Infomerics) for an exact match on the promoter's own legal name -- not a fuzzy or \"probably the same company\" guess, since attributing a rating to the wrong legal entity would itself be a serious error.":
        "Checked directly against every rating agency's public database (currently ICRA and Infomerics) for an exact match on the promoter's own legal name, not a fuzzy or \"probably the same company\" guess, since attributing a rating to the wrong legal entity would itself be a serious error.",
    "Pulled directly from 3 independent public company-registry mirror(s) by CIN -- an exact-identifier lookup, not a name-based guess -- and cross-checked against each other; any disagreement between them on the current director roster is called out under Gaps & Limitations rather than silently resolved.":
        "Pulled directly from 3 independent public company-registry mirror(s) by CIN, an exact-identifier lookup rather than a name-based guess, and cross-checked against each other; any disagreement between them on the current director roster is called out under Gaps & Limitations rather than silently resolved.",
    "Every entity below shares at least one concrete, named link with the promoter above -- a specific director in common, a shared registered office, or a filed subsidiary/associate/JV relationship -- rather than being inferred from a name or industry match.":
        "Every entity below shares at least one concrete, named link with the promoter above (a specific director in common, a shared registered office, or a filed subsidiary/associate/JV relationship) rather than being inferred from a name or industry match.",
    "Each of this promoter's directors, current or past, cross-referenced against how many of the group companies below name them as a shared director -- collapses the 299-entity list to the thing that actually matters here: how concentrated the group's leadership is around a small number of individuals, not a name-by-name read of every affiliated entity.":
        "Each of this promoter's directors, current or past, is cross-referenced against how many of the group companies below name them as a shared director. This collapses the 299-entity list to the thing that actually matters here: how concentrated the group's leadership is around a small number of individuals, not a name-by-name read of every affiliated entity.",
    "This page is generated directly from the same sources and gaps already listed earlier in this document -- it is a count, not a self-assessment.":
        "This page is generated directly from the same sources and gaps already listed earlier in this document. It is a count, not a self-assessment.",
    "This score rates how well-sourced and verified THIS DOCUMENT's own claims are -- source quality, completeness, cross-corroboration, recency, and re-check success -- it is NOT a rating of the underlying project's quality, safety, or investment merit.":
        "This score rates how well-sourced and verified this document's own claims are (source quality, completeness, cross-corroboration, recency, and re-check success). It is not a rating of the underlying project's quality, safety, or investment merit.",
    "It is a weighted average of six criteria computed below from this document's own sources and gaps -- informed by the structure of CRISIL's real-estate methodology (a small number of named factors rather than one flat checklist) but NOT a replica of any CRISIL formula: CRISIL does not publish a numeric weighting scheme for any of its three real-estate products, and the weights used here are this project's own calibration.":
        "It is a weighted average of six criteria computed below from this document's own sources and gaps, informed by the structure of CRISIL's real-estate methodology (a small number of named factors rather than one flat checklist) but not a replica of any CRISIL formula: CRISIL does not publish a numeric weighting scheme for any of its three real-estate products, and the weights used here are this project's own calibration.",
    "Checked directly against every rating agency's public database (currently ICRA and Infomerics) for an exact match on the promoter's own legal name -- not a fuzzy or \"probably the same company\" guess, since attributing a rating to the wrong legal entity would itself be a serious error. Every agency is checked regardless of whether an earlier one already found something, so that if two agencies rate the same entity, both ratings are shown here for comparison rather than silently reporting only one.":
        "Checked directly against every rating agency's public database (currently ICRA and Infomerics) for an exact match on the promoter's own legal name, not a fuzzy or \"probably the same company\" guess, since attributing a rating to the wrong legal entity would itself be a serious error. Every agency is checked regardless of whether an earlier one already found something, so that if two agencies rate the same entity, both ratings are shown here for comparison rather than silently reporting only one.",
    "Each row below reflects the actual order PDF already on file for that complaint (downloaded via the same document-retrieval mechanism used for project documents), classified by a small, named set of outcome keywords -- not a self-reported summary.":
        "Each row below reflects the actual order PDF already on file for that complaint (downloaded via the same document-retrieval mechanism used for project documents), classified by a small, named set of outcome keywords rather than a self-reported summary.",
    "Of 10 cited source(s) in this report, 3 (30%) come from a primary regulatory record, a document opened directly from it, or a live Google Maps route checked in this review -- the highest-confidence tier.":
        "Of 10 cited source(s) in this report, 3 (30%) come from a primary regulatory record, a document opened directly from it, or a live Google Maps route checked in this review: the highest-confidence tier.",
    "Of 8 cited source(s) in this report, 2 (25%) come from a primary regulatory record, a document opened directly from it, or a live Google Maps route checked in this review -- the highest-confidence tier.":
        "Of 8 cited source(s) in this report, 2 (25%) come from a primary regulatory record, a document opened directly from it, or a live Google Maps route checked in this review: the highest-confidence tier.",
    "facts that were sought but could not be confirmed, listed in full under \"Gaps & Limitations\" earlier in this document -- rather than filled in with an estimate.":
        "facts that were sought but could not be confirmed, listed in full under \"Gaps & Limitations\" earlier in this document, rather than filled in with an estimate.",
    "CIN / LLPIN (if incorporated -- note: N/A for an unincorporated Partnership)":
        "CIN / LLPIN (if incorporated; N/A for an unincorporated Partnership)",
    "See Litigation Status -- not a clean read": "See Litigation Status: not a clean read",
    "promoter_portfolio.totals.area_within_5km_lakh_sqft not available -- this pass's promoter_portfolio.json predates the geocoding-based 5km filter (build_promoter_portfolio's subject_project_partners_data/subject_reg_no params), or no subject location could be geocoded. Re-run the pipeline to compute it, rather than treating this as a permanent gap.":
        "Not available. This review's portfolio data predates the geocoding-based 5km filter, or no subject location could be geocoded. This can be computed on a future review rather than treated as a permanent gap.",
    "promoter_portfolio.totals.total_area_developed_lakh_sqft not available -- requires area figures aggregated across the promoter's other MahaRERA-registered projects, not yet computed this pass.":
        "Not available. This requires area figures aggregated across the promoter's other MahaRERA-registered projects, not yet computed this review.",
    "Not applicable this pass -- excluded rather than scored as a failure; remaining weights renormalized to still sum to 100%.":
        "Not applicable this review: excluded rather than scored as a failure; remaining weights renormalized to still sum to 100%.",
    "Complaint volume: 3 total filing(s) on record against this project -- a visible floor worth noting, not blended into zero.":
        "Complaint volume: 3 total filing(s) on record against this project, a visible floor worth noting, not blended into zero.",
    "Entity Rating, 7.5% each) -- each sub-metric independently banded AAA/AA/A/B/C/D.":
        "Entity Rating, 7.5% each), each sub-metric independently banded AAA/AA/A/B/C/D.",

    # --- shared across both fixtures (identical boilerplate text) ---
    "Available -- downloaded": "Downloaded",
    "Form B declares the standard 70% construction/land-cost escrow undertaking (Clause under Rule 5) -- exact bank/branch not disclosed in the documents reviewed.":
        "Form B declares the standard 70% construction/land-cost escrow undertaking (Clause under Rule 5); exact bank/branch is not disclosed in the documents reviewed.",
    "No public rating found for this exact legal entity name from any agency checked (ICRA, Infomerics). This is NOT itself a red flag -- these agencies only rate developers that sought a public rating (typically larger, listed, or NCD-issuing entities); most MahaRERA promoters are too small or private to ever be rated.":
        "No public rating found for this exact legal entity name from any agency checked (ICRA, Infomerics). This is not itself a red flag: these agencies only rate developers that sought a public rating (typically larger, listed, or NCD-issuing entities), and most MahaRERA promoters are too small or private to ever be rated.",
    "Portfolio is based on a name match against MahaRERA's own Promoters-tab search -- punctuation/suffix variants (e.g. 'Pvt Ltd', typos) may cause under-counting; cross-check manually if more projects are expected.":
        "Portfolio is based on a name match against MahaRERA's own Promoters-tab search. Punctuation or suffix variants (e.g. 'Pvt Ltd', typos) may cause under-counting; cross-check manually if more projects are expected.",
    "MahaRERA's Promoters-tab search has no Registered/Revoked toggle (unlike the Projects tab) -- 'lapsed_or_flagged_count' is a best-effort signal from each project's own status field, not an independently confirmed revocation count.":
        "MahaRERA's Promoters-tab search has no Registered/Revoked toggle (unlike the Projects tab); the flagged-project count is a best-effort signal from each project's own status field, not an independently confirmed revocation count.",
    "Mortgage lender is NOT tracked here across a promoter's portfolio: none of MahaRERA's structured project category APIs (projects, partners, professionals, sro_details, past_experiences, documents, complaints, appeals, spocs -- checked exhaustively) expose a bank/lender/mortgage/finance field anywhere. Lender identity is only ever recoverable as free text inside a project's own documents (see company_charter.py's mortgage_lender field, sourced from a per-project document read), which this deterministic, document-free portfolio scan does not open. Building cross-project lender tracking would require downloading and OCR/reading documents for every project in a promoter's portfolio -- a different, much heavier architecture, out of scope for this pass.":
        "Mortgage lender is not tracked here across a promoter's portfolio: none of MahaRERA's structured project category APIs (projects, partners, professionals, sro_details, past_experiences, documents, complaints, appeals, spocs, checked exhaustively) expose a bank, lender, mortgage, or finance field anywhere. Lender identity is only ever recoverable as free text inside a project's own documents, which this deterministic, document-free portfolio scan does not open. Building cross-project lender tracking would require downloading and reading documents for every project in a promoter's portfolio, a different, much heavier approach, out of scope for this review.",
    "Past-experience completion dates (on_time_rate_pct and friends) are self-reported by the promoter to MahaRERA, not independently verified against any external record -- treat this as the promoter's own claimed track record, not a confirmed one.":
        "Past-experience completion dates are self-reported by the promoter to MahaRERA, not independently verified against any external record; treat this as the promoter's own claimed track record, not a confirmed one.",
    "total_area_developed_lakh_sqft and area_within_5km_lakh_sqft are both summed from each portfolio project's own past_experiences.landArea (self-reported, sqm, converted to lakh sq ft), excluding the subject project's own entry -- if the promoter declared the SAME historical project's area under more than one of its current registrations, that area is counted once per declaration, not deduplicated (the same limitation already applies to on_time_count/delayed_count above).":
        "Total area developed and area within 5km figures are both summed from each portfolio project's own reported land area (self-reported, converted to lakh sq ft), excluding the subject project's own entry. If the promoter declared the same historical project's area under more than one of its current registrations, that area is counted once per declaration, not deduplicated.",
    "area_within_5km_lakh_sqft additionally requires geocoding: the subject project's own locality (from its partners category data, if supplied) is resolved via OpenStreetMap's public Nominatim API -- a free, no-API-key service, rate-limited to ~1 request/second per its usage policy. Each past_experiences entry's own address is geocoded by preference on any 6-digit Indian pincode found inside it (confirmed live: MahaRERA's own `address` field is often a full legal land description -- survey numbers, stray commas -- that Nominatim's free-form search fails on entirely, even when the correct locality name is embedded in it; the pincode alone resolves reliably), falling back to the raw address text, then to the portfolio project's district, if no pincode is present. An entry that still can't be geocoded is excluded from the 5km sum entirely (never guessed in or out), and a pincode covers a small area, not a point, so this figure is a reasonable estimate, not a surveyed one. If no subject location was supplied at all, or it can't be geocoded, area_within_5km_lakh_sqft stays None entirely.":
        "The area-within-5km figure additionally requires geocoding: the subject project's own locality is resolved via OpenStreetMap's public Nominatim API, a free, no-API-key service, rate-limited to roughly 1 request per second. Each portfolio entry's own address is geocoded by preference on any 6-digit Indian pincode found inside it, since MahaRERA's own address field is often a full legal land description (survey numbers, stray commas) that free-form search fails on, while the pincode alone resolves reliably, falling back to the raw address text and then to the project's district if no pincode is present. An entry that still cannot be geocoded is excluded from the 5km sum entirely, and a pincode covers a small area rather than a point, so this figure is a reasonable estimate, not a surveyed one. If no subject location was supplied at all, or it cannot be geocoded, this figure is left blank entirely.",
    "No public source discloses internal team headcounts by function (Liaisoning / Project Development / Sales & CRM) -- a structural limitation of publicly-available data, not something a future pass can close.":
        "No public source discloses internal team headcounts by function (Liaisoning / Project Development / Sales & CRM). This is a structural limitation of publicly available data, not something a future review can close.",
    "Debt-to-capital ratio, secured-debt ratio, and default occurrence are not disclosed by any source this pipeline checks (MahaRERA/ZaubaCorp don't carry balance-sheet debt structure) -- only resolvable when a credit-rating rationale or MCA financial filing states these figures.":
        "Debt-to-capital ratio, secured-debt ratio, and default occurrence are not disclosed by any source checked (MahaRERA and ZaubaCorp do not carry balance-sheet debt structure). This is only resolvable when a credit-rating rationale or MCA financial filing states these figures.",
    "RERA compliance history (registration timeliness, extension/penalty record) is not yet computed as a scored metric -- no data source has been wired in for this criterion yet; a pending build, not a permanent public-data gap.":
        "RERA compliance history (registration timeliness, extension/penalty record) is not yet computed as a scored metric. No data source has been wired in for this criterion yet; this is a pending build, not a permanent public-data gap.",
    "GST/TDS compliance is not yet computed as a scored metric -- no data source has been wired in for this criterion yet; a pending build, not a permanent public-data gap.":
        "GST/TDS compliance is not yet computed as a scored metric. No data source has been wired in for this criterion yet; this is a pending build, not a permanent public-data gap.",
    "Nearest metro station -- estimated via web search, not a live driving-route lookup.":
        "Nearest metro station, estimated via web search, not a live driving-route lookup.",

    # --- Pranami Bliss ---
    "Pranami Bliss (MahaRERA P51800077150) is a 234-unit residential redevelopment on Veera Desai Road, Azad Nagar, Andheri West, Mumbai, by Pranami Neev Realty Limited -- a Mumbai-market SPV (incorporated 2022) of the longer-established Pranami Group (Ranchi, founded 2002).":
        "Pranami Bliss (MahaRERA P51800077150) is a 234-unit residential redevelopment on Veera Desai Road, Azad Nagar, Andheri West, Mumbai, by Pranami Neev Realty Limited, a Mumbai-market SPV (incorporated 2022) of the longer-established Pranami Group (Ranchi, founded 2002).",
    "Amalgamated plot CTS No. 183(pt), Village Mauje Aambivali -- comprising Property No. 49 (Survey No. 133(part) / City Survey No. 630-B(part)) and Property No. 50 (Survey No. 132(part) / City Survey No. 183(part))":
        "Amalgamated plot CTS No. 183(pt), Village Mauje Aambivali, comprising Property No. 49 (Survey No. 133(part) / City Survey No. 630-B(part)) and Property No. 50 (Survey No. 132(part) / City Survey No. 183(part))",
    "None disclosed in the Title Report -- no acquisition/reservation affecting this land was found in the 30-year search":
        "None disclosed in the Title Report: no acquisition or reservation affecting this land was found in the 30-year search",
    "1,674.63 sq.m (same as total gross area -- no deductions disclosed in the Title Report)":
        "1,674.63 sq.m (same as total gross area; no deductions disclosed in the Title Report)",
    "Public Limited Company (per ZaubaCorp/InstaFinancials -- Company limited by shares, Non-government)":
        "Public Limited Company (per ZaubaCorp/InstaFinancials: company limited by shares, non-government)",
    "U70109MH2022PLC385473 -- confirmed identically across ZaubaCorp, Tofler, and InstaFinancials, and on the company's own PAN card":
        "U70109MH2022PLC385473, confirmed identically across ZaubaCorp, Tofler, and InstaFinancials, and on the company's own PAN card",
    "Not separately disclosed -- Form B's authorizing resolution (2 Feb 2024) does not itself restate a registered-office address distinct from the MCA record above":
        "Not separately disclosed. Form B's authorizing resolution (2 Feb 2024) does not itself restate a registered-office address distinct from the MCA record above.",
    "Not applicable -- no separate planning-stage address disclosed anywhere in this project's documents":
        "Not applicable: no separate planning-stage address is disclosed anywhere in this project's documents.",
    "Mr. Sundeep Poddar, Director -- signs both the Form B declaration and the Non-Encumbrance/Legal declarations on the company's behalf":
        "Mr. Sundeep Poddar, Director, signs both the Form B declaration and the Non-Encumbrance/Legal declarations on the company's behalf.",
    "Bijay Kumar Agarwal (DIN 00448678, appointed 2022-09-07), Nitish Kumar Agarwal (DIN 02750231, appointed 2022-06-27), Sundeep Poddar (DIN 05217062, appointed 2022-06-27) -- current directors, confirmed identically by ZaubaCorp, Tofler, and InstaFinancials (no cross-source disagreement, unlike some other promoters checked)":
        "Bijay Kumar Agarwal (DIN 00448678, appointed 2022-09-07), Nitish Kumar Agarwal (DIN 02750231, appointed 2022-06-27), and Sundeep Poddar (DIN 05217062, appointed 2022-06-27) are the current directors, confirmed identically by ZaubaCorp, Tofler, and InstaFinancials with no cross-source disagreement.",
    "Land owner of record: Maharashtra Housing and Area Development Authority (MHADA), via two lessee societies -- Azad Nagar Sai-Chhaya CHS Ltd. (Property 49) and Azad Nagar Himalaya CHS Ltd. (Property 50), merged into one society by Deputy Registrar order dated 6 March 2024.":
        "Land owner of record: Maharashtra Housing and Area Development Authority (MHADA), via two lessee societies: Azad Nagar Sai-Chhaya CHS Ltd. (Property 49) and Azad Nagar Himalaya CHS Ltd. (Property 50), merged into one society by Deputy Registrar order dated 6 March 2024.",
    "Two distinct matters, kept separate rather than blended into one verdict: (1) LAND TITLE litigation -- the Title Report's own 30-year search (1994-2023, Bombay High Court, Bombay City Civil Court Dindoshi, Small Causes Court Bandra, Debt Recovery Tribunal) found no litigation directly against Property 49 or 50, save for a Notice of Lis Pendens (20 Dec 2017) concerning an adjacent, unrelated society (Azad Nagar Krupa Sagar CHS) with no relief sought against this subject property. Form B additionally declares \"we have no litigation on the said Land.\" (2) RERA CONSUMER COMPLAINTS/APPEALS -- MahaRERA's own live record shows 0 complaints and 0 appeals against this project as of this pass. Neither the land itself nor the promoter's delivery/compliance record shows disclosed litigation risk at this time.":
        "Two distinct matters, kept separate rather than blended into one verdict. (1) Land title litigation: the Title Report's own 30-year search (1994-2023, Bombay High Court, Bombay City Civil Court Dindoshi, Small Causes Court Bandra, Debt Recovery Tribunal) found no litigation directly against Property 49 or 50, save for a Notice of Lis Pendens (20 Dec 2017) concerning an adjacent, unrelated society (Azad Nagar Krupa Sagar CHS) with no relief sought against this subject property. Form B additionally declares \"we have no litigation on the said Land.\" (2) RERA consumer complaints and appeals: MahaRERA's own live record shows 0 complaints and 0 appeals against this project as of this review. Neither the land itself nor the promoter's delivery/compliance record shows disclosed litigation risk at this time.",
    "Andheri West / Oshiwara residential-commercial corridor -- proximate to":
        "Andheri West / Oshiwara residential-commercial corridor, proximate to",
    "Governed by MHADA's redevelopment framework for the two amalgamated cessed/MHADA-lease buildings (Bldg. 49 -- Azad Nagar Sai-Chhaya CHS; Bldg. 50 -- Azad Nagar Himalaya CHS) on CTS 183(pt), Village Ambivali, under DCPR 2034 -- confirmed via the MHADA Intimation of Approval (IOA) chain (initial \"Zero FSI\" IOA dated 18-Sep-2023, amended 28-Mar-2025 for \"FSI Potential\").":
        "Governed by MHADA's redevelopment framework for the two amalgamated cessed/MHADA-lease buildings (Bldg. 49, Azad Nagar Sai-Chhaya CHS; Bldg. 50, Azad Nagar Himalaya CHS) on CTS 183(pt), Village Ambivali, under DCPR 2034. Confirmed via the MHADA Intimation of Approval (IOA) chain (initial \"Zero FSI\" IOA dated 18-Sep-2023, amended 28-Mar-2025 for \"FSI Potential\").",
    "The project's approvals moved from an initial \"Zero FSI\" IOA (18-Sep-2023 -- approving the redevelopment in principle without quoting specific FSI/BUA figures) to an amended IOA (28-Mar-2025) for \"FSI Potential\" -- also without the amendment letter itself quoting the specific area/FSI numbers being unlocked.":
        "The project's approvals moved from an initial \"Zero FSI\" IOA (18-Sep-2023, approving the redevelopment in principle without quoting specific FSI/BUA figures) to an amended IOA (28-Mar-2025) for \"FSI Potential,\" also without the amendment letter itself quoting the specific area/FSI numbers being unlocked.",
    "1,674.63 sq.m (per Title Report -- see Land Identification)":
        "1,674.63 sq.m (per Title Report; see Land Identification)",
    "Not stated in the documents reviewed -- the Zero-FSI IOA and its 28-Mar-2025 amendment both approve the redevelopment/amended plan in principle without quoting a specific BUA figure.":
        "Not stated in the documents reviewed. The Zero-FSI IOA and its 28-Mar-2025 amendment both approve the redevelopment in principle without quoting a specific BUA figure.",
    "Not applicable -- no mortgage was disclosed as currently existing (see mortgage_lender below); LLDPL holds a contractual RIGHT to mortgage the free-sale area under the Supplemental Agreements, not a disclosed existing charge.":
        "Not applicable. No mortgage was disclosed as currently existing; LLDPL holds a contractual right to mortgage the free-sale area under the Supplemental Agreements, not a disclosed existing charge.",
    "Not computed -- BUA figure unavailable from the documents reviewed (see fsi_interpretation).":
        "Not computed. BUA figure unavailable from the documents reviewed.",
    "No lender/bank named in either Non-Encumbrance declaration reviewed -- both are self-declarations by the promoter (\"free from all sort of encumbrances\", \"free from all sort of legal encumbrances\") with no financial institution identified.":
        "No lender or bank is named in either Non-Encumbrance declaration reviewed. Both are self-declarations by the promoter (\"free from all sort of encumbrances\", \"free from all sort of legal encumbrances\") with no financial institution identified.",
    "Registered 2024-07-23 (acknowledgement REA51800155212), currently \"Certificate Signed\" status, proposed completion 2027-04-15 (unchanged since original registration -- no extension on record).":
        "Registered 2024-07-23 (acknowledgement REA51800155212), currently \"Certificate Signed\" status, proposed completion 2027-04-15 (unchanged since original registration; no extension on record).",
    "Not independently confirmed from the documents reviewed this pass (see Gaps) -- MahaRERA's own project API does not expose a bank/account-number field.":
        "Not independently confirmed from the documents reviewed this review (see Gaps). MahaRERA's own project API does not expose a bank or account-number field.",
    "Form B declares no litigation on the land, holds a legal title report, and commits to the 15-Apr-2027 completion date -- signed by Sundeep Poddar, Director, under a board resolution dated 2-Feb-2024.":
        "Form B declares no litigation on the land, holds a legal title report, and commits to the 15-Apr-2027 completion date, signed by Sundeep Poddar, Director, under a board resolution dated 2-Feb-2024.",
    "specific to this micro-pocket was found -- the figures above are real-estate-aggregator averages":
        "specific to this micro-pocket was found. The figures above are real-estate-aggregator averages",
    "Price on request -- not disclosed": "Price on request (not disclosed)",
    "tracked by one aggregator) -- consistent with an actively developing":
        "tracked by one aggregator), consistent with an actively developing",
    "2027-04-15 -- unchanged since original registration, no extension on record":
        "2027-04-15, unchanged since original registration, no extension on record",
    "Per Pranami Group's own website (pranamigroup.com/about), the group was founded in 2002 by Bijay Kumar Agarwal (~23-24 years active as of 2026), operating primarily in Ranchi and Gurgaon, citing 24+ projects and 5M+ sq ft delivered -- self-reported by the group, not independently audited. Pranami Neev Realty Limited itself (the RERA-registered promoter entity for this specific project) was only incorporated 2022-06-27, consistent with being a newer Mumbai-market SPV of the longer-established group rather than the group's founding entity -- confirmed via an August 2023 announcement of a Rs. 225 crore Integrow investment partnership covering Pranami Group's newly-acquired Andheri/Ghatkopar developments, of which this project is one.":
        "Per Pranami Group's own website (pranamigroup.com/about), the group was founded in 2002 by Bijay Kumar Agarwal (roughly 23-24 years active as of 2026), operating primarily in Ranchi and Gurgaon, citing 24+ projects and 5M+ sq ft delivered; this is self-reported by the group, not independently audited. Pranami Neev Realty Limited itself (the RERA-registered promoter entity for this specific project) was only incorporated 2022-06-27, consistent with being a newer Mumbai-market SPV of the longer-established group rather than the group's founding entity. This is confirmed via an August 2023 announcement of a Rs. 225 crore Integrow investment partnership covering Pranami Group's newly-acquired Andheri/Ghatkopar developments, of which this project is one.",
    "Per-building sold/unsold carpet-area figures (not unit counts) are disclosed in the promoter's own quarterly Sold/Unsold Inventory filings -- see RERA Compliance's construction_progress note for the Jun-2025 to Dec-2025 trend.":
        "Per-building sold/unsold carpet-area figures (not unit counts) are disclosed in the promoter's own quarterly Sold/Unsold Inventory filings; see RERA Compliance for the Jun-2025 to Dec-2025 trend.",
    "Per-building/wing floor counts and unit-type breakdowns were not confirmed from the documents reviewed this pass -- see Gaps.":
        "Per-building/wing floor counts and unit-type breakdowns were not confirmed from the documents reviewed this review (see Gaps).",
    "No documents from the MahaRERA-listed library are absent -- all 60 listed documents downloaded successfully this pass.":
        "No documents from the MahaRERA-listed library are absent: all 60 listed documents downloaded successfully this review.",
    "No independent third-party market-research report (Knight Frank/ANAROCK/CBRE-style) specific to the Veera Desai Road/Azad Nagar micro-pocket was found -- the cited pricing/trend figures are real-estate-aggregator averages, not an audited report.":
        "No independent third-party market-research report (Knight Frank/ANAROCK/CBRE-style) specific to the Veera Desai Road/Azad Nagar micro-pocket was found. The cited pricing and trend figures are real-estate-aggregator averages, not an audited report.",
    "The exact list of Pranami Group's other projects and their current completion status could not be independently verified -- the group's own project-list webpage marks all entries \"archived\" with no visible per-project completion dates.":
        "The exact list of Pranami Group's other projects and their current completion status could not be independently verified. The group's own project-list webpage marks all entries \"archived\" with no visible per-project completion dates.",
    "Director Sundeep Poddar could not be conclusively matched to an independent director-profile record beyond the MCA-mirror data already cited -- a similarly-named \"Sandeep Kumar Poddar\" (different DIN) linked to unrelated companies was found and explicitly NOT assumed to be the same person.":
        "Director Sundeep Poddar could not be conclusively matched to an independent director-profile record beyond the MCA-mirror data already cited. A similarly named \"Sandeep Kumar Poddar\" (different DIN) linked to unrelated companies was found and explicitly not assumed to be the same person.",
    "No public credit rating was found for this exact legal entity from any agency checked (ICRA, Infomerics) -- not itself a red flag; most MahaRERA promoters are too small or private to be rated.":
        "No public credit rating was found for this exact legal entity from any agency checked (ICRA, Infomerics). This is not itself a red flag; most MahaRERA promoters are too small or private to be rated.",
    "promoter_portfolio's cross-project totals (area developed, area within 5km) are not available -- MahaRERA's own Promoters-tab name search matched only this one project under this exact promoter name, so portfolio-based Developer Score criteria 3/4 are unscored gaps, not zero.":
        "This promoter's cross-project totals (area developed, area within 5km) are not available. MahaRERA's own Promoters-tab name search matched only this one project under this exact promoter name, so the portfolio-based Developer Score criteria are unscored gaps, not zero.",
    "0 defaults -- IBBI shows no insolvency process against this CIN, and nothing in the credit-rating check (if one ran) states otherwise.":
        "0 defaults. IBBI shows no insolvency process against this CIN, and nothing in the credit-rating check states otherwise.",
    "1 of 4 core figures (FSI, land/built-up area, unit counts, pricing) have no unresolved gap against them -- flagged as an open gap: fsi, land_built_up_area, pricing":
        "1 of 4 core figures (FSI, land/built-up area, unit counts, pricing) have no unresolved gap against them. Flagged as open gaps: FSI, land/built-up area, pricing.",
    "Per-shareholder/promoter shareholding percentages are gated behind ZaubaCorp's paid report and were not available on this pass; only aggregate authorised/paid-up capital above is free. This is not itself a red flag for a private company -- detailed cap tables are rarely public for unlisted entities.":
        "Per-shareholder shareholding percentages are not publicly available for this private company (normal for unlisted entities); aggregate authorised/paid-up capital is shown above.",

    # --- IRA Insignia ---
    "None disclosed as a current acquisition/reservation in the Title Report itself -- a separate ~65 sq.ft strip was purchased from Chandresh Vastu CHS Ltd in 2019 for road access, already folded into the land assembly.":
        "None disclosed as a current acquisition or reservation in the Title Report itself. A separate approximately 65 sq. ft. strip was purchased from Chandresh Vastu CHS Ltd in 2019 for road access, already folded into the land assembly.",
    "Partnership Firm -- confirmed via the 4th character of its PAN (AAHFI1448M; 'F' denotes Firm, not 'C' for Company or 'P' for Individual/proprietorship), and independently corroborated by the Title Report's and Commencement Certificate's own repeated references to 'M/s. Ira Homes Partnership Firm'. No CIN or LLPIN exists for this entity -- partnership firms are not required to register with the MCA.":
        "Partnership Firm, confirmed via the 4th character of its PAN (AAHFI1448M; 'F' denotes Firm, not 'C' for Company or 'P' for Individual/proprietorship), and independently corroborated by the Title Report's and Commencement Certificate's own repeated references to 'M/s. Ira Homes Partnership Firm'. No CIN or LLPIN exists for this entity: partnership firms are not required to register with the MCA.",
    "Not applicable -- IRA Homes is an unregistered (non-LLP) partnership firm; no CIN/LLPIN exists to look up.":
        "Not applicable. IRA Homes is an unregistered (non-LLP) partnership firm; no CIN/LLPIN exists to look up.",
    "Not applicable -- a partnership firm has no board resolution; Form B is instead executed by partner Hiren Kantilal Vador under the firm's own authority.":
        "Not applicable. A partnership firm has no board resolution; Form B is instead executed by partner Hiren Kantilal Vador under the firm's own authority.",
    "Hiren Kantilal Vador, Partner -- signs Form B (notarized 2021-09-28). Ketan Kantilal Vador, Partner -- named on the Title Report, NA Tax Conversion letter, and Commencement Certificate as the firm's representative.":
        "Hiren Kantilal Vador, Partner, signs Form B (notarized 2021-09-28). Ketan Kantilal Vador, Partner, is named on the Title Report, NA Tax Conversion letter, and Commencement Certificate as the firm's representative.",
    "Hiren Kantilal Vador and Ketan Kantilal Vador, Partners of M/s. Ira Homes Partnership Firm -- named across the Title Report, PAN card, Form B, NA Tax Conversion letter, and Commencement Certificate consistently. No MCA/ROC filing exists to independently cross-check this roster (partnership, not a company/LLP), so this is based entirely on the firm's own project documents, not a third-party registry.":
        "Hiren Kantilal Vador and Ketan Kantilal Vador, Partners of M/s. Ira Homes Partnership Firm, are named across the Title Report, PAN card, Form B, NA Tax Conversion letter, and Commencement Certificate consistently. No MCA/ROC filing exists to independently cross-check this roster (partnership, not a company or LLP), so this is based entirely on the firm's own project documents, not a third-party registry.",
    "then reassigned in 2019 -- with Shah Builders/Kanti Ratanshi Shah as a CONFIRMING PARTY -- to M/s. Ira Homes Partnership Firm.":
        "then reassigned in 2019, with Shah Builders/Kanti Ratanshi Shah as a confirming party, to M/s. Ira Homes Partnership Firm.",
    "The project site pincode (421204) and the promoter's registered-office pincode (421201) differ -- both are Dombivli East localities and this is common for a project-site-vs-registered-office split, not necessarily an inconsistency, but it was not independently reconciled this pass.":
        "The project site pincode (421204) and the promoter's registered-office pincode (421201) differ. Both are Dombivli East localities, and this is common for a project-site-versus-registered-office split, not necessarily an inconsistency, but it was not independently reconciled this review.",
    "No MCA/ROC cross-check is possible -- IRA Homes is a partnership firm (PAN 4th character 'F'), not a company or LLP, and carries no CIN/LLPIN.":
        "No MCA/ROC cross-check is possible. IRA Homes is a partnership firm (PAN 4th character 'F'), not a company or LLP, and carries no CIN/LLPIN.",
    "Two distinct matters. (1) LAND TITLE: the Title Report's advocate opines title is clear and marketable, but the underlying 30-year Search Report itself shows real gaps -- records were unavailable/not indexed for 1995-2002 and 2017-2020, so the title's own supporting search could not verify those years, not a clean confirmed record. No specific litigation against the land itself is stated in the Title Report. (2) MAHARERA COMPLAINTS -- 3 on record, a materially different picture from a clean project: (a) Kanti Shah vs. Ira Homes (CC12400706) -- Shah, a predecessor developer under the 2007 Development Agreement who surrendered his rights in the 2019 settlement, alleged Ira Homes built ~9 floors beyond its sanctioned Commencement Certificate (constructing up to ~33 floors when only up to the 5th floor -- G+5 -- was sanctioned at the time), sold 80+ flats without permission, and misrepresented the project's sanctioned floor count on MahaRERA's own portal. MahaRERA DISMISSED this complaint ON JURISDICTIONAL GROUNDS ONLY -- Shah was held not to be an \"allottee\" under RERA (his flats were compensation under a private settlement, a civil matter outside MahaRERA's remit) -- NOT on the merits of whether the construction itself was authorized. MahaRERA separately noted the promoter obtained a revised building permission up to the 21st floor on 2025-02-17, stating the CC-violation issue \"does not survive as on date\" -- but this does not affirmatively establish the earlier construction (before that revision) WAS authorized; it was never adjudicated on the merits. (b) Amit Salaskar vs. Ira Homes (CC12500570) -- alleged construction \"stalled since 2021\" with no progress since booking, sought a full refund; MahaRERA rejected the delay/refund claim as premature (declared completion date is 2026-12-31, no formal allotment letter existed to trigger Section 18(1)) but ordered a refund of the amount paid, minus a standard 2% cancellation deduction. (c) Mohan Rajkumar Kumhar vs. Ira Homes (CC12602499, filed 2026-04-29) -- no order yet on record as of this pass; substance not yet known.":
        "Two distinct matters. (1) Land title: the Title Report's advocate opines title is clear and marketable, but the underlying 30-year Search Report itself shows real gaps. Records were unavailable or not indexed for 1995-2002 and 2017-2020, so the title's own supporting search could not verify those years; this is not a clean, confirmed record. No specific litigation against the land itself is stated in the Title Report. (2) MahaRERA complaints: 3 on record, a materially different picture from a clean project. (a) Kanti Shah vs. Ira Homes (CC12400706): Shah, a predecessor developer under the 2007 Development Agreement who surrendered his rights in the 2019 settlement, alleged Ira Homes built roughly 9 floors beyond its sanctioned Commencement Certificate (constructing up to roughly 33 floors when only up to the 5th floor, G+5, was sanctioned at the time), sold 80+ flats without permission, and misrepresented the project's sanctioned floor count on MahaRERA's own portal. MahaRERA dismissed this complaint on jurisdictional grounds only: Shah was held not to be an \"allottee\" under RERA (his flats were compensation under a private settlement, a civil matter outside MahaRERA's remit), not on the merits of whether the construction itself was authorized. MahaRERA separately noted the promoter obtained a revised building permission up to the 21st floor on 2025-02-17, stating the CC-violation issue \"does not survive as on date,\" but this does not affirmatively establish that the earlier construction, before that revision, was authorized; it was never adjudicated on the merits. (b) Amit Salaskar vs. Ira Homes (CC12500570): alleged construction \"stalled since 2021\" with no progress since booking, and sought a full refund. MahaRERA rejected the delay/refund claim as premature (declared completion date is 2026-12-31, and no formal allotment letter existed to trigger Section 18(1)), but ordered a refund of the amount paid, minus a standard 2% cancellation deduction. (c) Mohan Rajkumar Kumhar vs. Ira Homes (CC12602499, filed 2026-04-29): no order yet on record as of this review; substance not yet known.",
    "(boundary marker, per MahaRERA's own record -- no named landmark disclosed)":
        "(boundary marker per MahaRERA's own record; no named landmark disclosed)",
    "Governed by the Kalyan-Dombivli Municipal Corporation (KDMC) under the Maharashtra Regional and Town Planning Act, 1966 (Section 45) read with the Maharashtra Municipal Corporations Act, 1949 (Section 253) -- confirmed via the Commencement Certificate (Office No. KDMC/TPD/BP/27Village/2021-22/14, sanctioned 2021-07-30) and the accompanying sanctioned Layout/Building Plans of the same date.":
        "Governed by the Kalyan-Dombivli Municipal Corporation (KDMC) under the Maharashtra Regional and Town Planning Act, 1966 (Section 45), read with the Maharashtra Municipal Corporations Act, 1949 (Section 253). Confirmed via the Commencement Certificate (Office No. KDMC/TPD/BP/27Village/2021-22/14, sanctioned 2021-07-30) and the accompanying sanctioned Layout/Building Plans of the same date.",
    "across Buildings 1, 2, and 3 -- i.e. a low-rise sanction at that stage,":
        "across Buildings 1, 2, and 3, a low-rise sanction at that stage,",
    "that revised sanction itself was not among the documents reviewed this pass, so the current, post-revision FSI/floor count could not be independently confirmed -- only the original 2021 sanction was directly read.":
        "that revised sanction itself was not among the documents reviewed this review, so the current, post-revision FSI and floor count could not be independently confirmed; only the original 2021 sanction was directly read.",
    "Approx. 8,824.99 sq.m proposed built-up area per the 2021 Layout Plan's own area statement -- figures across the scanned plan set are fragmented and not fully reconcilable against a separate ~2,982-2,992 sq.m \"FSI combined\" figure also referenced in the same documents; treated as an approximate, not a precisely confirmed figure.":
        "Approx. 8,824.99 sq.m proposed built-up area per the 2021 Layout Plan's own area statement. Figures across the scanned plan set are fragmented and not fully reconcilable against a separate approximately 2,982-2,992 sq.m \"FSI combined\" figure also referenced in the same documents; treated as an approximate, not a precisely confirmed figure.",
    "Not precisely reconcilable from the documents reviewed -- see approved_bua's note on the two differing figures found in the same plan set.":
        "Not precisely reconcilable from the documents reviewed; see the note above on the two differing figures found in the same plan set.",
    "Not computed -- the underlying BUA figures could not be reconciled to a single confirmed number (see approved_bua).":
        "Not computed. The underlying BUA figures could not be reconciled to a single confirmed number.",
    "3 complaints on record (Kanti Shah -- dismissed on jurisdictional grounds; Amit Salaskar -- partly allowed, refund ordered minus 2% cancellation; Mohan Rajkumar Kumhar -- filed 2026-04-29, no order yet). 0 appeals.":
        "3 complaints on record (Kanti Shah: dismissed on jurisdictional grounds; Amit Salaskar: partly allowed, refund ordered minus 2% cancellation; Mohan Rajkumar Kumhar: filed 2026-04-29, no order yet). 0 appeals.",
    "Form B declares the standard title/encumbrance/RERA-compliance undertakings and commits to a 2026-12-31 completion date -- signed by Hiren Kantilal Vador, Partner, notarized 2021-09-28.":
        "Form B declares the standard title, encumbrance, and RERA-compliance undertakings and commits to a 2026-12-31 completion date, signed by Hiren Kantilal Vador, Partner, notarized 2021-09-28.",
    "Only 1 of 338 total units sold as of this pass, despite registration since 2021-10-23 (over 4 years) -- consistent with the Salaskar complaint's own allegation of construction \"stalled since 2021.\" The Kanti Shah order references a revised building permission obtained up to the 21st floor as of 2025-02-17, indicating construction has progressed further than the original 2021 low-rise sanction at some point, but current physical progress was not independently verified this pass beyond what these two complaint orders describe.":
        "Only 1 of 338 total units sold as of this review, despite registration since 2021-10-23 (over 4 years), consistent with the Salaskar complaint's own allegation of construction \"stalled since 2021.\" The Kanti Shah order references a revised building permission obtained up to the 21st floor as of 2025-02-17, indicating construction has progressed further than the original 2021 low-rise sanction at some point, but current physical progress was not independently verified this review beyond what these two complaint orders describe.",
    "MahaRERA (project registration); Kalyan-Dombivli Municipal Corporation -- KDMC (building/layout plan sanction, under DCPR); Tahsildar, Kalyan (NA land-use conversion).":
        "MahaRERA (project registration); Kalyan-Dombivli Municipal Corporation, KDMC (building/layout plan sanction, under DCPR); Tahsildar, Kalyan (NA land-use conversion).",
    "2026-12-31 -- unchanged since original registration, no extension on record":
        "2026-12-31, unchanged since original registration, no extension on record",
    "Estimated via web search, not a live driving-route lookup -- exact distance from Shankeshwar Nagar/Bhopar specifically not independently confirmed.":
        "Estimated via web search, not a live driving-route lookup; exact distance from Shankeshwar Nagar/Bhopar specifically was not independently confirmed.",
    "Located off Bhopar Road, Shankeshwar Nagar Phase-3, Dombivli East -- within the Kalyan-Dombivli Municipal Corporation limits.":
        "Located off Bhopar Road, Shankeshwar Nagar Phase-3, Dombivli East, within the Kalyan-Dombivli Municipal Corporation limits.",
    "No operational metro line serves Dombivli East directly as of this pass -- planned/upcoming extensions were referenced in general market research but not independently confirmed for this specific pocket.":
        "No operational metro line serves Dombivli East directly as of this review. Planned or upcoming extensions were referenced in general market research but not independently confirmed for this specific pocket.",
    "Dombivli East residential belt within Kalyan-Dombivli Municipal Corporation limits -- an affordable-to-mid-tier suburb":
        "Dombivli East residential belt within Kalyan-Dombivli Municipal Corporation limits, an affordable-to-mid-tier suburb",
    "plus planned metro extensions -- no source quantified":
        "plus planned metro extensions. No source quantified",
    "1BHK from INR 30L, 2BHK from INR 45L (possession projected Dec 2029 -- a notably long horizon)":
        "1BHK from INR 30L, 2BHK from INR 45L (possession projected Dec 2029, a notably long horizon)",
    "neither a conspicuous premium nor discount -- though this was not independently verified against a specific unit size and configuration this pass.":
        "neither a conspicuous premium nor discount, though this was not independently verified against a specific unit size and configuration this review.",
    "IRA Homes (the partnership firm itself, PAN AAHFI1448M) was formed 2018-06-18 per its own PAN card -- 8 years old as of this pass. Multiple listing sites (Dwello, ira-insignia.in) describe \"IRA\" as a brand of a broader \"Vador Group\" with a longer history, and a company called \"Vador Properties Private Limited\" (CIN U70102MH2010PTC210775, incorporated 2010) was found on ZaubaCorp -- but no source confirms this CIN is the same corporate lineage as the Ira Insignia promoter, so that longer history is NOT credited here; only the firm's own confirmed formation date is used, deliberately conservative rather than assuming an unconfirmed parent-group link.":
        "IRA Homes (the partnership firm itself, PAN AAHFI1448M) was formed 2018-06-18 per its own PAN card, 8 years old as of this review. Multiple listing sites (Dwello, ira-insignia.in) describe \"IRA\" as a brand of a broader \"Vador Group\" with a longer history, and a company called \"Vador Properties Private Limited\" (CIN U70102MH2010PTC210775, incorporated 2010) was found on ZaubaCorp. However, no source confirms this CIN is the same corporate lineage as the Ira Insignia promoter, so that longer history is not credited here; only the firm's own confirmed formation date is used, deliberately conservative rather than assuming an unconfirmed parent-group link.",
    "One third-party listing (99acres) instead shows 385 units across 3 towers/28 floors for this same project -- a discrepancy against MahaRERA's own 338-unit figure that could not be reconciled this pass; MahaRERA's own record is treated as authoritative here, but the discrepancy itself is flagged as a gap.":
        "One third-party listing (99acres) instead shows 385 units across 3 towers/28 floors for this same project, a discrepancy against MahaRERA's own 338-unit figure that could not be reconciled this review. MahaRERA's own record is treated as authoritative here, but the discrepancy itself is flagged as a gap.",
    "338 total (RERA record) -- see unit_summary_note on a conflicting 385-unit third-party figure":
        "338 total (RERA record); see the note above on a conflicting 385-unit third-party figure",
    "The Title Report's own 30-year search has real gaps -- records were unavailable/not indexed for 1995-2002 and 2017-2020 -- so the advocate's \"clear and marketable\" opinion rests on an incomplete underlying search for those years.":
        "The Title Report's own 30-year search has real gaps: records were unavailable or not indexed for 1995-2002 and 2017-2020, so the advocate's \"clear and marketable\" opinion rests on an incomplete underlying search for those years.",
    "The Kanti Shah complaint's core allegation (that construction proceeded materially beyond the sanctioned Commencement Certificate before a 2025-02-17 revision) was never adjudicated on the merits -- MahaRERA dismissed the complaint on jurisdictional grounds (Shah wasn't an allottee) rather than ruling on whether the construction itself was authorized at the time it occurred.":
        "The Kanti Shah complaint's core allegation, that construction proceeded materially beyond the sanctioned Commencement Certificate before a 2025-02-17 revision, was never adjudicated on the merits. MahaRERA dismissed the complaint on jurisdictional grounds (Shah wasn't an allottee) rather than ruling on whether the construction itself was authorized at the time it occurred.",
    "The revised building permission -- reportedly approving construction up to the 21st floor, obtained on 2025-02-17 per MahaRERA's own order -- was not available in the documents reviewed, so the current, post-revision sanctioned floor count and FSI could not be independently confirmed.":
        "The revised building permission (reportedly approving construction up to the 21st floor, obtained on 2025-02-17 per MahaRERA's own order) was not available in the documents reviewed, so the current, post-revision sanctioned floor count and FSI could not be independently confirmed.",
    "A third MahaRERA complaint (Mohan Rajkumar Kumhar, CC12602499, filed 2026-04-29) has no order on record yet -- its substance is not yet known.":
        "A third MahaRERA complaint (Mohan Rajkumar Kumhar, CC12602499, filed 2026-04-29) has no order on record yet; its substance is not yet known.",
    "No CIN/LLPIN exists for IRA Homes (an unregistered partnership firm) -- the MCA-mirror company-profile chain, IBBI insolvency check, and ZaubaCorp group-companies crosswalk this pipeline otherwise runs for corporate promoters could not run at all; there is no company/LLP registry record to check.":
        "No CIN or LLPIN exists for IRA Homes (an unregistered partnership firm), so the MCA-mirror company-profile chain, IBBI insolvency check, and ZaubaCorp group-companies crosswalk otherwise run for corporate promoters could not run at all; there is no company or LLP registry record to check.",
    "Whether \"Ira Homes\" is the same corporate lineage as the longer-established \"Vador Group\"/\"Vador Properties Private Limited\" (CIN U70102MH2010PTC210775, found via ZaubaCorp) could not be confirmed -- treated as an unconfirmed lead, not credited toward this promoter's track record.":
        "Whether \"Ira Homes\" is the same corporate lineage as the longer-established \"Vador Group\"/\"Vador Properties Private Limited\" (CIN U70102MH2010PTC210775, found via ZaubaCorp) could not be confirmed. It is treated as an unconfirmed lead, not credited toward this promoter's track record.",
    "A third-party listing (99acres) shows 385 units/3 towers/28 floors for this project, conflicting with MahaRERA's own live record of 338 total units -- not reconciled this pass.":
        "A third-party listing (99acres) shows 385 units across 3 towers/28 floors for this project, conflicting with MahaRERA's own live record of 338 total units; not reconciled this review.",
    "No independent news coverage, forum discussion, or review-site commentary corroborating (or contradicting) the 3 known MahaRERA complaints was found -- an explicit search gap (MahaRERA's own complaint records are not generally indexed by search engines), not evidence the underlying issues are resolved or absent.":
        "No independent news coverage, forum discussion, or review-site commentary corroborating or contradicting the 3 known MahaRERA complaints was found. This is an explicit search gap (MahaRERA's own complaint records are not generally indexed by search engines), not evidence the underlying issues are resolved or absent.",
    "Only 1 of 338 units sold as of this pass, over 4 years after registration -- consistent with the Salaskar complaint's allegation of stalled construction, but current physical construction progress was not independently site-verified this pass beyond what the two complaint orders describe.":
        "Only 1 of 338 units sold as of this review, over 4 years after registration, consistent with the Salaskar complaint's allegation of stalled construction. Current physical construction progress was not independently site-verified this review beyond what the two complaint orders describe.",
    "IRA Insignia (MahaRERA P51700031409) is a 338-unit residential project in Shankeshwar Nagar, Dombivli East, by IRA Homes -- an unregistered partnership firm (no CIN/LLPIN) formed 2018, apparently operating under a broader \"Vador Group\" brand whose exact corporate relationship to this promoter could not be confirmed.":
        "IRA Insignia (MahaRERA P51700031409) is a 338-unit residential project in Shankeshwar Nagar, Dombivli East, by IRA Homes, an unregistered partnership firm (no CIN/LLPIN) formed 2018, apparently operating under a broader \"Vador Group\" brand whose exact corporate relationship to this promoter could not be confirmed.",
    "Most notably, a predecessor developer's complaint (Kanti Shah) alleged the promoter constructed roughly 9 floors beyond its original sanctioned Commencement Certificate and misrepresented the project's sanctioned floor count on MahaRERA's own portal -- MahaRERA dismissed this complaint on jurisdictional grounds only (the complainant wasn't an allottee), never ruling on whether the construction itself was authorized at the time.":
        "Most notably, a predecessor developer's complaint (Kanti Shah) alleged the promoter constructed roughly 9 floors beyond its original sanctioned Commencement Certificate and misrepresented the project's sanctioned floor count on MahaRERA's own portal. MahaRERA dismissed this complaint on jurisdictional grounds only (the complainant wasn't an allottee), never ruling on whether the construction itself was authorized at the time.",
    "No CIN exists for the MCA-mirror/IBBI checks this pipeline otherwise runs to apply to.":
        "No CIN exists for the MCA-mirror or IBBI checks otherwise run for corporate promoters to apply to.",
    "All 70 listed documents downloaded successfully this pass -- none absent. Two files listed under MahaRERA's \"Others\" category (\"DOC1.pdf\", \"File1.pdf\") were opened and found to be effectively blank placeholders (a single page reading only \"Ira homes\", no letterhead/signature/substantive content) -- confirmed empty, not a download failure.":
        "All 70 listed documents downloaded successfully this review; none absent. Two files listed under MahaRERA's \"Others\" category (\"DOC1.pdf\", \"File1.pdf\") were opened and found to be effectively blank placeholders (a single page reading only \"Ira homes,\" no letterhead, signature, or substantive content), confirmed empty rather than a download failure.",
    "IBBI insolvency check did not run this pass -- nothing to count past defaults from.":
        "IBBI insolvency check did not run this review; there is nothing to count past defaults from.",
    "LLP with no independently confirmed intent to convert to Pvt Ltd -- conservatively scored at the \"not willing to convert\" band rather than assuming unconfirmed intent.":
        "LLP with no independently confirmed intent to convert to Pvt Ltd. Conservatively scored at the \"not willing to convert\" band rather than assuming unconfirmed intent.",
    "2 of 4 core figures (FSI, land/built-up area, unit counts, pricing) have no unresolved gap against them -- flagged as an open gap: fsi, land_built_up_area":
        "2 of 4 core figures (FSI, land/built-up area, unit counts, pricing) have no unresolved gap against them. Flagged as open gaps: FSI, land/built-up area.",
}

# Ordered (most-specific-first) regex substitutions for _externalize_prose
# -- targets the specific internal-process phrasing actually found in this
# Charter's own generated prose (mentions of this session's own tooling/
# bug-fix history, and "this pass"/"this session" as a research-cycle
# marker). Best-effort pattern matching over KNOWN phrasing, not a
# guaranteed-exhaustive rewrite -- new prose using different wording could
# still slip through un-cleaned.
_EXTERNAL_PROSE_SUBSTITUTIONS = (
    (r"\s*Specific schools/hospitals within a fixed radius were not independently verified (this pass|this review)\s*(\(see Gaps\))?\.?", ""),
    (r"\s*Not every document was individually read(\s*\([^)]*\))?;\s*the quarterly certificates/sale deeds not opened are lower-diligence-priority than the title/Form-B/encumbrance/IOA documents that were\.?", ""),
    (r"Allotment mechanics: Standard MahaRERA-governed sale under the Act -- no allotment mechanism beyond the statutory framework was separately disclosed in the documents reviewed\.?", ""),
    (r"confirmed on re-extraction after fixing this session's OCR pipeline", "independently re-verified"),
    (r"re-extraction after fixing this session's OCR pipeline", "independent re-verification"),
    (r"this session's OCR extraction pipeline \(Tesseract\)", "the OCR process used for this review"),
    (r"this session's OCR (extraction )?pipeline", "the OCR process used for this review"),
    (r"documents opened (in )?this pass", "documents reviewed"),
    (r"opened this pass", "reviewed"),
    (r"this research session", "this review"),
    (r"this session", "this review"),
    (r"this pass", "this review"),
    # "(per live RERA record)" etc. is an internal sourcing annotation --
    # useful context for an Internal reviewer (this figure was pulled fresh
    # from the API, not stale), but reads as an odd internal-process note
    # to an External/client reader, so it's dropped there rather than kept.
    (r"\s*\(per live RERA record\)", ""),
)

# Matches "(see gaps[0])", "(see fsi_interpretation)", "(see local_planning
# note on ...)" -- a raw facts.json field path (snake_case and/or dotted/
# bracketed) used as an internal QA cross-reference. Requires the token
# right after "see " to contain an underscore, dot, or bracket so genuine
# in-document references like "(see Sources)" or "(see the tier table
# above...)" are left untouched -- those start uppercase or are plain
# English words with no internal-identifier shape.
_INTERNAL_FIELD_ANNOTATION_RE = re.compile(
    r"\s*\(see (?=[a-z][a-z0-9_.\[\]]*[_.\[])[a-z][a-z0-9_]*(?:\.[a-z0-9_]+|\[\d+\])*[^)]*\)"
)


def _externalize_prose(facts: dict, text: str) -> str:
    """Best-effort cleanup of known internal-process phrasing in `text`
    for the External variant only -- Internal returns `text` verbatim.
    Preserves the original match's capitalization (so a sentence-initial
    "This pass..." becomes "This review...", not a lowercase "this
    review..." mid-sentence-looking fragment)."""
    if not text or facts.get("_doc_variant") != "external":
        return text

    def _replace_preserving_case(pattern: str, replacement: str, s: str) -> str:
        def _repl(m: re.Match) -> str:
            matched = m.group(0)
            if not replacement or not matched[:1].isupper():
                return replacement
            return replacement[0].upper() + replacement[1:]
        return re.sub(pattern, _repl, s, flags=re.IGNORECASE)

    for original, rewritten in _EXTERNAL_DASH_REWRITES.items():
        text = text.replace(original, rewritten)
    for pattern, replacement in _EXTERNAL_PROSE_SUBSTITUTIONS:
        text = _replace_preserving_case(pattern, replacement, text)
    text = _INTERNAL_FIELD_ANNOTATION_RE.sub("", text)
    # Last-resort fallback, after every curated rewrite above: those are
    # keyed to specific, hand-checked WHOLE sentences, so any text that
    # doesn't match one exactly -- new facts.json prose never seen before,
    # or an old registered sentence now split into a smaller fragment by
    # _set_paragraph_as_bullets -- would otherwise carry its " -- "/" — "
    # straight through and fail the External quality gate. A semicolon is
    # grammatical wherever a dash joins two related clauses or introduces
    # an explanation, and unlike a comma it can't create a comma splice --
    # the same reasoning the shared prose builder uses for this, just never
    # added here until it was actually confirmed to matter live.
    text = text.replace(" -- ", "; ").replace(" — ", "; ").replace("—", "; ")
    return text


def _externalized_facts_copy(facts: dict) -> dict:
    """Returns a deep copy of `facts` with every string value passed
    through _externalize_prose -- used once, at the top of _fill_template,
    only for doc_variant="external". A deep copy (not an in-place walk)
    is deliberate: _fill_template's caller (run_company_charter) persists
    its own `facts` dict to .facts.json after rendering, and that
    persisted copy must stay the real, un-rewritten internal content --
    the lossy External phrasing must never leak into it, however many
    times or in whatever order the two variants get rendered."""
    copied = copy.deepcopy(facts)

    def _walk(obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                if isinstance(v, str):
                    obj[k] = _externalize_prose(copied, v)
                else:
                    _walk(v)
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                if isinstance(v, str):
                    obj[i] = _externalize_prose(copied, v)
                else:
                    _walk(v)

    _walk(copied)
    copied["gaps"] = _external_gaps(copied.get("gaps", []) or [])
    return copied


# A gap that records a TOOLING failure rather than a fact about the project:
# a re-verification that could not run, an authentication error, a pending
# manual step, or a raw internal path/key. CLAUDE.md Section B: these "stay in
# full in the Internal document and are cut from the External one", and
# separately, no file path, JSON key or raw exception string may appear in
# either document -- least of all External.
_PROCESS_FAILURE_GAP_RE = re.compile(
    r"could not resolve authentication|verification could not run|"
    r"^\s*[a-z_]+(?:\.[a-z_]+)*\s*:|"          # a bare facts.json key used as a label
    r"\boutput/[\w.-]+/|\.json\b",
    re.IGNORECASE,
)
# The exception: a cross-corroboration gap wraps a REAL finding (this topic
# rests on a single source) around the same error text. Section B says to
# "preserve the finding in one consolidated line and drop the error text".
_SINGLE_SOURCE_GAP_RE = re.compile(r"^Cross-corroboration: the '([\w_]+)' topic is backed by only one source", re.IGNORECASE)


def _external_gaps(gaps: list) -> list:
    """Filters facts["gaps"] down to what belongs in the External document.

    Drops pure process failures outright. Consolidates every single-source
    cross-corroboration gap into ONE line naming the affected topics, which
    keeps the finding a reader can act on ("these topics rest on a single
    source") while dropping the tooling error each one currently carries."""
    kept, single_source_topics = [], []
    for gap in gaps:
        m = _SINGLE_SOURCE_GAP_RE.match(gap.strip())
        if m:
            single_source_topics.append(m.group(1).replace("_", " "))
            continue
        if _PROCESS_FAILURE_GAP_RE.search(gap):
            continue
        kept.append(gap)

    if single_source_topics:
        topics = ", ".join(single_source_topics[:-1]) + " and " + single_source_topics[-1] \
            if len(single_source_topics) > 1 else single_source_topics[0]
        kept.append(
            f"The following topics rest on a single source each, with no independent second source "
            f"confirming them: {topics}. Treat these with more caution than multi-sourced findings."
        )
    return kept


def _external_heading(facts: dict, heading_text: str) -> str:
    """Drops the phrase "Code-Computed" from a section heading for the
    External variant -- Internal keeps it unchanged. Removes it together
    with an adjacent comma so "(Code-Computed)" disappears entirely
    (parens included), "(ZaubaCorp, Code-Computed)" becomes
    "(ZaubaCorp)", and "(Code-Computed Director & Address Crosswalk)"
    becomes "(Director & Address Crosswalk)" -- the rest of a
    parenthetical (an org name, a description of what the section is)
    is real content, not internal-process language, and stays."""
    if facts.get("_doc_variant") != "external":
        return heading_text
    cleaned = re.sub(r",?\s*Code-Computed\s*,?", "", heading_text)
    cleaned = re.sub(r"\(\s*\)", "", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip()


def _topic_citation(facts: dict, topic: str) -> str | None:
    """Looks up facts['sources'] for the first entry tagged with `topic`
    and returns its bare citation label, e.g. "99acres, accessed
    2026-07-17" (no surrounding punctuation -- see _citation_text for how
    this gets wrapped for the doc's variant). Returns None if no source
    carries that topic -- never invents one just to fill a paragraph."""
    for s in facts.get("sources", []) or []:
        if s.get("topic") == topic:
            label = s.get("label") or s.get("ref") or "source"
            accessed = s.get("accessed_date")
            if accessed and accessed != "unknown":
                return f"{label}, accessed {accessed}"
            return label
    return None


def _cite_marker(*topics: str, facts: dict) -> str | None:
    """Same topic lookup as _cite, but returns just the bare marker
    string (e.g. "[1]", or "[1][2]" if the matching source's own label
    packs several semicolon-joined sources -- see _citation_text) instead
    of concatenating it onto a piece of text. Lets a caller attach the
    SAME marker to every bullet clause _split_into_bullet_clauses
    produces from one field, rather than the old bug: gluing it once onto
    the whole multi-sentence blob BEFORE splitting, so only the last
    resulting bullet ever carried a citation and every earlier one carried
    none (see _set_paragraph_as_bullets's `citation` param, which this
    feeds)."""
    for topic in topics:
        citation = _citation_text(facts, _topic_citation(facts, topic))
        if citation:
            return citation
    return None


def _cite(text: str, *topics: str, facts: dict) -> str:
    """Appends the first matching topic citation (tried in the given
    order) to `text`, or returns `text` unchanged if none of the topics
    has a matching source -- never invents a citation just to fill a
    paragraph. Only safe for a single, un-split paragraph; a field that
    will be rendered via _set_paragraph_as_bullets must use _cite_marker
    plus that function's `citation` param instead, so the marker reaches
    every bullet clause, not just the tail of the pre-split text."""
    if not text:
        return text
    marker = _cite_marker(*topics, facts=facts)
    return f"{text} {marker}" if marker else text


def _compute_authenticity_summary(facts: dict) -> dict:
    tier_counts = {}
    for src in facts.get("sources", []):
        tier = _classify_source_tier(src.get("ref", ""))
        tier_counts[tier] = tier_counts.get(tier, 0) + 1

    total_sources = sum(tier_counts.values())
    primary_tiers = (
        "Primary regulatory record (MahaRERA/MCA, or a document opened from it)",
        "Credit rating agency (CRISIL/ICRA/CARE/India Ratings)",
        "Government legal/insolvency record (IBBI, NCLT, NCDRC, MahaREAT judgments)",
        "Live Google Maps verification",
    )
    primary_count = sum(c for t, c in tier_counts.items() if t in primary_tiers)

    return {
        "tier_counts": tier_counts,
        "total_sources": total_sources,
        "primary_count": primary_count,
        "total_gaps": len(facts.get("gaps", [])),
    }


# ---------------------------------------------------------------------------
# Cross-corroboration enforcement -- turns the Cross-Corroboration score
# criterion below (which only ever *measures* what fraction of topics have
# 2+ sources) into something that actually *acts* on a single-sourced topic,
# instead of letting it sit silently inside a percentage. For every topic
# backed by exactly one source, this makes ONE bounded attempt (never
# retried further, same policy as deep_research._resolve_gaps) to find a
# genuinely independent second source via an agentic web-search pass. A
# topic that gets a real second source is upgraded in place; a topic that
# doesn't gets an explicit, named gap -- so a reader always knows exactly
# which specific findings in this Charter rest on only one source, rather
# than having to infer it from the Cross-Corroboration percentage alone.
# ---------------------------------------------------------------------------

_SECOND_SOURCE_SYSTEM_PROMPT = """You are trying to find a SECOND, INDEPENDENT source that \
corroborates a claim already sourced from one place. Use web_search to look for a different \
publisher/site that confirms the same claim -- not a mirror, syndication, or re-post of the same \
original source, and not the same claim phrased differently by the same underlying publisher.

Your FINAL reply must be ONLY a single raw JSON object -- no prose, no markdown fences -- matching \
exactly this shape: {"found": true, "label": "publisher name", "ref": "one-line description -- url", \
"published_date": "YYYY-MM-DD or 'unknown'", "accessed_date": "YYYY-MM-DD"} if a genuinely \
independent second source was found, or {"found": false, "reason": "one sentence"} if not."""


def _find_single_source_topics(facts: dict) -> dict:
    """Groups facts['sources'] by topic and returns {topic: source} for
    every topic currently backed by exactly one source."""
    by_topic = {}
    for src in facts.get("sources", []):
        topic = src.get("topic")
        if not topic:
            continue
        by_topic.setdefault(topic, []).append(src)
    return {topic: srcs[0] for topic, srcs in by_topic.items() if len(srcs) == 1}


def _attempt_second_source(topic: str, existing_source: dict) -> dict:
    """One bounded attempt to find a second, independent source for a claim
    currently backed by only one. Returns {"found": True, "source": {...}}
    or {"found": False, "reason": ...} -- a missing ANTHROPIC_API_KEY or any
    other failure to even run the attempt surfaces honestly as a reason
    here rather than crashing the caller, same policy as
    deep_research._verify_claim elsewhere in this codebase."""
    prompt = (
        f"Existing source: {existing_source.get('label', '')} -- {existing_source.get('ref', '')}\n"
        f"Topic: {topic}\n"
        f"Find one genuinely independent second source that corroborates this."
    )
    try:
        result = deep_research._run_agentic_pass(prompt, _SECOND_SOURCE_SYSTEM_PROMPT, label="second_source_verify")
    except Exception as e:
        return {"found": False, "reason": f"verification could not run: {e}"}
    if not isinstance(result, dict) or "found" not in result:
        return {"found": False, "reason": "second-source search returned an unrecognized response"}
    if result.get("found"):
        return {
            "found": True,
            "source": {
                "label": result.get("label", ""),
                "ref": result.get("ref", ""),
                "topic": topic,
                "published_date": result.get("published_date") or "unknown",
                "accessed_date": result.get("accessed_date") or datetime.now().strftime("%Y-%m-%d"),
            },
        }
    return {"found": False, "reason": result.get("reason", "no independent second source found")}


def verify_cross_corroboration(facts: dict) -> dict:
    """The code-enforced version of 'verify every material fact from 2+
    independent angles': finds every topic currently backed by exactly one
    source, attempts to upgrade each to two sources, and explicitly names
    the ones that can't be -- rather than leaving single-sourced topics
    invisible inside the Cross-Corroboration score's percentage. Call this
    after all other sources have been assembled (credit rating, IBBI,
    judgments, Maps, etc.), so the single-source topics found here are the
    real final set, not a mid-assembly snapshot."""
    single_source_topics = _find_single_source_topics(facts)
    gaps = list(facts.get("gaps", []))
    attempted = 0
    upgraded = 0
    for topic, existing_source in single_source_topics.items():
        attempted += 1
        attempt = _attempt_second_source(topic, existing_source)
        if attempt["found"]:
            facts.setdefault("sources", []).append(attempt["source"])
            upgraded += 1
        else:
            gaps.append(
                f"Cross-corroboration: the '{topic}' topic is backed by only one source "
                f"({existing_source.get('label', 'unknown')}); one independent-second-source "
                f"retry attempt did not find a genuinely separate corroborating source "
                f"({attempt['reason']}). Treat this specific finding with more caution than "
                f"multi-sourced ones in this Charter."
            )
    facts["gaps"] = gaps
    facts.setdefault("_verification_stats", {})["cross_corroboration"] = {
        "single_source_topics_found": len(single_source_topics),
        "attempted": attempted,
        "upgraded_to_two_sources": upgraded,
    }
    return facts


# ---------------------------------------------------------------------------
# Composite Documentation Confidence Score -- a score of how well-sourced and
# verified THIS DOCUMENT's own claims are (source quality, completeness,
# cross-corroboration, recency, re-check success), NOT a rating of the
# underlying project's quality, safety, or investment merit. A project with a
# genuinely bad track record whose problems are thoroughly documented can
# score HIGH here; a genuinely fine project with little public paper trail to
# verify against can score LOW. Six criteria, each computed directly from
# data already in `facts`, none of them a model self-assessment. CRISIL's own
# real-estate methodology (researched, not guessed -- see project notes) does
# NOT publish a numeric weighting formula for any of its three products
# (project grading, developer grading, or credit rating); it relies on a
# committee's qualitative judgment over a factor tree. So there is no
# published formula to replicate here -- the weights below are this project's
# own calibration, informed by CRISIL's STRUCTURE (a small number of named
# top-level buckets rather than one flat checklist; funding/verification
# treated as their own distinct dimension; recency of evidence mattering)
# but not a copy of anything CRISIL discloses. Say so plainly wherever this
# score is presented, so nobody mistakes "CRISIL-informed" for "CRISIL's
# formula".
# ---------------------------------------------------------------------------

_TIER_WEIGHTS = {
    "Primary regulatory record (MahaRERA/MCA, or a document opened from it)": 100,
    "Credit rating agency (CRISIL/ICRA/CARE/India Ratings)": 95,
    "Government legal/insolvency record (IBBI, NCLT, NCDRC, MahaREAT judgments)": 90,
    "Live Google Maps verification": 90,
    "Corporate-registry mirror": 75,
    "Homebuyer-advocacy / watchdog body": 60,
    "Real-estate aggregator / listing site": 50,
    "News / press coverage": 40,
    "Crowd-sourced / user-editable reference (Wikipedia, etc.)": 30,
    "Developer/promoter's own website (self-published)": 25,
    "Social media / user-generated content": 20,
    "Other / unclassified source": 30,
}

_DOC_CONFIDENCE_WEIGHTS = {
    "source_tier_quality": 0.20,
    "primary_tier_density": 0.15,
    "completeness_rate": 0.05,  # was 0.15 -- 0.10 moved to financial_figures_confirmed (Step 5)
    "cross_corroboration": 0.15,
    "recency_legal": 0.05,
    "recency_other": 0.05,
    "verification_rate": 0.25,
    "financial_figures_confirmed": 0.10,  # was 0.00/TBD -- implemented in Step 5
}

_RECENCY_WINDOW_MONTHS = 18  # this project's own calibration, not a disclosed CRISIL threshold
_RECENCY_WINDOW_MONTHS_LEGAL = 3  # legal/litigation sources are held to a much tighter freshness bar

# The four core investment-case figures financial_figures_confirmed checks
# for, and the gap-text markers used to detect each is still an open,
# unconfirmed gap rather than something read from a primary document.
_FINANCIAL_FIGURE_MARKERS = {
    "fsi": (r"\bfsi\b",),
    "land_built_up_area": (r"built-up area", r"built up area", r"\bbua\b"),
    "unit_counts": (r"unit breakdown", r"unit count", r"unit mix", r"number of units"),
    "pricing": (r"\bpricing\b", r"per-sq\.?ft", r"per sq\.?ft", r"\bprice\b"),
}

# ---------------------------------------------------------------------------
# Groundwork constants for the upcoming Developer Score / risk-flagging
# feature -- every number that currently lives only in the design memo, moved
# here so later steps import from one place instead of re-hardcoding it.
# Establishing the constants only: nothing below is wired into scoring or
# flagging logic yet (see _compute_documentation_confidence_score, which
# still computes a single "recency" criterion and has no
# "financial_figures_confirmed" criterion -- that wiring, and the
# corresponding _DOC_CONFIDENCE_WEIGHTS key split above, are a later step).
# ---------------------------------------------------------------------------

_FLAG_THRESHOLDS = {
    "complaint_monitor": 15, "complaint_imminent": 40,
    "appeal_monitor": 5, "appeal_imminent": 15,
    "credit_rating_min_units": 500,  # below this, "no rating found" isn't flagged
    # GST filing-pattern breakpoints -- shared between _score_gst_compliance
    # and _classify_flags (same reasoning as the complaint/appeal pair
    # above: a project already flagged for a bad GST pattern can't silently
    # score AAA, and vice versa).
    "gst_late_pct_monitor": 15, "gst_late_pct_imminent": 40,
    "gst_recent_delays_monitor": 1, "gst_recent_delays_imminent": 3,
}
# The Developer Score's 3-bucket / 9-sub-metric AAA-D framework (see
# _DEVELOPER_SCORE_STRUCTURE below): each sub-metric is scored independently
# against its own AAA-D band, converted to an even 0-100 tier-equivalent,
# then combined by its FIXED structural weight -- unlike every other score
# in this Charter, missing data is never redistributed/renormalized here
# (see _compute_developer_score's own note).
_DEVELOPER_SCORE_TIER_SCORES = {"AAA": 100.0, "AA": 83.3, "A": 66.7, "B": 50.0, "C": 33.3, "D": 16.7}
# Thresholds sit at the midpoint between adjacent tier-equivalents, used to
# map a combined composite (which can land between two tiers once several
# criteria are averaged) back onto a single overall letter grade.
_DEVELOPER_SCORE_TIER_THRESHOLDS = (("AAA", 91.65), ("AA", 75.0), ("A", 58.35), ("B", 41.65), ("C", 25.0))  # D below C
_DATA_AUTHENTICITY_BANDS = {"High": 70, "Moderate": 45}  # Limited below Moderate


def _band_label(score: float) -> str:
    if score >= _DATA_AUTHENTICITY_BANDS["High"]:
        return "High"
    if score >= _DATA_AUTHENTICITY_BANDS["Moderate"]:
        return "Moderate"
    return "Limited"


# Shared red/amber/green convention so a grade, a band, and a flag all read
# the same way wherever they appear -- fills for KPI cards, text colors for
# flag lines and headline score runs.
_FILL_RED = "F8CBAD"
_FILL_AMBER = "FFE699"
_FILL_GREEN = "C6E0B4"
_FILL_NEUTRAL = "D9E2F3"
_TEXT_RED = "C00000"
_TEXT_AMBER = "BF8F00"
_TEXT_GREEN = "375623"


def _grade_fill(grade: str) -> str:
    return {
        "AAA": _FILL_GREEN, "AA": _FILL_GREEN,
        "A": _FILL_AMBER, "B": _FILL_AMBER,
        "C": _FILL_RED, "D": _FILL_RED,
    }.get(grade, _FILL_NEUTRAL)


def _band_fill(band: str) -> str:
    return {"High": _FILL_GREEN, "Moderate": _FILL_AMBER, "Limited": _FILL_RED}.get(band, _FILL_NEUTRAL)


def _color_run(run, hex_color: str) -> None:
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor.from_string(hex_color)


def _months_since(date_str: str) -> float | None:
    """Parses a YYYY-MM-DD string and returns months elapsed since then, or
    None for anything unparseable -- including the literal "unknown" the
    schema explicitly allows, which must never be silently coerced into a
    fake date."""
    if not date_str or date_str.strip().lower() == "unknown":
        return None
    try:
        parsed = datetime.strptime(date_str.strip()[:10], "%Y-%m-%d")
    except ValueError:
        return None
    return (datetime.now() - parsed).days / 30.44


def _compute_documentation_confidence_score(facts: dict, authenticity_summary: dict) -> dict:
    """Returns {overall, band, criteria: {name: {score, weight, note}}} -- a
    score of how well-sourced and verified THIS DOCUMENT's claims are, not a
    rating of the underlying project (see the module note above). Every
    criterion is either a plain ratio over data already in `facts`, or
    explicitly marked not-applicable and excluded (with the remaining
    weights renormalized to still sum to 100%) rather than silently scored
    as 0 -- a project with, say, no distinct source topics yet shouldn't be
    penalized on Cross-Corroboration as if it had failed that check."""
    sources = facts.get("sources", [])
    total_sources = authenticity_summary["total_sources"]
    total_gaps = authenticity_summary["total_gaps"]
    criteria = {}

    # 1. Primary-Tier Density
    if total_sources:
        criteria["primary_tier_density"] = {
            "score": 100 * authenticity_summary["primary_count"] / total_sources,
            "note": f"{authenticity_summary['primary_count']} of {total_sources} sources are top-tier (primary regulatory, credit rating, government legal record, or live Maps verification).",
        }

    # 2. Weighted Source-Tier Quality
    if total_sources:
        weighted_sum = sum(
            _TIER_WEIGHTS.get(tier, _TIER_WEIGHTS["Other / unclassified source"]) * count
            for tier, count in authenticity_summary["tier_counts"].items()
        )
        criteria["source_tier_quality"] = {
            "score": weighted_sum / total_sources,
            "note": "Weighted average trust level across all cited source tiers (see the tier table above for this document's own calibration of each tier's weight).",
        }

    # 3. Completeness Rate
    if total_sources + total_gaps:
        criteria["completeness_rate"] = {
            "score": 100 * total_sources / (total_sources + total_gaps),
            "note": f"{total_sources} confirmed source(s) versus {total_gaps} item(s) left as an explicit, unresolved gap.",
        }

    # 4. Independent Verification Rate -- combines the URL-based re-check
    # (_verify_material_claims) and the document-grounding re-check
    # (_check_document_grounding); N/A if neither ever ran (nothing sourced
    # from a URL or a checkable local document this pass). Zero attempts is
    # NOT just renormalized away silently, unlike every other N/A criterion
    # here -- it's a distinct, worse signal (nothing in this Charter was
    # independently re-checked at all) that gets its own override below,
    # after `overall`/`band` are otherwise computed.
    v = facts.get("_verification_stats", {})
    url_stats = v.get("url", {"attempted": 0, "confirmed": 0})
    doc_stats = v.get("document", {"attempted": 0, "confirmed": 0})
    total_attempted = url_stats["attempted"] + doc_stats["attempted"]
    if total_attempted:
        total_confirmed = url_stats["confirmed"] + doc_stats["confirmed"]
        criteria["verification_rate"] = {
            "score": 100 * total_confirmed / total_attempted,
            "note": f"{total_confirmed} of {total_attempted} independently-checkable claims were confirmed on re-check (web-sourced claims re-verified via search; local-document claims cross-checked against the actual extracted text).",
        }

    # 5. Cross-Corroboration -- % of distinct topics backed by >=2 sources.
    topics = [s.get("topic") for s in sources if s.get("topic")]
    if topics:
        topic_counts = {}
        for t in topics:
            topic_counts[t] = topic_counts.get(t, 0) + 1
        multi_backed = sum(1 for c in topic_counts.values() if c >= 2)
        criteria["cross_corroboration"] = {
            "score": 100 * multi_backed / len(topic_counts),
            "note": f"{multi_backed} of {len(topic_counts)} distinct topics cited are backed by 2 or more independent sources, not just one.",
        }

    # 6. Recency -- % of dated sources (published_date, falling back to
    # accessed_date for a live-record source with no publish date of its
    # own) accessed/published within _RECENCY_WINDOW_MONTHS. Sources with
    # no parseable date at all (both fields "unknown" or missing) are
    # excluded from the denominator, not counted as stale -- there's
    # nothing to assess. Split legal vs. other (not one blended "recency"
    # criterion) because a stale legal/insolvency source is a materially
    # different risk than a stale pricing/press source -- a legal record
    # going 18 months without a re-check is a real gap; a market-pricing
    # figure that old is a much smaller one. "Legal" is judged by the
    # source's own topic (litigation/insolvency/appeal-adjudication ones)
    # or, failing that, its classified tier being the government
    # legal/insolvency record tier -- everything else is "other". Each
    # half is independently omittable (renormalized away) if that half
    # simply has no dated sources this pass, same as every other criterion
    # here.
    _LEGAL_TOPICS = {"legal_documents", "litigation", "insolvency_status", "appeal_judgments"}
    _LEGAL_TIER = "Government legal/insolvency record (IBBI, NCLT, NCDRC, MahaREAT judgments)"
    legal_ages, other_ages = [], []
    for s in sources:
        age = _months_since(s.get("published_date", ""))
        if age is None:
            age = _months_since(s.get("accessed_date", ""))
        if age is None:
            continue
        is_legal = s.get("topic") in _LEGAL_TOPICS or _classify_source_tier(s.get("ref", "")) == _LEGAL_TIER
        (legal_ages if is_legal else other_ages).append(age)

    if legal_ages:
        fresh = sum(1 for a in legal_ages if a <= _RECENCY_WINDOW_MONTHS_LEGAL)
        criteria["recency_legal"] = {
            "score": 100 * fresh / len(legal_ages),
            "note": f"{fresh} of {len(legal_ages)} dated legal/litigation/insolvency sources are within this document's tighter {_RECENCY_WINDOW_MONTHS_LEGAL}-month freshness window for legal sources (this project's own calibration, not a disclosed external standard).",
        }
    if other_ages:
        fresh = sum(1 for a in other_ages if a <= _RECENCY_WINDOW_MONTHS)
        criteria["recency_other"] = {
            "score": 100 * fresh / len(other_ages),
            "note": f"{fresh} of {len(other_ages)} other dated sources (pricing, corporate identity, distances, etc.) are within this document's {_RECENCY_WINDOW_MONTHS}-month freshness window.",
        }

    # 7. Financial Figures Confirmed -- % of the four core investment-case
    # figures (FSI, land/built-up area, unit counts, pricing) that are NOT
    # named in facts["gaps"] -- i.e. confirmed from a primary document
    # rather than left an open gap. Unlike every other criterion above,
    # this one always has something to say (facts["gaps"] is always at
    # least an empty list), so it's never skipped/renormalized away.
    gaps_blob = " ".join(facts.get("gaps", [])).lower()
    unconfirmed_figures = [name for name, patterns in _FINANCIAL_FIGURE_MARKERS.items() if any(re.search(p, gaps_blob) for p in patterns)]
    confirmed_count = len(_FINANCIAL_FIGURE_MARKERS) - len(unconfirmed_figures)
    note = f"{confirmed_count} of {len(_FINANCIAL_FIGURE_MARKERS)} core figures (FSI, land/built-up area, unit counts, pricing) have no unresolved gap against them"
    if unconfirmed_figures:
        note += f" -- flagged as an open gap: {', '.join(unconfirmed_figures)}"
    criteria["financial_figures_confirmed"] = {
        "score": 100 * confirmed_count / len(_FINANCIAL_FIGURE_MARKERS),
        "note": note,
    }

    applicable_weight = sum(_DOC_CONFIDENCE_WEIGHTS[k] for k in criteria)
    if applicable_weight == 0:
        overall = 0.0
    else:
        overall = sum(criteria[k]["score"] * _DOC_CONFIDENCE_WEIGHTS[k] for k in criteria) / applicable_weight
        for k in criteria:
            criteria[k]["weight"] = round(100 * _DOC_CONFIDENCE_WEIGHTS[k] / applicable_weight, 1)

    band = _band_label(overall)
    result = {
        "overall": round(overall, 1),
        "band": band,
        "criteria": criteria,
        "skipped_criteria": sorted(set(_DOC_CONFIDENCE_WEIGHTS) - set(criteria)),
    }

    # Zero-verification safeguard: 0 independently-checkable claims is a
    # materially worse signal than "N/A, renormalize away" -- nothing in
    # this Charter was ever re-checked against an independent source, so
    # the band is capped at "Moderate" regardless of what every other
    # criterion computes (restrains a High down to Moderate; a Limited
    # stays Limited -- same "never worsen, only restrain" cap pattern as
    # the Developer Score's imminent-flag override).
    if total_attempted == 0:
        result["verification_warning"] = "0 claims were independently re-checked this pass."
        if band == "High":
            result["band"] = "Moderate"

    return result


# ---------------------------------------------------------------------------
# Flag classification -- reads the already-assembled facts dict (nothing new
# fetched here) and sorts every material risk signal into exactly one of
# three urgency tiers: imminent (act before proceeding), structural (a
# standing characteristic, raise directly with the developer), or monitor
# (re-check on a future pass). Ported from this project's design memo and
# validated against Company_Charter_GodrejParkGreens_P52100019639.facts.json,
# whose known-correct split (4 imminent / 5 structural / 5 monitor) is
# asserted in test_classify_flags.py.
#
# Two rules below go beyond what the memo states outright, both found
# necessary to reproduce that validated split and flagged here rather than
# silently invented:
#   1. "Near-sellout despite an extended completion date -> imminent" --
#      the memo's rule list has no explicit slot for this, but the reference
#      Charter's 4th Imminent flag is exactly this signal (a project that is
#      ~99% sold against a completion date already pushed back twice is a
#      real, high-urgency exposure for almost the entire buyer base). Gated
#      on _NEAR_SELLOUT_PCT below, a local threshold (not in _FLAG_THRESHOLDS,
#      since the memo's dict has no key for it).
#   2. Two specific gap texts (the one about complaint SUBSTANCE being
#      unconfirmable, and the one about no independent market comparable
#      being confirmable) are promoted to structural rather than falling to
#      the "any remaining gap -> monitor" default. Both describe a STANDING
#      limitation of the data source/market (not closable by a future
#      documents pass), versus every other unmatched gap here, which is
#      closable by more work. That distinction is judgment, not a clean
#      keyword rule -- _STRUCTURAL_GAP_PHRASES below is deliberately narrow
#      (calibrated to the one validated example, not claimed to generalize)
#      rather than pretending a broader heuristic was validated.
# ---------------------------------------------------------------------------

_NEAR_SELLOUT_PCT = 90  # this project's own calibration, not from the memo's threshold dict

_KNOWN_LISTED_GROUP_MARKERS = (
    "godrej", "dlf", "oberoi", "prestige", "sobha", "brigade", "lodha", "macrotech",
    "puravankara", "sunteck", "kolte", "mahindra lifespace", "shapoorji", "tata housing",
)

_FSI_AREA_GAP_MARKERS = ("fsi", "bua", "built-up area", "built up area")

# Deliberately narrow, matched to this one validated Charter's own wording --
# see the module note above.
_STRUCTURAL_GAP_PHRASES = ("complaint substance", "genuinely independent comparable")

_ESCROW_GAP_MARKERS = ("escrow", "collection account")


def _first_int(text: str, pattern: str) -> int | None:
    match = re.search(pattern, text or "", re.IGNORECASE)
    if not match:
        return None
    try:
        return int(match.group(1).replace(",", ""))
    except ValueError:
        return None


def _parse_complaint_appeal_counts(facts: dict) -> tuple[int | None, int | None]:
    """Returns (complaint_count, appeal_count). Reads rera_core_fields'
    structured total_complaints_count/total_appeals_count companions first
    -- exact by construction, no regex involved -- and only falls back to
    parsing litigations_per_record's prose for older facts.json files that
    predate those fields. The prose-parsing fallback tries the specific
    phrasing first (e.g. "67 total complaint-category filings ..., 18
    appeals on record ..."), then a plainer "N complaints"/"N appeals"
    match (e.g. "0 complaints, 0 appeals on MahaRERA's own record ...") --
    a genuinely clean 0/0 record must parse just as confidently as a high
    one, or it silently drops out of the Developer Score's legal_compliance
    pillar. Either return value is None only if nothing above could
    confidently resolve it -- never a guessed number."""
    core = facts.get("rera_core_fields", {}) or {}
    complaint_count = core.get("total_complaints_count")
    appeal_count = core.get("total_appeals_count")
    if isinstance(complaint_count, int) and isinstance(appeal_count, int):
        return complaint_count, appeal_count

    text = core.get("litigations_per_record", "") or ""
    if not isinstance(complaint_count, int):
        complaint_count = _first_int(text, r"([\d,]+)\s+total complaint")
        if complaint_count is None:
            complaint_count = _first_int(text, r"([\d,]+)\s+complaints?\b")
    if not isinstance(appeal_count, int):
        appeal_count = _first_int(text, r"([\d,]+)\s+appeals?\s+on record")
        if appeal_count is None:
            appeal_count = _first_int(text, r"([\d,]+)\s+appeals?\b")
    return complaint_count, appeal_count


def _parse_total_units_and_sold(facts: dict) -> tuple[int | None, int | None]:
    """Returns (total_units, sold_units). Reads rera_core_fields'
    structured units_total/units_sold companions first, falling back to
    regexing total_building_units's prose only for older facts.json files
    that predate those fields. Either is None if not confidently
    resolvable -- never a guessed number."""
    core = facts.get("rera_core_fields", {}) or {}
    total = core.get("units_total")
    sold = core.get("units_sold")
    if isinstance(total, int) and isinstance(sold, int):
        return total, sold

    text = core.get("total_building_units", "") or ""
    if not isinstance(total, int):
        total = _first_int(text, r"([\d,]+)\s+total units")
    if not isinstance(sold, int):
        sold = _first_int(text, r"([\d,]+)\s+sold")
    return total, sold


def _matches_known_listed_group(promoter_name: str) -> bool:
    name = (promoter_name or "").lower()
    return any(marker in name for marker in _KNOWN_LISTED_GROUP_MARKERS)


def _credit_rating_applies(facts: dict, total_units: int | None) -> bool:
    """A missing credit rating is only worth flagging for a project large
    enough, or a promoter well-known enough, that a rating would plausibly
    exist -- small LLPs are not expected to have one (see
    lookup_credit_rating's own module note)."""
    if total_units is not None and total_units > _FLAG_THRESHOLDS["credit_rating_min_units"]:
        return True
    promoter_name = ((facts.get("corporate_identity", {}) or {}).get("promoter_name") or {}).get("value", "")
    return _matches_known_listed_group(promoter_name)


def _classify_ibbi_hit(note_text: str) -> tuple[str, str]:
    """Reads an IBBI status_text/note for whether a positive hit is against
    the exact promoter entity and whether it's open or closed/settled.
    Returns (severity, reason) where severity is "imminent", "structural",
    or "unclear" -- this text is freeform, so anything not confidently
    readable is treated as "unclear" (caller flags it structural and says
    so explicitly) rather than guessed at."""
    text = (note_text or "").lower()
    is_affiliated_not_itself = "affiliated" in text or "not " in text and " itself" in text
    is_closed = any(term in text for term in ("closure", "settlement", "settled", "12a", "withdrawn"))
    if is_closed or is_affiliated_not_itself:
        return "structural", "the IBBI record reads as either a closed/settled matter or tied to an affiliated (not the exact promoter) entity"
    if "open" in text or "ongoing" in text or "pending" in text:
        return "imminent", "the IBBI record reads as an open matter against the exact promoter entity"
    return "unclear", "whether this IBBI hit is open/closed or against the exact promoter entity could not be confidently determined from the freeform status text"


def _classify_flags(facts: dict) -> dict:
    """Sorts every material risk signal already present in `facts` into
    {"imminent": [...], "structural": [...], "monitor": [...]} -- each item
    is {"text": str, "field": str}. Never fetches anything new; purely a
    read-and-classify pass over data this Charter already assembled."""
    imminent, structural, monitor = [], [], []

    core = facts.get("rera_core_fields", {}) or {}
    corp = facts.get("corporate_identity", {}) or {}
    land = facts.get("land_identification", {}) or {}

    # 1. RERA registration number
    reg_no = (core.get("registration_number") or "").strip()
    if not reg_no:
        imminent.append({
            "text": "No MahaRERA registration number could be resolved for this project.",
            "field": "rera_core_fields.registration_number",
        })

    # 2. CTS/plot number
    cts_value = (land.get("survey_cts_plot_numbers") or {}).get("value", "") or ""
    if not cts_value.strip() or cts_value.strip().lower().startswith(("not found", "not disclosed", "not confirmed", "not available")):
        net_area = (land.get("net_area") or {}).get("value", "") or ""
        blocks_entirely = not net_area.strip() or net_area.strip().lower().startswith(("not found", "not disclosed", "not confirmed", "not available"))
        target = imminent if blocks_entirely else structural
        target.append({
            "text": "No CTS/plot number could be confirmed for this project's land" + (
                " -- and with no area figure available either, land verification is blocked entirely." if blocks_entirely
                else ", though area figures are otherwise available."
            ),
            "field": "land_identification.survey_cts_plot_numbers",
        })

    # 2b. CTS mismatch between a pre-RERA carryover intake and this
    # Charter's own RERA-sourced record -- see run_cts_land_lookup's
    # carryover path. A real discrepancy here, not an absence, so it's
    # always imminent regardless of whether check #2 above also fired.
    if facts.get("cts_mismatch_note"):
        imminent.append({"text": facts["cts_mismatch_note"], "field": "cts_mismatch_note"})

    # 3. CIN/LLPIN
    cin_llpin_text = (corp.get("cin_llpin") or {}).get("value", "") or ""
    if not extract_cin(cin_llpin_text) and not extract_llpin(cin_llpin_text):
        structural.append({
            "text": "No CIN or LLPIN could be confirmed for the promoter.",
            "field": "corporate_identity.cin_llpin",
        })

    # 4 & 8. IBBI insolvency + credit rating -- handled together since,
    # when NEITHER ran this pass at all, the memo's reference Charter
    # reports them as a single combined "not checked" structural note
    # rather than two near-duplicate items.
    total_units, sold_units = _parse_total_units_and_sold(facts)
    promoter_name = (corp.get("promoter_name") or {}).get("value", "")
    ibbi_check = facts.get("ibbi_insolvency_check")
    credit_check = facts.get("credit_rating_check")

    if ibbi_check is None and credit_check is None:
        if _credit_rating_applies(facts, total_units):
            structural.append({
                "text": "Credit rating and IBBI insolvency status were not checked in this Charter pass -- a process gap, not a finding.",
                "field": "ibbi_insolvency_check / credit_rating_check",
            })
    else:
        if ibbi_check is not None and ibbi_check.get("found_process") is True:
            note_text = ibbi_check.get("status_text") or ibbi_check.get("note") or ""
            severity, reason = _classify_ibbi_hit(note_text)
            if severity == "imminent":
                imminent.append({"text": f"IBBI insolvency record found against the exact promoter entity, appearing open: {reason}.", "field": "ibbi_insolvency_check.status_text"})
            else:
                prefix = "IBBI insolvency record found" if severity == "structural" else "IBBI insolvency record found, but"
                structural.append({"text": f"{prefix} -- {reason}; shown here rather than suppressed.", "field": "ibbi_insolvency_check.status_text"})
        elif ibbi_check is not None and ibbi_check.get("found_process") is None:
            monitor.append({"text": f"IBBI insolvency check could not run this pass: {ibbi_check.get('note', 'reason not recorded')}.", "field": "ibbi_insolvency_check.note"})

        if credit_check is not None:
            promoter_result = (credit_check.get("promoter") or {})
            if promoter_result.get("found"):
                ratings_text = " ".join(
                    item.get("rating", "") for agency in promoter_result.get("ratings", []) for item in agency.get("instruments", [])
                ).lower()
                sub_investment_grade = bool(re.search(r"\b(bb|b|c|d)\b", ratings_text)) and "bbb" not in ratings_text
                if "downgraded" in ratings_text or sub_investment_grade:
                    imminent.append({
                        "text": "Credit rating found is below investment grade or was downgraded versus a prior run.",
                        "field": "credit_rating_check.promoter.ratings",
                    })
            elif _credit_rating_applies(facts, total_units):
                structural.append({
                    "text": "No public credit rating was found for the promoter, despite this project's scale/promoter profile making one plausible.",
                    "field": "credit_rating_check.promoter.note",
                })

    # 5. Complaint count vs thresholds
    complaint_count, appeal_count = _parse_complaint_appeal_counts(facts)
    if complaint_count is not None:
        if complaint_count > _FLAG_THRESHOLDS["complaint_imminent"]:
            imminent.append({"text": f"Complaint volume: {complaint_count} total filings against this project -- above the imminent threshold ({_FLAG_THRESHOLDS['complaint_imminent']}).", "field": "rera_core_fields.litigations_per_record"})
        elif complaint_count > _FLAG_THRESHOLDS["complaint_monitor"]:
            monitor.append({"text": f"Complaint volume: {complaint_count} total filings against this project -- above the monitor threshold ({_FLAG_THRESHOLDS['complaint_monitor']}) but not yet at the imminent level.", "field": "rera_core_fields.litigations_per_record"})
        elif complaint_count >= 1:
            structural.append({"text": f"Complaint volume: {complaint_count} total filing(s) on record against this project -- a visible floor worth noting, not blended into zero.", "field": "rera_core_fields.litigations_per_record"})

    # 6. Appeal count vs thresholds -- same pattern
    if appeal_count is not None:
        if appeal_count > _FLAG_THRESHOLDS["appeal_imminent"]:
            imminent.append({"text": f"Appeal volume: {appeal_count} appeals on record -- above the imminent threshold ({_FLAG_THRESHOLDS['appeal_imminent']}).", "field": "rera_core_fields.litigations_per_record"})
        elif appeal_count > _FLAG_THRESHOLDS["appeal_monitor"]:
            monitor.append({"text": f"Appeal volume: {appeal_count} appeals on record -- above the monitor threshold ({_FLAG_THRESHOLDS['appeal_monitor']}) but not yet at the imminent level.", "field": "rera_core_fields.litigations_per_record"})
        elif appeal_count >= 1:
            structural.append({"text": f"Appeal volume: {appeal_count} appeal(s) on record -- a visible floor worth noting, not blended into zero.", "field": "rera_core_fields.litigations_per_record"})

    # 6b. GST filing pattern -- only present when a human has supplied
    # gst_filing_input.json (see run_gst_compliance_check); reuses the same
    # thresholds _score_gst_compliance scores against, so a project flagged
    # here can't silently score AAA there. A bad RECENT pattern is exactly
    # what should prompt raising this directly with the developer, which is
    # why delays_last_12_months (not lifetime late_pct alone) drives the
    # imminent/structural split.
    gst_check = facts.get("gst_compliance_check")
    if gst_check and gst_check.get("found"):
        gst_summary = gst_check["summary"]
        gst_late_pct = gst_summary.get("late_pct", 0)
        gst_delays_recent = gst_summary.get("delays_last_12_months", 0)
        gst_text = (
            f"GST filing pattern for GSTIN {gst_check.get('gstin', '')}: {gst_late_pct}% of rated periods filed "
            f"late, {gst_delays_recent} delayed/unfiled period(s) due in the trailing 12 months."
        )
        if gst_delays_recent > _FLAG_THRESHOLDS["gst_recent_delays_imminent"] or gst_late_pct > _FLAG_THRESHOLDS["gst_late_pct_imminent"]:
            imminent.append({"text": gst_text + " Above the threshold for raising directly with the developer.", "field": "gst_compliance_check.summary"})
        elif gst_delays_recent > _FLAG_THRESHOLDS["gst_recent_delays_monitor"] or gst_late_pct > _FLAG_THRESHOLDS["gst_late_pct_monitor"]:
            structural.append({"text": gst_text, "field": "gst_compliance_check.summary"})
        elif gst_delays_recent > 0 or gst_late_pct > 0:
            monitor.append({"text": gst_text, "field": "gst_compliance_check.summary"})

    # 7. Promoter portfolio total -- always structural, never colour-scored
    portfolio = facts.get("promoter_portfolio")
    if portfolio:
        portfolio_total_complaints = (portfolio.get("totals", {}) or {}).get("total_complaints")
        text = f"Promoter portfolio: {portfolio.get('projects_analyzed', 0)} MahaRERA-registered project(s) tied to this promoter."
        if portfolio_total_complaints and complaint_count is not None and portfolio_total_complaints > 0:
            concentration_pct = round(100 * complaint_count / portfolio_total_complaints, 1)
            text += f" This project alone accounts for {complaint_count} of {portfolio_total_complaints} total complaints across that portfolio ({concentration_pct}%)."
        structural.append({"text": text, "field": "promoter_portfolio.totals"})

    # Near-sellout despite an extended completion date -- see module note
    # above (a memo-gap-filling addition, not stated in the original rules).
    proposed_completion_text = (core.get("proposed_completion_date") or "").lower()
    if total_units and sold_units and total_units > 0 and "extend" in proposed_completion_text:
        sold_pct = 100 * sold_units / total_units
        if sold_pct >= _NEAR_SELLOUT_PCT:
            imminent.append({
                "text": f"Completion timeline has been extended, against a project that is {round(sold_pct, 1)}% sold -- nearly the entire buyer base is exposed to this delay.",
                "field": "rera_core_fields.proposed_completion_date",
            })

    # 10, 11, 12, 13. Mortgage lender change, escrow gap, FSI/area figures,
    # and the remaining gaps -- monitor by default.
    if facts.get("mortgage_lender_history_note"):
        monitor.append({"text": facts["mortgage_lender_history_note"], "field": "mortgage_lender_history_note"})

    for i, gap_text in enumerate(facts.get("gaps", [])):
        lowered = gap_text.lower()
        # gap_number is this gap's stable 1-based identity, assigned here so a
        # flag and its Gaps & Sources entry can be cross-referenced without
        # either renderer having to re-derive it (CLAUDE.md Section B, "Flags
        # summarise, Gaps explain"). It indexes facts["gaps"] and does not
        # renumber per document variant -- see _rendered_gap_numbers.
        entry = {"text": gap_text, "field": f"gaps[{i}]", "gap_number": i + 1}
        if any(marker in lowered for marker in _ESCROW_GAP_MARKERS):
            structural.append(entry)
        elif any(marker in lowered for marker in _FSI_AREA_GAP_MARKERS):
            imminent.append(entry)
        elif any(phrase in lowered for phrase in _STRUCTURAL_GAP_PHRASES):
            structural.append(entry)
        else:
            monitor.append(entry)

    return {"imminent": imminent, "structural": structural, "monitor": monitor}


def _rendered_gap_numbers(facts: dict) -> set:
    """The stable gap numbers (see _classify_flags) that THIS document variant
    actually prints under Gaps & Sources.

    Internal prints every gap. External prints only those material enough to
    have also earned an Imminent or Structural flag -- see _fill_template's own
    note where that list is built. _append_flag_list consults this before
    emitting a "(Gap N)" pointer, so External can never point at a number it
    does not print. Numbers stay stable either way: External's list reads
    "Gap 3.", "Gap 7." rather than renumbering to 1, 2, because the number is
    the gap's identity, not its position in one variant's list."""
    gaps = facts.get("gaps", []) or []
    if facts.get("_doc_variant") != "external":
        return set(range(1, len(gaps) + 1))

    flags = _classify_flags(facts)
    return {
        item["gap_number"]
        for severity in ("imminent", "structural")
        for item in flags.get(severity, [])
        if item.get("gap_number")
    }


def _flag_headline(facts: dict, item: dict) -> tuple:
    """Returns (text_to_render, points_at_a_gap) for one Overview & Flags item.

    CLAUDE.md Section B, "Flags summarise, Gaps explain": a flag whose full
    explanation already appears under Gaps & Sources is reduced to a
    one-sentence headline plus a "(Gap N)" pointer, instead of restating the
    whole thing. In the Pranami Bliss run that duplication ran to ~1.5 pages.

    A flag with no gap entry to point at -- a CTS mismatch, a missing
    registration number -- keeps its full text, because nothing else in the
    document states it. Same when this variant does not print that gap at all
    (External drops monitor-only gaps): a headline whose detail appears
    nowhere would just lose information, so the full text stays. The second
    return value tells the caller which of those two happened, so it does not
    have to infer it by comparing strings.

    KNOWN LIMITATION, deliberate and deferred: "one sentence" only compresses a
    gap that HAS more than one. On the 2026-08 Pranami Bliss data 11 of 17 gaps
    are written as a single dense sentence, so their headline comes back byte-
    identical to the gap entry and the duplication this function exists to
    remove survives for those 11. Shortening further (a clause-boundary cut or
    a character budget) was considered and rejected: both produce grammatical
    fragments, which CLAUDE.md's "Bullets and grammar" rule specifically warns
    against, and the real defect is upstream -- a gap should not be one
    350-character sentence in the first place. The per-finding research stage
    rewrites gaps into properly structured multi-sentence entries, at which
    point this compresses naturally with no change here. Do not "fix" this by
    truncating."""
    text = item.get("text", "")
    number = item.get("gap_number")
    if not number or number not in _rendered_gap_numbers(facts):
        return text, False

    # A model-written headline, if run_editorial_passes cached one. This is what
    # closes the duplication the KNOWN LIMITATION above describes: "first
    # sentence" compresses nothing when the gap IS one sentence. A miss falls
    # through to first-sentence, so no headline means today's behaviour exactly.
    written = facts.get("_flag_headlines", {}).get(number)
    if written:
        return f"{written} (Gap {number})", True

    clauses = _split_into_bullet_clauses(text)
    return f"{clauses[0] if clauses else text} (Gap {number})", True


# ---------------------------------------------------------------------------
# Developer Score -- a composite read on the PROMOTER's own standing,
# scored against a fixed 3-bucket / 9-sub-metric industry rubric:
#   Operational Strength (50%): Team Strength, Influence in Micromarket
#     (area within 5km), Past Experience - Area, Track Record -- 12.5% each.
#   Financial Strength (20%): Financial Strength (debt structure) -- one
#     sub-metric, the whole bucket.
#   Governance Strength (30%): RERA Compliance, GST/TDS Compliance, Cases
#     (Past Defaults), Entity Rating -- 7.5% each.
# Each sub-metric independently banded AAA/AA/A/B/C/D. Distinct from
# _compute_documentation_confidence_score above (which scores this
# DOCUMENT's sourcing, not the promoter).
#
# Unlike every other score/count in this Charter, this rubric does NOT
# renormalize when data is missing: each sub-metric's weight is fixed
# (12.5%/20%/7.5%) whether or not it actually scores, and an unscored
# sub-metric's weight is simply never redistributed to the others or
# excluded from the composite's denominator -- so incomplete public
# disclosure structurally lowers the composite, rather than a promoter
# with less available data being scored purely on the strength of
# whatever little IS known. Three sub-metrics (Team Strength; Financial
# Strength's debt ratios; RERA/GST-TDS Compliance, not yet built) have no
# wired-in data source today -- included for completeness, not left out,
# so a reader always sees why they're missing rather than wondering if
# they were forgotten.
# ---------------------------------------------------------------------------


def _tier_from_score(score: float) -> str:
    for tier, threshold in _DEVELOPER_SCORE_TIER_THRESHOLDS:
        if score >= threshold:
            return tier
    return "D"


def _score_track_record_years(facts: dict) -> dict:
    """Criterion 1: years this promoter (or its parent group, if this is a
    group-affiliated SPV) has been in the real-estate industry. Reads
    developer_track_record.years_in_industry -- a field the deep-research
    promoter-profile step is expected to populate with a sourced start
    date; not yet wired into every pipeline run, so this is commonly N/A
    today, not a parsing failure."""
    years = (facts.get("developer_track_record") or {}).get("years_in_industry")
    if not isinstance(years, (int, float)):
        return {"score": None, "tier": None, "reason": "Years in the industry not confirmed this pass; needs a sourced start date for the promoter or its parent group from the deep-research profile step."}
    if years > 20:
        tier = "AAA"
    elif years >= 17:
        tier = "AA"
    elif years >= 13:
        tier = "A"
    elif years >= 9:
        tier = "B"
    elif years >= 5:
        tier = "C"
    else:
        tier = "D"
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": f"{years} years in the industry"}


def _score_team_strength(facts: dict) -> dict:
    """Criterion 2: strength/experience of the Liaisoning, Project
    Development, and Sales & CRM teams. No public source (MahaRERA,
    ZaubaCorp, or general web search) discloses a promoter's internal
    headcount by function -- this is a permanent, structural gap, not a
    pending build, unless a specific internal data feed is wired in."""
    return {"score": None, "tier": None, "reason": "No public source discloses internal team headcounts by function (Liaisoning / Project Development / Sales & CRM) -- a structural limitation of publicly-available data, not something a future pass can close."}


def _score_past_area_developed(facts: dict) -> dict:
    """Criterion 3: total area (lakh sq ft) this promoter has developed and
    delivered in the past, across its MahaRERA-registered portfolio.
    Reads promoter_portfolio.totals.total_area_developed_lakh_sqft --
    requires per-project area figures aggregated across the promoter's
    other registered projects, not yet computed by every pipeline run."""
    area = ((facts.get("promoter_portfolio") or {}).get("totals") or {}).get("total_area_developed_lakh_sqft")
    if not isinstance(area, (int, float)):
        return {"score": None, "tier": None, "reason": "promoter_portfolio.totals.total_area_developed_lakh_sqft not available -- requires area figures aggregated across the promoter's other MahaRERA-registered projects, not yet computed this pass."}
    if area > 120:
        tier = "AAA"
    elif area >= 81:
        tier = "AA"
    elif area >= 51:
        tier = "A"
    elif area >= 21:
        tier = "B"
    elif area >= 6:
        tier = "C"
    else:
        tier = "D"
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": f"~{area} lakh sq ft developed across the promoter's MahaRERA-registered portfolio"}


def _score_area_within_5km(facts: dict) -> dict:
    """Criterion 4: this promoter's own total developed area (lakh sq ft)
    within a 5km radius of THIS project -- a read on promoter influence in
    the micro-market, not a competitor/all-developer figure. Reads
    promoter_portfolio.totals.area_within_5km_lakh_sqft, geo-filtered from
    the same portfolio data as criterion 3."""
    area = ((facts.get("promoter_portfolio") or {}).get("totals") or {}).get("area_within_5km_lakh_sqft")
    if not isinstance(area, (int, float)):
        return {"score": None, "tier": None, "reason": "promoter_portfolio.totals.area_within_5km_lakh_sqft not available -- this pass's promoter_portfolio.json predates the geocoding-based 5km filter (build_promoter_portfolio's subject_project_partners_data/subject_reg_no params), or no subject location could be geocoded. Re-run the pipeline to compute it, rather than treating this as a permanent gap."}
    if area > 50:
        tier = "AAA"
    elif area >= 21:
        tier = "AA"
    elif area >= 6:
        tier = "A"
    elif area >= 2:
        tier = "B"
    elif area >= 1:
        tier = "C"
    else:
        tier = "D"
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": f"~{area} lakh sq ft developed by this promoter within 5km of this project"}


def _score_financial_strength_debt(facts: dict) -> dict:
    """Criterion 5: a combined 1-50 point score across % debt in total
    capital employed, % secured debt in total debt, and past-default
    occurrence (LOWER points = stronger, per the rubric). Requires actual
    balance-sheet debt structure -- MahaRERA and ZaubaCorp's public profile
    (authorised/paid-up capital only) don't disclose this; only resolvable
    when a credit-rating agency's rationale or an MCA financial filing
    states these ratios explicitly, which is uncommon for the private,
    unrated companies that make up most MahaRERA promoters."""
    points = (facts.get("developer_track_record") or {}).get("financial_strength_points")
    if not isinstance(points, (int, float)):
        return {"score": None, "tier": None, "reason": "Debt-to-capital ratio, secured-debt ratio, and default occurrence are not disclosed by any source this pipeline checks (MahaRERA/ZaubaCorp don't carry balance-sheet debt structure) -- only resolvable when a credit-rating rationale or MCA financial filing states these figures."}
    if points <= 8:
        tier = "AAA"
    elif points <= 17:
        tier = "AA"
    elif points <= 26:
        tier = "A"
    elif points <= 34:
        tier = "B"
    elif points <= 42:
        tier = "C"
    else:
        tier = "D"
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": f"{points} combined debt/default points (lower is stronger)"}


def _score_past_default_count(facts: dict) -> dict:
    """Criterion 6: number of past default events. Derived from the IBBI
    insolvency check (a clean IBBI record scores 0 defaults) cross-checked
    against the credit-rating agency's own note text when a rating exists
    -- IBBI alone only catches insolvency proceedings that reached the
    Board, not every default event, so a rating note that itself flags
    "default" language blocks this from being scored 0 without a human
    reading the raw text. Never guesses an exact count beyond 0/unclear."""
    ibbi_check = facts.get("ibbi_insolvency_check")
    if ibbi_check is None or ibbi_check.get("found_process") is None:
        return {"score": None, "tier": None, "reason": "IBBI insolvency check did not run this pass -- nothing to count past defaults from."}
    if ibbi_check["found_process"] is not False:
        return {"score": None, "tier": None, "reason": "IBBI returned an insolvency record against this CIN -- see the Insolvency Check section for the raw detail; an exact default count isn't machine-countable from that text, so this is left unscored rather than guessed."}
    credit_check = facts.get("credit_rating_check") or {}
    note_text = ((credit_check.get("promoter") or {}).get("note") or "").lower()
    if "default" in note_text and "no instance of default" not in note_text:
        return {"score": None, "tier": None, "reason": "IBBI shows no insolvency process, but the credit-rating check's own note mentions \"default\" in a way this checker can't confidently classify as zero -- see the Credit Rating Check section for the raw text rather than guessing a count."}
    return {"score": _DEVELOPER_SCORE_TIER_SCORES["AAA"], "tier": "AAA", "note": "0 defaults -- IBBI shows no insolvency process against this CIN, and nothing in the credit-rating check (if one ran) states otherwise."}


def _score_entity_rating(facts: dict) -> dict:
    """Criterion 7: entity/organization type. Reads
    corporate_identity.organization_type -- a Private or Public Limited
    Company scores AAA outright; an LLP or Partnership's stated
    "willingness to convert to Pvt Ltd" can't be independently verified
    from public filings, so both are conservatively scored at the "not
    willing to convert" band rather than crediting unconfirmed intent."""
    org_type = ((facts.get("corporate_identity") or {}).get("organization_type") or {}).get("value", "")
    text = org_type.lower()
    if not text:
        return {"score": None, "tier": None, "reason": "corporate_identity.organization_type not confirmed this pass."}
    if "llp" in text or "liability partnership" in text:
        return {"score": _DEVELOPER_SCORE_TIER_SCORES["D"], "tier": "D", "note": "LLP with no independently confirmed intent to convert to Pvt Ltd -- conservatively scored at the \"not willing to convert\" band rather than assuming unconfirmed intent."}
    if "partnership" in text:
        return {"score": _DEVELOPER_SCORE_TIER_SCORES["D"], "tier": "D", "note": "Partnership with no independently confirmed intent to convert to Pvt Ltd -- conservatively scored at the \"not willing to convert\" band rather than assuming unconfirmed intent."}
    if "limited" in text:
        return {"score": _DEVELOPER_SCORE_TIER_SCORES["AAA"], "tier": "AAA", "note": f"Corporatized entity ({org_type.split('(')[0].strip()})."}
    return {"score": None, "tier": None, "reason": f"organization_type ('{org_type}') did not match a recognized entity-type category (Private/Public Limited, LLP, or Partnership)."}


def _score_rera_compliance(facts: dict) -> dict:
    """Governance sub-metric: RERA compliance history for THIS project --
    whether its RERA-registered completion commitment has ever been
    extended, plus how many complaints/appeals are on record against it.
    Combines both into a single 0(best)-100(worst) "compliance friction"
    point score, banded onto the AAA-D scale (LOWER points = stronger,
    same convention as _score_financial_strength_debt). Reuses
    _FLAG_THRESHOLDS' own complaint/appeal breakpoints so a project that's
    already flagged in Overview & Flags for complaint/appeal volume can't
    silently score AAA here -- the two systems read the same underlying
    numbers.

    Deliberately does NOT read complaint-order OUTCOME text (dismissed /
    allowed / etc.) into this score: summarize_complaint_outcomes' own
    docstring is explicit that its keyword classification is a heuristic
    for a quick tally, "not a substitute for a human reading the actual
    order for anything that matters to a real decision" -- a complaint
    dismissed purely on jurisdictional grounds, with the underlying
    allegation never actually ruled on, would be indistinguishable from
    one dismissed on the merits. Raw complaint/appeal COUNT and extension
    status are the only signals reliable enough to score automatically;
    the substance of any complaint still belongs in Litigation Status and
    the flag list, read by a human, not folded into this number.

    Requires complaint_count and appeal_count to both be confidently
    parseable -- _parse_completion_slippage's own "never extended" default
    is always well-defined by construction (see its docstring), so
    extension status alone never blocks scoring."""
    complaint_count, appeal_count = _parse_complaint_appeal_counts(facts)
    if complaint_count is None or appeal_count is None:
        return {"score": None, "tier": None, "reason": "This project's complaint/appeal counts are not confidently parseable from rera_core_fields; both are needed to score RERA compliance friction."}

    _, was_extended = _parse_completion_slippage(facts)
    points = 25 if was_extended else 0
    if complaint_count > _FLAG_THRESHOLDS["complaint_imminent"]:
        points += 45
    elif complaint_count > _FLAG_THRESHOLDS["complaint_monitor"]:
        points += 30
    elif complaint_count >= 1:
        points += 15
    if appeal_count > _FLAG_THRESHOLDS["appeal_imminent"]:
        points += 45
    elif appeal_count > _FLAG_THRESHOLDS["appeal_monitor"]:
        points += 30
    elif appeal_count >= 1:
        points += 15

    if points == 0:
        tier = "AAA"
    elif points <= 20:
        tier = "AA"
    elif points <= 40:
        tier = "A"
    elif points <= 60:
        tier = "B"
    elif points <= 80:
        tier = "C"
    else:
        tier = "D"
    extension_note = "completion date extended" if was_extended else "no completion extension"
    note = f"{points} compliance-friction points (lower is stronger): {extension_note}, {complaint_count} complaint(s), {appeal_count} appeal(s) on record for this project."
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": note}


def _score_gst_compliance(facts: dict) -> dict:
    """Governance sub-metric: GST return-filing compliance (GSTR-1/GSTR-3B
    due-date and delay pattern -- see gst_compliance.py). TDS is explicitly
    out of scope (deferred per an earlier scoping decision), hence "GST
    Compliance" not "GST/TDS Compliance" in _DEVELOPER_SCORE_STRUCTURE.

    Requires facts["gst_compliance_check"] to have been populated by
    run_gst_compliance_check -- itself gated on a human-supplied
    output/<reg_no>/gst_filing_input.json, since there is no automated GST
    portal scrape here at all: the portal's own "Search Taxpayer" filing
    table sits behind a CAPTCHA solved fresh per lookup (see
    gst_compliance.py's own module docstring), the same hard constraint
    already documented for MahaRERA/Maha Bhulekh elsewhere in this file.
    Unscored (not a permanent gap, just nothing supplied this pass) when
    that file was never dropped in.

    Banded on the same "friction points, lower is stronger" convention as
    _score_rera_compliance, and reuses _FLAG_THRESHOLDS' own GST breakpoints
    so a project already flagged in Overview & Flags for a bad filing
    pattern can't silently score AAA here -- the two systems read the same
    underlying numbers. Weights three signals: overall late-filing rate,
    the single worst delay on record, and -- weighted highest, since it's
    what should actually drive an "ask the developer" conversation today --
    how many periods were late or still unfiled within the trailing 12
    months."""
    check = facts.get("gst_compliance_check")
    if not check or not check.get("found"):
        return {
            "score": None, "tier": None,
            "reason": "No GST filing data available this pass. Requires a human-supplied output/<reg_no>/"
            "gst_filing_input.json (GSTIN + filing dates), since the GST portal's own CAPTCHA makes live "
            "scraping unautomatable here; this is a pending-input gap, not a permanent one.",
        }

    summary = check["summary"]
    if summary["on_time"] + summary["late"] == 0:
        return {"score": None, "tier": None, "reason": "GST filing input contained no period with both a resolvable due date and a recorded filing outcome to score against."}

    late_pct = summary["late_pct"]
    points = 45 if late_pct > _FLAG_THRESHOLDS["gst_late_pct_imminent"] \
        else 30 if late_pct > _FLAG_THRESHOLDS["gst_late_pct_monitor"] \
        else 15 if late_pct > 0 else 0

    worst_delay = summary.get("worst_delay_days")
    if worst_delay is not None:
        points += 30 if worst_delay > 60 else 15 if worst_delay > 20 else 0

    delays_recent = summary.get("delays_last_12_months", 0)
    points += 45 if delays_recent > _FLAG_THRESHOLDS["gst_recent_delays_imminent"] \
        else 25 if delays_recent > _FLAG_THRESHOLDS["gst_recent_delays_monitor"] else 0

    if points == 0:
        tier = "AAA"
    elif points <= 20:
        tier = "AA"
    elif points <= 40:
        tier = "A"
    elif points <= 60:
        tier = "B"
    elif points <= 80:
        tier = "C"
    else:
        tier = "D"

    note = (
        f"{points} compliance-friction points (lower is stronger): {late_pct}% of rated periods filed late, "
        f"worst delay {worst_delay if worst_delay is not None else 0} day(s), {delays_recent} delayed/unfiled "
        f"period(s) due in the trailing 12 months."
    )
    return {"score": _DEVELOPER_SCORE_TIER_SCORES[tier], "tier": tier, "note": note}


# Three fixed-weight buckets, each holding a fixed number of sub-metrics
# that split the bucket's weight EQUALLY -- but that split is structural,
# not a renormalization: a sub-metric's weight (12.5% / 20% / 7.5%) is the
# same whether or not it actually has data this pass. When a sub-metric is
# unscored, its weight simply contributes nothing to the composite -- it is
# NOT redistributed to the other sub-metrics, and the composite is NOT
# divided by the sum of only the available weights. This is a deliberate
# choice (over the equal-weight-renormalized approach used elsewhere in
# this Charter): a promoter with less publicly-verifiable data structurally
# scores lower, even if everything that IS available is top-tier, rather
# than have missing disclosure quietly get "filled in" by inflating what
# little is known.
_DEVELOPER_SCORE_STRUCTURE = (
    ("Operational Strength", 50.0, (
        ("team_strength", "Team Strength", _score_team_strength),
        ("area_within_5km", "Influence in Micromarket", _score_area_within_5km),
        ("past_area_developed", "Past Experience - Area", _score_past_area_developed),
        ("track_record_years", "Track Record", _score_track_record_years),
    )),
    ("Financial Strength", 20.0, (
        ("financial_strength_debt", "Financial Strength", _score_financial_strength_debt),
    )),
    ("Governance Strength", 30.0, (
        ("rera_compliance", "RERA Compliance", _score_rera_compliance),
        ("gst_compliance", "GST Compliance", _score_gst_compliance),
        ("past_default_count", "Cases (Past Defaults)", _score_past_default_count),
        ("entity_rating", "Entity Rating", _score_entity_rating),
    )),
)


def _compute_developer_score(facts: dict, flags: dict) -> dict:
    """Returns {"composite": 0-100, "grade": one of AAA/AA/A/B/C/D,
    "criteria": {name: {"score", "tier", "weight", "bucket",
    "display_name", "note"} or {"score": None, "tier": None, "weight":
    <fixed %>, "bucket", "display_name", "reason"}}}. All nine sub-metrics
    (see _DEVELOPER_SCORE_STRUCTURE) always appear in `criteria`, each
    carrying its FIXED structural weight regardless of whether it actually
    scored -- see _DEVELOPER_SCORE_STRUCTURE's own note on why missing
    weight is never redistributed or renormalized away."""
    criteria = {}
    composite = 0.0
    for bucket_name, bucket_weight, metrics in _DEVELOPER_SCORE_STRUCTURE:
        sub_weight = round(bucket_weight / len(metrics), 2)
        for key, display_name, score_fn in metrics:
            result = score_fn(facts)
            result["weight"] = sub_weight
            result["bucket"] = bucket_name
            result["display_name"] = display_name
            criteria[key] = result
            if result["score"] is not None:
                composite += result["score"] * sub_weight / 100.0

    any_scored = any(v["score"] is not None for v in criteria.values())
    grade = _tier_from_score(composite) if any_scored else "D"
    if flags.get("imminent") and grade in ("AAA", "AA"):
        # Hard cap: an imminent-tier flag means this can never be graded
        # better than A, regardless of how strong the composite otherwise
        # looks -- but a composite that already bands to A or below is left
        # as is, since the cap only restrains a grade that would otherwise
        # be too generous, never softens one that's already lower.
        grade = "A"

    return {"composite": round(composite, 1), "grade": grade, "criteria": criteria}


def _parse_completion_slippage(facts: dict) -> tuple[str, bool]:
    """Returns (display text, was_extended). Reads rera_core_fields'
    structured completion_date_current/completion_date_original companions
    first -- an empty completion_date_original means "never extended" by
    construction (see _SYSTEM_PROMPT), not "unknown". Falls back to
    regexing proposed_completion_date's prose (e.g. "2028-03-30 per the
    live RERA record (originally 2024-03-30; ...)" -> ("2024-03-30 -> 2028-
    03-30", True)) only for older facts.json files that predate those
    fields -- never a guessed date either way."""
    core = facts.get("rera_core_fields", {}) or {}
    current = core.get("completion_date_current")
    original = core.get("completion_date_original")
    if current:
        if original and original != current:
            return f"{original} -> {current}", True
        return current, False

    text = core.get("proposed_completion_date", "") or ""
    current_match = re.match(r"(\d{4}-\d{2}-\d{2})", text)
    original_match = re.search(r"originally\s+(\d{4}-\d{2}-\d{2})", text, re.IGNORECASE)
    if current_match and original_match and current_match.group(1) != original_match.group(1):
        return f"{original_match.group(1)} -> {current_match.group(1)}", True
    return (current_match.group(1) if current_match else "Not confirmed"), False


def _land_title_summary(facts: dict) -> tuple[str, str]:
    """Returns (short display text, fill color) for the land-title portion
    of litigation_status.value -- green if that text reads as no title
    litigation/dispute/encumbrance found (checked against several
    projects' worth of actual phrasing, not just one), amber if
    litigation_status exists but doesn't read clean, neutral only if
    litigation_status itself is empty. Never invents a status beyond what
    that text actually says."""
    text = ((facts.get("litigation_status") or {}).get("value", "") or "")
    if not text.strip():
        return "Not confirmed this pass", _FILL_NEUTRAL
    lowered = text.lower()
    clean_signals = (
        "no litigation affecting the property",
        "not shown to be under title dispute",
        "no litigation directly against",
        "clear, marketable, encumbrance-free title",
        "no litigation on the land",
    )
    if any(signal in lowered for signal in clean_signals):
        return "Clear (see Litigation Status for source)", _FILL_GREEN
    return "See Litigation Status -- not a clean read", _FILL_AMBER


def _append_executive_summary_kpis(doc, facts: dict, flags: dict) -> None:
    """Appends the four material-findings cards a reviewer needs inside 60
    seconds -- Completion Slippage, Units Sold, Litigation Load, Land Title
    -- extracted from facts this Charter already established (nothing new
    computed here), each shaded by real severity. Deliberately distinct
    from the Developer Score/Data Authenticity/Flags/Project Status
    scorecard, which lives in Overview & Flags -- this card row is about
    what makes THIS deal specifically notable, not the generic scorecard."""
    total_units, sold_units = _parse_total_units_and_sold(facts)
    complaint_count, appeal_count = _parse_complaint_appeal_counts(facts)
    completion_text, was_extended = _parse_completion_slippage(facts)
    land_title_text, land_title_fill = _land_title_summary(facts)

    if total_units and sold_units:
        sold_pct = round(100 * sold_units / total_units, 1)
        units_text = f"{sold_units:,} / {total_units:,} ({sold_pct}%)"
        units_fill = _FILL_RED if (was_extended and sold_pct >= _NEAR_SELLOUT_PCT) else _FILL_NEUTRAL
    else:
        units_text, units_fill = "Not confirmed", _FILL_NEUTRAL

    if complaint_count is not None and appeal_count is not None:
        litigation_text = f"{complaint_count} complaints / {appeal_count} appeals"
        # Colored off this card's own numbers against their own thresholds
        # -- not off flags["imminent"] in general, which can be entirely
        # unrelated (e.g. an FSI/BUA gap) and would otherwise paint a
        # genuinely clean 0/0 record red.
        if complaint_count > _FLAG_THRESHOLDS["complaint_imminent"] or appeal_count > _FLAG_THRESHOLDS["appeal_imminent"]:
            litigation_fill = _FILL_RED
        elif complaint_count > _FLAG_THRESHOLDS["complaint_monitor"] or appeal_count > _FLAG_THRESHOLDS["appeal_monitor"]:
            litigation_fill = _FILL_AMBER
        else:
            litigation_fill = _FILL_GREEN if complaint_count == 0 and appeal_count == 0 else _FILL_NEUTRAL
    else:
        litigation_text, litigation_fill = "Not confirmed", _FILL_NEUTRAL

    cards = (
        ("Completion Slippage", completion_text, _FILL_RED if was_extended else _FILL_NEUTRAL),
        ("Units Sold", units_text, units_fill),
        ("Litigation Load", litigation_text, litigation_fill),
        ("Land Title", land_title_text, land_title_fill),
    )

    kpi_table = doc.add_table(rows=2, cols=len(cards))
    _set_table_borders(kpi_table)
    header_cells = kpi_table.rows[0].cells
    value_cells = kpi_table.rows[1].cells
    for i, (label, value, fill) in enumerate(cards):
        header_cells[i].text = label
        value_cells[i].text = _externalize_prose(facts, value)
        _shade_cell(header_cells[i], fill)
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
        for para in value_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "The cards above are the material findings from this Charter's own research -- see Overview & "
        "Flags below for the full flag detail behind each one."
    )


def _append_overview_section(doc, facts: dict, flags: dict) -> None:
    """The first substantive thing a reader sees after the title block: a
    deal snapshot table, the headline KPI scorecard (Developer Score, Data
    Authenticity, flag counts, project status -- each shaded red/amber/
    green by its own grade or band), and the full Imminent/Structural/
    Monitor flag lists from _classify_flags -- each item shown with both
    its plain-English text and the exact facts.json field it came from,
    never just a bare label, and colored by severity (red/amber/plain) so
    risk is visible at a glance. Nothing here is a new data point; it's a
    new, urgency-ordered arrangement of facts this Charter already
    established elsewhere."""
    heading_style = doc.paragraphs[4].style  # a known Heading 1 in the template -- removed from the doc later, but only after every _append_* call finishes

    heading_para = doc.add_paragraph("Overview & Flags")
    heading_para.style = heading_style

    core = facts.get("rera_core_fields", {}) or {}
    ci = facts.get("corporate_identity", {}) or {}
    promoter_name = (ci.get("promoter_name") or {}).get("value", "")

    def _header_row(cells):
        for cell in cells:
            _shade_cell(cell, "D9E2F3")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    doc.add_paragraph("Deal snapshot")
    snapshot_table = doc.add_table(rows=1, cols=2)
    _set_table_borders(snapshot_table)
    header_cells = snapshot_table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Value"
    _header_row(header_cells)
    for label, key in (
        ("Project", "project_name"),
        ("Promoter", None),
        ("MahaRERA Registration", "registration_number"),
        ("Project Status", "project_status"),
        ("Proposed Completion Date", "proposed_completion_date"),
        ("Total Units", "total_building_units"),
        ("Litigations on Record", "litigations_per_record"),
    ):
        value = promoter_name if key is None else core.get(key, "")
        row = snapshot_table.add_row()
        row.cells[0].text = label
        cell = row.cells[1]
        cell.text = ""
        run = cell.paragraphs[0].add_run(_externalize_prose(facts, str(value or "")))
        if label == "Proposed Completion Date" and _parse_completion_slippage(facts)[1]:
            # Only colored when this project's own text documents a real
            # slippage (original date vs. the extended one(s)) -- an
            # un-extended date is a plain fact, not a discrepancy, and
            # must not be colored red just because this field sometimes is.
            _color_run(run, _TEXT_RED)

    doc.add_paragraph()
    developer_score = facts.get("developer_score", {}) or {}
    # Computed here rather than read blind: _append_authenticity_page (which
    # normally populates this key) runs LATER in the same _fill_template pass,
    # in the Diligence Appendix. Reading the key directly at this point showed
    # "N/A (N/A/100)" on the Internal document, while the External document --
    # built second, from the same mutated facts dict -- showed the real score.
    # Same document, two different answers for one figure. This makes the
    # scorecard self-sufficient and order-independent.
    doc_confidence = _ensure_documentation_confidence(facts)
    imminent_count = len(flags.get("imminent", []))
    grade = developer_score.get("grade")
    band = doc_confidence.get("band")

    kpi_table = doc.add_table(rows=2, cols=4)
    _set_table_borders(kpi_table)
    header_cells = kpi_table.rows[0].cells
    for i, label in enumerate(("Developer Score", "Data Authenticity", "Flags (Imminent / Structural / Monitor)", "Project Status")):
        header_cells[i].text = label
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    value_cells = kpi_table.rows[1].cells
    value_cells[0].text = f"{grade or 'N/A'} ({developer_score.get('composite', 'N/A')}/100)"
    value_cells[1].text = f"{band or 'N/A'} ({doc_confidence.get('overall', 'N/A')}/100)"
    value_cells[2].text = f"{imminent_count} / {len(flags.get('structural', []))} / {len(flags.get('monitor', []))}"
    value_cells[3].text = _externalize_prose(facts, str(core.get("project_status", "")))
    for cell, fill in (
        (header_cells[0], _grade_fill(grade)),
        (header_cells[1], _band_fill(band)),
        (header_cells[2], _FILL_RED if imminent_count else _FILL_GREEN),
        (header_cells[3], _FILL_NEUTRAL),
    ):
        _shade_cell(cell, fill)
    for cell in value_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Every figure above is drawn directly from the underlying facts already researched for this "
        "project -- see the flag lists below for detail and the Diligence Appendix for the per-pillar/"
        "per-check breakdown behind the Developer Score and Data Authenticity figures."
    )

    def _append_flag_list(title: str, items: list, text_color: str | None) -> None:
        doc.add_paragraph()
        list_heading = doc.add_paragraph(f"{title} ({len(items)})")
        for run in list_heading.runs:
            run.bold = True
            if text_color:
                _color_run(run, text_color)
        if not items:
            # CLAUDE.md Section B: a genuinely-empty section collapses to a
            # single plain "Nothing found" line, not a padded placeholder.
            doc.add_paragraph("Nothing found.")
            return
        for item in items:
            line = doc.add_paragraph()
            _apply_bullet_hanging_indent(line)
            headline, points_at_a_gap = _flag_headline(facts, item)
            # "(see gaps[0])" / "(see rera_core_fields.litigations_per_record)"
            # are raw internal facts.json paths -- meaningless (and
            # unprofessional-looking) to a client, so External drops the
            # annotation entirely rather than trying to reword it. A flag
            # carrying a "(Gap N)" pointer drops it in BOTH variants: the
            # pointer already says where the detail lives, in language a
            # reader can actually follow.
            if points_at_a_gap:
                raw_text = f"• {headline}"
            else:
                raw_text = _variant(facts, f"• {headline} (see {item['field']})", f"• {headline}")
            rendered = _annotate_flag_citations(facts, _externalize_prose(facts, raw_text))
            run = line.add_run(rendered)
            if text_color:
                _color_run(run, text_color)

    _append_flag_list("Imminent Red Flags -- act on these before proceeding", flags.get("imminent", []), _TEXT_RED)
    _append_flag_list("Structural Flags -- standing characteristics, raise directly with the developer", flags.get("structural", []), _TEXT_AMBER)
    _append_flag_list("Monitor Flags -- re-check on a future pass", flags.get("monitor", []), None)


def _append_developer_score_section(doc, facts: dict) -> None:
    """Renders facts["developer_score"] (see _compute_developer_score) as a
    per-sub-metric table against the 3-bucket / 9-sub-metric AAA-D
    industry rubric -- silently does nothing if it was never computed
    (defensive only; the normal pipeline always sets it via
    _append_overview_section before this is called)."""
    developer_score = facts.get("developer_score")
    if not developer_score:
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_external_heading(facts, "Developer Score (Code-Computed)"))
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            f"Composite: {developer_score['composite']}/100 -- Grade {developer_score['grade']}. Scored "
            "against a fixed 3-bucket rubric -- Operational Strength (50%: Team Strength, Influence in "
            "Micromarket, Past Experience - Area, Track Record, 12.5% each), Financial Strength (20%: debt "
            "structure), Governance Strength (30%: RERA Compliance, GST/TDS Compliance, Cases/Past Defaults, "
            "Entity Rating, 7.5% each) -- each sub-metric independently banded AAA/AA/A/B/C/D. Every sub-"
            "metric's weight below is fixed and always shown, even when it couldn't be scored this pass: "
            "unscored weight is never redistributed to the others or excluded from the composite's "
            "denominator, so incomplete public disclosure structurally lowers the composite rather than "
            "being scored purely on the strength of whatever little is known. An imminent-tier flag (see "
            "Overview & Flags) caps the grade at A regardless of the composite, unless the composite alone "
            "already bands lower than that."
        ),
        external_text=(
            f"Composite: {developer_score['composite']}/100 (Grade {developer_score['grade']}), scored "
            "across Operational Strength (50%), Financial Strength (20%), and Governance Strength (30%). "
            "A grade below A signals either weak fundamentals or gaps in public disclosure that could not "
            "be scored."
        ),
    )

    # External drops the Weight column -- the bucket/composite explanation
    # above already covers how scoring works, and a numeric weight per row
    # reads as internal-methodology detail rather than a finding. Internal
    # keeps all 6 columns.
    show_weight = facts.get("_doc_variant") != "external"
    columns = ("Bucket", "Sub-metric", "Tier", "Score", "Weight", "Note / Reason") if show_weight \
        else ("Bucket", "Sub-metric", "Tier", "Score", "Note / Reason")
    note_col = 5 if show_weight else 4

    table = doc.add_table(rows=1, cols=len(columns))
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, label in enumerate(columns):
        header_cells[i].text = label
        _shade_cell(header_cells[i], "D9E2F3")
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for bucket_name, _bucket_weight, metrics in _DEVELOPER_SCORE_STRUCTURE:
        for key, _display_name, _fn in metrics:
            criterion = developer_score.get("criteria", {}).get(key, {})
            row = table.add_row()
            row.cells[0].text = bucket_name
            row.cells[1].text = criterion.get("display_name", key.replace("_", " ").title())
            if show_weight:
                row.cells[4].text = f"{criterion.get('weight', '')}%"
            if criterion.get("score") is None:
                row.cells[2].text = "N/A"
                row.cells[3].text = "N/A"
                row.cells[note_col].text = _externalize_prose(facts, criterion.get("reason", ""))
            else:
                row.cells[2].text = criterion["tier"]
                row.cells[3].text = str(criterion["score"])
                row.cells[note_col].text = _externalize_prose(facts, criterion.get("note", ""))


def _append_counterparty_summary(doc, facts: dict) -> None:
    """A few short, one-line-each summaries of the credit rating,
    insolvency, company-registration, group-affiliation, and Developer
    Score checks -- the key figure and nothing more, each pointing to the
    Diligence Appendix for the full table. Lets Counterparty answer "how
    strong is this promoter" at a glance without repeating the full detail
    twice. Silently omits any check that never ran, same as every other
    _append_*_section here; adds no heading of its own if nothing ran at
    all (so an empty Counterparty subsection never appears)."""
    heading_style = doc.paragraphs[9].style  # "Land Identification" -- a known Heading 2
    heading_added = False

    def _ensure_heading():
        nonlocal heading_added
        if not heading_added:
            heading_para = doc.add_paragraph("Credit Rating, Insolvency & Group Affiliation (Summary)")
            heading_para.style = heading_style
            heading_added = True

    # A plain colon reads better than " -- " in these short one-liners, but
    # each has a dynamic value (agency list, status text, grade...) sitting
    # right where a literal dash would be, so a fixed dict/regex substitution
    # can't reach it -- this separator is built once and spliced in instead.
    # Internal keeps " -- " exactly as before; External gets ":".
    _sep = _variant_sep(facts)

    credit_check = facts.get("credit_rating_check")
    if credit_check:
        _ensure_heading()
        promoter_result = credit_check.get("promoter") or {}
        if promoter_result.get("found"):
            agencies = ", ".join(r["agency"] for r in promoter_result.get("ratings", []))
            doc.add_paragraph(f"Credit rating: found ({agencies or 'agency unspecified'}){_sep} see Diligence Appendix for the full instrument/rating detail.")
        else:
            doc.add_paragraph(f"Credit rating: {promoter_result.get('note', 'not found')} (full detail in Diligence Appendix).")

    ibbi_check = facts.get("ibbi_insolvency_check")
    if ibbi_check and ibbi_check.get("found_process") is not None:
        _ensure_heading()
        if ibbi_check["found_process"] is False:
            # No pointer to the Appendix on a clean result: that section now
            # collapses to "Nothing found." (see _append_ibbi_check_section),
            # so promising "full detail" there would send the reader to an
            # empty page. This line is the verdict CLAUDE.md's "Say it once"
            # rule asks for, and on a clean check it is the whole story.
            # The raw status string is still glossed in place, per "Gloss raw
            # status strings".
            doc.add_paragraph(
                f"IBBI insolvency status: clean{_sep} the register returns "
                f"\"{ibbi_check.get('status_text', '')}\", meaning no insolvency process is open against this promoter."
            )
        else:
            doc.add_paragraph(f"IBBI insolvency status: a record was found against this CIN{_sep} see Diligence Appendix for the raw detail (not auto-classified).")

    profile_check = facts.get("company_profile_check")
    if profile_check and profile_check.get("found"):
        _ensure_heading()
        doc.add_paragraph(
            f"Company registration: confirmed via ZaubaCorp{_sep} {profile_check.get('status', 'status unknown')}, "
            f"incorporated {profile_check.get('incorporation_date', 'unknown')} (directors and full detail in Diligence Appendix)."
        )

    group_check = facts.get("group_companies_check")
    if group_check and group_check.get("found") and group_check.get("companies"):
        _ensure_heading()
        doc.add_paragraph(
            f"Group affiliation: {len(group_check['companies'])} linked entit(y/ies) via shared directors or "
            "registered office (full crosswalk in Diligence Appendix)."
        )

    developer_score = facts.get("developer_score")
    if developer_score:
        _ensure_heading()
        doc.add_paragraph(
            f"Developer Score: {developer_score['composite']}/100{_sep} Grade {developer_score['grade']} "
            "(per-pillar breakdown in Diligence Appendix)."
        )


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_table_borders(table) -> None:
    """Applies grid borders directly via XML rather than a named table
    style -- this template (built with docx-js, not Word) has no "Table
    Grid" style defined, confirmed live (table.style = "Table Grid" raises
    KeyError even though the existing template tables render with visible
    grid lines -- they use explicit borders, not a named style)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


def _append_complaint_outcomes_section(doc, facts: dict) -> None:
    """Appends a page summarizing real, per-complaint outcomes extracted
    from downloaded order PDFs (see api_client.download_complaint_orders +
    summarize_complaint_outcomes) -- code-computed, not model-authored, same
    reasoning as the Document Library table. Silently does nothing if
    complaint_outcomes_summary was never set (e.g. no auth token was
    available to download the orders this pass)."""
    summary = facts.get("complaint_outcomes_summary")
    if not summary or not summary.get("per_complaint"):
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_variant(facts, "Complaint Order Outcomes (Code-Extracted)", "Complaint Order Outcomes"))
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "Each row below reflects the actual order PDF already on file for that complaint (downloaded "
            "via the same document-retrieval mechanism used for project documents), classified by a small, "
            "named set of outcome keywords -- not a self-reported summary. An outcome of \"not determinable\" "
            "means the extracted text didn't match any of those keywords; it is not a claim that the "
            "complaint was resolved favourably or unfavourably, only that the automated classification "
            "could not tell from the text available."
        ),
        external_text=(
            "Each outcome below is classified directly from the complaint's own order PDF, not "
            "self-reported. \"Not determinable\" means the classifier couldn't tell from the text; "
            "it is not a claim the complaint was resolved either way."
        ),
    )

    table = doc.add_table(rows=1, cols=2)
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Outcome"
    header_cells[1].text = "Count"
    for cell in header_cells:
        _shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for outcome, count in sorted(summary["outcome_counts"].items(), key=lambda kv: -kv[1]):
        row = table.add_row()
        row.cells[0].text = outcome.replace("_", " ")
        row.cells[1].text = str(count)

    doc.add_paragraph()
    doc.add_paragraph(
        f"{len(summary['per_complaint'])} complaint order(s) were available and read this pass. "
        f"Per-complaint detail (registration number, outcome, source order filename) is retained in "
        f"this Charter's .facts.json rather than repeated in full here."
    )


def _add_rating_comparison_table(doc, rating_result: dict, facts: dict) -> None:
    """Renders every agency's rating(s) for one entity into a SINGLE
    Agency | Instrument | Rating table -- deliberately one table, not one
    per agency, so that if two agencies rate the same entity differently
    (or agree), a reader sees both rows side by side instead of having to
    flip between separate mini-sections to compare them."""
    ratings = rating_result.get("ratings", [])
    for r in ratings:
        citation = _citation_text(facts, _clean_source_label(r["url"]) or r["url"])
        doc.add_paragraph(f"Match found ({r['agency']}): {r['company_name']} {citation}")

    table = doc.add_table(rows=1, cols=3)
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Agency"
    header_cells[1].text = "Instrument"
    header_cells[2].text = "Rating"
    for cell in header_cells:
        _shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r in ratings:
        for item in r.get("instruments", []):
            row = table.add_row()
            row.cells[0].text = r["agency"]
            row.cells[1].text = item["instrument"]
            row.cells[2].text = item["rating"]

    not_found = rating_result.get("not_found_agencies", [])
    if not_found:
        doc.add_paragraph(f"No public rating found from: {', '.join(not_found)}.")


def _append_credit_rating_section(doc, facts: dict) -> None:
    """Appends a section reporting the code-computed credit-rating check
    on the promoter's exact legal name across every agency checked (see
    lookup_credit_rating) -- if more than one agency rates the same
    entity, all of their ratings are shown together for direct comparison,
    rather than just the first match found. Silently does nothing if
    credit_rating_check was never set (e.g. no promoter name was available
    to check against)."""
    check = facts.get("credit_rating_check")
    if not check:
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_external_heading(facts, "Credit Rating Check (Code-Computed)"))
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "Checked directly against every rating agency's public database (currently ICRA and Infomerics) "
            "for an exact match on the promoter's own legal name -- not a fuzzy or \"probably the same "
            "company\" guess, since attributing a rating to the wrong legal entity would itself be a serious "
            "error. Every agency is checked regardless of whether an earlier one already found something, so "
            "that if two agencies rate the same entity, both ratings are shown here for comparison rather "
            "than silently reporting only one. A promoter having no public rating anywhere is the ordinary "
            "case, not a red flag: these agencies only rate developers that sought a public rating (typically "
            "larger, listed, or NCD-issuing entities)."
        ),
        external_text="Checked against ICRA and Infomerics for a rating under the promoter's exact legal name.",
    )

    promoter_result = check.get("promoter", {})
    if promoter_result.get("found"):
        _add_rating_comparison_table(doc, promoter_result, facts)
    else:
        doc.add_paragraph(promoter_result.get("note", "No result recorded."))

    parent_result = check.get("parent_group")
    if parent_result:
        doc.add_paragraph()
        doc.add_paragraph(
            "The following is a SEPARATE check on a distinct parent/group entity -- it describes that "
            "entity's own credit standing, not the specific promoter/SPV named above; the two must not "
            "be conflated."
        )
        if parent_result.get("found"):
            _add_rating_comparison_table(doc, parent_result, facts)
        else:
            doc.add_paragraph(parent_result.get("note", "No result recorded."))


def _append_ibbi_check_section(doc, facts: dict) -> None:
    """Appends a section reporting the code-computed IBBI insolvency check
    on the promoter's CIN (see lookup_ibbi_insolvency_status). Silently
    does nothing if ibbi_insolvency_check was never set (e.g. no CIN was
    extractable this pass)."""
    check = facts.get("ibbi_insolvency_check")
    if not check or check.get("found_process") is None:
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_external_heading(facts, "Insolvency Check -- IBBI Corporate Debtor Master Data (Code-Computed)"))
    heading_para.style = heading_style
    if check["found_process"] is False:
        # CLAUDE.md Section B: a check that ran and came back clear produces no
        # sentence -- no scope-of-search paragraph, no named source, and above
        # all no citation, since "never attach a citation marker to an absence".
        # This section used to spend a cited paragraph establishing that there
        # was nothing to report, which is the exact defect the rule targets.
        # The score carve-out does not rescue it: the clean IBBI result IS the
        # stated basis of the Cases/Past Defaults sub-metric, but that basis
        # belongs to the sub-metric's own note, and Section B is explicit that
        # "a sentence explaining the figure elsewhere in narrative prose is
        # not [a finding], and is deleted". So the whole section collapses to
        # the empty-section form: heading plus one bare line.
        doc.add_paragraph("Nothing found.")
        return

    _variant_paragraph(
        doc, facts,
        internal_text=(
            "Checked directly against the Insolvency and Bankruptcy Board of India's public Corporate "
            "Debtor Master Data, by the promoter's own CIN -- an exact-identifier lookup, not a name-based "
            "guess."
        ),
        external_text="Checked against IBBI's public Corporate Debtor Master Data by the promoter's own CIN.",
    )
    # Reached only when found_process is True -- a real insolvency signal, which
    # is a finding and stays in full.
    doc.add_paragraph(
        "Result: this CIN returned something other than the standard \"no process\" result. The "
        "raw extracted page content is reproduced below verbatim for a human to read directly -- "
        "this checker was not validated against a real active/past insolvency case, so it does not "
        "attempt to summarize or classify this content itself:"
    )
    doc.add_paragraph(check.get("status_text", ""))
    doc.add_paragraph(f"Source: {_citation_text(facts, _clean_source_label(check.get('url', '')) or check.get('url', ''))}")


def _clean_scraped_address(address: str) -> str:
    """Registry-scraped addresses (ZaubaCorp/Tofler/InstaFinancials) come
    from concatenating separate HTML fields with no space normalization --
    real example seen in this Charter's own data: 'Flat 2207,Floor 22, Wing
    B,Ashok Tower-B ... Parel East,   , Mumbai'. Fixes both variants
    (a data-quality bug, not a content/scope difference): collapses a
    stray empty ", ," segment first, then inserts the missing space after
    every remaining comma."""
    if not address:
        return address
    address = re.sub(r"\s*,\s*,", ",", address)
    address = re.sub(r",(?=\S)", ", ", address)
    return re.sub(r"\s+", " ", address).strip()


def _append_company_profile_section(doc, facts: dict) -> None:
    """Appends a section reporting the code-computed company registration
    profile merged from the MCA-mirror chain (ZaubaCorp -> Tofler ->
    InstaFinancials; see _run_mca_profile_chain). Silently does nothing if
    company_profile_check was never set or found no record."""
    check = facts.get("company_profile_check")
    if not check or not check.get("found"):
        return

    heading_style = doc.paragraphs[4].style
    sources_used = check.get("sources_used") or ["ZaubaCorp"]
    # .title() was clobbering the intentional internal capitals in these
    # brand names ("zaubacorp" -> "Zaubacorp" instead of "ZaubaCorp") --
    # _MCA_SOURCE_DISPLAY_NAMES already has the correct display form.
    sources_label = ", ".join(_MCA_SOURCE_DISPLAY_NAMES.get(s, s.replace(".com", "").replace(".in", "").title()) for s in sources_used)

    doc.add_page_break()
    # Which registries were cross-checked is methodology detail, not
    # something a CEO needs in a section heading.
    heading_para = doc.add_paragraph(_variant(
        facts,
        _external_heading(facts, f"Company Registration Profile ({sources_label}, Code-Computed)"),
        "Company Registration Profile",
    ))
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            f"Pulled directly from {len(sources_used)} independent public company-registry mirror(s) by "
            "CIN -- an exact-identifier lookup, not a name-based guess -- and cross-checked against each "
            "other; any disagreement between them on the current director roster is called out under "
            "Gaps & Limitations rather than silently resolved."
        ),
        external_text=f"Cross-checked across {len(sources_used)} independent company registries by CIN.",
    )
    profile_citation = _citation_text(facts, _clean_source_label(check.get("url", "")))
    _sep = _variant_sep(facts)
    doc.add_paragraph(
        f"{check['name']} ({check['cin']}){_sep} Status: {check.get('status', 'unknown')}; "
        f"Class: {check.get('class_of_company', 'unknown')}; "
        f"Category: {check.get('company_category', 'unknown')}; ROC: {check.get('roc', 'unknown')}"
        + (f" {profile_citation}" if profile_citation else "")
    )
    doc.add_paragraph(f"Incorporated: {check.get('incorporation_date', 'unknown')}")
    doc.add_paragraph(f"Registered address: {_clean_scraped_address(check.get('registered_address', 'unknown'))}")
    doc.add_paragraph(
        f"Authorised capital: {check.get('authorised_capital') or 'not publicly available'}; "
        f"Paid-up capital: {check.get('paid_up_capital') or 'not publicly available'}"
    )
    doc.add_paragraph(check.get("shareholding_note", ""))

    def _add_director_table(title: str, directors: list):
        if not directors:
            return
        doc.add_paragraph()
        doc.add_paragraph(title)
        cols = list(directors[0].keys())
        table = doc.add_table(rows=1, cols=len(cols))
        _set_table_borders(table)
        header_cells = table.rows[0].cells
        for i, col in enumerate(cols):
            header_cells[i].text = col
            _shade_cell(header_cells[i], "D9E2F3")
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
        for director in directors:
            row = table.add_row()
            for i, col in enumerate(cols):
                # `or ""` rather than a .get default: a director record can
                # carry an explicit None (e.g. a DIN the registry never
                # returned), and python-docx raises TypeError on None.
                row.cells[i].text = director.get(col) or ""

    _add_director_table("Current Directors & Key Managerial Personnel", check.get("current_directors", []))
    _add_director_table("Past Directors & Key Managerial Personnel", check.get("past_directors", []))
    doc.add_paragraph()
    doc.add_paragraph(f"Source: {_citation_text(facts, _clean_source_label(check.get('url', '')) or check.get('url', ''))}")


def _build_director_company_links(facts: dict) -> list:
    """Cross-references company_profile_check's current + past directors
    against group_companies_check's per-company "shared director: NAME
    (designation)" basis strings, returning one row per named director --
    {name, role, tenure, link_count}, sorted by link_count descending.
    Built entirely from data this Charter already fetched; no new lookup.
    A director who shows 0 links here genuinely has none among the
    companies found, not a parsing gap."""
    from collections import Counter

    cp = facts.get("company_profile_check") or {}
    gc = facts.get("group_companies_check") or {}

    link_counts = Counter()
    for company in gc.get("companies") or []:
        for basis in company.get("basis", []):
            m = re.match(r"shared director:\s*(.+?)\s*\(", basis)
            if m:
                link_counts[" ".join(m.group(1).split())] += 1

    past_sorted = sorted(cp.get("past_directors") or [], key=lambda d: d.get("Cessation") or "", reverse=True)
    rows, seen = [], set()
    for status, directors in (("Current", cp.get("current_directors") or []), ("Past", past_sorted)):
        for d in directors:
            name = " ".join((d.get("Director Name") or "").split())
            if not name or name in seen:
                continue
            seen.add(name)
            appointment, cessation = d.get("Appointment Date"), d.get("Cessation")
            if status == "Current":
                tenure = f"since {appointment}" if appointment and appointment != "-" else "tenure not confirmed"
            else:
                tenure = f"until {cessation}" if cessation and cessation != "-" else "cessation date not confirmed"
            rows.append({
                "name": name, "role": f"{status} {d.get('Designation', '')}".strip(),
                "tenure": tenure, "link_count": link_counts.get(name, 0),
            })
    rows.sort(key=lambda r: -r["link_count"])
    return rows


def _render_director_hub_diagram(promoter_name: str, rows: list):
    """Renders a small hub-and-spoke diagram (promoter at the center, each
    director with at least one group-company link as a spoke sized by that
    count) to an in-memory PNG buffer -- returns None if no director has
    any links (nothing meaningful to draw). The 299 linked companies
    themselves are deliberately NOT drawn as individual nodes, only counts
    per director -- a graph with that many nodes would be unreadable, and
    the point of this diagram is what a reviewer can read in 60 seconds."""
    linked = [r for r in rows if r["link_count"] > 0]
    if not linked:
        return None

    import io
    import math
    import textwrap

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(7, 5.6))
    ax.axis("off")
    ax.set_xlim(-1.55, 1.55)
    ax.set_ylim(-1.4, 1.4)

    max_count = max(r["link_count"] for r in linked)
    n = len(linked)
    # Offset the first spoke away from due-east/due-west so no line runs
    # straight through the horizontal midline the center label sits on.
    angle_offset = math.pi / n
    for i, r in enumerate(linked):
        angle = angle_offset + 2 * math.pi * i / n
        x, y = math.cos(angle), math.sin(angle)
        weight = r["link_count"] / max_count
        ax.plot([0, x], [0, y], color="#BF8F00", linewidth=1 + 3 * weight, zorder=1)
        ax.scatter([x], [y], s=400 + 2200 * weight, color="#FFE699", edgecolor="#BF8F00", zorder=2)
        ax.text(x * 1.22, y * 1.22, f"{r['name']}\n({r['link_count']} linked)", ha="center", va="center", fontsize=7.5, zorder=4)

    # Drawn after the spokes (not before) so the center disc fully covers
    # every line's inner end rather than the lines drawing over its edge.
    ax.scatter([0], [0], s=4200, color="#2E74B5", zorder=3)
    ax.text(0, 0, textwrap.fill(promoter_name, width=14), ha="center", va="center", color="white", fontsize=7.5, fontweight="bold", zorder=4)

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _append_group_companies_section(doc, facts: dict) -> None:
    """Appends a section listing companies linked to the promoter via
    shared directors or a shared registered office (see
    find_group_companies_by_cin), led by a director relationship map (a
    compact table plus a hub-and-spoke diagram) so a reviewer can read how
    concentrated the group's leadership is without scanning all 299 rows.
    Silently does nothing if group_companies_check was never set or found
    nothing."""
    check = facts.get("group_companies_check")
    if not check or not check.get("found") or not check.get("companies"):
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_external_heading(facts, "Group / Affiliated Companies (Code-Computed Director & Address Crosswalk)"))
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "Every entity below shares at least one concrete, named link with the promoter above -- a "
            "specific director in common, a shared registered office, or a filed subsidiary/associate/JV "
            "relationship -- rather than being inferred from a name or industry match. The strength of each "
            "link should be judged from the basis given, not assumed uniform across the list."
        ),
        external_text=(
            "Each entity below shares a concrete link with the promoter above: a shared director, "
            "registered office, or a filed subsidiary/associate/JV relationship."
        ),
    )

    # A relationship whose counterparty identity is paywalled is still a real,
    # reportable relationship -- it just can't be named (see
    # find_group_companies_by_cin's own note). Stated explicitly so it isn't
    # silently absent from a list a reader would otherwise take as complete.
    undisclosed = check.get("undisclosed_relationship_counts") or {}
    for relationship, count in sorted(undisclosed.items()):
        doc.add_paragraph(
            f"In addition, {count} {relationship} relationship(s) appear on the registry record but their "
            f"counterparty identities are withheld behind that source's paid tier, so they cannot be named "
            f"here. Their existence is confirmed; who they are is not."
        )

    director_rows = _build_director_company_links(facts)
    if any(r["link_count"] for r in director_rows):
        sub_heading = doc.add_paragraph("Director Relationship Map")
        for run in sub_heading.runs:
            run.bold = True
        _variant_paragraph(
            doc, facts,
            internal_text=(
                "Each of this promoter's directors, current or past, cross-referenced against how many of "
                "the group companies below name them as a shared director -- collapses the 299-entity list "
                "to the thing that actually matters here: how concentrated the group's leadership is around "
                "a small number of individuals, not a name-by-name read of every affiliated entity."
            ),
            external_text=(
                "Shows how concentrated the group's leadership is around a small number of individuals, "
                "based on shared directorships across the 299 linked entities."
            ),
        )
        dir_table = doc.add_table(rows=1, cols=3)
        _set_table_borders(dir_table)
        header_cells = dir_table.rows[0].cells
        for i, label in enumerate(("Director", "Role at This Promoter", "Group Companies Linked (of 299)")):
            header_cells[i].text = label
            _shade_cell(header_cells[i], "D9E2F3")
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
        for r in director_rows:
            row = dir_table.add_row()
            row.cells[0].text = r["name"]
            row.cells[1].text = f"{r['role']} ({r['tenure']})"
            row.cells[2].text = str(r["link_count"])
        _remove_fully_empty_rows(dir_table)

        promoter_name = ((facts.get("corporate_identity") or {}).get("promoter_name") or {}).get("value") or "This promoter"
        diagram_buf = _render_director_hub_diagram(promoter_name, director_rows)
        if diagram_buf is not None:
            from docx.shared import Inches

            doc.add_paragraph()
            doc.add_picture(diagram_buf, width=Inches(6))

        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Company Name"
    header_cells[1].text = "CIN / Identifier"
    header_cells[2].text = "Basis for Link"
    for cell in header_cells:
        _shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for company in sorted(check["companies"], key=lambda c: -len(c.get("basis", []))):
        row = table.add_row()
        row.cells[0].text = company.get("name", "")
        row.cells[1].text = company.get("cin", "")
        row.cells[2].text = "; ".join(company.get("basis", []))
    _remove_fully_empty_rows(table)

    doc.add_paragraph()
    doc.add_paragraph(f"Source: {_citation_text(facts, _clean_source_label(check.get('url', '')) or check.get('url', ''))}")


def _append_cts_land_record_section(doc, facts: dict) -> None:
    """Appends a section reporting the CTS -> Property Card lookup (see
    run_cts_land_lookup / mahabhumi.fetch_property_card). Silently does
    nothing if cts_land_record_check was never set (the ordinary case --
    it only runs when a human has supplied cts_lookup_input.json) or found
    no result."""
    check = facts.get("cts_land_record_check")
    if not check or not check.get("found"):
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph("Land Record Check -- Maha Bhulekh Property Card (Code-Assisted, Human-Verified)")
    heading_para.style = heading_style
    doc.add_paragraph(
        "Fetched directly from bhulekh.mahabhumi.gov.in, Maharashtra's official land-records portal, by "
        "exact CTS number -- every field/office/village selection leading up to this was an exact match, "
        "never a guess, and a human read and solved the site's own CAPTCHA to reveal this record."
    )

    fields = check.get("fields") or {}
    if fields:
        table = doc.add_table(rows=1, cols=2)
        _set_table_borders(table)
        header_cells = table.rows[0].cells
        header_cells[0].text = "Field"
        header_cells[1].text = "Value"
        for cell in header_cells:
            _shade_cell(cell, "D9E2F3")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for k, v in fields.items():
            row = table.add_row()
            row.cells[0].text = k
            row.cells[1].text = v
    else:
        # This lookup's result-page structure was never independently
        # confirmed live (see mahabhumi.py's module note) -- when no
        # structured fields could be pulled out, the full raw page text is
        # shown instead of silently dropping the result.
        doc.add_paragraph(
            "This lookup's result page could not be parsed into structured fields this pass -- the raw "
            "page text is reproduced below for a human to read directly:"
        )
        doc.add_paragraph((check.get("raw_text") or "")[:4000])

    doc.add_paragraph()
    doc.add_paragraph(f"Source: {_citation_text(facts, _clean_source_label(check.get('url', '')) or check.get('url', ''))}")


def _append_appeal_judgments_section(doc, facts: dict) -> None:
    """Appends a section listing any appeal-level judgments matched against
    this project's own appeals.json (see search_maharera_judgments +
    cross_reference_appeals). Silently does nothing if the field was never
    set or nothing matched -- the common case, since most appeals are still
    pending and have no published judgment yet."""
    matched = facts.get("appeal_judgments_found")
    if not matched:
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph("Appeal-Level Judgments Found (Code-Matched)")
    heading_para.style = heading_style
    doc.add_paragraph(
        "Each row below is a real judgment from MahaRERA's public Orders/Judgments search, matched to "
        "one of this project's own appeal records by an exact ID (the search result's Complainant No. "
        "against this project's own complaintRegistrationNo) -- not a name-based guess. The underlying "
        "judgment PDF has been downloaded and saved alongside this project's other documents."
    )

    table = doc.add_table(rows=1, cols=4)
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, text in enumerate(("Complainant No.", "Respondent", "Uploaded", "Judgment PDF")):
        header_cells[i].text = text
        _shade_cell(header_cells[i], "D9E2F3")
        for p in header_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for item in matched:
        row = table.add_row()
        row.cells[0].text = item.get("complainant_no") or ""
        row.cells[1].text = item.get("respondent_name") or ""
        row.cells[2].text = item.get("uploaded_date") or ""
        row.cells[3].text = item.get("saved_filename") or "(save failed -- see facts.json)"


def _append_review_authenticity_section(doc, facts: dict) -> None:
    """Appends a section reporting the code-computed review-authenticity
    heuristic triage (see run_review_authenticity_triage). Silently does
    nothing if review_authenticity_triage was never set -- the ordinary
    case, since this pipeline has no automated review-fetching mechanism;
    it only runs when a caller supplies pre-collected reviews (see
    run_company_charter's `reviews` param and --reviews-file)."""
    triage = facts.get("review_authenticity_triage")
    if not triage:
        return

    heading_style = doc.paragraphs[4].style

    doc.add_page_break()
    heading_para = doc.add_paragraph(_external_heading(facts, "Review Authenticity Triage (Code-Computed Heuristics)"))
    heading_para.style = heading_style
    doc.add_paragraph(
        "Five heuristic checks over reviews supplied for this run -- a red-flag surfacing aid for a "
        "human to look at, NOT a fabrication verdict. True fabrication-probability scoring would need "
        "platform-internal data (poster IP/device clustering, account history) this project has no "
        "access to. There is no automated review-fetching mechanism in this pipeline; the reviews "
        "analyzed here were supplied directly to this run."
    )
    doc.add_paragraph(f"Total reviews analyzed: {triage.get('total_reviews_analyzed', 0)}")

    polarization = triage.get("rating_polarization", {})
    if polarization.get("total"):
        doc.add_paragraph(
            f"Rating polarization: {polarization['pct_extreme']}% of ratings sit at the extremes "
            f"(1 or 5 stars) rather than the middle -- a heavily bimodal split is a directional red "
            f"flag (fabricated reviews skew to the extremes), not proof on its own."
        )

    bursts = triage.get("burst_clusters", [])
    if bursts:
        doc.add_paragraph(f"Review bursts detected ({len(bursts)}): {len(bursts)} cluster(s) of 5+ reviews posted within a 7-day window:")
        for b in bursts:
            doc.add_paragraph(f"  - {b['count']} reviews between {b['start_date']} and {b['end_date']}")

    dupes = triage.get("near_duplicate_pairs", [])
    if dupes:
        doc.add_paragraph(
            f"Near-duplicate review text: {len(dupes)} pair(s) of reviews with near-identical "
            f"(templated/copy-pasted) language -- see facts.json for the full text of each pair."
        )

    one_hit = triage.get("one_hit_wonder_reviewers", [])
    if one_hit:
        doc.add_paragraph(
            f"One-hit-wonder reviewers: {len(one_hit)} reviewer(s) whose only review anywhere is this "
            f"one -- a documented pattern for accounts that exist solely to post a single review."
        )

    claims = triage.get("claim_cross_reference", [])
    if claims:
        doc.add_paragraph(
            "Possession-delay claims cross-referenced against this project's own confirmed "
            "proposed_completion_date:"
        )
        table = doc.add_table(rows=1, cols=3)
        _set_table_borders(table)
        header_cells = table.rows[0].cells
        for i, text in enumerate(("Claimed year", "Project record year", "Verdict")):
            header_cells[i].text = text
            _shade_cell(header_cells[i], "D9E2F3")
            for p in header_cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
        for item in claims:
            row = table.add_row()
            row.cells[0].text = str(item.get("claimed_year", ""))
            row.cells[1].text = str(item.get("project_year", "")) if item.get("project_year") is not None else "n/a"
            row.cells[2].text = item.get("verdict", "")


def _ensure_documentation_confidence(facts: dict) -> dict:
    """Returns the Documentation Confidence Score, computing and caching it on
    `facts` the first time it is asked for.

    Exists because two places need this figure -- the Overview & Flags
    scorecard near the top of the document, and the Documentation Authenticity
    & Confidence Summary in the Diligence Appendix -- and the appendix one used
    to be the only writer. Whichever renders first now populates it, so the two
    can never disagree, and a caller that renders only the scorecard still gets
    a real number instead of N/A.

    Idempotent: recomputing from the same facts yields the same score, so a
    cached value is returned as-is rather than recomputed per call site."""
    cached = facts.get("documentation_confidence_score")
    if cached:
        return cached
    summary = _compute_authenticity_summary(facts)
    confidence = _compute_documentation_confidence_score(facts, summary)
    facts["documentation_confidence_score"] = confidence
    return confidence


def _append_authenticity_page(doc, facts: dict) -> None:
    """Appends a new, code-computed section (not model-authored) summarizing
    what tier each cited source falls into and how many claims remain
    explicit gaps -- lets a reader judge this report's reliability from the
    same underlying data already visible in the Sources/Gaps sections,
    rather than trusting a self-assessment."""
    summary = _compute_authenticity_summary(facts)
    # Routed through the shared helper (which caches onto `facts`) rather than
    # computing independently, so this page and the Overview & Flags scorecard
    # can never print two different numbers for the same score. The helper also
    # keeps it available to callers building a separate output from the same
    # facts dict -- e.g. report.py's PDF.
    confidence = _ensure_documentation_confidence(facts)

    # This template's heading styles were created via docx-js, not Word --
    # python-docx's add_heading() looks up a style named "Heading 1" and
    # doesn't find a matching one (confirmed live: raises KeyError even
    # though paragraph.style.name reports "Heading 1" when reading an
    # existing heading). Reusing the actual style OBJECT already applied to
    # an existing Heading-1 paragraph in this same document sidesteps the
    # name-lookup mismatch entirely.
    heading_style = doc.paragraphs[4].style  # a known Heading 1 in the template -- removed from the doc later, but only after every _append_* call finishes

    doc.add_page_break()
    heading_para = doc.add_paragraph("Documentation Authenticity & Confidence Summary")
    heading_para.style = heading_style
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "This page is generated directly from the same sources and gaps already listed earlier in "
            "this document -- it is a count, not a self-assessment. A report author claiming its own "
            "work is \"reliable\" would just be another unverified claim; this page instead classifies "
            "every cited source by tier so a reader can judge confidence from the same underlying data."
        ),
        external_text="Reflects how well this Charter's own claims are sourced and cross-verified.",
    )

    from docx.shared import Pt

    _sep = _variant_sep(facts)
    score_para = doc.add_paragraph()
    score_run = score_para.add_run(f"Documentation Confidence Score: {confidence['overall']}/100{_sep} {confidence['band']}")
    score_run.bold = True
    score_run.font.size = Pt(14)
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "This score rates how well-sourced and verified THIS DOCUMENT's own claims are -- source "
            "quality, completeness, cross-corroboration, recency, and re-check success -- it is NOT a "
            "rating of the underlying project's quality, safety, or investment merit. A project with a "
            "genuinely poor track record whose problems are thoroughly documented can score HIGH here; a "
            "genuinely sound project with little public paper trail to verify against can score LOW. It "
            "is a weighted average of six criteria computed below from this document's own sources and "
            "gaps -- informed by the structure of CRISIL's real-estate methodology (a small number of "
            "named factors rather than one flat checklist) but NOT a replica of any CRISIL formula: CRISIL "
            "does not publish a numeric weighting scheme for any of its three real-estate products, and "
            "the weights used here are this project's own calibration."
        ),
        external_text=(
            "This score measures how well this report's own claims are documented and verified; it "
            "does not rate the project itself. A well-documented project with real problems can score "
            "HIGH here; a genuinely sound project with a thin public paper trail can score LOW."
        ),
    )

    score_table = doc.add_table(rows=1, cols=4)
    _set_table_borders(score_table)
    score_headers = score_table.rows[0].cells
    for i, text in enumerate(("Criterion", "Score", "Weight", "What it means")):
        score_headers[i].text = text
        _shade_cell(score_headers[i], "D9E2F3")
        for p in score_headers[i].paragraphs:
            for run in p.runs:
                run.bold = True

    _CRITERION_LABELS = {
        "source_tier_quality": "Weighted Source-Tier Quality",
        "primary_tier_density": "Primary-Tier Density",
        "completeness_rate": "Completeness Rate",
        "cross_corroboration": "Cross-Corroboration",
        "recency": "Recency",
        "verification_rate": "Independent Verification Rate",
    }
    for key, label in _CRITERION_LABELS.items():
        row = score_table.add_row()
        if key in confidence["criteria"]:
            c = confidence["criteria"][key]
            row.cells[0].text = label
            row.cells[1].text = f"{round(c['score'])}/100"
            row.cells[2].text = f"{c['weight']}%"
            row.cells[3].text = _externalize_prose(facts, c["note"])
        else:
            row.cells[0].text = label
            row.cells[1].text = "N/A"
            row.cells[2].text = "excluded"
            row.cells[3].text = _externalize_prose(facts, "Not applicable this pass -- excluded rather than scored as a failure; remaining weights renormalized to still sum to 100%.")

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    _set_table_borders(table)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Source Tier"
    header_cells[1].text = "Count"
    for cell in header_cells:
        _shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    tier_order = [t for t, _ in _SOURCE_TIERS] + ["Other / unclassified source"]
    for tier in tier_order:
        count = summary["tier_counts"].get(tier, 0)
        if count == 0:
            continue
        row = table.add_row()
        row.cells[0].text = tier
        row.cells[1].text = str(count)

    total = summary["total_sources"]
    primary = summary["primary_count"]
    gaps = summary["total_gaps"]
    pct_primary = round(100 * primary / total) if total else 0

    doc.add_paragraph()
    _variant_paragraph(
        doc, facts,
        internal_text=(
            f"Of {total} cited source(s) in this report, {primary} ({pct_primary}%) come from a primary "
            f"regulatory record, a document opened directly from it, or a live Google Maps route checked "
            f"in this session -- the highest-confidence tier. The remainder are corporate-registry "
            f"mirrors, real-estate aggregator listings, press, or social-media corroboration, all "
            f"explicitly labelled by tier above rather than presented as equivalent to a primary source. "
            f"Separately, {gaps} item(s) in this report are recorded as explicit, unresolved gaps -- facts "
            f"that were sought but could not be confirmed, listed in full under \"Gaps & Limitations\" "
            f"earlier in this document -- rather than filled in with an estimate."
        ),
        external_text=(
            f"{primary} of {total} cited sources ({pct_primary}%) are primary-regulatory or "
            "live-verified; the rest are registry mirrors, aggregators, or press. "
            f"{gaps} item(s) remain open gaps (see Gaps & Sources)."
        ),
    )
    _variant_paragraph(
        doc, facts,
        internal_text=(
            "Recommended reading of this summary: treat primary-regulatory-tier and live-Maps findings as "
            "confirmed; treat aggregator/press/social-tier findings as directional and worth an independent "
            "check before any financial or legal decision; treat every listed gap as genuinely open, not as "
            "an implicit \"probably fine.\""
        ),
        external_text=(
            "Treat primary/regulatory findings as confirmed, aggregator/press findings as directional, "
            "and every listed gap as genuinely open."
        ),
    )


def _diff_mortgage_lender(facts: dict, prior_facts: dict | None) -> str | None:
    """Compares this run's disclosed mortgage_lender against the immediately
    prior Company Charter run for the same project (see
    run_archive.archive_previous_charter). Returns a plain-English note if
    the disclosed lender changed, was newly added, or was no longer
    confirmed -- None if there's no prior run to compare against or nothing
    changed. This is the only historical-diffing this project can do for
    lender identity: MahaRERA's own API exposes no lender/mortgage field at
    all (confirmed by checking every category JSON), so there is nothing to
    diff at the promoter-portfolio level, only same-project run-over-run."""
    if not prior_facts:
        return None
    current = ((facts.get("fsi_metrics", {}) or {}).get("mortgage_lender") or {}).get("value", "").strip()
    prior = ((prior_facts.get("fsi_metrics", {}) or {}).get("mortgage_lender") or {}).get("value", "").strip()
    if current == prior:
        return None
    if current and prior:
        return f"Mortgage lender changed since the prior Charter run: \"{prior}\" -> \"{current}\"."
    if current and not prior:
        return f"Mortgage lender newly disclosed since the prior Charter run: \"{current}\" (not disclosed previously)."
    return (
        f"Mortgage lender \"{prior}\" was disclosed in the prior Charter run but is not "
        "confirmed in this one -- this does not necessarily mean it was paid off, only that "
        "it wasn't re-confirmed from a document opened this pass."
    )


# ---------------------------------------------------------------------------
# Phase 2 checks run concurrently, not sequentially -- confirmed safe: each
# of the 5 checks below hits a different external site (ICRA, Infomerics,
# IBBI, ZaubaCorp x2, MahaRERA Orders) via its own plain requests.get/post
# call, and none of them share module-level session state (the one
# requests.Session() in this file, inside _maharera_orders_search_once, is
# a local variable created fresh per call) or depend on each other's
# results -- so running them in a thread pool is safe and shortens the
# wall-clock cost of this step from the SUM of all 5 calls to roughly the
# SLOWEST single one. Each wrapper below swallows its own exceptions into
# the same honest {"found": False, "note": ...}-shaped result the
# sequential code already produced, so a future's .result() never raises
# and one check's failure still can't affect any other's.
# ---------------------------------------------------------------------------

def _safe_credit_rating(promoter_name: str) -> dict:
    try:
        return lookup_credit_rating(promoter_name)
    except Exception as e:
        return {"found": False, "note": f"Credit rating lookup could not run this pass: {e}"}


def _safe_ibbi_check(identifier: str) -> dict:
    try:
        return lookup_ibbi_insolvency_status(identifier)
    except Exception as e:
        return {"found_process": None, "note": f"IBBI lookup could not run this pass: {e}"}


def _safe_company_profile(identifier: str, promoter_name: str = "") -> dict:
    try:
        return _run_mca_profile_chain(identifier, promoter_name)
    except Exception as e:
        return {"found": False, "note": f"MCA-mirror company-profile chain could not run this pass: {e}"}


def _safe_group_companies(identifier: str) -> dict:
    try:
        return find_group_companies_by_cin(identifier)
    except Exception as e:
        return {"found": False, "note": f"ZaubaCorp group-companies crosswalk could not run this pass: {e}"}


def _safe_judgments_search(project_name: str) -> tuple:
    """Returns (judgments, error_note) instead of raising, matching the
    other _safe_* wrappers -- error_note is None on success."""
    try:
        return search_maharera_judgments(project_name), None
    except Exception as e:
        return [], f"MahaRERA Orders/Judgments search for appeal-level outcomes could not run this pass: {e}"


# ---------------------------------------------------------------------------
# Standalone promoter intake -- runs the same promoter-side checks
# run_company_charter runs internally (_safe_company_profile/_safe_ibbi_check/
# _safe_group_companies/_safe_credit_rating), but from a bare CIN handed to us
# directly, with no RERA project/reg_no required at all. Exists for the case
# where a CIN reaches this pipeline before any RERA number does (a promoter's
# company registration predates their RERA filing by definition).
#
# Deliberately writes its result under the SAME facts.json keys
# run_company_charter already uses for these four checks (see its own
# "results" handling above: facts["company_profile_check"],
# facts["ibbi_insolvency_check"], facts["group_companies_check"],
# facts["credit_rating_check"]) -- so a later run_company_charter pass for
# the same promoter can absorb this file with a plain dict update, never a
# reshape. This does duplicate a small amount of the source/gap-shaping
# logic already in run_company_charter rather than refactoring it out into a
# shared helper -- deliberate: it keeps this addition from touching a single
# existing line in that function, so it stays trivially revertable (delete
# this function and promoter_intake.py, nothing else changes) until this
# standalone path has actually proven useful.
# ---------------------------------------------------------------------------

def run_promoter_intake(cin: str, company_name: str = "", output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Runs the MCA-mirror company-profile chain, IBBI insolvency check, and
    ZaubaCorp group-companies crosswalk for a bare CIN -- plus a credit
    rating check if `company_name` is given (rating lookups are name-based,
    not CIN-based). Persists to output/_pending/<CIN>/promoter_profile.json
    and returns the same dict.

    `company_name` is optional but improves two of these checks: Tofler's
    resolve step inside the company-profile chain needs it (skipped as a
    gap without it, same as run_company_charter's own behavior), and the
    credit rating check is skipped entirely without it since it has no CIN
    equivalent to search by."""
    cin = cin.strip()
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        futures = {
            "profile": executor.submit(_safe_company_profile, cin, company_name),
            "ibbi": executor.submit(_safe_ibbi_check, cin),
            "group": executor.submit(_safe_group_companies, cin),
        }
        if company_name:
            futures["rating"] = executor.submit(_safe_credit_rating, company_name)
        results = {key: future.result() for key, future in futures.items()}

    record = {
        "cin": cin,
        "company_name": company_name,
        "generated_at": datetime.now().isoformat(),
        "sources": [],
        "gaps": [],
    }
    accessed_date = datetime.now().strftime("%Y-%m-%d")

    profile_result = results["profile"]
    record["company_profile_check"] = profile_result
    if profile_result.get("found"):
        sources_used = profile_result.get("sources_used") or ["zaubacorp.com"]
        record["sources"].append({
            "label": f"MCA-mirror company registration profile ({', '.join(sources_used)})",
            "ref": f"{profile_result['name']} -- {profile_result['url']}",
            "topic": "company_profile",
            "published_date": "unknown",
            "accessed_date": accessed_date,
        })
        for conflict in profile_result.get("roster_conflicts") or []:
            record["gaps"].append(conflict)

    ibbi_result = results["ibbi"]
    record["ibbi_insolvency_check"] = ibbi_result
    if ibbi_result.get("found_process") is not None:
        record["sources"].append({
            "label": "IBBI Corporate Debtor Master Data",
            "ref": f"CIN {cin} -- {ibbi_result.get('url', '')}",
            "topic": "insolvency_status",
            "published_date": "unknown",
            "accessed_date": accessed_date,
        })

    group_result = results["group"]
    record["group_companies_check"] = group_result
    if group_result.get("found") and group_result.get("companies"):
        record["sources"].append({
            "label": "ZaubaCorp director/address crosswalk",
            "ref": f"CIN {cin} -- {group_result.get('url', '')} -- {len(group_result['companies'])} linked entit(y/ies)",
            "topic": "group_companies",
            "published_date": "unknown",
            "accessed_date": accessed_date,
        })

    if "rating" in results:
        rating_result = results["rating"]
        record["credit_rating_check"] = {"promoter": rating_result}
        for agency_rating in rating_result.get("ratings", []):
            record["sources"].append({
                "label": f"{agency_rating['agency']} credit rating",
                "ref": f"{agency_rating['company_name']} -- {agency_rating['url']}",
                "topic": "credit_rating",
                "published_date": "unknown",
                "accessed_date": accessed_date,
            })

    out_dir = os.path.join(output_dir, "_pending", cin)
    os.makedirs(out_dir, exist_ok=True)
    with open(os.path.join(out_dir, "promoter_profile.json"), "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return record


# ---------------------------------------------------------------------------
# Phase 3: attaching a pending CIN/CTS-only case to a RERA number once one
# becomes known for that same promoter/plot. This is an explicit action
# (never auto-matched/guessed -- same "a human confirms the link, this
# never fuzzy-matches" policy as every other identifier-linking decision in
# this file) that copies the pending case's already-fetched records into
# the reg_no-keyed output tree under fixed, predictable filenames
# (promoter_profile_carryover.json / land_record_carryover.json) --
# run_company_charter/run_cts_land_lookup check for those specific
# filenames (see below) to reuse this data instead of re-fetching, and
# run_cts_land_lookup's carryover path also runs the CTS cross-check that
# promotes a genuine mismatch to an imminent flag (see _classify_flags).
# ---------------------------------------------------------------------------

def attach_rera_number(case_id: str, reg_no: str, output_dir: str = config.OUTPUT_ROOT) -> dict:
    """Copies whatever promoter_profile.json/land_record.json exist under
    output/_pending/<case_id>/ into output/<reg_no>/ as *_carryover.json --
    copies, not moves, so the original pending case stays intact if the
    same CIN/CTS case later needs attaching to a DIFFERENT reg_no too (a
    promoter's CIN, in particular, can legitimately apply to more than one
    of their projects).

    Returns {"attached": bool, "case_id", "reg_no", "had_promoter_profile",
    "had_land_record", "note"} -- attached is False (with a note, not an
    exception) if the pending case directory doesn't exist or has neither
    file, since a typo'd case_id here should surface as an honest "nothing
    to attach" rather than silently doing nothing indistinguishable from
    success."""
    pending_dir = os.path.join(output_dir, "_pending", case_id)
    if not os.path.isdir(pending_dir):
        return {
            "attached": False, "case_id": case_id, "reg_no": reg_no,
            "had_promoter_profile": False, "had_land_record": False,
            "note": f"No pending case directory found at {pending_dir}",
        }

    promoter_profile_path = os.path.join(pending_dir, "promoter_profile.json")
    land_record_path = os.path.join(pending_dir, "land_record.json")
    had_promoter_profile = os.path.exists(promoter_profile_path)
    had_land_record = os.path.exists(land_record_path)

    if not had_promoter_profile and not had_land_record:
        return {
            "attached": False, "case_id": case_id, "reg_no": reg_no,
            "had_promoter_profile": False, "had_land_record": False,
            "note": f"{pending_dir} exists but has neither promoter_profile.json nor land_record.json",
        }

    project_dir = os.path.join(output_dir, reg_no)
    os.makedirs(project_dir, exist_ok=True)
    if had_promoter_profile:
        shutil.copy2(promoter_profile_path, os.path.join(project_dir, "promoter_profile_carryover.json"))
    if had_land_record:
        shutil.copy2(land_record_path, os.path.join(project_dir, "land_record_carryover.json"))

    return {
        "attached": True, "case_id": case_id, "reg_no": reg_no,
        "had_promoter_profile": had_promoter_profile, "had_land_record": had_land_record,
        "note": f"Copied from {pending_dir} into {project_dir}",
    }


def _load_promoter_carryover(facts: dict, output_dir: str, reg_no: str) -> bool:
    """If output/<reg_no>/promoter_profile_carryover.json exists (written by
    attach_rera_number), loads its 4 checks directly into `facts` and
    returns True -- callers use this to skip re-running the equivalent live
    CIN-based checks. Straight dict assignment, not a merge: the carryover
    record was built with the exact same keys/source shape those checks
    themselves produce (see run_promoter_intake's own docstring), so
    nothing here needs reshaping. Returns False (facts untouched) when no
    carryover file exists -- the ordinary case for every project that had
    its RERA number from the start."""
    carryover_path = os.path.join(output_dir, reg_no, "promoter_profile_carryover.json")
    if not os.path.exists(carryover_path):
        return False

    with open(carryover_path, "r", encoding="utf-8") as f:
        carryover = json.load(f)
    for key in ("company_profile_check", "ibbi_insolvency_check", "group_companies_check", "credit_rating_check"):
        if key in carryover:
            facts[key] = carryover[key]
    facts.setdefault("sources", []).extend(carryover.get("sources", []))
    facts.setdefault("gaps", []).extend(carryover.get("gaps", []))
    return True


def run_company_charter(
    reg_no: str,
    category_data: dict,
    documents_manifest: list,
    documents_dir: str,
    research_data: dict | None = None,
    output_dir: str = config.OUTPUT_ROOT,
    complaint_orders_manifest: list | None = None,
    complaint_orders_dir: str | None = None,
    reviews: list | None = None,
    review_source_label: str | None = None,
    promoter_portfolio: dict | None = None,
    pre_built_facts: dict | None = None,
    pipeline_start_time: float | None = None,
) -> tuple[str, dict]:
    """Returns (out_path, facts) -- facts is the complete, code-and-model
    -assembled Charter data (same content as the .facts.json written
    alongside the docx), so callers can build a separate output (e.g.
    report.py's PDF) from the same source data without re-reading it from
    disk.

    `pre_built_facts`: skips the `_run_charter_pass` API call entirely and
    uses this dict as the model-authored layer instead (still schema-shaped
    like `_CHARTER_FACTS_SCHEMA`) -- for when that call's own auth isn't
    available (no ANTHROPIC_API_KEY configured) and an agentic Claude Code
    session has produced the equivalent JSON directly instead, following
    the exact same system/user prompt _run_charter_pass would have sent
    (see _SYSTEM_PROMPT and this function's own user_prompt construction
    below). Every downstream step (document grounding, professional team,
    developer score, company profile/IBBI/credit checks) still runs
    unchanged -- none of those ever called the API to begin with.

    `pipeline_start_time`: a `time.time()` timestamp from the start of the
    FULL pipeline (main.py's scrape + deep research, not just this
    function), so the Internal document's version-log line can report the
    true end-to-end run time. Defaults to this function's own start when
    omitted (e.g. a standalone `python company_charter.py <REG_NO>` run against
    already-scraped data), which then reports charter-generation time only."""
    _charter_start_time = time.time()
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at {TEMPLATE_PATH}")

    extracted_docs, doc_library_status = _select_documents_for_extraction(documents_manifest, documents_dir)

    if pre_built_facts is not None:
        facts = pre_built_facts
    else:
        research_context = ""
        if research_data:
            research_context = (
                "\n\nAlready-confirmed deep-research findings for this project (reuse these, don't "
                "re-research from scratch):\n" + json.dumps({k: research_data.get(k) for k in deep_research.RESEARCH_KEYS})
            )

        user_prompt = (
            f"MahaRERA registration: {reg_no}\n"
            f"RERA project data (JSON): {json.dumps(category_data.get('projects') or {})}\n"
            f"Promoter/partner data (JSON): {json.dumps(category_data.get('partners') or {})}\n"
            f"Extracted document text (label -> text, high-priority documents only):\n"
            f"{json.dumps(extracted_docs)[:_MAX_TOTAL_DOC_CHARS]}\n"
            f"Full document library ({len(doc_library_status)} entries, for reference/completeness):\n"
            f"{json.dumps([d['document_name'] for d in doc_library_status])}"
            f"{research_context}\n\n"
            f"Produce the complete Company Charter facts as the raw JSON object described above."
        )

        facts = _run_charter_pass(user_prompt)
    facts = _verify_material_claims(facts)
    facts = _check_document_grounding(facts, extracted_docs, category_data, documents_manifest)
    facts["document_library"] = doc_library_status  # always the full, code-computed list -- not model-generated

    # Also code-computed, never model-authored (see summarize_professionals'
    # own note for the real mis-reporting this replaces): MahaRERA's
    # professionals category already names every professional on record with
    # their registration number, so it is read directly rather than
    # paraphrased out of the documents.
    facts["professional_team"] = summarize_professionals(category_data)
    if facts["professional_team"]:
        # Overwrite the model's own prose line for the same fact, which was
        # observed under-reporting it badly ("Engineer and CA firms ... not
        # individually named") while the structured data named all of them.
        # Built from the code-computed list so the two can never disagree.
        facts.setdefault("local_planning", {})["professionals_of_record"] = "; ".join(
            f"{p['role']}: {p['name']}"
            + (f" ({p['registration_label']} {p['registration_number']})" if p["registration_number"] else "")
            for p in facts["professional_team"]
        )

    if promoter_portfolio is not None:
        # Also code-computed, never model-authored -- same reasoning as
        # document_library above: the promoter's cross-project track record
        # (see promoter_portfolio.build_promoter_portfolio) is already a
        # deterministic MahaRERA-only computation by the time it reaches
        # here, so it's attached as-is rather than re-derived or
        # paraphrased by the model.
        facts["promoter_portfolio"] = promoter_portfolio
    if complaint_orders_manifest and complaint_orders_dir:
        # Also code-computed, never model-authored -- the same reasoning as
        # document_library above: a real per-complaint outcome breakdown
        # should never be re-derived or paraphrased by the model when we
        # can just compute it directly from the actual downloaded orders.
        facts["complaint_outcomes_summary"] = summarize_complaint_outcomes(complaint_orders_manifest, complaint_orders_dir)

    promoter_name_for_rating = (facts.get("corporate_identity", {}).get("promoter_name") or {}).get("value", "")

    # Prefer a CIN when both are present in the text (the common case for a
    # company); fall back to an LLPIN for LLP promoters (e.g. Trimity Realty
    # LLP), which have no CIN at all. Confirmed live that every check below
    # accepts either format identically -- IBBI's URL slot and ZaubaCorp's
    # redirect-by-identifier trick both work the same way for a CIN or an
    # LLPIN (see lookup_ibbi_insolvency_status / lookup_company_by_cin's own
    # module notes) -- so one identifier drives all three checks either way.
    identity_text = (facts.get("corporate_identity", {}).get("cin_llpin") or {}).get("value", "")
    corp_identifier = extract_cin(identity_text) or extract_llpin(identity_text)
    identifier_label = "CIN" if extract_cin(identity_text) else "LLPIN"

    project_name_for_judgments = (facts.get("rera_core_fields", {}) or {}).get("project_name", "")

    used_promoter_carryover = _load_promoter_carryover(facts, output_dir, reg_no)

    # All 5 checks below hit different external sites and don't depend on
    # each other -- run them concurrently rather than one after another
    # (see the module note above _safe_credit_rating). Each is gated on the
    # same input-availability condition the sequential version used, so a
    # check is simply never submitted (not run at all, not run-and-discarded)
    # when e.g. no CIN/LLPIN was extractable -- or, now, when a carryover
    # profile above already answered it.
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        futures = {}
        if promoter_name_for_rating and not used_promoter_carryover:
            futures["rating"] = executor.submit(_safe_credit_rating, promoter_name_for_rating)
        if corp_identifier and not used_promoter_carryover:
            futures["ibbi"] = executor.submit(_safe_ibbi_check, corp_identifier)
            futures["profile"] = executor.submit(_safe_company_profile, corp_identifier, promoter_name_for_rating)
            futures["group"] = executor.submit(_safe_group_companies, corp_identifier)
        if project_name_for_judgments:
            futures["judgments"] = executor.submit(_safe_judgments_search, project_name_for_judgments)

        results = {key: future.result() for key, future in futures.items()}

    # From here on, everything is pure in-memory processing of already-fetched
    # results (facts assignment, source bookkeeping) -- fast enough that
    # doing it sequentially costs nothing worth parallelizing further.
    if "rating" in results:
        rating_result = results["rating"]
        facts["credit_rating_check"] = {"promoter": rating_result}
        # One source entry PER agency that found a rating -- both agencies
        # checking the same entity is exactly the comparison case this is
        # meant to surface, so each gets credited as its own source rather
        # than collapsing into one generic "credit rating" entry.
        for agency_rating in rating_result.get("ratings", []):
            facts.setdefault("sources", []).append({
                "label": f"{agency_rating['agency']} credit rating",
                "ref": f"{agency_rating['company_name']} -- {agency_rating['url']}",
                "topic": "credit_rating",
                "published_date": "unknown",
                "accessed_date": datetime.now().strftime("%Y-%m-%d"),
            })

    if "ibbi" in results:
        ibbi_result = results["ibbi"]
        facts["ibbi_insolvency_check"] = ibbi_result
        if ibbi_result.get("found_process") is not None:
            facts.setdefault("sources", []).append({
                "label": "IBBI Corporate Debtor Master Data",
                "ref": f"{identifier_label} {corp_identifier} -- {ibbi_result.get('url', '')}",
                "topic": "insolvency_status",
                "published_date": "unknown",
                "accessed_date": datetime.now().strftime("%Y-%m-%d"),
            })

    if "profile" in results:
        profile_result = results["profile"]
        facts["company_profile_check"] = profile_result
        if profile_result.get("found"):
            sources_used = profile_result.get("sources_used") or ["zaubacorp.com"]
            facts.setdefault("sources", []).append({
                "label": f"MCA-mirror company registration profile ({', '.join(sources_used)})",
                "ref": f"{profile_result['name']} -- {profile_result['url']}",
                "topic": "company_profile",
                "published_date": "unknown",
                "accessed_date": datetime.now().strftime("%Y-%m-%d"),
            })
            # Director-roster disagreements between the MCA mirrors queried
            # above (e.g. one lists a director the others don't) are
            # genuine data-quality signal for a reader, not internal
            # process narration -- surfaced via the existing gaps
            # convention so they render in both Charter variants (see
            # _merge_director_rosters).
            for conflict in profile_result.get("roster_conflicts") or []:
                facts.setdefault("gaps", []).append(conflict)

    if "group" in results:
        group_result = results["group"]
        facts["group_companies_check"] = group_result
        if group_result.get("found") and group_result.get("companies"):
            facts.setdefault("sources", []).append({
                "label": "ZaubaCorp director/address crosswalk",
                "ref": f"{identifier_label} {corp_identifier} -- {group_result.get('url', '')} -- {len(group_result['companies'])} linked entit(y/ies)",
                "topic": "group_companies",
                "published_date": "unknown",
                "accessed_date": datetime.now().strftime("%Y-%m-%d"),
            })

    if "judgments" in results:
        judgments, judgments_error = results["judgments"]
        if judgments_error:
            facts.setdefault("gaps", []).append(judgments_error)
        appeals_data = category_data.get("appeals") or []
        matched_judgments = cross_reference_appeals(judgments, appeals_data)
        matched_judgments = _save_appeal_judgment_pdfs(reg_no, matched_judgments, output_dir)
        facts["appeal_judgments_found"] = matched_judgments
        if matched_judgments:
            facts.setdefault("sources", []).append({
                "label": "MahaRERA Orders/Judgments",
                "ref": f"{project_name_for_judgments} -- {_MAHARERA_ORDERS_URL} -- {len(matched_judgments)} matched judgment(s)",
                "topic": "appeal_judgments",
                "published_date": "unknown",
                "accessed_date": datetime.now().strftime("%Y-%m-%d"),
            })

    # No automated review-fetching mechanism exists in this pipeline (see
    # run_review_authenticity_triage's own module note) -- this only runs
    # when a caller supplies pre-collected reviews, e.g. via --reviews-file.
    if reviews:
        facts["review_authenticity_triage"] = run_review_authenticity_triage(reviews, facts)
        facts.setdefault("sources", []).append({
            "label": review_source_label or "User reviews (authenticity triage)",
            "ref": f"{len(reviews)} review(s) analyzed for authenticity signals",
            "topic": "review_authenticity",
            "published_date": "unknown",
            "accessed_date": datetime.now().strftime("%Y-%m-%d"),
        })

    facts = run_cts_land_lookup(facts, reg_no, output_dir)
    facts = run_gst_compliance_check(facts, reg_no, output_dir)

    land = facts.get("land_identification", {})
    origin_locality = (land.get("village_locality") or {}).get("value", "")
    origin_district = (land.get("mandal_taluka_district") or {}).get("value", "")
    if origin_locality:
        origin = f"{origin_locality}, {origin_district}, Maharashtra" if origin_district else f"{origin_locality}, Maharashtra"
        facts = _refine_distances_with_maps(facts, origin)

    # Runs last, after every other step above has had its chance to add a
    # source -- otherwise a topic that only looks single-sourced mid-assembly
    # (e.g. before the credit-rating/IBBI checks ran) would get a spurious
    # gap for a corroboration problem that a later step already fixed.
    facts = verify_cross_corroboration(facts)

    # The template is a fixed reference asset checked into the real repo
    # output root regardless of --output-dir; generated charters follow
    # whatever output_dir the caller actually asked for.
    out_dir = os.path.join(output_dir, "company_charters")
    os.makedirs(out_dir, exist_ok=True)
    project_name = (facts.get("rera_core_fields", {}).get("project_name") or reg_no).strip().replace(" ", "_")

    # Snapshot the prior Charter run (if any) before this run's write
    # overwrites it, so lender/mortgage disclosure can be diffed run-over-run.
    prior_charter_archive_dir = run_archive.archive_previous_charter(reg_no, output_dir)
    prior_charter_facts = run_archive.load_prior_charter_facts(prior_charter_archive_dir)
    lender_history_note = _diff_mortgage_lender(facts, prior_charter_facts)
    if lender_history_note:
        facts["mortgage_lender_history_note"] = lender_history_note

    # Cross-run source-trust bookkeeping (see _record_source_hits_and_promote's
    # own module note) -- always recorded regardless of doc_variant; only
    # rendered in the Internal document below.
    facts["source_promotion_notes"] = _record_source_hits_and_promote(facts, reg_no)

    # Per-finding deep research (CLAUDE.md Section B, "Deep research on every
    # finding"). Runs here, after every check above has had its say and before
    # either document renders, so the enriched text reaches both variants and
    # is persisted to .facts.json.
    #
    # Normalization runs first so a relocated finding is researched in the
    # field it now belongs to. The clean-check scrubber deliberately does NOT
    # run first: _collect_findings already skips clean-check clauses on its
    # own, and scrubbing here would mean restoring afterwards, which would
    # overwrite the very enrichment this stage just wrote into those fields.
    # Never fatal: a failed call leaves that finding's original text in place.
    _normalize_misfiled_facts(facts)

    # Editorial judgements a model does better than a keyword table, computed
    # once here for both variants. Every consumer falls back to its existing
    # deterministic path on a miss, so a failure leaves output unchanged rather
    # than degraded.
    try:
        _editorial = run_editorial_passes(facts)
        if any(_editorial.values()):
            print(f"[OK] Editorial passes: {_editorial['clean_checks']} clause verdict(s), "
                  f"{_editorial['citations']} citation match(es), {_editorial['headlines']} headline(s).")
    except Exception as e:
        print(f"[WARN] Editorial passes failed ({e}) -- falling back to deterministic rules.")

    try:
        _research_summary = run_finding_research(facts)
        print(f"[OK] Per-finding research: {_research_summary['enriched']} of "
              f"{_research_summary['findings_seen']} finding(s) resolved in depth"
              + (f", {_research_summary['kept_original']} kept as-is"
                 if _research_summary["kept_original"] else "") + ".")
    except Exception as e:
        # Broad on purpose, same policy as every other API-dependent stage:
        # findings keep their original text and the Charter still builds.
        print(f"[WARN] Per-finding research failed ({e}) -- findings keep their original text.")

    # Two documents from the same underlying facts: Internal (today's
    # existing behavior -- inline "(label)" citations, "(Code-Computed)"
    # labels, verbatim prose) for the team's own diligence use, and
    # External (numbered "[N]" citations resolving to a generic-language
    # Sources list, no internal-process labels/phrasing) for sharing
    # outside it. Internal renders FIRST and on the real `facts` dict --
    # it's the one that computes developer_score/documentation_confidence_
    # score etc. and must be what gets persisted to .facts.json below.
    # External renders from an internal deep copy (see
    # _externalized_facts_copy) and never touches the real `facts` dict,
    # however many times or in whatever order the two variants render.
    # Version log inputs -- computed here, right before rendering, so the
    # cost/call count reflect every API call this run made (charter_pass,
    # verification passes, document grounding, etc. all log to the same
    # deep_research._USAGE_LOG). Only passed into the Internal render below
    # -- External never receives these kwargs, so it stays unaffected.
    _usage_total = deep_research.usage_summary()["total"]
    _run_elapsed_seconds = time.time() - (pipeline_start_time or _charter_start_time)

    out_path = os.path.join(out_dir, f"Company_Charter_{project_name}_{reg_no}_Internal.docx")
    _fill_template(
        reg_no, facts, out_path, doc_variant="internal",
        elapsed_seconds=_run_elapsed_seconds, cost_usd=_usage_total["cost_usd"], api_calls=_usage_total["calls"],
    )
    external_out_path = os.path.join(out_dir, f"Company_Charter_{project_name}_{reg_no}_External.docx")
    _fill_template(reg_no, facts, external_out_path, doc_variant="external")

    # Final stage: both documents are re-read and audited against the CLAUDE.md
    # rules they were written under, before the PDFs are produced. Advisory --
    # it reports and writes a review file, and never blocks delivery. See
    # run_claude_md_document_review for why.
    facts["claude_md_review"] = run_claude_md_document_review(
        {"internal": out_path, "external": external_out_path},
        output_dir=output_dir, reg_no=reg_no,
    )

    internal_pdf_path = _convert_docx_to_pdf(out_path)
    _convert_docx_to_pdf(external_out_path)

    # The page drops clean checks; the record keeps them (CLAUDE.md Section B:
    # the scope of what was checked "stay[s] in the facts file"). Restore
    # before persisting, so a later pass reading this file still sees what was
    # actually searched rather than inheriting a hollowed-out record.
    _restore_clean_checks(facts, facts.pop("_pre_scrub_narrative", None))
    _pre_sanitize_gaps = facts.pop("_pre_sanitize_gaps", None)
    if _pre_sanitize_gaps:
        facts["gaps"] = _pre_sanitize_gaps

    facts_path = os.path.join(out_dir, f"Company_Charter_{project_name}_{reg_no}.facts.json")
    with open(facts_path, "w", encoding="utf-8") as f:
        json.dump(facts, f, indent=2, ensure_ascii=False)

    # PDF is the actual final deliverable format -- fall back to the .docx
    # path only if conversion itself failed (no Word installed, COM error).
    return internal_pdf_path or out_path, facts


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a Company Charter docx for an already-scraped project.")
    parser.add_argument("reg_no", help="MahaRERA registration number whose output/ folder already exists.")
    parser.add_argument("--output-dir", default=config.OUTPUT_ROOT)
    parser.add_argument(
        "--reviews-file",
        default=None,
        help=(
            "Optional path to a JSON file containing a list of pre-collected reviews "
            "({rating, date, text, reviewer_name, reviewer_review_count}) to run the "
            "review-authenticity heuristic triage against. There is no automated "
            "review-fetching mechanism in this pipeline -- reviews must be collected "
            "and saved to this file separately."
        ),
    )
    parser.add_argument(
        "--reviews-source-label",
        default=None,
        help="Label for the reviews source shown in the Charter's Sources list (e.g. 'MouthShut.com'). Defaults to a generic label.",
    )
    args = parser.parse_args()

    project_out_dir = os.path.join(args.output_dir, args.reg_no)
    raw_dir = os.path.join(project_out_dir, "raw")
    if not os.path.isdir(raw_dir):
        print(f"[ERROR] No raw data found at {raw_dir} -- run `python main.py {args.reg_no}` first.")
        return 1

    category_data = finalize_report.load_category_data(raw_dir)
    manifest_path = os.path.join(project_out_dir, "documents_manifest.json")
    documents_manifest = json.load(open(manifest_path, encoding="utf-8")) if os.path.exists(manifest_path) else []
    documents_dir = os.path.join(project_out_dir, "documents")

    research_path = os.path.join(project_out_dir, "research", "deep_research.json")
    research_data = json.load(open(research_path, encoding="utf-8")) if os.path.exists(research_path) else None

    reviews = None
    if args.reviews_file:
        if not os.path.exists(args.reviews_file):
            print(f"[ERROR] --reviews-file not found: {args.reviews_file}")
            return 1
        reviews = json.load(open(args.reviews_file, encoding="utf-8"))

    print(f"[..] Generating Company Charter for {args.reg_no} (model={MODEL})")
    out_path, _facts = run_company_charter(
        args.reg_no, category_data, documents_manifest, documents_dir, research_data, args.output_dir,
        reviews=reviews, review_source_label=args.reviews_source_label,
    )
    print(f"[OK] Company Charter written to {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
