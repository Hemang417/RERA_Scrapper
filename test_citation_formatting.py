"""
Regression tests for two real bugs found by inspecting the Pranami Bliss and
IRA Insignia External Company Charters and fixed in company_charter.py:

1. _clean_source_label/_citation_text split a source string on every ";",
   including semicolons that sit INSIDE a source's own trailing "(...)"
   annotation -- e.g. "Title Report - RERA.pdf (First Schedule; land-title
   chain; 30-year litigation search)" is ONE source with a 3-part
   annotation, not 3 separate sources. The naive split tore it into 3 fake
   citations, one left with a dangling, unmatched ")" ("30-year litigation
   search)"). Fixed with _split_outside_parens, a paren-depth-aware split.

2. _fill_variable_paragraphs cloned the paragraph STYLE for citations past
   the template's pre-built slot_count, but a bullet's numbering (w:numPr)
   lives separately from its style and was never copied -- so the 9th+
   Source in a Charter's numbered list silently lost its bullet and
   indent. Fixed by capturing the last real slot's numPr once and
   re-applying it to every overflow paragraph.

Run directly: python test_citation_formatting.py
"""

import docx

import company_charter as cc

_TEMPLATE_PATH = cc.TEMPLATE_PATH


def test_split_outside_parens_ignores_semicolons_inside_parens():
    """The core utility: a semicolon inside "(...)" is not a separator."""
    assert cc._split_outside_parens("a; b (c; d); e", ";") == ["a", " b (c; d)", " e"]
    assert cc._split_outside_parens("no separators here", ";") == ["no separators here"]
    assert cc._split_outside_parens("x; y; z", ";") == ["x", " y", " z"]
    print("test_split_outside_parens_ignores_semicolons_inside_parens: PASS")


def test_clean_source_label_keeps_multi_part_annotation_as_one_source():
    """Real regression case: a single document with a 3-part parenthetical
    annotation must stay ONE cleaned label, not fragment into 3 -- and
    critically, the last fragment must never end up with a dangling ")"."""
    raw = "output/P51800077150/documents/Title Report - RERA.pdf (First Schedule; land-title chain; 30-year litigation search)"
    cleaned = cc._clean_source_label(raw)
    # The whole annotation must survive attached to ONE filename -- this is
    # the exact opposite of the bug, where it fragmented into "Title
    # Report - RERA.pdf (First Schedule", "land-title chain", and a
    # dangling "30-year litigation search)" with no opening paren at all.
    assert cleaned == "Title Report - RERA.pdf (First Schedule; land-title chain; 30-year litigation search)"
    assert cleaned.count("(") == 1 and cleaned.count(")") == 1
    print("test_clean_source_label_keeps_multi_part_annotation_as_one_source: PASS")


def test_clean_source_label_still_splits_genuinely_separate_sources():
    """Guards the other direction -- real multi-source strings (joined by
    "; " with NO shared enclosing parens) must still split correctly, so
    the paren-aware fix doesn't regress the original multi-citation case."""
    raw = "output/P51700031409/documents/A.pdf; output/P51700031409/documents/B.pdf; https://example.com/page"
    cleaned = cc._clean_source_label(raw)
    parts = cleaned.split("; ")
    assert parts == ["A.pdf", "B.pdf", "example.com"]
    print("test_clean_source_label_still_splits_genuinely_separate_sources: PASS")


def test_citation_text_does_not_refragment_a_parenthetical_annotation():
    """_citation_text re-splits the ALREADY-cleaned label on ";" a second
    time (to register each real source separately) -- it must use the same
    paren-aware split, or a cleaned label carrying a multi-part annotation
    would fragment again at this second stage."""
    facts = {"_doc_variant": "external", "_citation_registry": {"order": [], "index": {}}}
    cleaned_label = "Title Report - RERA.pdf (First Schedule; land-title chain; 30-year litigation search)"
    marker = cc._citation_text(facts, cleaned_label)
    assert marker == "[1]"  # ONE citation registered, not 3
    assert facts["_citation_registry"]["order"] == ["Title Report"]
    print("test_citation_text_does_not_refragment_a_parenthetical_annotation: PASS")


def test_fill_variable_paragraphs_preserves_bullet_past_template_slots():
    """The Sources list template only has 8 pre-built bulleted slots.
    Every item past that (9, 10, 11, ...) must still render as a bulleted,
    indented list item -- not a plain unbulleted paragraph."""
    doc = docx.Document(_TEMPLATE_PATH)
    texts = [f"[{i}] test source {i}" for i in range(1, 12)]  # 11 items, 3 past the 8 slots
    cc._fill_variable_paragraphs(doc, 69, 8, texts)

    # doc.paragraphs builds fresh Paragraph wrappers on every access, so
    # capture index+paragraph together in one pass rather than calling
    # .index() against a second, unrelated fresh list.
    all_paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(all_paragraphs) if p.text.strip().startswith("[1] test source"))

    for offset in range(11):
        para = all_paragraphs[start + offset]
        assert para.text.strip() == f"[{offset + 1}] test source {offset + 1}"
        pPr = para._p.pPr
        assert pPr is not None and pPr.numPr is not None, (
            f"item {offset + 1} (index {start + offset}) lost its bullet numbering"
        )
    print("test_fill_variable_paragraphs_preserves_bullet_past_template_slots: PASS")


if __name__ == "__main__":
    test_split_outside_parens_ignores_semicolons_inside_parens()
    test_clean_source_label_keeps_multi_part_annotation_as_one_source()
    test_clean_source_label_still_splits_genuinely_separate_sources()
    test_citation_text_does_not_refragment_a_parenthetical_annotation()
    test_fill_variable_paragraphs_preserves_bullet_past_template_slots()
    print("\nAll tests passed.")
