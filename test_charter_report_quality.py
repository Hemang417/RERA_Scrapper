"""
Regression tests for charter_report.verify_charter_report_quality -- the
Company Charter quality gate. Each test proves the gate actually detects
the specific bug it claims to catch (not just that it returns something),
and that the real, current Pranami documents pass clean.

Run directly: python test_charter_report_quality.py
"""

import json
import os

import docx

import charter_report as cr

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_quality_gate")


def _scratch_path(name):
    os.makedirs(_SCRATCH, exist_ok=True)
    return os.path.join(_SCRATCH, name)


def _cleanup():
    import shutil
    shutil.rmtree(_SCRATCH, ignore_errors=True)


def test_catches_dash_artifact():
    path = _scratch_path("dash.docx")
    doc = docx.Document()
    doc.add_paragraph("This has a bad -- dash artifact.")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("dash artifact" in v for v in violations), violations
    print("test_catches_dash_artifact: PASS")


def test_catches_redundant_attribution_only_alongside_a_citation():
    """The "(per X)" phrasing is legitimate on its own (a value with no
    separate source at all) -- only a violation once a bracketed citation
    is ALSO present, duplicating what the citation already says."""
    path = _scratch_path("redundant.docx")
    doc = docx.Document()
    doc.add_paragraph("Public Limited Company (per ZaubaCorp) [1]")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("redundant inline attribution" in v for v in violations), violations

    path_clean = _scratch_path("not_redundant.docx")
    doc = docx.Document()
    doc.add_paragraph("Certificate Signed / New (per live RERA record)")
    doc.save(path_clean)
    violations_clean = cr.verify_charter_report_quality(path_clean)
    assert not any("redundant inline attribution" in v for v in violations_clean), violations_clean
    print("test_catches_redundant_attribution_only_alongside_a_citation: PASS")


def test_catches_bare_domain_mention():
    path = _scratch_path("bare_domain.docx")
    doc = docx.Document()
    doc.add_paragraph("Per the group's own website (example.com/about), founded in 2002.")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("bare-domain mention" in v for v in violations), violations
    print("test_catches_bare_domain_mention: PASS")


def test_catches_manual_page_break():
    """A manual page-break character is exactly how the blank-page bug
    came back once -- every major-section heading must use
    page_break_before instead (see _Builder.heading)."""
    path = _scratch_path("manual_break.docx")
    doc = docx.Document()
    doc.add_paragraph("Some text")
    doc.add_page_break()
    doc.add_paragraph("More text")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("manual page-break character" in v for v in violations), violations
    print("test_catches_manual_page_break: PASS")


def test_catches_missing_page_number_fields():
    path = _scratch_path("no_page_numbers.docx")
    doc = docx.Document()
    doc.add_paragraph("Some text")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("PAGE field" in v for v in violations), violations
    assert any("NUMPAGES field" in v for v in violations), violations
    print("test_catches_missing_page_number_fields: PASS")


def test_catches_orphan_citation():
    path = _scratch_path("orphan_citation.docx")
    doc = docx.Document()
    doc.add_paragraph("Some claim [5].")
    doc.save(path)
    violations = cr.verify_charter_report_quality(path)
    assert any("[5]" in v and "missing from References" in v for v in violations), violations
    print("test_catches_orphan_citation: PASS")


def test_real_documents_pass_clean():
    """The actual current Pranami Bliss Internal and External documents,
    rebuilt fresh, must pass the gate with zero violations -- this is the
    gate build_charter_report itself already runs and raises on, so this
    test is really just confirming that invariant stays true."""
    import charter_research_prep as prep
    import run_charter_pipeline as pipeline

    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        facts = json.load(f)

    minimal_research = {"directors": {}, "group": {}}
    for variant in ("internal", "external"):
        out = _scratch_path(f"real_{variant}.docx")
        cr.build_charter_report("P51800077150", json.loads(json.dumps(facts)), minimal_research, out, variant, "31 July 2026")
        violations = cr.verify_charter_report_quality(out)
        assert violations == [], (variant, violations)
    print("test_real_documents_pass_clean: PASS")


if __name__ == "__main__":
    test_catches_dash_artifact()
    test_catches_redundant_attribution_only_alongside_a_citation()
    test_catches_bare_domain_mention()
    test_catches_manual_page_break()
    test_catches_missing_page_number_fields()
    test_catches_orphan_citation()
    test_real_documents_pass_clean()
    _cleanup()
    print("\nAll tests passed.")
