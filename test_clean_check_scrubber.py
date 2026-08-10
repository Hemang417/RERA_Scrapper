"""
Tests for CLAUDE.md Section B's "A clean check produces no sentence" rule --
the deterministic layer (_scrub_clean_checks) and the External gate
(_verify_external_document_quality's cited-absence check).

The rule's hard part, and what most of these tests pin, is what must NOT be
deleted. These two sentences are grammatically identical:
    "No litigation is disclosed against the promoter."           -> delete
    "Not found among the documents reviewed: an FSI certificate." -> KEEP
The first says a RISK is absent; the second says EVIDENCE is absent, which is
a gap, and Section B says gaps are never compressed or deleted. Shape alone
cannot tell them apart, so the scrubber requires both a reviewed field and a
risk noun. Deleting a real finding or a gap is far worse than leaving a clean
check on the page, and these tests encode that asymmetry.

Run directly: python test_clean_check_scrubber.py
"""

import json
import os

import docx

import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_clean_checks")


def _facts() -> dict:
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        return json.load(f)


def _all_text(path: str) -> str:
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _build(variant: str) -> str:
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    cc._fill_template("P51800077150", _facts(), out, doc_variant=variant)
    return out


# --- the detector -----------------------------------------------------------

def test_detects_a_bare_clean_check():
    for text in (
        "No litigation is disclosed against the promoter, the project, or the underlying land in any source reviewed.",
        "No discrepancy found.",
        "No complaints, appeals, or warrants are on record for this project.",
    ):
        assert cc._is_clean_check_clause(text), text
    print("test_detects_a_bare_clean_check: PASS")


def test_a_gap_is_not_a_clean_check():
    """Section B: "We checked and it is clear" is deleted; "we looked and could
    not establish this" is kept in full. Same grammar, opposite treatment --
    what differs is that EVIDENCE is missing, not a risk."""
    for text in (
        "Not found among the documents reviewed: a standalone sanctioned-FSI certificate.",
        "No public rating found for this exact legal entity name from any agency checked.",
        "MahaRERA's Promoters-tab search has no Registered/Revoked toggle, unlike the Projects tab.",
    ):
        assert not cc._is_clean_check_clause(text), text
    print("test_a_gap_is_not_a_clean_check: PASS")


def test_an_absence_carrying_a_finding_survives():
    """"...found nothing except a Notice of Lis Pendens" reports a real
    finding. A clause naming a specific dated event is not a bare nothing."""
    for text in (
        "The 30-year search found nothing against either property except an incidental Notice of Lis Pendens naming an adjoining society.",
        "The two 6 March 2024 orders are administrative Deputy Registrar actions, not litigation or a tribunal proceeding.",
    ):
        assert not cc._is_clean_check_clause(text), text
    print("test_an_absence_carrying_a_finding_survives: PASS")


def test_a_terse_value_needs_its_field_for_risk_context():
    """fsi_metrics.mortgage_lender's whole value is "None disclosed." -- only
    readable as a clean check because of the field it sits in."""
    assert not cc._is_clean_check_clause("None disclosed.")
    assert cc._is_clean_check_clause("None disclosed.", risk_context=True)
    print("test_a_terse_value_needs_its_field_for_risk_context: PASS")


def test_strips_a_clean_check_tail_but_keeps_the_leading_finding():
    clause = (
        "An independent advocate's Title Report (18 April 2024) opines the land title is clear, "
        "marketable and free of encumbrances; no litigation was found against the promoter."
    )
    out = cc._strip_absence_tail(clause)
    assert "opines the land title is clear" in out, out
    assert "no litigation was found" not in out, out
    print("test_strips_a_clean_check_tail_but_keeps_the_leading_finding: PASS")


def test_the_leading_segment_is_never_dropped():
    """A bullet that opens mid-thought is worse than one that runs long."""
    clause = "No litigation was found; the Title Report confirms this."
    assert cc._strip_absence_tail(clause).startswith("No litigation was found")
    print("test_the_leading_segment_is_never_dropped: PASS")


# --- the scrubber ------------------------------------------------------------

def test_scrub_is_reversible_so_the_record_keeps_what_the_page_drops():
    """Section B: the scope of what was checked "stay[s] in the facts file;
    they do not reach the page". run_company_charter restores before
    persisting, so this must round-trip exactly."""
    facts = _facts()
    before = json.dumps(facts, sort_keys=True)
    changed = cc._scrub_clean_checks(facts)
    assert changed, "fixture must contain clean checks for this test's premise"
    assert json.dumps(facts, sort_keys=True) != before, "scrub changed nothing"
    cc._restore_clean_checks(facts, changed)
    assert json.dumps(facts, sort_keys=True) == before, "restore was not lossless"
    print(f"test_scrub_is_reversible_so_the_record_keeps_what_the_page_drops: PASS ({len(changed)} field(s))")


def test_scrub_is_idempotent():
    """Both variants render from one dict, so a second pass must be a no-op."""
    facts = _facts()
    cc._scrub_clean_checks(facts)
    once = json.dumps(facts, sort_keys=True)
    assert not cc._scrub_clean_checks(facts), "second pass reported further changes"
    assert json.dumps(facts, sort_keys=True) == once
    print("test_scrub_is_idempotent: PASS")


def test_scrub_leaves_gaps_and_score_inputs_alone():
    facts = _facts()
    before_gaps = list(facts.get("gaps", []))
    before_score = json.dumps(facts.get("developer_score"), sort_keys=True)
    cc._scrub_clean_checks(facts)
    assert facts["gaps"] == before_gaps, "gaps must never be scrubbed"
    assert json.dumps(facts.get("developer_score"), sort_keys=True) == before_score
    print("test_scrub_leaves_gaps_and_score_inputs_alone: PASS")


def test_scrubbing_never_moves_the_developer_score():
    """The score reads structured counts, not this prose. Pinned because a
    scrubber that silently re-graded a promoter would be a serious defect."""
    facts = _facts()
    flags = cc._classify_flags(facts)
    before = cc._compute_developer_score(facts, flags)
    cc._scrub_clean_checks(facts)
    after = cc._compute_developer_score(facts, cc._classify_flags(facts))
    assert before["composite"] == after["composite"], (before["composite"], after["composite"])
    assert before["grade"] == after["grade"], (before["grade"], after["grade"])
    print(f"test_scrubbing_never_moves_the_developer_score: PASS (composite {after['composite']}, grade {after['grade']})")


# --- the gate ----------------------------------------------------------------

def test_gate_flags_a_citation_attached_to_an_absence():
    assert cc._is_cited_absence("No litigation is disclosed against the promoter[3].")
    print("test_gate_flags_a_citation_attached_to_an_absence: PASS")


def test_gate_ignores_an_uncited_absence_and_a_cited_finding():
    assert not cc._is_cited_absence("No litigation is disclosed against the promoter.")
    assert not cc._is_cited_absence("The Title Report records a Notice of Lis Pendens dated 20 December 2017[4].")
    print("test_gate_ignores_an_uncited_absence_and_a_cited_finding: PASS")


def test_gate_exempts_a_numbered_gap_entry():
    """Gaps & Sources entries are open unknowns, and Section B keeps them in
    full -- including their citations."""
    assert not cc._is_cited_absence("Gap 4. No confirmation of the escrow account could be obtained[7].")
    print("test_gate_exempts_a_numbered_gap_entry: PASS")


# --- end to end --------------------------------------------------------------

def test_target_clean_checks_are_gone_and_findings_survive():
    """The worked example from the reshape spec, checked in both variants."""
    for variant in ("internal", "external"):
        blob = _all_text(_build(variant))
        for gone in (
            "No litigation is disclosed",
            "records for this project are empty",
            "No discrepancy found",
            "no insolvency process is recorded",
        ):
            assert gone not in blob, f"{variant}: clean check survived: {gone!r}"
        for kept in (
            "Lis Pendens",                          # a real finding
            "contractually permit the developer",   # a live right is a finding
            "opines the land title is clear",       # finding, tail stripped
            "standalone sanctioned-FSI certificate",  # a gap
        ):
            assert kept in blob, f"{variant}: content that must survive was deleted: {kept!r}"
    print("test_target_clean_checks_are_gone_and_findings_survive: PASS")


if __name__ == "__main__":
    test_detects_a_bare_clean_check()
    test_a_gap_is_not_a_clean_check()
    test_an_absence_carrying_a_finding_survives()
    test_a_terse_value_needs_its_field_for_risk_context()
    test_strips_a_clean_check_tail_but_keeps_the_leading_finding()
    test_the_leading_segment_is_never_dropped()
    test_scrub_is_reversible_so_the_record_keeps_what_the_page_drops()
    test_scrub_is_idempotent()
    test_scrub_leaves_gaps_and_score_inputs_alone()
    test_scrubbing_never_moves_the_developer_score()
    test_gate_flags_a_citation_attached_to_an_absence()
    test_gate_ignores_an_uncited_absence_and_a_cited_finding()
    test_gate_exempts_a_numbered_gap_entry()
    test_target_clean_checks_are_gone_and_findings_survive()
    print("\nAll tests passed.")
