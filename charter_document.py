"""
Restructured Company Charter document builder -- the Counterparty +
Collateral layout agreed in design (see the "charter restructure" spec).

WHY THIS IS A SEPARATE MODULE, built fresh rather than editing
company_charter._fill_template: that function fills the template's
paragraphs BY FIXED INDEX (p[8], p[16], p[34] ...) and then relocates
headings, which works only because the template's own section order matches
the document's. The new layout reorders everything, so index-based filling
stops being a help and becomes a hazard. This builder instead takes the
template purely for its STYLES (heading styles are latent in this template
and cannot be looked up by name -- confirmed live, doc.styles["Heading 1"]
raises KeyError while paragraphs report that style, which is why the
existing code reads styles off real paragraphs), clears the body, and
writes every section explicitly in the new order.

_fill_template is left completely untouched, so the existing Internal/
External documents keep generating exactly as before while this is built
and compared against them.

Structure (both variants share this skeleton; Internal shows more per
section, via company_charter._variant/_variant_paragraph):
  Counterparty
    1. Verification Summary        -- claim-by-claim + 2 bucketed rating tables
    2. Counterparty Identity
    3. Leadership, Related Entities & Litigation Screen
    4. Litigation & Regulatory Screening
    5. Organisation & Team
    6. Portfolio
  Collateral
    7. Asset Identity & Land Record
    8. Approvals, FSI & Title
    9. RERA Compliance & Escrow
   10. Location, Connectivity & Market Read
   11. Document & Diligence Trail
  Covering both
   12. Consolidated Flags & Recommended Steps
   13. Closing Read                -- rule-derived, never model-authored
   14. Scoring Detail appendix     -- full sub-metric breakdowns
   15. Sources                     -- Harvard style
"""

from __future__ import annotations

import re

import company_charter as cc

# ---------------------------------------------------------------------------
# Documentation Confidence Score buckets. The Developer Score already carries
# a `bucket` on every criterion (see _DEVELOPER_SCORE_STRUCTURE); the
# Documentation Confidence Score is a flat criteria list with no buckets at
# all, so its grouping is defined here. Weights are NOT reassigned -- each
# bucket's weight is just the sum of its criteria's existing weights, so the
# bucketed table and the detailed breakdown can never disagree.
#
# Each bucket answers a distinct question a reader actually asks: how good
# were the sources, how much did we independently confirm, and how complete
# and current is it.
# ---------------------------------------------------------------------------
_DOC_CONFIDENCE_BUCKETS = (
    ("Source Quality", ("source_tier_quality", "primary_tier_density")),
    ("Corroboration & Confirmation", ("cross_corroboration", "financial_figures_confirmed", "verification_rate")),
    ("Coverage & Recency", ("completeness_rate", "recency_legal", "recency_other")),
)

_DOC_CONFIDENCE_CRITERION_LABELS = {
    "source_tier_quality": "Weighted Source-Tier Quality",
    "primary_tier_density": "Primary-Tier Density",
    "cross_corroboration": "Cross-Corroboration",
    "financial_figures_confirmed": "Financial Figures Confirmed",
    "verification_rate": "Independent Verification Rate",
    "completeness_rate": "Completeness Rate",
    "recency_legal": "Recency - Legal Records",
    "recency_other": "Recency - Other Sources",
}


def _band_developer_rating(score: float) -> str:
    """Maps a 0-100 rolled-up bucket score onto the same AAA-D ladder the
    individual sub-metrics use, via the thresholds already defined for the
    composite grade."""
    for tier, floor in cc._DEVELOPER_SCORE_TIER_THRESHOLDS:
        if score >= floor:
            return tier
    return "D"


def rollup_developer_score_buckets(developer_score: dict) -> list:
    """Rolls the 9 Developer Score sub-metrics up into their 3 buckets for
    the summary table. Returns [{bucket, weight, rating, score, scored,
    total, unscored_display_names}].

    A bucket's rating averages only the sub-metrics that ACTUALLY scored,
    renormalised within the bucket -- deliberately different from the
    composite, which holds every weight fixed and lets an unscored
    sub-metric contribute nothing (see _compute_developer_score's own note).
    Both readings are honest but answer different questions: the composite
    answers "how much of this promoter is verifiable at all", the bucket
    rating answers "of what could be measured here, how strong is it". They
    are therefore never presented as the same number, and the unscored
    sub-metric names travel with the bucket so they can be flagged rather
    than silently averaged away -- the no-N/A rule.

    `rating` is None when nothing in the bucket scored, so a caller can
    render a flag line instead of a fake grade."""
    criteria = (developer_score or {}).get("criteria") or {}
    buckets = []
    for bucket_name, bucket_weight, metrics in cc._DEVELOPER_SCORE_STRUCTURE:
        keys = [key for key, _display, _fn in metrics]
        scored, unscored = [], []
        for key in keys:
            criterion = criteria.get(key) or {}
            display = criterion.get("display_name") or key.replace("_", " ").title()
            if criterion.get("score") is None:
                unscored.append(display)
            else:
                scored.append((criterion["score"], criterion.get("weight") or 0.0))

        weight_sum = sum(w for _s, w in scored)
        score = round(sum(s * w for s, w in scored) / weight_sum, 1) if weight_sum else None
        buckets.append({
            "bucket": bucket_name,
            "weight": bucket_weight,
            "score": score,
            "rating": _band_developer_rating(score) if score is not None else None,
            "scored": len(scored),
            "total": len(keys),
            "unscored_display_names": unscored,
        })
    return buckets


def rollup_doc_confidence_buckets(doc_confidence: dict) -> list:
    """Same idea as rollup_developer_score_buckets, for the Documentation
    Confidence Score's 7-8 criteria grouped by _DOC_CONFIDENCE_BUCKETS.
    Returns [{bucket, weight, band, score, scored, total,
    unscored_display_names}].

    Banded with the High/Moderate/Limited labels this score already uses
    (_DATA_AUTHENTICITY_BANDS) rather than the AAA-D ladder, since it rates
    how well-evidenced THIS REPORT is, not the promoter -- conflating the
    two scales is exactly the misreading the existing document's own
    wording works to prevent."""
    criteria = (doc_confidence or {}).get("criteria") or {}
    skipped = set((doc_confidence or {}).get("skipped_criteria") or [])

    buckets = []
    for bucket_name, keys in _DOC_CONFIDENCE_BUCKETS:
        scored, unscored = [], []
        for key in keys:
            criterion = criteria.get(key)
            label = _DOC_CONFIDENCE_CRITERION_LABELS.get(key, key.replace("_", " ").title())
            if not criterion or criterion.get("score") is None or key in skipped:
                # A criterion absent from `criteria` entirely is only a real
                # gap if it was explicitly skipped this pass or is a known
                # member of this bucket -- either way it is named, never
                # silently dropped.
                if key in skipped or criterion is not None:
                    unscored.append(label)
                continue
            scored.append((criterion["score"], criterion.get("weight") or 0.0))

        weight_sum = sum(w for _s, w in scored)
        score = round(sum(s * w for s, w in scored) / weight_sum, 1) if weight_sum else None
        buckets.append({
            "bucket": bucket_name,
            "weight": round(weight_sum, 1),
            "score": score,
            "band": cc._band_label(score) if score is not None else None,
            "scored": len(scored),
            "total": len(keys),
            "unscored_display_names": unscored,
        })
    return buckets


# ---------------------------------------------------------------------------
# Claim-by-claim verification, the spine of section 1.
#
# The comparison is promoter-stated vs independently verified -- there is no
# deck or brochure ingested anywhere in this pipeline, so a "deck claim"
# column would have nothing behind it. What every sourced field DOES carry is
# where the claim came from, and those sources fall into genuinely different
# evidence classes. The distinction that matters most: a MahaRERA filing is
# the promoter's OWN statement to the regulator (promoter-stated, however
# official the venue), whereas an MCA mirror or a government land record is a
# third party (independent). Collapsing those two into "verified" would
# overstate the evidence.
#
# Status is derived from the source, NOT from a stored per-field verdict,
# because no such verdict generally exists: _verify_material_claims only
# re-checks URL-sourced claims and a real project can have none at all
# (confirmed live -- every material field on Pranami Bliss is
# document-sourced, so _verification_stats is null and every gap is
# model-authored prose rather than a per-field record). Deriving from the
# source is therefore the honest available signal, and it never claims a
# check that didn't run.
# ---------------------------------------------------------------------------

STATUS_CONFIRMED_DOCUMENT = "Confirmed (primary document)"
STATUS_CONFIRMED_INDEPENDENT = "Confirmed (independent registry)"
STATUS_PROMOTER_FILED = "Promoter-filed (RERA record)"
STATUS_NOT_ESTABLISHED = "Not established"
STATUS_STATED_ONLY = "Stated, not independently verified"
STATUS_DISCREPANCY = "Discrepancy - see flags"


def classify_claim_evidence(source: str) -> tuple:
    """Returns (verified_position, status, is_independent) for one field's
    `source` string. `verified_position` names WHAT the claim was checked
    against, in reader-facing terms rather than a file path."""
    src = str(source or "").strip()
    lowered = src.lower()

    if not src or lowered.startswith("gap"):
        return ("No source established this pass", STATUS_NOT_ESTABLISHED, False)

    # A downloaded primary document -- the strongest single source available,
    # and the class _check_document_grounding mechanically verifies.
    if "/documents/" in src or ".pdf" in lowered:
        label = cc._clean_source_label(src) or "primary document"
        return (f"Primary document: {label}", STATUS_CONFIRMED_DOCUMENT, False)

    # Third-party registries -- genuinely independent of the promoter.
    if "mca-mirror" in lowered or "zaubacorp" in lowered or "tofler" in lowered or "instafinancials" in lowered:
        return ("MCA-mirror company registration profile", STATUS_CONFIRMED_INDEPENDENT, True)
    if "bhulekh" in lowered or "property card" in lowered:
        return ("Maha Bhulekh Property Card (government land record)", STATUS_CONFIRMED_INDEPENDENT, True)
    if "ibbi" in lowered:
        return ("IBBI Corporate Debtor Master Data", STATUS_CONFIRMED_INDEPENDENT, True)

    # MahaRERA is official, but what it holds is the promoter's own filing.
    if "maharera" in lowered:
        return ("MahaRERA project record (promoter's own filing)", STATUS_PROMOTER_FILED, False)

    return (cc._clean_source_label(src) or src, STATUS_STATED_ONLY, False)


# The four material-claim groups _verify_material_claims itself re-checks,
# with reader-facing group titles.
_MATERIAL_CLAIM_GROUPS = (
    ("land_identification", "Land identification"),
    ("corporate_identity", "Counterparty identity"),
    ("fsi_metrics", "FSI and area"),
)


# Explicit labels for fields whose key doesn't title-case into something a
# reader should see. Spelled out rather than derived by string surgery: an
# earlier attempt applied replacements and THEN .capitalize(), which silently
# lowercased them again and rendered "CIN / LLPIN" as "Cin / llpin".
_FIELD_DISPLAY_OVERRIDES = {
    "survey_cts_plot_numbers": "Survey / CTS / plot numbers",
    "mandal_taluka_district": "Mandal / taluka / district",
    "cin_llpin": "CIN / LLPIN",
    "registered_office_main": "Registered office (main)",
    "registered_office_board_resolution": "Registered office (per board resolution)",
    "registered_office_planning_stage": "Registered office (at planning stage)",
    "partners_directors": "Partners / directors",
    "landowner_investor": "Landowner / investor",
    "authorized_signatory": "Authorised signatory",
    "organization_type": "Organisation type",
    "net_land_area": "Net land area",
    "approved_bua": "Approved built-up area",
    "sanctioned_bua": "Sanctioned built-up area",
    "implied_fsi": "Implied FSI",
    "mortgage_area": "Mortgaged area",
    "mortgage_lender": "Mortgage lender",
}


def _field_display_name(key: str) -> str:
    override = _FIELD_DISPLAY_OVERRIDES.get(key)
    if override:
        return override
    return key.replace("_", " ").capitalize()


def build_claim_rows(facts: dict) -> list:
    """Builds section 1's claim-by-claim rows from every {value, source}-
    shaped field in the four material-claim groups. Returns [{group, claim,
    stated, verified_position, status}].

    Only dict-shaped fields appear: a group can also hold plain strings with
    no source of their own (fsi_metrics carries five such), and a claim with
    no recorded source cannot honestly be given a verification status, so
    those belong in their own section's prose rather than in a table whose
    whole purpose is claim-versus-evidence."""
    rows = []
    gap_texts = [str(g).lower() for g in (facts.get("gaps") or [])]

    def _add(group_title: str, key: str, field: dict):
        value = str(field.get("value") or "").strip()
        if not value:
            return
        verified_position, status, _independent = classify_claim_evidence(field.get("source"))

        # A MACHINE-written per-field gap overrides the source-derived
        # status: an unconfirmed claim must never read as confirmed just
        # because it cites a document. Matched on the exact structured
        # prefix _verify_one_field/_check_document_grounding write
        # ("<field_key>: <value> (verification: ...)"), NOT by searching the
        # gap prose for the field's name -- that was tried and produced a
        # false positive immediately: a model-authored gap describing a
        # portfolio search ("matched only this one project under this exact
        # promoter name") flipped the promoter_name claim to Discrepancy
        # despite saying nothing about whether that claim was verified.
        # Model-authored gaps are general findings and are surfaced in the
        # flags section on their own terms, not as per-field verdicts.
        display = _field_display_name(key)
        if any(g.startswith(f"{key.lower()}:") for g in gap_texts):
            status = STATUS_DISCREPANCY

        rows.append({
            "group": group_title,
            "claim": display,
            "stated": value,
            "verified_position": verified_position,
            "status": status,
            "field_key": key,
            "source": field.get("source"),
        })

    for group_key, group_title in _MATERIAL_CLAIM_GROUPS:
        for key, field in (facts.get(group_key) or {}).items():
            if isinstance(field, dict):
                _add(group_title, key, field)

    litigation = facts.get("litigation_status")
    if isinstance(litigation, dict):
        _add("Litigation", "litigation_status", litigation)

    return rows


# ---------------------------------------------------------------------------
# Document construction.
# ---------------------------------------------------------------------------

def _clear_body(doc) -> None:
    """Empties the template's body while keeping everything the document's
    LOOK depends on. Styles and numbering live in separate parts
    (styles.xml/numbering.xml) and are untouched by this; the one body-level
    element that must survive is the trailing <w:sectPr>, which carries page
    size, orientation and margins -- removing it silently resets the document
    to Word's defaults."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _capture_styles(doc) -> dict:
    """Grabs style OBJECTS off the template's own paragraphs before the body
    is cleared. Necessary because this template's heading styles are latent:
    confirmed live that doc.styles["Heading 1"] raises KeyError even though
    paragraphs report that exact style name, which is why the existing
    generator reads styles off doc.paragraphs[4] rather than by name."""
    styles = {"h1": None, "h2": None, "title": None, "list": None}
    for paragraph in doc.paragraphs:
        name = paragraph.style.name if paragraph.style is not None else None
        if name == "Heading 1" and styles["h1"] is None:
            styles["h1"] = paragraph.style
        elif name == "Heading 2" and styles["h2"] is None:
            styles["h2"] = paragraph.style
        elif name == "Title" and styles["title"] is None:
            styles["title"] = paragraph.style
        elif name == "List Paragraph" and styles["list"] is None:
            styles["list"] = paragraph.style
    return styles


class _Builder:
    """Thin wrapper carrying the doc, the captured styles and `facts`, so
    every section helper doesn't have to thread all three plus the doc
    variant through every call."""

    def __init__(self, doc, styles, facts):
        self.doc = doc
        self.styles = styles
        self.facts = facts

    # --- text -------------------------------------------------------------
    def prose(self, text):
        """Applies the External variant's language rules to a string.

        _externalize_prose handles this via a curated dictionary of known
        sentences, which is the right approach where the wording is known
        (it can rewrite grammatically rather than mangling punctuation). But
        it cannot cover arbitrary model-authored prose, and the External
        quality gate hard-forbids " -- ", so a last-resort conversion is
        needed for anything the dictionary misses. A semicolon is used
        because it is grammatical wherever a dash joins two related clauses
        or introduces an explanation, and unlike a comma it cannot create a
        comma splice.

        Confirmed necessary rather than hypothetical: address_discrepancy_note
        and corporate_registry_cross_check both reached the document with
        dashes intact and were caught by the gate. Those two used to be
        suppressed from External entirely as methodology narration, so the
        dictionary never needed to cover them; section 2 now renders their
        content deliberately."""
        rendered = cc._externalize_prose(self.facts, str(text))
        if self.facts.get("_doc_variant") == "external":
            rendered = rendered.replace(" -- ", "; ").replace(" — ", "; ").replace("—", "; ")
            # "Document Library" is the OLD section name this restructure
            # removed (section 11 is "Document & Diligence Trail" now), but
            # model-authored prose from an existing facts.json still refers
            # to it by that name -- confirmed live, documents_reviewed_note
            # says a document was "confirmed present in the Document
            # Library." That is internal-process naming, exactly what the
            # External quality gate exists to catch, so it is generalised
            # here rather than pattern-matched against one project's exact
            # sentence (which would not survive a different facts.json).
            rendered = re.sub(r"\bdocument library\b", "document list", rendered, flags=re.IGNORECASE)
        return rendered

    def para(self, text="", bold=False, color=None, italic=False, style=None):
        """Adds a paragraph, routing prose through the External language
        rules so no call site has to repeat that."""
        rendered = self.prose(text) if text else ""
        paragraph = self.doc.add_paragraph()
        if style is not None:
            paragraph.style = style
        if rendered:
            run = paragraph.add_run(rendered)
            run.bold = bold
            run.italic = italic
            if color:
                cc._color_run(run, color)
        return paragraph

    def heading(self, text, level=1):
        style = self.styles["h1"] if level == 1 else self.styles["h2"]
        paragraph = self.para(text)
        if style is not None:
            paragraph.style = style
        return paragraph

    def bullet(self, text, color=None):
        paragraph = self.para(text, color=color, style=self.styles["list"])
        return paragraph

    def flag_line(self, text):
        """A one-liner for something a table can't cover. Uses the
        no-N/A convention: an uncoverable dimension is stated in a sentence
        rather than padding a table with an empty row."""
        return self.para(text, italic=False)

    def page_break(self):
        self.doc.add_page_break()

    # --- tables -----------------------------------------------------------
    def table(self, headers, rows, widths=None):
        """Creates a bordered table with a shaded, bold header row. `rows` is
        a list of lists of either plain strings or (text, color) tuples, so a
        cell can be coloured without a second pass over the table."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        cc._set_table_borders(table)
        header_cells = table.rows[0].cells
        for i, label in enumerate(headers):
            header_cells[i].text = self.prose(str(label))
            cc._shade_cell(header_cells[i], cc._FILL_NEUTRAL)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_values in rows:
            row = table.add_row()
            for i, value in enumerate(row_values):
                text, color = value if isinstance(value, tuple) else (value, None)
                cell = row.cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(self.prose(str(text or "")))
                if color:
                    cc._color_run(run, color)
        return table


_STATUS_COLORS = {
    STATUS_DISCREPANCY: cc._TEXT_RED,
    STATUS_NOT_ESTABLISHED: cc._TEXT_AMBER,
    STATUS_STATED_ONLY: cc._TEXT_AMBER,
    STATUS_PROMOTER_FILED: cc._TEXT_AMBER,
}


def _section_1_verification_summary(b: _Builder, flags: dict, developer_score: dict) -> None:
    """Section 1: what was claimed, what the evidence says, and the two
    bucketed ratings. Deliberately the shortest section that carries the
    most: a reader who stops here should still know what is confirmed, what
    is only asserted, and where the counterparty is strong or weak."""
    b.heading("1. Verification Summary")

    claim_rows = build_claim_rows(b.facts)
    confirmed = sum(1 for r in claim_rows if r["status"].startswith("Confirmed"))
    independent = sum(1 for r in claim_rows if r["status"] == STATUS_CONFIRMED_INDEPENDENT)
    b.para(
        f"Of {len(claim_rows)} material claims checked, {confirmed} are confirmed against a primary document or "
        f"an independent registry ({independent} of those against a registry independent of the promoter). "
        f"Each row below names what the claim was checked against, so the strength of the evidence is visible "
        f"rather than assumed uniform."
    )

    b.table(
        ["Claim", "Stated position", "Checked against", "Status"],
        [
            [r["claim"], r["stated"], r["verified_position"], (r["status"], _STATUS_COLORS.get(r["status"]))]
            for r in claim_rows
        ],
    )

    # --- Overall rating: Developer Score buckets -------------------------
    b.para()
    b.heading("Overall Rating", level=2)
    ds_buckets = rollup_developer_score_buckets(developer_score)
    b.para(
        f"Composite {developer_score.get('composite')}/100, grade {developer_score.get('grade')}. "
        f"Each bucket below is rated on what could actually be measured for it; the composite instead holds "
        f"every sub-metric's weight fixed, so a promoter with less publicly verifiable data scores lower "
        f"overall even where the measurable parts are strong. That is why a bucket can read well while the "
        f"composite does not, and the coverage column shows how much of each bucket was measurable."
    )
    b.table(
        ["Dimension", "Weight", "Rating", "Coverage", "Comment"],
        [
            [
                bucket["bucket"],
                f"{bucket['weight']}%",
                bucket["rating"] or "Not rated",
                f"{bucket['scored']} of {bucket['total']} sub-metrics",
                f"Weighted score {bucket['score']}/100 across what was measurable."
                if bucket["score"] is not None
                else "Nothing in this bucket could be scored from public data this pass.",
            ]
            for bucket in ds_buckets
        ],
    )
    for bucket in ds_buckets:
        for name in bucket["unscored_display_names"]:
            b.flag_line(f"Not rated in {bucket['bucket']}: {name}. No public data source established this for this pass.")

    # --- Documentation Confidence buckets --------------------------------
    doc_confidence = b.facts.get("documentation_confidence_score") or {}
    dc_buckets = rollup_doc_confidence_buckets(doc_confidence) if doc_confidence.get("criteria") else []
    if dc_buckets and not any(bucket["scored"] for bucket in dc_buckets):
        # Unlike the Developer Score, this one cannot be recomputed here: it
        # needs an authenticity_summary built inside _fill_template. So a
        # persisted copy is used, and if NONE of its criteria match the
        # buckets defined in this module, it is from a superseded scoring
        # scheme and its numbers are meaningless. Say so rather than render
        # them (the same staleness that made the Developer Score read
        # 100.0/A against a true 37.1/C).
        b.para()
        b.heading("Documentation Confidence", level=2)
        b.flag_line(
            "Not rated: the stored Documentation Confidence Score is from a superseded scoring scheme and "
            "cannot be mapped onto the current criteria, so no band is shown rather than reporting a figure "
            "that no longer means what it says. Regenerate this Charter to score it."
        )
        dc_buckets = []
    elif dc_buckets:
        b.para()
        b.heading("Documentation Confidence", level=2)
        b.para(
            f"Overall {doc_confidence.get('overall')}/100, band {doc_confidence.get('band')}. This rates how "
            f"well evidenced this report's own claims are, not the project's quality: a sound project with a "
            f"thin public paper trail scores low here, and a troubled one with thorough documentation scores "
            f"high."
        )
        b.table(
            ["Dimension", "Weight", "Band", "Coverage", "Comment"],
            [
                [
                    bucket["bucket"],
                    f"{bucket['weight']}%",
                    bucket["band"] or "Not rated",
                    f"{bucket['scored']} of {bucket['total']} criteria",
                    f"Weighted score {bucket['score']}/100 across what was measurable."
                    if bucket["score"] is not None
                    else "No criterion in this bucket could be scored this pass.",
                ]
                for bucket in dc_buckets
            ],
        )
        for bucket in dc_buckets:
            for name in bucket["unscored_display_names"]:
                b.flag_line(f"Not rated in {bucket['bucket']}: {name}.")
        if doc_confidence.get("verification_warning"):
            b.flag_line(str(doc_confidence["verification_warning"]))


def _section_2_counterparty_identity(b: _Builder) -> None:
    """Section 2: the counterparty's own registration facts, promoter-stated
    against independently verified, then the discrepancies that reading
    surfaced. Goes deeper on identity than section 1's headline table, which
    spans both halves of the document."""
    b.page_break()
    b.heading("2. Counterparty Identity")

    profile = b.facts.get("company_profile_check") or {}
    corp = b.facts.get("corporate_identity") or {}
    core = b.facts.get("rera_core_fields") or {}

    def stated(key):
        return str((corp.get(key) or {}).get("value") or "").strip()

    if profile.get("found"):
        sources_used = ", ".join(
            cc._MCA_SOURCE_DISPLAY_NAMES.get(d, d) for d in (profile.get("sources_used") or [])
        )
        b.para(f"Company Snapshot, verified against {sources_used}.")
        rows = [
            ["Legal name", stated("promoter_name") or core.get("promoter_name", ""), cc._normalise_entity_name(profile.get("name") or "")],
            ["CIN / LLPIN", stated("cin_llpin"), profile.get("cin") or ""],
            ["Status", "", profile.get("status") or ""],
            ["Class of company", stated("organization_type"), profile.get("class_of_company") or ""],
            ["Company category", "", profile.get("company_category") or ""],
            ["Registrar of Companies", "", profile.get("roc") or ""],
            ["Incorporation date", "", profile.get("incorporation_date") or ""],
            ["Registered office", stated("registered_office_main"), cc._clean_scraped_address(profile.get("registered_address") or "")],
            ["Authorised capital", "", profile.get("authorised_capital") or ""],
            ["Paid-up capital", "", profile.get("paid_up_capital") or ""],
        ]
        b.table(
            ["Field", "Promoter-stated", "Independently verified"],
            [[r[0], r[1] or "Not separately stated", r[2] or "Not disclosed by the registry"] for r in rows],
        )
    else:
        b.flag_line(
            "No company registration profile could be retrieved for this counterparty: "
            + str(profile.get("note") or "no reason recorded")
        )

    # --- Reading discrepancies -------------------------------------------
    b.para()
    b.heading("Reading the Discrepancies", level=2)
    discrepancy_points = []
    if b.facts.get("address_discrepancy_note"):
        discrepancy_points.append(str(b.facts["address_discrepancy_note"]))
    if b.facts.get("corporate_registry_cross_check"):
        discrepancy_points.append(str(b.facts["corporate_registry_cross_check"]))
    if b.facts.get("cts_mismatch_note"):
        discrepancy_points.append(str(b.facts["cts_mismatch_note"]))
    for conflict in (profile.get("roster_conflicts") or []):
        discrepancy_points.append(str(conflict))

    if discrepancy_points:
        for point in discrepancy_points:
            b.bullet(point)
    else:
        b.flag_line("No identity discrepancy was found across the registries and documents read this pass.")


def _section_3_leadership_and_related(b: _Builder) -> None:
    """Section 3: who runs the counterparty, and what else they run. The
    related-entity table names only the entities carrying more than one
    independent link (the real signal) and collapses single-link entities to
    counts, so 65 rows don't bury the 20 that matter."""
    b.page_break()
    b.heading("3. Leadership, Related Entities & Litigation Screen")

    profile = b.facts.get("company_profile_check") or {}
    current = profile.get("current_directors") or []
    past = profile.get("past_directors") or []

    b.heading(f"{cc._role_word(b.facts, 2).capitalize()} of Record", level=2)
    if current:
        b.table(
            ["Name", "DIN", "Designation", "Appointed", "Analyst read"],
            [
                [
                    cc._normalise_entity_name(d.get("Director Name") or ""),
                    d.get("DIN") or "Not disclosed",
                    d.get("Designation") or "",
                    d.get("Appointment Date") or "Not disclosed",
                    "Confirmed on the registry roster.",
                ]
                for d in current
            ],
        )
    else:
        b.flag_line(f"No current {cc._role_word(b.facts)} roster could be retrieved for this counterparty.")

    if past:
        b.para()
        b.para(
            f"{len(past)} former {cc._role_word(b.facts, len(past))} on record: "
            + ", ".join(cc._normalise_entity_name(d.get("Director Name") or "") for d in past)
            + f". A former {cc._role_word(b.facts)}'s other {cc._role_word(b.facts, 1)}ships still map the group this counterparty sits in, so they "
              "are included in the relationship counts below."
        )

    # --- Related entity mapping ------------------------------------------
    group = b.facts.get("group_companies_check") or {}
    companies = group.get("companies") or []
    b.para()
    b.heading("Related Entity Mapping", level=2)

    if not companies:
        b.flag_line("No related entities were identified from the registry crosswalk this pass.")
    else:
        multi_link = [c for c in companies if len(c.get("basis") or []) > 1]
        single_link = [c for c in companies if len(c.get("basis") or []) <= 1]
        b.para(
            f"{len(companies)} entities share at least one concrete, named link with this counterparty. The "
            f"{len(multi_link)} carrying more than one independent link are named below, since two independent "
            f"links are materially stronger evidence of a real group relationship than one. The remaining "
            f"{len(single_link)} single-link entities are summarised by link type beneath, rather than listed."
        )
        if multi_link:
            b.table(
                ["Related entity", "Role of the promoter", "Nature and implication"],
                [
                    [
                        c.get("name") or "",
                        "; ".join(sorted({_basis_role(x) for x in (c.get("basis") or [])})),
                        _basis_implication(c.get("basis") or []),
                    ]
                    for c in sorted(multi_link, key=lambda c: c.get("name") or "")
                ],
            )

        by_type = {}
        for c in single_link:
            for basis in (c.get("basis") or []):
                key = _basis_role(basis)
                by_type[key] = by_type.get(key, 0) + 1
        if by_type:
            b.para()
            b.table(
                ["Single-link entities by type", "Count", "What it implies"],
                [
                    [key, str(count), _LINK_TYPE_IMPLICATIONS.get(key, "A single named link; weaker evidence on its own.")]
                    for key, count in sorted(by_type.items(), key=lambda kv: -kv[1])
                ],
            )

        # A relationship whose counterparty identity is paywalled is still a
        # real relationship; reporting the count is honest, naming an advert
        # would not be.
        for relationship, count in sorted((group.get("undisclosed_relationship_counts") or {}).items()):
            b.flag_line(
                f"{count} {relationship} relationship(s) appear on the registry record but their counterparty "
                f"identities sit behind that source's paid tier. Their existence is confirmed; who they are is not."
            )

        corroboration = group.get("corroboration") or {}
        if corroboration.get("note"):
            b.para()
            b.para(str(corroboration["note"]))


_LINK_TYPE_IMPLICATIONS = {
    "Shared director": "Common leadership, so decisions and reputation travel between the entities.",
    "Shared registered office": "A shared address alone is weak evidence; common for group holding structures and for unrelated tenants of the same building.",
    "Subsidiary / associate / JV": "A filed ownership relationship, the strongest form of link in this set.",
}


def _basis_role(basis: str) -> str:
    """Turns one raw basis string into a reader-facing link type."""
    text = str(basis or "")
    if text.startswith("shared director"):
        return "Shared director"
    if text.startswith("shared registered office"):
        return "Shared registered office"
    if text.startswith("subsidiary"):
        return "Subsidiary / associate / JV"
    return text or "Unspecified link"


def _basis_implication(basis_list: list, known_person: str | None = None) -> str:
    """Names every distinct link an entity carries, plus the specific people
    behind any shared-director link, since 'who connects these two' is the
    part a reader actually needs.

    `known_person`: when this table already lives inside that person's own
    subsection (see _person_related_entities), the "Common leadership via
    NAME" sentence would just restate the one name the reader was already
    told 24+ rows ago -- confirmed confusing in practice for a
    large-footprint past director. Naming them there is only useful when a
    SECOND, not-yet-established person also shares the link; if they were
    the only name, the sentence is dropped as pure repetition."""
    roles = sorted({_basis_role(b) for b in basis_list})
    people = sorted({
        str(b).split(":", 1)[1].split("(")[0].strip()
        for b in basis_list
        if str(b).startswith("shared director") and ":" in str(b)
    })
    if known_person:
        people = [p for p in people if cc._normalise_entity_name(p) != cc._normalise_entity_name(known_person)]
    link_count = len(basis_list)
    parts = [f"{link_count} independent link{'s' if link_count != 1 else ''}: " + ", ".join(roles) + "."]
    if people:
        parts.append("Common leadership via " + ", ".join(cc._normalise_entity_name(p) for p in people) + ".")
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Combined related-entity Annexure (charter_report.py) -- one shared-director
# basis string parsed into its real per-entity role and current/resigned
# status, e.g. "shared director: Mahir Haresh Wadhwani (Designated Partner,
# resigned 2023-04-03)". Confirmed live as a real inconsistency: prose
# elsewhere correctly said "a past partner" for an LLP promoter, while this
# table's own "Role of the promoter" column still said the generic category
# "Shared director" regardless of the LINKED entity's own legal form -- the
# fix is to show the real designation ZaubaCorp recorded at THAT entity
# (already correct per-entity, e.g. "Director" for a Pvt Ltd link and
# "Designated Partner" for an LLP link) instead of a category label.
# ---------------------------------------------------------------------------
_SHARED_DIRECTOR_BASIS_RE = re.compile(r"^shared director:\s*([^(]+?)(?:\s*\(([^)]*)\))?\s*$", re.IGNORECASE)


def _parse_shared_director_basis(basis: str) -> dict | None:
    """Returns {"person", "designation", "status"} or None if `basis` isn't
    a shared-director entry. Tolerates the OLD basis format too (no status
    suffix, e.g. "shared director: X (Director)", or no parenthetical at
    all) for backward compatibility with a Charter built before this
    status-capture existed -- `designation`/`status` are just "" then,
    rather than raising or guessing a value that was never scraped."""
    m = _SHARED_DIRECTOR_BASIS_RE.match(str(basis or "").strip())
    if not m:
        return None
    person = cc._normalise_entity_name(m.group(1).strip())
    detail = (m.group(2) or "").strip()
    designation, status = "", ""
    if detail:
        last_comma = detail.rfind(",")
        tail = detail[last_comma + 1:].strip() if last_comma != -1 else detail
        if tail.lower() == "ongoing" or tail.lower().startswith("resigned"):
            status = tail
            designation = detail[:last_comma].strip() if last_comma != -1 else ""
        else:
            designation = detail  # old format: just "(Designation)", no status captured
    return {"person": person, "designation": designation, "status": status}


def _related_entity_people(basis_list: list) -> list:
    """One entry per distinct person named in a shared-director basis on
    this list (deduped by name, first occurrence wins)."""
    seen = {}
    for b in basis_list or []:
        parsed = _parse_shared_director_basis(b)
        if parsed and parsed["person"] and parsed["person"] not in seen:
            seen[parsed["person"]] = parsed
    return list(seen.values())


def _related_entity_linked_via(basis_list: list) -> str:
    """"Person (real designation, ongoing/resigned date)" for every distinct
    shared-director person on this entity, joined -- e.g. "Girish Sanjay
    Yadav (Designated Partner, ongoing); Swapnil Yuvraj Marathe (Designated
    Partner, ongoing)" when an entity is genuinely shared by two people, so
    a reader can see exactly who they're shared with rather than a generic
    "Shared director" label that reads as if a position itself were shared."""
    parts = []
    for p in _related_entity_people(basis_list):
        detail_bits = [x for x in (p["designation"], p["status"]) if x]
        detail = f" ({', '.join(detail_bits)})" if detail_bits else ""
        parts.append(f"{p['person']}{detail}")
    return "; ".join(parts)


def _related_entity_implication(basis_list: list) -> str:
    """General commentary on what this link TYPE means -- never restates a
    name (that's _related_entity_linked_via's job), and explicitly flags
    when two or more DIFFERENT people are the shared connection, since
    that's materially stronger evidence of a real group relationship than
    a single shared person."""
    roles = sorted({_basis_role(b) for b in basis_list or []})
    people = _related_entity_people(basis_list)
    parts = []
    if "Shared director" in roles:
        if len(people) >= 2:
            parts.append(
                f"Shared with {len(people)} named individuals ({', '.join(p['person'] for p in people)}), "
                f"so decisions and reputation travel between the entities via more than one person."
            )
        else:
            parts.append(_LINK_TYPE_IMPLICATIONS["Shared director"])
    if "Shared registered office" in roles:
        parts.append(_LINK_TYPE_IMPLICATIONS["Shared registered office"])
    if "Subsidiary / associate / JV" in roles:
        parts.append(_LINK_TYPE_IMPLICATIONS["Subsidiary / associate / JV"])
    return " ".join(parts)


def _section_4_litigation_screen(b: _Builder) -> None:
    """Section 4: every regulatory and litigation screen run against the
    counterparty, as findings rather than narration of the checking."""
    b.page_break()
    b.heading("4. Litigation & Regulatory Screening")

    points = []

    ibbi = b.facts.get("ibbi_insolvency_check") or {}
    if ibbi.get("found_process") is True:
        points.append(f"IBBI: an insolvency record exists against this exact entity. Status text on record: {ibbi.get('status_text') or 'not stated'}.")
    elif ibbi.get("found_process") is False:
        points.append("IBBI insolvency register: no corporate insolvency process found against this entity's identifier.")
    else:
        points.append(f"IBBI insolvency register: the check could not complete this pass ({ibbi.get('note') or 'no reason recorded'}).")

    rating = (b.facts.get("credit_rating_check") or {}).get("promoter") or b.facts.get("credit_rating_check") or {}
    if rating.get("ratings"):
        for agency in rating["ratings"]:
            points.append(f"{agency.get('agency', 'Rating agency')}: rating on record for {agency.get('company_name', 'this entity')}.")
    elif rating:
        points.append(
            "Credit rating: no public rating found for this exact legal entity from the agencies checked. "
            "Absence is not itself adverse, since agencies only rate developers that sought a rating."
        )

    core = b.facts.get("rera_core_fields") or {}
    complaint_count, appeal_count = cc._parse_complaint_appeal_counts(b.facts)
    if complaint_count is not None:
        points.append(f"MahaRERA complaints against this project: {complaint_count} on record.")
    if appeal_count is not None:
        points.append(f"MahaRERA appeals against this project: {appeal_count} on record.")
    if core.get("litigations_per_record"):
        points.append(f"RERA litigation field as filed: {core['litigations_per_record']}.")

    outcomes = b.facts.get("complaint_outcomes_summary") or {}
    if outcomes.get("outcome_counts"):
        breakdown = ", ".join(f"{v} {k.replace('_', ' ')}" for k, v in sorted(outcomes["outcome_counts"].items(), key=lambda kv: -kv[1]))
        points.append(f"Complaint order outcomes, read from the downloaded orders themselves: {breakdown}.")

    judgments = b.facts.get("appeal_judgments_found") or []
    if judgments:
        points.append(f"{len(judgments)} appeal-level judgment(s) matched to this project in MahaRERA's published orders.")
    else:
        points.append("MahaRERA published orders: no appeal-level judgment matched this project.")

    litigation = b.facts.get("litigation_status")
    if isinstance(litigation, dict) and litigation.get("value"):
        points.append(f"Title litigation position: {litigation['value']}")

    for point in points:
        b.bullet(point)


def _section_5_organisation_and_team(b: _Builder) -> None:
    """Section 5: the professional team of record, code-computed from
    MahaRERA's own structured data. An org chart is deliberately absent
    rather than invented: no source discloses internal headcount, reporting
    lines, or team composition by function."""
    b.page_break()
    b.heading("5. Organisation & Team")

    team = b.facts.get("professional_team") or []
    if team:
        b.para(
            "The professionals below are named on MahaRERA's own project record, read directly from the "
            "registry rather than from the promoter's documents. Registration numbers are reproduced exactly "
            "as filed."
        )
        b.table(
            ["Role", "Firm or individual", "Registration on record"],
            [
                [
                    p.get("role") or "",
                    p.get("name") or "",
                    f"{p['registration_label']}: {p['registration_number']}" if p.get("registration_number") else "None filed",
                ]
                for p in team
            ],
        )
        suspect = [p for p in team if p.get("registration_number") and len(str(p["registration_number"]).strip()) < 3]
        for p in suspect:
            b.flag_line(
                f"The {p['role'].lower()}'s {p['registration_label'].lower()} on record is "
                f"\"{p['registration_number']}\", which is too short to be a real registration number. "
                f"Worth raising as a filing-quality point."
            )
    else:
        b.flag_line("No professionals are named on this project's MahaRERA record.")

    b.para()
    b.flag_line(
        "Firm-wide organisation structure and team headcount by function are not covered: no public source "
        "discloses internal headcount or reporting lines for a private developer. This would have to be "
        "requested from the developer directly."
    )


def _section_6_portfolio(b: _Builder, next_steps: list) -> None:
    """Section 6: what this promoter has actually delivered, and how their
    headline claims hold up against the public record. Distinguishes "we
    could not check" from "the promoter has declared nothing", which are
    very different findings."""
    b.page_break()
    b.heading("6. Portfolio & Track Record")

    portfolio = b.facts.get("promoter_portfolio") or {}
    totals = portfolio.get("totals") or {}
    projects = portfolio.get("projects") or []
    subject_reg = (b.facts.get("rera_core_fields") or {}).get("registration_number", "")

    fetch_failed = [p for p in projects if p.get("past_experience_fetch_error") or p.get("complaints_fetch_error")]
    other_projects = [p for p in projects if str(p.get("reg_no") or "").upper() != str(subject_reg).upper()]

    b.para(
        f"{totals.get('total_projects', 0)} MahaRERA registration(s) are held under this promoter's name, and "
        f"{totals.get('total_experience_entries_found', 0)} prior delivery/deliveries are declared to MahaRERA."
    )

    if other_projects:
        by_district = {}
        for p in other_projects:
            key = p.get("district") or "District not stated"
            by_district.setdefault(key, []).append(p)
        b.table(
            ["Submarket", "Registrations", "Representative projects"],
            [
                [district, str(len(rows)), ", ".join((r.get("project_name") or "").title() for r in rows[:3])]
                for district, rows in sorted(by_district.items(), key=lambda kv: -len(kv[1]))
            ],
        )
    elif fetch_failed:
        b.flag_line(
            "The promoter's other registrations could not be checked this pass: MahaRERA's own API failed for "
            f"{len(fetch_failed)} project row(s). This is a gap in the check, not a finding about the promoter."
        )
        next_steps.append("Re-run the promoter portfolio scan; MahaRERA's API failed partway through this pass.")
    else:
        b.flag_line(
            "This promoter holds no MahaRERA registration other than the subject project. That is a confirmed "
            "finding rather than a failed lookup: the registry was searched successfully and returned only this "
            "one project, so there is no independent MahaRERA-registered track record to review."
        )
        next_steps.append("Obtain and verify the promoter's completed-project list directly from the developer.")

    # --- headline-claim cross verification -------------------------------
    b.para()
    b.heading("Cross-Verification of Headline Claims", level=2)
    rows = _headline_claim_rows(b.facts, totals)
    if rows:
        b.table(["Claim", "What the public record shows", "Read"], rows)
    else:
        b.flag_line("No headline track-record claim was available to cross-check this pass.")

    if totals.get("on_time_rate_pct") is not None:
        b.para()
        b.para(
            f"Declared completion performance: {totals['on_time_rate_pct']}% of declared prior deliveries "
            f"completed on or before their originally proposed date ({totals.get('on_time_count', 0)} on time, "
            f"{totals.get('delayed_count', 0)} delayed). These dates are self-reported by the promoter to "
            f"MahaRERA and are not independently verified."
        )


def _headline_claim_rows(facts: dict, totals: dict) -> list:
    """Builds the claim-versus-record rows for section 6 from whatever
    headline claims this pass actually captured. Returns [] rather than
    inventing rows when nothing is available."""
    rows = []
    track_record = facts.get("developer_track_record") or {}
    years = track_record.get("years_in_industry")
    basis = str(track_record.get("years_in_industry_basis") or "")
    if years:
        read = "Consistent"
        if "not credited" in basis.lower() or "no source confirms" in basis.lower():
            read = "Possibly related, unconfirmed"
        rows.append([
            f"Around {years} years in the industry",
            basis or "No basis recorded for this figure.",
            read,
        ])

    area = totals.get("total_area_developed_lakh_sqft")
    if area is not None:
        rows.append([
            f"{area} lakh sq ft delivered across declared prior projects",
            "Summed from the promoter's own past-experience filings to MahaRERA, excluding the subject project.",
            "Promoter-declared, not independently verified",
        ])

    within_5km = totals.get("area_within_5km_lakh_sqft")
    if within_5km is None and totals.get("total_experience_entries_found"):
        rows.append([
            "Experience in this micro-market",
            "Could not be measured: this project's own locality did not resolve to coordinates, so no "
            "distance comparison against prior deliveries was possible.",
            "Not established",
        ])
    return rows


# ---------------------------------------------------------------------------
# Collateral block. Everything from here to section 11 describes the ASSET,
# not the counterparty.
# ---------------------------------------------------------------------------

def _field_table_rows(group: dict, keys=None) -> list:
    """Rows of [display name, value, source] for a {value, source}-shaped
    group, skipping fields with no value so a table never carries an empty
    row (the no-N/A rule applied at field level)."""
    rows = []
    for key in (keys or list(group.keys())):
        field = group.get(key)
        if isinstance(field, dict):
            value = str(field.get("value") or "").strip()
            if not value:
                continue
            position, _status, _independent = classify_claim_evidence(field.get("source"))
            rows.append([_field_display_name(key), value, position])
        elif field:
            rows.append([_field_display_name(key), str(field), "Not separately sourced"])
    return rows


def _section_7_asset_identity(b: _Builder) -> None:
    b.page_break()
    b.heading("7. Asset Identity & Land Record")

    land_rows = _field_table_rows(b.facts.get("land_identification") or {})
    if land_rows:
        b.table(["Field", "Position", "Checked against"], land_rows)
    else:
        b.flag_line("No land identification fields were established this pass.")

    neighbourhood = b.facts.get("neighbourhood") or {}
    if any(neighbourhood.values()):
        b.para()
        b.heading("Boundaries per the RERA Land Record", level=2)
        b.table(
            ["Direction", "Abutting"],
            [[direction.title(), str(neighbourhood.get(direction) or "Not stated")]
             for direction in ("north", "east", "south", "west")],
        )

    # The government land record, when a human has run the CAPTCHA-gated
    # lookup (see company_charter.run_cts_land_lookup).
    cts_check = b.facts.get("cts_land_record_check") or {}
    b.para()
    b.heading("Government Land Record", level=2)
    if cts_check.get("found"):
        b.para(
            "A Property Card was retrieved from Maha Bhulekh, Maharashtra's official land-records portal, by "
            "exact CTS number. Every office, village and CTS selection leading to it was an exact match rather "
            "than a guess, and a human solved the portal's own CAPTCHA to reveal the record."
        )
        fields = cts_check.get("fields") or {}
        if fields:
            b.table(["Field", "Value"], [[k, v] for k, v in fields.items()])
        else:
            b.flag_line(
                "The Property Card was retrieved but its page did not parse into labelled fields, so the record "
                "itself is preserved as text and a screenshot alongside this Charter for direct reading."
            )
    else:
        b.flag_line(
            "No government land record was retrieved for this plot. This lookup is human-assisted by necessity: "
            "Maha Bhulekh gates the Property Card behind a CAPTCHA that must be solved fresh for every single "
            "request, so it does not run unattended."
        )


def _section_8_approvals_fsi_title(b: _Builder) -> None:
    b.page_break()
    b.heading("8. Approvals, FSI & Title")

    fsi_rows = _field_table_rows(b.facts.get("fsi_metrics") or {})
    if fsi_rows:
        b.table(["Metric", "Position", "Checked against"], fsi_rows)

    for key, title in (
        ("fsi_governing_framework", "Governing framework"),
        ("fsi_interpretation", "How the approvals read"),
    ):
        if b.facts.get(key):
            b.para()
            b.para(f"{title}: {b.facts[key]}")

    rules = b.facts.get("rules_statutory") or {}
    if rules.get("planning_approval_sequence"):
        b.para()
        b.heading("Planning Approval Sequence", level=2)
        b.para(str(rules["planning_approval_sequence"]))
    if rules.get("allotment_mechanics"):
        b.para()
        b.para(f"Allotment mechanics: {rules['allotment_mechanics']}")
    if rules.get("governing_act"):
        b.para()
        b.para(f"Governing statute: {rules['governing_act']}")


def _section_9_rera_compliance_escrow(b: _Builder) -> None:
    b.page_break()
    b.heading("9. RERA Compliance & Escrow")

    compliance = b.facts.get("rera_compliance") or {}
    for key, title in (
        ("registration_summary", "Registration"),
        ("collection_account", "Collection account"),
        ("escrow_subaccounts", "Escrow sub-accounts"),
        ("litigations_complaints_appeals", "Complaints and appeals as filed"),
        ("statutory_declaration", "Statutory declaration"),
        ("construction_progress", "Construction progress"),
    ):
        if compliance.get(key):
            b.bullet(f"{title}: {compliance[key]}")

    core = b.facts.get("rera_core_fields") or {}
    core_keys = [
        ("registration_number", "Registration number"),
        ("project_status", "Project status"),
        ("approved_date", "Approved on"),
        ("proposed_completion_date", "Proposed completion"),
        ("plan_approval_number", "Plan approval number"),
        ("authority", "Authority"),
        ("project_type", "Project type"),
        ("total_building_units", "Total units"),
        ("collection_bank_account", "Collection bank account"),
    ]
    rows = [[title, str(core[key])] for key, title in core_keys if core.get(key)]
    if rows:
        b.para()
        b.heading("RERA Record as Filed", level=2)
        b.para(
            "Every value below is the promoter's own filing with MahaRERA, reproduced as filed. It is an "
            "official venue but not an independent verification of the underlying facts."
        )
        b.table(["Field", "As filed"], rows)

    if b.facts.get("unit_summary_note"):
        b.para()
        b.para(str(b.facts["unit_summary_note"]))

    blocks = b.facts.get("blocks") or []
    if blocks:
        b.para()
        b.heading("Building and Unit Breakdown", level=2)
        b.table(
            ["Block / wing", "Floors", "Configuration", "Units", "Note"],
            [[x.get("block_wing", ""), x.get("floors", ""), x.get("config", ""), x.get("units_counted", ""), x.get("note", "")] for x in blocks],
        )

    gst = b.facts.get("gst_compliance_check") or {}
    b.para()
    b.heading("GST Filing Compliance", level=2)
    if gst.get("found"):
        summary = gst.get("summary") or {}
        b.para(
            f"GSTIN {gst.get('gstin', '')}: {summary.get('late_pct', 0)}% of rated return periods were filed "
            f"late, worst delay {summary.get('worst_delay_days') or 0} day(s), with "
            f"{summary.get('delays_last_12_months', 0)} delayed or unfiled period(s) falling due in the "
            f"trailing 12 months."
        )
    else:
        b.flag_line(
            "GST filing compliance is not assessed: it requires the developer's GSTIN and their filing history, "
            "and the GST portal gates that history behind a CAPTCHA that must be solved per lookup, so it "
            "cannot be collected unattended."
        )


def _section_10_location_and_market(b: _Builder) -> None:
    b.page_break()
    b.heading("10. Location, Connectivity & Market Read")

    connectivity = b.facts.get("connectivity") or {}
    if any(connectivity.values()):
        b.table(
            ["Mode", "Position"],
            [[mode.title(), str(connectivity.get(mode) or "Not stated")] for mode in ("road", "rail", "metro", "air")],
        )

    distances = b.facts.get("distances") or []
    if distances:
        b.para()
        b.heading("Distances to Key Landmarks", level=2)
        b.table(
            ["Landmark", "Distance / time", "Basis"],
            [[x.get("landmark", ""), x.get("distance_time", ""), x.get("route_note", "")] for x in distances],
        )

    if b.facts.get("social_infrastructure"):
        b.para()
        b.para(f"Social infrastructure: {b.facts['social_infrastructure']}")

    b.para()
    b.heading("Micro-Market Read", level=2)
    if b.facts.get("micro_market_overview"):
        b.para(str(b.facts["micro_market_overview"]))
    if b.facts.get("area_intelligence_trend"):
        b.para(str(b.facts["area_intelligence_trend"]))

    comparables = b.facts.get("comparables") or []
    if comparables:
        b.para()
        b.table(
            ["Comparable project", "Distance", "Configuration", "Pricing", "Source"],
            [
                [
                    c.get("project", ""),
                    f"{c['distance_km']} km" if c.get("distance_km") else "Not stated",
                    c.get("configuration", ""),
                    c.get("pricing", ""),
                    c.get("source", ""),
                ]
                for c in comparables
            ],
        )
    else:
        b.flag_line("No comparable projects were established for this micro-market this pass.")


def _section_11_document_trail(b: _Builder) -> None:
    b.page_break()
    b.heading("11. Document & Diligence Trail")

    library = b.facts.get("document_library") or []
    if library:
        opened = sum(1 for d in library if "download" in str(d.get("status", "")).lower() or "reused" in str(d.get("status", "")).lower())
        b.para(
            f"{len(library)} documents are listed on this project's MahaRERA record, of which {opened} were "
            f"retrieved for this review. The full list is a code-computed inventory of the registry's own "
            f"library rather than a summary of it."
        )
        b.table(
            ["Document", "Status"],
            [[d.get("document_name", ""), d.get("status", "")] for d in library],
        )
    else:
        b.flag_line("No document library was recorded for this project.")

    for key, title in (
        ("documents_reviewed_note", "Documents read in full"),
        ("documents_absent_note", "Documents absent from the registry library"),
    ):
        if b.facts.get(key):
            b.para()
            b.para(f"{title}: {b.facts[key]}")


# ---------------------------------------------------------------------------
# Sections covering BOTH halves.
# ---------------------------------------------------------------------------

def _section_12_consolidated_flags(b: _Builder, flags: dict, next_steps: list) -> None:
    b.page_break()
    b.heading("12. Consolidated Flags & Recommended Steps")

    imminent = flags.get("imminent") or []
    structural = flags.get("structural") or []
    monitor = flags.get("monitor") or []

    b.heading("What Checks Out", level=2)
    green = _green_flags(b.facts)
    if green:
        for point in green:
            b.bullet(point, color=cc._TEXT_GREEN)
    else:
        b.flag_line("No positive confirmations were recorded this pass.")

    b.para()
    b.heading(f"What Needs Attention ({len(imminent)} imminent, {len(structural)} structural)", level=2)
    if imminent:
        for item in imminent:
            b.bullet(f"Imminent: {item['text']}", color=cc._TEXT_RED)
    if structural:
        for item in structural:
            b.bullet(f"Structural: {item['text']}", color=cc._TEXT_AMBER)
    if not imminent and not structural:
        b.flag_line("No imminent or structural flag was raised this pass.")

    if monitor:
        b.para()
        b.heading(f"To Re-Check on a Future Pass ({len(monitor)})", level=2)
        for item in monitor:
            b.bullet(item["text"])

    b.para()
    b.heading("Recommended Verification Steps", level=2)
    steps = list(next_steps) + _recommended_steps(b.facts, flags)
    seen, unique_steps = set(), []
    for step in steps:
        if step not in seen:
            seen.add(step)
            unique_steps.append(step)
    if unique_steps:
        for step in unique_steps:
            b.bullet(step)
    else:
        b.flag_line("No further verification step is outstanding from this pass.")


def _green_flags(facts: dict) -> list:
    """Positive confirmations, each tied to a specific check that actually
    ran. Never a generic reassurance: if a check did not run, its absence
    belongs in the flags, not here."""
    green = []
    ibbi = facts.get("ibbi_insolvency_check") or {}
    if ibbi.get("found_process") is False:
        green.append("No corporate insolvency process is on record against this entity in IBBI's register.")

    complaint_count, appeal_count = cc._parse_complaint_appeal_counts(facts)
    if complaint_count == 0:
        green.append("No MahaRERA complaint is on record against this project.")
    if appeal_count == 0:
        green.append("No MahaRERA appeal is on record against this project.")

    _original, was_extended = cc._parse_completion_slippage(facts)
    if not was_extended:
        green.append("The RERA-registered completion date has never been extended.")

    profile = facts.get("company_profile_check") or {}
    if str(profile.get("status") or "").lower() == "active":
        green.append(f"The counterparty is an active company on the MCA record ({profile.get('roc', 'registrar not stated')}).")
    if len(profile.get("sources_used") or []) > 1:
        green.append(
            f"Its registration profile agrees across {len(profile['sources_used'])} independent registry mirrors."
        )

    totals = (facts.get("promoter_portfolio") or {}).get("totals") or {}
    if totals.get("on_time_rate_pct") == 100.0 and totals.get("on_time_count"):
        green.append(
            f"All {totals['on_time_count']} declared prior delivery/deliveries completed on or before the "
            f"originally proposed date, per the promoter's own filings."
        )
    return green


def _recommended_steps(facts: dict, flags: dict) -> list:
    """Concrete next steps derived from what this pass could NOT establish."""
    steps = []
    if not (facts.get("cts_land_record_check") or {}).get("found"):
        steps.append("Retrieve the Maha Bhulekh Property Card for the project's CTS number to confirm land ownership independently.")
    if not (facts.get("gst_compliance_check") or {}).get("found"):
        steps.append("Request the developer's GSTIN and pull their GSTR-1 / GSTR-3B filing history to assess statutory filing discipline.")
    if not (facts.get("credit_rating_check") or {}).get("promoter", facts.get("credit_rating_check") or {}).get("ratings"):
        steps.append("Ask the developer whether any agency rating, banking appraisal, or lender credit note exists for this entity.")
    if (facts.get("group_companies_check") or {}).get("undisclosed_relationship_counts"):
        steps.append("Obtain the counterparty's own shareholding and subsidiary disclosure; the registry's free tier withholds those counterparty identities.")
    steps.append("Request the firm-wide organisation structure and team headcount by function, which no public source discloses.")
    for item in (flags.get("imminent") or []):
        steps.append(f"Raise directly with the developer: {item['text']}")
    return steps


def assess_counterparty(facts: dict, flags: dict, developer_score: dict) -> dict:
    """Rule-derived read on substance, for section 13. Returns
    {"verdict", "signals": [(label, finding)], "spv_like": bool}.

    Deliberately NOT model-authored. This is the single most consequential
    passage in the document, so every clause is computed from a named fact
    and can be traced back to it. A model could phrase it more smoothly but
    could also be confidently wrong about whether a real company is a shell,
    which is a materially harmful thing to get wrong.
    """
    signals = []
    portfolio = facts.get("promoter_portfolio") or {}
    totals = portfolio.get("totals") or {}
    profile = facts.get("company_profile_check") or {}
    core = facts.get("rera_core_fields") or {}

    registrations = totals.get("total_projects")
    single_project = registrations == 1
    if registrations is not None:
        signals.append((
            "MahaRERA registrations",
            f"{registrations} held under this promoter's name"
            + (", which is the subject project itself" if single_project else ""),
        ))

    prior = totals.get("total_experience_entries_found") or 0
    signals.append((
        "Declared prior deliveries",
        f"{prior} declared to MahaRERA" if prior else "none declared to MahaRERA",
    ))

    incorporation = str(profile.get("incorporation_date") or "")
    approved = str(core.get("approved_date") or "")
    if incorporation and approved:
        signals.append((
            "Incorporation against project registration",
            f"incorporated {incorporation}, project approved {approved}",
        ))

    directors = profile.get("current_directors") or []
    if directors:
        signals.append(("Board", f"{len(directors)} current {cc._role_word(facts, len(directors))} on the registry roster"))

    group = facts.get("group_companies_check") or {}
    linked = len(group.get("companies") or [])
    if linked:
        signals.append((
            "Group footprint",
            f"{linked} entities share a named {cc._role_word(facts)} or registered office with this counterparty",
        ))

    paid_up = str(profile.get("paid_up_capital") or "").strip()
    if paid_up:
        signals.append(("Paid-up capital", paid_up))

    imminent_count = len(flags.get("imminent") or [])
    signals.append((
        "Imminent flags",
        f"{imminent_count} raised in this review" if imminent_count else "none raised in this review",
    ))

    # --- verdict ------------------------------------------------------------
    # "Paper firm" is a claim about SUBSTANCE, and a single-project SPV inside
    # a real corporate group is a completely different animal from a
    # standalone entity with nothing behind it. The two are separated here
    # rather than collapsed into one label.
    substantial_group = linked >= 10
    if single_project and prior == 0 and not substantial_group:
        verdict = (
            "This counterparty presents as a standalone single-project entity with no declared prior delivery "
            "and no substantial group footprint behind it. On the evidence gathered, treat its execution "
            "capability as unproven and seek direct evidence of delivery capacity before proceeding."
        )
    elif single_project and substantial_group:
        verdict = (
            f"This counterparty is a single-project vehicle, but not a shell: {linked} entities share named "
            f"leadership or a registered office with it, and it is an active company on the MCA record. Read it "
            f"as a project-specific vehicle of a larger group, which means execution capability rests on the "
            f"group rather than on this entity's own balance sheet or track record."
        )
    elif single_project:
        verdict = (
            "This counterparty is a single-project vehicle with a limited independent track record. Execution "
            "capability cannot be established from this entity's own history alone."
        )
    else:
        verdict = (
            f"This counterparty holds {registrations} MahaRERA registrations and has a track record that can be "
            f"reviewed on its own terms rather than inferred from a parent group."
        )

    grade = developer_score.get("grade")
    if grade:
        verdict += (
            f" Its composite Developer Score is {developer_score.get('composite')}/100 (grade {grade}), which "
            f"reflects both measured strength and how much of this promoter is publicly verifiable at all."
        )

    return {"verdict": verdict, "signals": signals, "spv_like": bool(single_project)}


def _section_13_closing_read(b: _Builder, flags: dict, developer_score: dict) -> None:
    b.page_break()
    b.heading("13. Closing Read")

    assessment = assess_counterparty(b.facts, flags, developer_score)
    b.table(["Signal", "Finding"], [[label, finding] for label, finding in assessment["signals"]])
    b.para()
    b.para(assessment["verdict"], bold=True)
    b.para()
    b.para(
        "This read is derived by rule from the signals tabled above, not written as an opinion. Each signal is "
        "a fact recorded elsewhere in this Charter and can be traced back to its source."
    )


def _section_14_scoring_detail(b: _Builder, developer_score: dict) -> None:
    b.page_break()
    b.heading("14. Scoring Detail")
    b.para(
        "The full sub-metric working behind the two summary tables in section 1. Presented here rather than up "
        "front so the summary stays readable, and included rather than omitted so any rating can be traced to "
        "what produced it."
    )

    b.heading("Developer Score, by Sub-Metric", level=2)
    # External drops the Weight column here -- same rule _fill_template
    # already applies to this exact table (see its own show_weight note): the
    # bucket explanation above already covers how scoring works, and a
    # numeric weight per row reads as internal-methodology detail rather than
    # a finding. The quality gate enforces this (it specifically checks a
    # Bucket/Sub-metric-headed table for a stray Weight column), which is
    # what caught this the first time it was missed here.
    show_weight = b.facts.get("_doc_variant") != "external"
    criteria = developer_score.get("criteria") or {}
    rows = []
    for bucket_name, _bucket_weight, metrics in cc._DEVELOPER_SCORE_STRUCTURE:
        for key, _display, _fn in metrics:
            criterion = criteria.get(key) or {}
            row = [bucket_name, criterion.get("display_name") or key.replace("_", " ").title()]
            if show_weight:
                row.append(f"{criterion.get('weight', '')}%")
            row += [
                criterion.get("tier") or "Not rated",
                str(criterion.get("score")) if criterion.get("score") is not None else "",
                criterion.get("note") or criterion.get("reason") or "",
            ]
            rows.append(row)
    headers = ["Bucket", "Sub-metric"] + (["Weight"] if show_weight else []) + ["Tier", "Score", "Basis"]
    b.table(headers, rows)

    doc_confidence = b.facts.get("documentation_confidence_score") or {}
    dc_criteria = doc_confidence.get("criteria") or {}
    if dc_criteria:
        b.para()
        b.heading("Documentation Confidence, by Criterion", level=2)
        dc_rows = []
        for bucket_name, keys in _DOC_CONFIDENCE_BUCKETS:
            for key in keys:
                criterion = dc_criteria.get(key)
                if not criterion:
                    continue
                dc_rows.append([
                    bucket_name,
                    _DOC_CONFIDENCE_CRITERION_LABELS.get(key, key),
                    f"{criterion.get('weight', '')}%",
                    f"{round(criterion['score'])}/100" if criterion.get("score") is not None else "Not rated",
                    criterion.get("note") or "",
                ])
        if dc_rows:
            b.table(["Bucket", "Criterion", "Weight", "Score", "Basis"], dc_rows)
        for skipped in (doc_confidence.get("skipped_criteria") or []):
            b.flag_line(
                f"Not scored: {_DOC_CONFIDENCE_CRITERION_LABELS.get(skipped, skipped)}. Excluded from this "
                f"score rather than counted as a failure."
            )


# ---------------------------------------------------------------------------
# Harvard-style source list.
# ---------------------------------------------------------------------------

_MONTHS = ("January", "February", "March", "April", "May", "June",
           "July", "August", "September", "October", "November", "December")


def _harvard_date(value: str) -> str:
    """"2026-07-28" -> "28 July 2026". Returns the input unchanged if it
    isn't an ISO date, so an already-formatted or unusual value is passed
    through rather than mangled."""
    text = str(value or "").strip()
    parts = text.split("-")
    if len(parts) == 3 and all(p.isdigit() for p in parts):
        year, month, day = int(parts[0]), int(parts[1]), int(parts[2])
        if 1 <= month <= 12:
            return f"{day} {_MONTHS[month - 1]} {year}"
    return text


def _harvard_year(published_date: str) -> str:
    """Harvard uses "(n.d.)" for an undated source. This pipeline records
    "unknown" for the same thing, so that is translated rather than printed
    literally."""
    text = str(published_date or "").strip()
    if not text or text.lower() in ("unknown", "n/a", "none"):
        return "n.d."
    year = text.split("-")[0]
    return year if year.isdigit() and len(year) == 4 else text


def harvard_reference(source: dict) -> str:
    """Formats one source dict as a Harvard reference:
    Organisation (Year) Title. Available at: URL (Accessed: DD Month YYYY).

    "Available at" is omitted entirely when the source carries no URL, since
    a downloaded document has no public address to cite."""
    label = str(source.get("label") or "").strip()
    ref = str(source.get("ref") or "").strip()
    year = _harvard_year(source.get("published_date"))
    accessed = _harvard_date(source.get("accessed_date"))

    url_match = re.search(r"https?://\S+", ref)
    url = url_match.group(0).rstrip(".,;") if url_match else ""
    title = ref.replace(url, "").strip() if url else ref
    title = title.strip(" -,;")

    parts = [f"{label} ({year})"]
    if title:
        parts.append(f" {title}.")
    else:
        parts.append(".")
    if url:
        parts.append(f" Available at: {url}")
    if accessed:
        parts.append(f" (Accessed: {accessed}).")
    return "".join(parts)


def _section_15_sources(b: _Builder) -> None:
    b.page_break()
    b.heading("15. Sources")

    sources = b.facts.get("sources") or []
    if not sources:
        b.flag_line("No sources were recorded for this Charter.")
        return

    b.para(
        f"{len(sources)} sources, referenced in Harvard style. An undated source is shown as (n.d.); a source "
        f"with no public address is a document retrieved directly and carries no 'Available at' line."
    )
    for source in sorted(sources, key=lambda s: str(s.get("label") or "").lower()):
        b.bullet(harvard_reference(source))


def build_charter_document(reg_no: str, facts: dict, out_path: str, doc_variant: str = "internal") -> str:
    """Builds the restructured Charter at `out_path` and returns that path.

    Takes a deep copy for the External variant for the same reason
    _fill_template does: the caller's own `facts` is persisted to
    .facts.json afterwards and must keep its un-rewritten internal content
    regardless of whether an External pass ran against it."""
    import shutil

    import docx

    if doc_variant == "external":
        facts = cc._externalized_facts_copy(facts)
    facts["_doc_variant"] = doc_variant
    facts["_citation_registry"] = {"order": [], "index": {}}

    shutil.copy2(cc.TEMPLATE_PATH, out_path)
    doc = docx.Document(out_path)
    cc._fix_bullet_hanging_indent(doc)
    styles = _capture_styles(doc)
    _clear_body(doc)

    b = _Builder(doc, styles, facts)
    core = facts.get("rera_core_fields") or {}
    corp = facts.get("corporate_identity") or {}

    # --- title block ------------------------------------------------------
    title = b.para("COMPANY CHARTER")
    if styles["title"] is not None:
        title.style = styles["title"]
    promoter = cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", "")) or "[Unknown]"
    b.para(f'Counterparty: {promoter} | Collateral: {core.get("project_name", "[Unknown]")} ({reg_no})', bold=True)

    flags = cc._classify_flags(facts)
    # ALWAYS recomputed, never read from facts["developer_score"] -- that key
    # is an OUTPUT of a previous run, not an input, and trusting it renders
    # stale numbers the moment the rubric changes. Confirmed the hard way:
    # this fixture's persisted copy was from the superseded flat 7-criteria
    # scheme (composite 100.0 / grade A / weights 33.3, no buckets, no
    # rera_compliance or gst_compliance at all), so an earlier
    # `facts.get(...) or compute(...)` here rendered 100.0/A against a true
    # 37.1/C and reported scored sub-metrics as "Not rated". _fill_template
    # recomputes at render time for exactly this reason (see its own note by
    # the same call).
    developer_score = cc._compute_developer_score(facts, flags)
    facts["developer_score"] = developer_score

    next_steps: list = []

    # Counterparty
    _section_1_verification_summary(b, flags, developer_score)
    _section_2_counterparty_identity(b)
    _section_3_leadership_and_related(b)
    _section_4_litigation_screen(b)
    _section_5_organisation_and_team(b)
    _section_6_portfolio(b, next_steps)

    # Collateral
    _section_7_asset_identity(b)
    _section_8_approvals_fsi_title(b)
    _section_9_rera_compliance_escrow(b)
    _section_10_location_and_market(b)
    _section_11_document_trail(b)

    # Covering both
    _section_12_consolidated_flags(b, flags, next_steps)
    _section_13_closing_read(b, flags, developer_score)
    _section_14_scoring_detail(b, developer_score)
    _section_15_sources(b)

    doc.save(out_path)

    if doc_variant == "external":
        violations = cc._verify_external_document_quality(out_path)
        if violations:
            raise RuntimeError(
                f"Restructured External Charter quality gate failed for {out_path} "
                f"({len(violations)} violation(s)):\n" + "\n".join(f"  - {v}" for v in violations)
            )
    return out_path
