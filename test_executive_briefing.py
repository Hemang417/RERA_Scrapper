"""
Tests for executive_briefing.py -- the 8-page narrative companion to the
15-section Company Charter.

Run directly: python test_executive_briefing.py
"""

import json
import os
import shutil

import docx

import charter_document as cd
import company_charter as cc
import executive_briefing as eb

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_briefing")


def _load_facts():
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        return json.load(f)


def _all_text(doc):
    """Quote Highlight and Key Statistic Card widgets are rendered as table
    cells (see executive_briefing._quote_highlight/_key_statistic_cards),
    which python-docx does not surface via doc.paragraphs -- only the body's
    own top-level paragraphs are. Any assertion over the document's full text
    must include table cell text or it will silently miss widget content."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_builds_against_real_data_with_eight_logical_pages():
    """A page_break is inserted before each of the 7 pages after the cover,
    so a correctly built briefing has exactly 7 explicit page breaks (8
    logical pages). Word may reflow a page's own content across more than
    one physical page if it overflows -- that's expected and not a bug, the
    company-report-generator skill's own framework says to adjust page count
    when content requires it -- but the number of *authored* page breaks
    must stay fixed at 7."""
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    out = os.path.join(_SCRATCH, "briefing.docx")
    eb.build_executive_briefing("P51800077150", facts, out, "30 July 2026")
    assert os.path.exists(out)

    doc = docx.Document(out)
    page_breaks = sum(
        1 for p in doc.paragraphs for run in p.runs
        for br in run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
        if br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
    )
    assert page_breaks == 7, page_breaks
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_builds_against_real_data_with_eight_logical_pages: PASS")


def test_never_trusts_a_stale_persisted_developer_score():
    """Same bug class as charter_document.py: facts["developer_score"] is an
    OUTPUT of a prior run. The briefing must recompute it via
    cc._compute_developer_score, exactly like the Charter does, so the two
    documents can never disagree."""
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    facts["developer_score"] = {
        "composite": 100.0, "grade": "A",
        "criteria": {"track_record_years": {"score": 100.0, "tier": "AAA", "weight": 33.3}},
    }
    out = os.path.join(_SCRATCH, "stale.docx")
    eb.build_executive_briefing("P51800077150", facts, out, "30 July 2026")

    text = _all_text(docx.Document(out))
    assert "100.0/100" not in text, "stale persisted composite was rendered"
    assert "37.1/100" in text, text[:400]
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_never_trusts_a_stale_persisted_developer_score: PASS")


def test_stale_doc_confidence_is_flagged_not_rendered():
    """Mirrors charter_document.py's own staleness guard: this score cannot
    be recomputed standalone, so if none of a persisted score's criteria map
    onto the current buckets, it must be flagged rather than rendered."""
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    facts["documentation_confidence_score"] = {
        "overall": 88.0, "band": "High",
        "criteria": {"some_retired_criterion": {"score": 88.0, "weight": 100.0}},
    }
    out = os.path.join(_SCRATCH, "stale_dc.docx")
    eb.build_executive_briefing("P51800077150", facts, out, "30 July 2026")

    text = _all_text(docx.Document(out))
    assert "88.0" not in text, "a superseded Documentation Confidence figure was rendered"
    assert "superseded scoring scheme" in text, text[:600]
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_stale_doc_confidence_is_flagged_not_rendered: PASS")


def test_verdict_matches_the_charter_word_for_word():
    """The whole point of this briefing is that it narrates the Charter's
    own rule-derived verdict rather than authoring a new one. The quote pulled
    onto pages 2 (Executive Summary) and 8 (Conclusion) must be byte-identical
    to charter_document.assess_counterparty's own output, computed the same
    way from the same facts."""
    facts = _load_facts()
    flags = cc._classify_flags(facts)
    developer_score = cc._compute_developer_score(facts, flags)
    expected_verdict = cd.assess_counterparty(facts, flags, developer_score)["verdict"]

    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, "verdict.docx")
    eb.build_executive_briefing("P51800077150", dict(facts), out, "30 July 2026")
    text = _all_text(docx.Document(out))
    assert f"“{expected_verdict}”" in text, (expected_verdict, text[:2000])
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_verdict_matches_the_charter_word_for_word: PASS")


if __name__ == "__main__":
    test_builds_against_real_data_with_eight_logical_pages()
    test_never_trusts_a_stale_persisted_developer_score()
    test_stale_doc_confidence_is_flagged_not_rendered()
    test_verdict_matches_the_charter_word_for_word()
    print("\nAll tests passed.")
