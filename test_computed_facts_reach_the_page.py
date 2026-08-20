"""Stops code-computed facts being built, tested, and then rendered nowhere.

THE FAILURE THIS FILE EXISTS FOR. On 19 August 2026 a review of this repo
found three capabilities that were fully built, unit-tested and verified
against live sources, and consumed by NOTHING:

    company_charter.summarise_charges   0 production callers
    group_entities.build_entity_graph   0 production callers
    facts["promoter_identity_check"]    written, never read

Between them: Rs 90.35 crore of open secured borrowing, a 65-entity group
graph, and the promoter's PAN. All correct, all invisible to the reader.
Every unit test passed the whole time, because a unit test proves a function
works, not that anything calls it.

So these tests render the real saved facts through _fill_template and assert
the numbers actually appear ON THE PAGE. They are deliberately end-of-chain:
they do not care how a value is computed, only that a reader would see it.

The second thing they pin is the opposite direction, and it matters more: a
value that failed verification must NOT reach the page. A PAN is a national
join key -- print a wrong one and every downstream search returns another
company's records. So "unverified" must render as a gap the reader can act
on, never as a fact.

No API calls: rendering from a saved facts.json is the established no-network
pattern (see test_state_labels.py and test_executive_ready.py). Renders are
slow, so each one is asserted against several times rather than re-run.

Run directly: python test_computed_facts_reach_the_page.py
"""

import copy
import io
import json
import os
import re
import shutil

import docx

import company_charter as cc
import group_entities as ge

_FACTS_PATH = os.path.join(
    "output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json"
)
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_computed_facts")
_REG_NO = "P51800077150"

_VERIFIED_PAN = {
    "pan": "AANCP0234D",
    "holder_type": "Company",
    "incorporation_date": "2022-06-27",
    "source_document": "PAN Card",
    "verified": True,
    "status": "verified",
    "checks": ["holder-type character 'C' = Company",
               "5th character 'P' matches promoter name initial"],
    "notes": [],
    "unverified_candidates": [],
}


def _base_facts():
    with open(_FACTS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def _render(facts, variant="internal", name="out"):
    """A fresh deep copy per render: _fill_template mutates what it is given
    (same reason test_state_labels does this)."""
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{name}_{variant}.docx")
    cc._fill_template(_REG_NO, copy.deepcopy(facts), out, doc_variant=variant)
    return docx.Document(out)


def _all_cells(document):
    return [c.text for t in document.tables for r in t.rows for c in r.cells]


def _all_text(document):
    return "\n".join([p.text for p in document.paragraphs] + _all_cells(document))


# --- the PAN reaches the page ---------------------------------------------

def test_a_verified_pan_reaches_the_identity_table_in_both_variants():
    """The whole point of extracting it. Also pins WHERE it lands: Section B
    allows an identifier like this in the identity tables and nowhere else,
    so it must be in a table cell, not in prose."""
    facts = _base_facts()
    facts["promoter_identity_check"] = copy.deepcopy(_VERIFIED_PAN)

    internal = _render(facts, "internal", "pan")
    cells = _all_cells(internal)
    row = [c for c in cells if "AANCP0234D" in c]
    assert row, "the verified PAN never reached the Internal document"
    assert any("Permanent Account Number" in c for c in cells), \
        "the PAN row label must expand the initialism (Section B, no jargon)"
    assert "(PAN Card)" in row[0], row[0]

    external = _render(facts, "external", "pan")
    ext_cells = _all_cells(external)
    assert any("AANCP0234D" in c for c in ext_cells), \
        "the verified PAN never reached the External document"
    print("test_a_verified_pan_reaches_the_identity_table_in_both_variants: PASS")


def test_the_external_pan_citation_is_generic_not_an_internal_artifact_name():
    """Section C: an External citation must describe the source in
    client-facing language, never name an internal file or category."""
    facts = _base_facts()
    facts["promoter_identity_check"] = copy.deepcopy(_VERIFIED_PAN)
    external = _render(facts, "external", "pancite")
    text = _all_text(external)
    assert "promoter identity filing" in text, \
        "the PAN source did not resolve to its generic External description"
    # ...and the raw label never leaks into the External sources list.
    sources = [p.text for p in external.paragraphs if p.text.strip().startswith("[")]
    assert not any(s.strip().endswith("(PAN Card)") for s in sources), sources
    print("test_the_external_pan_citation_is_generic_not_an_internal_artifact_name: PASS")


# --- and, more importantly, an unverified one does NOT --------------------

def test_an_unverified_pan_never_reaches_the_page_and_becomes_a_gap():
    """The rule that matters most here. A PAN that could not be tied to this
    promoter is either a scan problem or somebody else's number on the file.
    Printing it would send every downstream corporate, tax and charge search
    to the wrong entity, and each of those would look perfectly consistent."""
    facts = _base_facts()
    facts["promoter_identity_check"] = {
        "pan": None, "holder_type": None, "incorporation_date": None,
        "source_document": None, "verified": False, "status": "unverified_candidate",
        "checks": [], "notes": [],
        "unverified_candidates": [
            {"pan": "AANCZ0234D", "document": "PAN Card",
             "reason": "5th character 'Z' does not match the promoter name's initial 'P'"}
        ],
    }
    cc._record_promoter_identity_gap(facts)

    gaps = " ".join(facts.get("gaps") or [])
    assert "could not be matched to the promoter" in gaps, facts.get("gaps")
    assert "AANCZ0234D" not in gaps, "the rejected number must not be printed even in the gap"

    document = _render(facts, "internal", "panbad")
    text = _all_text(document)
    assert "AANCZ0234D" not in text, "an unverified PAN reached the document"
    assert "Permanent Account Number" in text, "the gap explaining the failure is missing"
    print("test_an_unverified_pan_never_reaches_the_page_and_becomes_a_gap: PASS")


def test_no_pan_card_on_file_produces_no_row_and_no_boilerplate_gap():
    """Karnataka and Telangana do not publish the card. Section B's
    clean-check rule means that absence must NOT put a line in every
    Charter -- it is an authority-level fact, already in the acquisition
    notes, not a per-project finding."""
    facts = _base_facts()
    before = len(facts.get("gaps") or [])
    facts["promoter_identity_check"] = {
        "pan": None, "verified": False, "status": "no_card", "unverified_candidates": [],
        "notes": ["No PAN card was found in this authority's document library for this project, "
                  "so the promoter's PAN could not be read from the RERA filing."],
    }
    cc._record_promoter_identity_gap(facts)
    assert len(facts.get("gaps") or []) == before, \
        "a missing PAN card added boilerplate to the gap list"

    document = _render(facts, "internal", "panabsent")
    cells = _all_cells(document)
    assert not any("Permanent Account Number" in c for c in cells), \
        "an empty PAN row was rendered instead of being omitted"
    print("test_no_pan_card_on_file_produces_no_row_and_no_boilerplate_gap: PASS")


def test_a_card_that_could_not_be_read_is_a_gap_not_a_silence():
    """"The promoter filed no PAN card" and "we could not read the card they
    filed" are different findings. Only the second one is true here, and a
    reader who sees nothing would conclude the first."""
    facts = _base_facts()
    facts["promoter_identity_check"] = {
        "pan": None, "verified": False, "status": "unreadable", "unverified_candidates": [],
        "notes": ["PAN Card: the file could not be read as text or by OCR."],
    }
    cc._record_promoter_identity_gap(facts)
    gaps = " ".join(facts.get("gaps") or [])
    assert "could not be read from the filed copy" in gaps, facts.get("gaps")
    print("test_a_card_that_could_not_be_read_is_a_gap_not_a_silence: PASS")


def test_gap_text_carries_no_filenames_keys_or_error_strings():
    """Section B forbids a file path, module name, function, JSON key or raw
    exception string in EITHER document. These gaps are written as plain
    prose so they satisfy that by construction, rather than depending on the
    process-text sanitizer to catch them."""
    for check in (
        {"pan": None, "verified": False, "status": "unreadable",
         "notes": ["PAN Card: the file could not be read as text or by OCR."],
         "unverified_candidates": []},
        {"pan": None, "verified": False, "status": "unverified_candidate", "notes": [],
         "unverified_candidates": [{"pan": "AANCZ0234D", "document": "PAN Card", "reason": "mismatch"}]},
    ):
        facts = {"promoter_identity_check": check}
        cc._record_promoter_identity_gap(facts)
    text = " ".join(facts.get("gaps") or [])
    for banned in (".json", ".py", ".pdf", "promoter_identity", "Traceback",
                   "_check", "output/", "facts[", "None"):
        assert banned not in text, f"{banned!r} leaked into a reader-facing gap: {text}"
    print("test_gap_text_carries_no_filenames_keys_or_error_strings: PASS")


def test_no_card_and_unreadable_card_are_told_apart_by_status_not_prose():
    """Regression. The first version of _record_promoter_identity_gap decided
    between these two by searching the notes for "could not be read" -- but
    the NO-CARD note also ends "...so the promoter's PAN could not be read
    from the RERA filing", so a state that simply does not publish the card
    got the unreadable-card gap on every single project. Caught by the test
    above; fixed by giving the result an explicit machine-readable status and
    keying on that. Prose is for the reader, never for control flow."""
    no_card_note = ("No PAN card was found in this authority's document library for this "
                    "project, so the promoter's PAN could not be read from the RERA filing.")
    assert "could not be read" in no_card_note, "the collision this test guards is gone"

    quiet = {"promoter_identity_check": {"status": "no_card", "verified": False,
                                         "notes": [no_card_note], "unverified_candidates": []}}
    cc._record_promoter_identity_gap(quiet)
    assert not quiet.get("gaps"), quiet.get("gaps")

    loud = {"promoter_identity_check": {"status": "unreadable", "verified": False,
                                        "notes": [no_card_note], "unverified_candidates": []}}
    cc._record_promoter_identity_gap(loud)
    assert loud.get("gaps"), "an unreadable card produced no gap"
    print("test_no_card_and_unreadable_card_are_told_apart_by_status_not_prose: PASS")


def test_every_status_the_extractor_emits_is_handled():
    """Anti-drift: a new status added to promoter_identity must be given an
    explicit decision here, not fall through to silence by accident."""
    import promoter_identity as pi
    source = io.open("promoter_identity.py", "r", encoding="utf-8").read()
    declared = set(re.findall(r'out\["status"\] = "(\w+)"', source))
    declared.add("no_card")  # the initial value
    handled = {"verified", "unverified_candidate", "unreadable", "ocr_unavailable", "no_card"}
    assert declared <= handled, f"unhandled status(es): {sorted(declared - handled)}"
    print("test_every_status_the_extractor_emits_is_handled: PASS")


# --- MCA charges reach the page -------------------------------------------

_LIVE_CHARGES = [
    {"charge_id": "100857390", "creation_date": "2023-10-31", "closure_date": None,
     "is_open": True, "amount": "3,491,110.00", "charge_holder": "HDFC BANK LIMITED"},
    {"charge_id": "100878097", "creation_date": "2024-01-29", "closure_date": None,
     "is_open": True, "amount": "300,000,000.00", "charge_holder": "CATALYST TRUSTEESHIP LIMITED"},
    {"charge_id": "100939871", "creation_date": "2024-03-26", "closure_date": None,
     "is_open": True, "amount": "300,000,000.00", "charge_holder": "CATALYST TRUSTEESHIP LIMITED"},
    {"charge_id": "100940922", "creation_date": "2024-03-26", "closure_date": None,
     "is_open": True, "amount": "300,000,000.00", "charge_holder": "CATALYST TRUSTEESHIP LIMITED"},
]


def test_the_profile_merge_carries_charges_through():
    """REGRESSION, and the most expensive bug found in this pass.

    lookup_company_by_cin parsed the charge table correctly all along, but
    _run_mca_profile_chain rebuilt its result from _SCALAR_PROFILE_FIELDS
    plus three explicitly-named keys, and `charges` was not one of them. So
    every caller saw an empty list: four open charges totalling Rs 90.35
    crore, parsed and then dropped by the merge. A Charter built on that
    would have reported a promoter with no secured borrowing at all.
    """
    original = cc._MCA_PROFILE_CHAIN
    cc._MCA_PROFILE_CHAIN = [
        ("zaubacorp.com", lambda cin, name: {
            "found": True, "name": "Test Co", "charges": copy.deepcopy(_LIVE_CHARGES),
            "past_directors": [], "shareholding_note": "n", "url": "https://example.test",
            "current_directors": [],
        }),
        ("tofler.in", lambda cin, name: {"found": True, "name": "Test Co", "current_directors": []}),
    ]
    try:
        merged = cc._run_mca_profile_chain("U70109MH2022PLC385473", "Test Co")
    finally:
        cc._MCA_PROFILE_CHAIN = original

    assert merged.get("charges"), "the merge dropped charges again"
    assert len(merged["charges"]) == 4, merged["charges"]
    summary = cc.summarise_charges(merged["charges"])
    assert summary["total_open_amount"] == 903491110.0, summary
    print("test_the_profile_merge_carries_charges_through: PASS")


def test_open_charges_and_lenders_reach_the_page():
    """Rs 90.35 crore and both lender names must be visible to a reader, in
    both variants -- this is the only independent check on the promoter's
    declared mortgage, which the RERA record states as an area with no
    lender named."""
    facts = _base_facts()
    facts["company_profile_check"]["charges"] = copy.deepcopy(_LIVE_CHARGES)

    for variant in ("internal", "external"):
        document = _render(facts, variant, "charges")
        text = _all_text(document)
        assert "90.35 crore" in text, f"[{variant}] the open charge total never reached the page"
        assert "HDFC BANK LIMITED" in text, f"[{variant}] lender missing"
        assert "CATALYST TRUSTEESHIP LIMITED" in text, f"[{variant}] lender missing"
        assert "100940922" in text, f"[{variant}] per-charge detail missing"
    print("test_open_charges_and_lenders_reach_the_page: PASS")


def test_unreadable_amounts_report_a_count_and_never_a_zero():
    """"Rs 0 of secured borrowing" and "the amounts could not be read" are
    opposite claims about a company's debt. summarise_charges returns None
    rather than 0.0 for exactly this reason, and the page must honour it."""
    facts = _base_facts()
    facts["company_profile_check"]["charges"] = [
        {"charge_id": "1", "creation_date": "2024-01-01", "closure_date": None,
         "is_open": True, "amount": None, "charge_holder": "A BANK"},
    ]
    document = _render(facts, "internal", "chargesunreadable")
    text = _all_text(document)
    assert "could not be read" in text, text[:400]
    assert "Rs 0" not in text, "an unreadable amount rendered as zero borrowing"
    assert "A BANK" in text
    print("test_unreadable_amounts_report_a_count_and_never_a_zero: PASS")


def test_a_promoter_with_no_charges_gets_the_empty_section_form():
    """Section B: a check that ran and came back clear produces no sentence
    -- heading plus one bare line, no scope-of-search paragraph, no source,
    no citation on the absence."""
    facts = _base_facts()
    facts["company_profile_check"]["charges"] = []
    document = _render(facts, "internal", "chargesnone")
    paragraphs = [p.text.strip() for p in document.paragraphs]
    idx = next((i for i, t in enumerate(paragraphs) if "Secured Borrowing" in t), None)
    assert idx is not None, "the section vanished entirely"
    assert paragraphs[idx + 1] == "Nothing found.", paragraphs[idx + 1:idx + 3]
    print("test_a_promoter_with_no_charges_gets_the_empty_section_form: PASS")


def test_a_charge_lookup_that_never_ran_asserts_nothing():
    """charges=None means the check did not happen (an older run, or a
    failed fetch). Printing "Nothing found." there would claim a clean
    register on the strength of a search that never occurred."""
    facts = _base_facts()
    facts["company_profile_check"].pop("charges", None)
    document = _render(facts, "internal", "chargesabsent")
    text = _all_text(document)
    assert "Secured Borrowing" not in text, "a section was rendered for a check that never ran"
    print("test_a_charge_lookup_that_never_ran_asserts_nothing: PASS")


# --- the pan-India state footprint ----------------------------------------

def test_a_declared_out_of_state_project_reaches_the_page():
    """The case this whole section exists for. Pranami holds ONE MahaRERA
    registration, and told MahaRERA about a completed Rs 128 crore mall in
    Ranchi. A Maharashtra-only read of this promoter shows a single project
    and no track record; the declared past project is the only thing that
    says otherwise, and it was being geocoded purely to decide the project
    was more than 5km from Mumbai, then discarded."""
    facts = _base_facts()
    facts["state_footprint"] = ge.state_footprint(
        {"confirmed": [{"name": "Mall of Ranchi Private Limited", "cin": "U93000JH2020PTC014638"}]},
        [{"projectName": "Mall of Ranchi", "address": "Ratu Road Ranchi Jharkhand 835222"}],
    )
    for variant in ("internal", "external"):
        document = _render(facts, variant, "footprint")
        text = _all_text(document)
        assert "Jharkhand" in text, f"[{variant}] the out-of-state project never reached the page"
        assert "Mall of Ranchi" in text, f"[{variant}] the project was not named"
    print("test_a_declared_out_of_state_project_reaches_the_page: PASS")


def test_incorporation_and_operation_are_reported_separately():
    """Incorporation is not operation. A shell registered in Maharashtra
    says nothing about where the group builds, so merging the two counts
    into one "states" number would overstate the footprint."""
    footprint = ge.state_footprint(
        {"confirmed": [{"name": "A Ltd", "cin": "U45200MH2005PTC111111"},
                       {"name": "B Ltd", "cin": "U45200MH2007PTC222222"}]},
        [{"projectName": "Ranchi Mall", "address": "Ranchi Jharkhand 835222"}],
    )
    assert [r["state"] for r in footprint["incorporated_in"]] == ["Maharashtra"]
    assert footprint["incorporated_in"][0]["count"] == 2
    assert [r["state"] for r in footprint["built_in"]] == ["Jharkhand"]
    print("test_incorporation_and_operation_are_reported_separately: PASS")


def test_llps_are_counted_out_loud_not_quietly_omitted():
    """20 of Pranami's 65 linked entities are partnerships, and an LLP
    identifier encodes no state at all. Dropping them silently would make
    the footprint look more complete than it is."""
    footprint = ge.state_footprint({"confirmed": [
        {"name": "Real Co", "cin": "U45200MH2005PTC111111"},
        {"name": "Some LLP", "cin": "AAM-0112"},
        {"name": "Other LLP", "cin": "ABZ-6154"},
    ]}, [{"address": "Mumbai Maharashtra 400058"}])
    assert footprint["unmapped_entities"] == 2, footprint
    assert any("partnerships" in l for l in footprint["limitations"]), footprint["limitations"]
    print("test_llps_are_counted_out_loud_not_quietly_omitted: PASS")


def test_an_unrecognised_state_code_is_named_never_guessed():
    """A wrong state here sends a sweep to the wrong regulator, which then
    reports a clean record. "PN" is real and live in this promoter's own
    group graph, and it is not a code this map recognises."""
    footprint = ge.state_footprint({"confirmed": [
        {"name": "Bord Systems India Private Limited", "cin": "U72900PN2021FTC204054"},
    ]})
    assert footprint["unrecognised_codes"] == ["PN"], footprint
    assert not footprint["incorporated_in"], "an unknown code was assigned to a state anyway"
    assert any("not recognised" in l for l in footprint["limitations"]), footprint["limitations"]
    print("test_an_unrecognised_state_code_is_named_never_guessed: PASS")


def test_the_state_comes_from_the_address_name_before_the_pin_code():
    """Name first because it is unambiguous. The PIN fallback is coarse and
    only covers ranges that do not straddle two states."""
    assert ge.state_from_address("Ratu Road Ranchi Jharkhand 835222") == "Jharkhand"
    assert ge.state_from_address("somewhere 835222") == "Jharkhand"      # PIN fallback
    assert ge.state_from_address("Andheri West 400058") == "Maharashtra"
    assert ge.state_from_address("no state and no pin here") is None
    assert ge.state_from_address("") is None
    # An alias the formal name would miss.
    assert ge.state_from_address("Bhubaneswar Orissa") == "Odisha"
    print("test_the_state_comes_from_the_address_name_before_the_pin_code: PASS")


def test_cin_state_extraction_handles_llpins_and_junk():
    assert ge.state_from_cin("U93000JH2020PTC014638") == ("Jharkhand", "JH")
    assert ge.state_from_cin("U72900PN2021FTC204054") == (None, "PN")
    assert ge.state_from_cin("AAM-0112") == (None, None)
    assert ge.state_from_cin(None) == (None, None)
    assert ge.state_from_cin("") == (None, None)
    print("test_cin_state_extraction_handles_llpins_and_junk: PASS")


def test_no_footprint_data_renders_no_section():
    """An empty footprint must not produce a heading with nothing under it,
    and must never imply the group operates in one state only."""
    facts = _base_facts()
    facts["state_footprint"] = ge.state_footprint({}, [])
    document = _render(facts, "internal", "fpempty")
    assert "State Footprint" not in _all_text(document)
    print("test_no_footprint_data_renders_no_section: PASS")


# --- the pipeline stages are actually reachable ---------------------------

def test_charge_movement_reaches_the_page_but_only_when_something_moved():
    """A charge list is a snapshot; a lender asks what MOVED. Real movement
    must be printed -- and "no change since last time" must not be, because
    Section B's clean-check rule deletes a sentence that only establishes
    there is nothing to report."""
    facts = _base_facts()
    facts["company_profile_check"]["charges"] = copy.deepcopy(_LIVE_CHARGES)
    facts["charge_movement"] = {
        "checked": True,
        "changes": [{"type": "satisfied",
                     "text": "SATISFIED: the charge of Rs 3,491,110.00 to HDFC BANK LIMITED now "
                             "carries a closure date of 2026-09-30."}],
        "note": "Read from an MCA mirror, which lags the Registrar.",
    }
    text = _all_text(_render(facts, "internal", "movement"))
    assert "SATISFIED" in text, "charge movement never reached the page"
    assert "lags the Registrar" in text, "the mirror caveat was dropped"

    quiet = _base_facts()
    quiet["company_profile_check"]["charges"] = copy.deepcopy(_LIVE_CHARGES)
    quiet["charge_movement"] = {"checked": True, "changes": [], "still_open": [],
                                "note": "Read from an MCA mirror, which lags the Registrar."}
    quiet_text = _all_text(_render(quiet, "internal", "movementquiet"))
    assert "lags the Registrar" not in quiet_text,         "a no-change note was printed; Section B deletes a sentence that reports nothing"
    print("test_charge_movement_reaches_the_page_but_only_when_something_moved: PASS")


def test_the_group_sweep_is_opt_in_and_says_nothing_when_off():
    """It queries several state portals in sequence, so it is off unless
    asked for. Off must mean the section is ABSENT, not present-and-empty --
    an empty coverage table would imply authorities were searched."""
    import company_charter as charter

    assert charter._safe_group_rera_sweep({"found": True}, "X", enabled=False) == {}
    facts = _base_facts()
    facts["group_rera_sweep"] = {}
    assert "Group Projects on Other State Registers" not in _all_text(
        _render(facts, "internal", "sweepoff"))
    print("test_the_group_sweep_is_opt_in_and_says_nothing_when_off: PASS")


def test_every_stage_this_session_added_is_reachable_from_the_pipeline():
    """Anti-drift guard, and the reason this file exists: each of these was
    at some point computed correctly and rendered nowhere. A stage must be
    (a) written into facts by run_company_charter and (b) consumed by a
    render section -- neither half is enough on its own."""
    import inspect

    import company_charter as charter

    source = inspect.getsource(charter)
    for key, section in (
        ("promoter_identity_check", "_append"),          # identity table row
        ("charge_movement", "_append_secured_borrowing_section"),
        ("state_footprint", "_append_state_footprint_section"),
        ("group_rera_sweep", "_append_group_rera_sweep_section"),
    ):
        assert f'facts["{key}"] =' in source, f"{key} is never written into facts"
        assert f'facts.get("{key}")' in source, f"{key} is written but never read back"
        assert section in source, f"{section} vanished"
    # ...and each render section is actually registered to run.
    for section in ("_append_secured_borrowing_section", "_append_state_footprint_section",
                    "_append_group_rera_sweep_section"):
        assert f"lambda: {section}(doc, facts)" in source,             f"{section} exists but is never called during a render"
    print("test_every_stage_this_session_added_is_reachable_from_the_pipeline: PASS")


def _cleanup():
    shutil.rmtree(_SCRATCH, ignore_errors=True)


if __name__ == "__main__":
    try:
        test_a_verified_pan_reaches_the_identity_table_in_both_variants()
        test_the_external_pan_citation_is_generic_not_an_internal_artifact_name()
        test_an_unverified_pan_never_reaches_the_page_and_becomes_a_gap()
        test_no_pan_card_on_file_produces_no_row_and_no_boilerplate_gap()
        test_a_card_that_could_not_be_read_is_a_gap_not_a_silence()
        test_gap_text_carries_no_filenames_keys_or_error_strings()
        test_no_card_and_unreadable_card_are_told_apart_by_status_not_prose()
        test_every_status_the_extractor_emits_is_handled()
        test_the_profile_merge_carries_charges_through()
        test_open_charges_and_lenders_reach_the_page()
        test_unreadable_amounts_report_a_count_and_never_a_zero()
        test_a_promoter_with_no_charges_gets_the_empty_section_form()
        test_a_charge_lookup_that_never_ran_asserts_nothing()
        test_a_declared_out_of_state_project_reaches_the_page()
        test_incorporation_and_operation_are_reported_separately()
        test_llps_are_counted_out_loud_not_quietly_omitted()
        test_an_unrecognised_state_code_is_named_never_guessed()
        test_the_state_comes_from_the_address_name_before_the_pin_code()
        test_cin_state_extraction_handles_llpins_and_junk()
        test_no_footprint_data_renders_no_section()
        test_charge_movement_reaches_the_page_but_only_when_something_moved()
        test_the_group_sweep_is_opt_in_and_says_nothing_when_off()
        test_every_stage_this_session_added_is_reachable_from_the_pipeline()
        print("\nAll tests passed.")
    finally:
        _cleanup()
