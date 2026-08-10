"""
Tests for the three facts-placement corrections: society amalgamation orders
belong under Land Identification not Litigation, the two mortgage fields
collapse to one row, and RERA Core Data's landowner entry is a pointer rather
than a second copy of the ownership chain.

All three are CLAUDE.md Section B rules -- "Ruled-out items survive, in the
right section", "Two table rows must not both exist to report the same
absence", and "Say it once" -- and unlike the clean-check scrubber these are
NOT reversed before the facts file is persisted: a relocation moves text to
the field it always belonged in, so nothing is lost and the corrected record
is the better one.

Run directly: python test_facts_normalization.py
"""

import copy
import json
import os

import docx

import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_normalization")


def _facts() -> dict:
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        return json.load(f)


def _all_text(variant: str) -> str:
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    cc._fill_template("P51800077150", _facts(), out, doc_variant=variant)
    d = docx.Document(out)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _rows(variant: str, label_contains: str) -> list:
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    cc._fill_template("P51800077150", _facts(), out, doc_variant=variant)
    hits = []
    for t in docx.Document(out).tables:
        for r in t.rows:
            if label_contains.lower() in r.cells[0].text.lower():
                hits.append((r.cells[0].text.strip(), r.cells[1].text.strip()))
    return hits


# --- 1. merge orders are land assembly, not litigation ----------------------

def _unnormalized() -> dict:
    """A facts dict in the pre-normalization state these transformations expect.

    Built inline rather than read from the fixture on purpose. The pipeline
    PERSISTS normalized and research-enriched facts back to that same file, so
    a test asserting "the fixture still holds the un-normalized form" is
    asserting a transient property of live data, not the behaviour of the code.
    It passed until the first real re-run and then broke, which is exactly the
    failure mode worth designing out."""
    return {
        "litigation_status": {
            # Sentence order mirrors the real fixture, because it is load-bearing:
            # the Form B clean check sits immediately BEFORE the merge-order
            # sentence, and the period closing "...said Land.'" falls inside a
            # quote so the splitter fuses the two. That fusion is why Form B
            # survived scrubbing until the merge orders were moved out.
            "value": (
                "No litigation is disclosed against the promoter. The Title Report's search found nothing "
                "except a Notice of Lis Pendens dated 20 December 2017 naming an adjoining society. "
                "The promoter's Form B declaration states 'We have no litigation on the said Land.' "
                "The two 6 March 2024 orders sometimes referenced as a 'merge order' are administrative "
                "Deputy Registrar of Co-operative Societies actions (amalgamating Azad Nagar Sai Chhaya CHS "
                "into Azad Nagar Himalaya CHS under the Maharashtra Co-operative Societies Act, 1960), not "
                "litigation or a tribunal proceeding."
            ),
            "source": "Title Report - RERA.pdf",
        },
        "land_identification": {},
        "rera_core_fields": {"promoter_land_owner_investor": "MHADA (original owner) leased the buildings to the societies."},
        "corporate_identity": {"landowner_investor": {"value": "MHADA (original owner) leased the buildings to the societies.", "source": "x.pdf"}},
    }


def test_merge_orders_move_out_of_litigation_status():
    facts = _unnormalized()
    assert "Deputy Registrar" in facts["litigation_status"]["value"], "fixture premise"
    assert cc._relocate_merge_orders(facts) is True

    assert "Deputy Registrar" not in facts["litigation_status"]["value"]
    moved = facts["land_identification"]["land_assembly"]["value"]
    assert "Deputy Registrar of Co-operative Societies" in moved
    # Reframed on arrival: the denial only made sense in the section it left.
    assert "not litigation or a tribunal proceeding" not in moved, moved
    # No stray quote from the preceding sentence's closing "...said Land.'"
    assert not moved.startswith(("'", '"')), moved
    print("test_merge_orders_move_out_of_litigation_status: PASS")


def test_relocation_is_idempotent():
    facts = _unnormalized()
    assert cc._relocate_merge_orders(facts) is True
    before = json.dumps(facts, sort_keys=True)
    assert cc._relocate_merge_orders(facts) is False
    assert json.dumps(facts, sort_keys=True) == before
    print("test_relocation_is_idempotent: PASS")


def test_relocating_first_lets_the_scrubber_reach_the_form_b_declaration():
    """Order matters, which is why _fill_template normalizes before scrubbing.
    The Form B "we have no litigation" declaration is a clean check, but the
    sentence splitter merges it with the merge-order sentence (the preceding
    period sits inside a quote), so it survived scrubbing until the merge
    orders were moved out from under it."""
    facts = _unnormalized()
    cc._relocate_merge_orders(facts)
    cc._scrub_clean_checks(facts)
    remaining = facts["litigation_status"]["value"]
    assert "no litigation on the said Land" not in remaining, remaining
    assert "Lis Pendens" in remaining, "the real finding must survive"
    print("test_relocating_first_lets_the_scrubber_reach_the_form_b_declaration: PASS")


def test_merge_orders_appear_exactly_once_in_the_document():
    """Asserts PLACEMENT rather than a global phrase count.

    A count over the whole document is the wrong test twice over: the
    per-finding research stage rewrites this passage, and the Sources list
    legitimately names the Deputy Registrar orders as source documents. What
    actually matters is that the orders describe land assembly and no longer
    sit under Litigation Status being denied."""
    for variant in ("internal", "external"):
        rows = _rows(variant, "Land assembly")
        assert len(rows) == 1, f"{variant}: expected one Land assembly row, got {len(rows)}"
        assert "Deputy Registrar of Co-operative Societies" in rows[0][1], rows[0][1]

        blob = _all_text(variant)
        assert "not litigation or a tribunal proceeding" not in blob, variant
    print("test_merge_orders_appear_exactly_once_in_the_document: PASS")


# --- 2. one mortgage row, carrying the finding ------------------------------

def test_mortgage_renders_as_a_single_row():
    for variant in ("internal", "external"):
        rows = _rows(variant, "mortgage")
        assert len(rows) == 1, f"{variant}: expected one mortgage row, got {[r[0] for r in rows]}"
        label, value = rows[0]
        assert label == "Mortgage / charge on the land", label
        # "A live right is a finding" -- the unexercised contractual permission
        # is the reason this row exists at all.
        assert "contractually permit the developer" in value, value
    print("test_mortgage_renders_as_a_single_row: PASS")


def test_merged_mortgage_prefers_the_finding_and_dedupes_a_shared_opening():
    facts = {"sources": []}
    metrics = {
        "mortgage_area": "None disclosed. Agreements permit a charge on the free-sale area.",
        "mortgage_lender": {"value": "None disclosed. Agreements permit a charge on the free-sale area.", "source": ""},
    }
    text, _ = cc._merged_mortgage_value(facts, metrics)
    assert text.count("Agreements permit a charge") == 1, text
    print("test_merged_mortgage_prefers_the_finding_and_dedupes_a_shared_opening: PASS")


def test_merged_mortgage_survives_both_fields_being_absent():
    text, citation = cc._merged_mortgage_value({"sources": []}, {})
    assert text and citation is None, (text, citation)
    print("test_merged_mortgage_survives_both_fields_being_absent: PASS")


def test_mortgage_lender_field_is_not_destroyed():
    """_diff_mortgage_lender compares this run's lender against the prior run,
    so the field has to keep existing -- the merge is a render-time concern."""
    facts = _facts()
    cc._normalize_misfiled_facts(facts)
    assert isinstance(facts["fsi_metrics"].get("mortgage_lender"), dict)
    print("test_mortgage_lender_field_is_not_destroyed: PASS")


# --- 3. the landowner chain is stated once ----------------------------------

def test_rera_core_landowner_becomes_a_pointer():
    facts = _unnormalized()
    assert cc._point_rera_landowner_at_identity_table(facts) is True
    pointer = facts["rera_core_fields"]["promoter_land_owner_investor"]
    assert pointer.startswith("See the Corporate/Promoter Identity table")
    assert "MHADA" not in pointer, "a pointer must not restate the chain"
    # idempotent
    assert cc._point_rera_landowner_at_identity_table(facts) is False
    print("test_rera_core_landowner_becomes_a_pointer: PASS")


def test_the_ownership_chain_is_not_told_twice():
    for variant in ("internal", "external"):
        blob = _all_text(variant)
        assert blob.count("MHADA (original owner)") <= 1, variant
        assert "See the Corporate/Promoter Identity table" in blob, variant
    print("test_the_ownership_chain_is_not_told_twice: PASS")


def test_pointer_is_skipped_when_there_is_nothing_to_point_at():
    """Never leave a pointer to a row that carries no content."""
    facts = {"rera_core_fields": {"promoter_land_owner_investor": "Some chain."}, "corporate_identity": {}}
    assert cc._point_rera_landowner_at_identity_table(facts) is False
    assert facts["rera_core_fields"]["promoter_land_owner_investor"] == "Some chain."
    print("test_pointer_is_skipped_when_there_is_nothing_to_point_at: PASS")


# --- the summary pointer must resolve ---------------------------------------

def test_clean_insolvency_summary_does_not_promise_appendix_detail():
    """The Diligence Appendix section collapses to "Nothing found." on a clean
    check, so the summary line must not send the reader there."""
    blob = _all_text("internal")
    assert "IBBI insolvency status: clean" in blob
    idx = blob.index("IBBI insolvency status: clean")
    line = blob[idx:blob.index("\n", idx)]
    assert "Diligence Appendix" not in line, line
    assert "ASSIGNMENT NOT APPROVED YET" in line, "the raw status must still be glossed in place"
    print("test_clean_insolvency_summary_does_not_promise_appendix_detail: PASS")


if __name__ == "__main__":
    test_merge_orders_move_out_of_litigation_status()
    test_relocation_is_idempotent()
    test_relocating_first_lets_the_scrubber_reach_the_form_b_declaration()
    test_merge_orders_appear_exactly_once_in_the_document()
    test_mortgage_renders_as_a_single_row()
    test_merged_mortgage_prefers_the_finding_and_dedupes_a_shared_opening()
    test_merged_mortgage_survives_both_fields_being_absent()
    test_mortgage_lender_field_is_not_destroyed()
    test_rera_core_landowner_becomes_a_pointer()
    test_the_ownership_chain_is_not_told_twice()
    test_pointer_is_skipped_when_there_is_nothing_to_point_at()
    test_clean_insolvency_summary_does_not_promise_appendix_detail()
    print("\nAll tests passed.")
