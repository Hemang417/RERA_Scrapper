"""
Executive Briefing -- an 8-page narrative companion to the 15-section Company
Charter (charter_document.py), in the Cover / Executive Summary / Context /
Primary Analysis / Supporting Analysis / Key Data / Implications / Conclusion
flow, using Key Statistic Card / Quote Highlight / Donut Chart widgets.

This is a SEPARATE document, not a replacement. The Charter stays the single
evidentiary source of truth -- every claim-by-claim table, sub-metric score
and source citation lives there. This briefing narrates only what the Charter
already established (its rule-derived closing verdict, its rule-classified
flags, its recomputed scores) for a reader who wants the headline read
without the full appendix. It performs no new interpretation of raw data:
every sentence traces to a value or a paragraph the Charter's own functions
already computed.

Deliberately built as its own module rather than an extra mode on the shared
prose-and-table builder (charter_report.py's _Builder, and the one
charter_document.py used to carry before it was retired): that shape assumes a
running prose-and-table document inherited from the firm's docx template. A
card/quote/chart layout
is a different visual grammar (cards need a table used as a layout grid, a
chart needs a rendered image dropped in) and forcing both shapes through one
Builder would complicate the class for both callers. This module still reuses
every fact-shaping function from charter_document.py and company_charter.py
directly -- rollups, flag classification, the closing-verdict rule engine --
so the numbers in this briefing and in the Charter can never disagree.

No Integrow branding: the brand assets (logo, design-standards.md,
brand-colours.md, voice-and-examples.md) referenced by the
company-report-generator skill do not exist in this installation (confirmed
by directory listing). Plain formatting, matching the External Charter PDF.
"""

import io

import docx
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import charter_document as cd
import company_charter as cc

_CARD_BORDER = "BFBFBF"
_CARD_FILL = "F2F2F2"


# --- low-level widgets ------------------------------------------------------

def _set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _heading(doc, text, size=20, color="1F3864", space_after=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, size=size, bold=True, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _eyebrow(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    _set_run(run, size=10, bold=True, color="7F7F7F")
    p.paragraph_format.space_after = Pt(2)
    return p


def _body(doc, text, size=11, italic=False, space_after=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _key_statistic_cards(doc, cards):
    """Renders a row of "Key Statistic Card" widgets: a single-row table,
    one column per card, each a big bold figure over a small label. `cards`
    is a list of (figure, label, accent_hex_or_None)."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc._set_table_borders(table)
    for i, (figure, label, accent) in enumerate(cards):
        cell = table.rows[0].cells[i]
        cc._shade_cell(cell, _CARD_FILL)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(figure)
        _set_run(run, size=22, bold=True, color=accent or "1F3864")
        label_p = cell.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_p.add_run(label)
        _set_run(label_run, size=9, color="595959")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def _quote_highlight(doc, text, attribution=None):
    """A pulled-quote block: larger italic text with a left accent border,
    approximated in docx via an indented, shaded single-cell table (python-docx
    has no native pull-quote element)."""
    table = doc.add_table(rows=1, cols=1)
    cc._set_table_borders(table)
    cell = table.rows[0].cells[0]
    cc._shade_cell(cell, "EDEDF7")
    cell.paragraphs[0].paragraph_format.left_indent = Inches(0.15)
    run = cell.paragraphs[0].add_run(f"“{text}”")
    _set_run(run, size=13, italic=True, color="1F3864")
    if attribution:
        attr_p = cell.add_paragraph()
        attr_run = attr_p.add_run(f"— {attribution}")
        _set_run(attr_run, size=9, color="595959")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def _donut_chart(doc, labels, values, colors, title, width_inches=3.2):
    """Renders a donut chart with matplotlib and drops it in as an image --
    python-docx has no native chart object. Slices with a zero value are
    dropped rather than drawn as invisible slivers."""
    normalised_colors = [c if c.startswith("#") else f"#{c}" for c in colors]
    pairs = [(l, v, c) for l, v, c in zip(labels, values, normalised_colors) if v]
    if not pairs:
        _body(doc, f"{title}: no data to chart this pass.", italic=True)
        return
    plot_labels, plot_values, plot_colors = zip(*pairs)
    fig, ax = plt.subplots(figsize=(3.6, 3.0), dpi=200)
    wedges, _texts, autotexts = ax.pie(
        plot_values, colors=plot_colors, startangle=90,
        wedgeprops={"width": 0.4, "edgecolor": "white"},
        autopct=lambda pct: f"{pct:.0f}%" if pct >= 5 else "",
        pctdistance=0.8,
    )
    for t in autotexts:
        t.set_fontsize(8)
        t.set_color("white")
    fig.legend(wedges, plot_labels, loc="lower center", ncol=1, fontsize=7, frameon=False,
               bbox_to_anchor=(0.5, -0.05))
    ax.set_title(title, fontsize=10, color="#1F3864", pad=6)
    ax.axis("equal")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", transparent=True)
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_inches))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    cc._set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        cc._shade_cell(header_cells[i], cc._FILL_NEUTRAL)
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            text, color = value if isinstance(value, tuple) else (value, None)
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(text or ""))
            run.font.size = Pt(10)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def _page_footer_note(doc, page_label):
    p = doc.add_paragraph()
    run = p.add_run(
        f"Executive Briefing -- {page_label}. Summarises findings already established in the full Company "
        f"Charter; no new analysis was performed for this document."
    )
    _set_run(run, size=8, italic=True, color="A6A6A6")


# --- page builders -----------------------------------------------------------

def _page_1_cover(doc, facts, reg_no, generated_on):
    core = facts.get("rera_core_fields") or {}
    corp = facts.get("corporate_identity") or {}
    promoter = cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", "")) or "[Unknown]"
    project = core.get("project_name", "[Unknown]")

    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EXECUTIVE BRIEFING")
    _set_run(run, size=32, bold=True, color="1F3864")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Counterparty & Collateral Due-Diligence Summary")
    _set_run(run, size=14, color="595959")

    doc.add_paragraph()
    detail = doc.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = detail.add_run(f"Counterparty: {promoter}\nCollateral: {project} ({reg_no})")
    _set_run(run, size=12, bold=True)

    doc.add_paragraph()
    generated = doc.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = generated.add_run(f"Generated {generated_on}")
    _set_run(run, size=10, color="7F7F7F")

    for _ in range(6):
        doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "This briefing is a summary companion to the full 15-section Company Charter. It narrates the "
        "Charter's already-verified findings; the Charter itself carries the claim-by-claim evidence, "
        "full sub-metric scoring and source list."
    )
    _set_run(run, size=9, italic=True, color="A6A6A6")


def _page_2_executive_summary(doc, facts, claim_rows, developer_score, doc_confidence_display, verdict_text):
    doc.add_page_break()
    _eyebrow(doc, "Page 1 of 7")
    _heading(doc, "Executive Summary")

    confirmed = sum(1 for r in claim_rows if r["status"].startswith("Confirmed"))
    independent = sum(1 for r in claim_rows if r["status"] == cd.STATUS_CONFIRMED_INDEPENDENT)
    imminent_count = len(cc._classify_flags(facts).get("imminent") or [])

    if doc_confidence_display["stale"]:
        dc_figure, dc_label, dc_sentence = "Not rated", "Documentation Confidence (superseded scheme)", (
            "documentation confidence could not be rated this pass, since the stored score is from a "
            "superseded scoring scheme"
        )
    else:
        overall, band = doc_confidence_display["overall"], doc_confidence_display["band"]
        dc_figure, dc_label = f"{overall}/100", f"Documentation Confidence ({band})"
        dc_sentence = f"documentation confidence in this pass is {overall}/100 ({band})"

    _body(
        doc,
        f"Of {len(claim_rows)} material claims checked against this counterparty and this project, "
        f"{confirmed} are confirmed against a primary document or an independent registry, {independent} of "
        f"those against a registry independent of the promoter. The Developer Score composite is "
        f"{developer_score.get('composite')}/100 (grade {developer_score.get('grade')}); {dc_sentence}."
    )

    _key_statistic_cards(doc, [
        (f"{developer_score.get('composite')}/100", f"Developer Score (grade {developer_score.get('grade')})",
         cc._TEXT_RED if developer_score.get("grade") in ("C", "D") else cc._TEXT_GREEN),
        (dc_figure, dc_label, None),
        (f"{confirmed}/{len(claim_rows)}", "Claims confirmed", None),
        (str(imminent_count), "Imminent flags", cc._TEXT_RED if imminent_count else cc._TEXT_GREEN),
    ])

    _quote_highlight(doc, verdict_text, attribution="Closing Read, derived by rule (see page 3)")
    _page_footer_note(doc, "Executive Summary")


def _page_3_context(doc, facts):
    doc.add_page_break()
    _eyebrow(doc, "Page 2 of 7")
    _heading(doc, "Context & Background")

    core = facts.get("rera_core_fields") or {}
    corp = facts.get("corporate_identity") or {}
    land = facts.get("land_identification") or {}

    _heading(doc, "The Counterparty", size=13, color="404040", space_after=4)
    _simple_table(doc, ["Field", "As Recorded"], [
        ["Promoter", cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", ""))],
        ["Organisation type", (corp.get("organization_type") or {}).get("value", "Not established")],
        ["CIN / LLPIN", (corp.get("cin_llpin") or {}).get("value", "Not established")],
        ["Registered office", (corp.get("registered_office_main") or {}).get("value", "Not established")],
    ])

    _heading(doc, "The Collateral", size=13, color="404040", space_after=4)
    _simple_table(doc, ["Field", "As Recorded"], [
        ["Project", core.get("project_name", "Not established")],
        ["Authority", core.get("authority", "Not established")],
        ["Registration status", core.get("project_status", "Not established")],
        ["Plan approval", core.get("plan_approval_number", "Not established")],
        ["Proposed completion", core.get("proposed_completion_date", "Not established")],
        ["Village / locality", (land.get("village_locality") or {}).get("value", "Not established")],
    ])
    _page_footer_note(doc, "Context & Background")


def _page_4_primary_analysis(doc, assessment):
    doc.add_page_break()
    _eyebrow(doc, "Page 3 of 7")
    _heading(doc, "Primary Analysis: Genuine Developer or Single-Project Vehicle?")
    _body(
        doc,
        "The single most consequential question in this diligence pass is whether the counterparty is a "
        "going concern with a track record, or a single-project vehicle standing behind one registration. "
        "The read below is derived by rule from named signals, not written as an opinion; each row traces "
        "back to a fact recorded in the full Charter."
    )
    _simple_table(doc, ["Signal", "Finding"], [[label, finding] for label, finding in assessment["signals"]])
    _quote_highlight(doc, assessment["verdict"])
    _page_footer_note(doc, "Primary Analysis")


def _page_5_supporting_analysis(doc, developer_score, dc_buckets, dc_stale):
    doc.add_page_break()
    _eyebrow(doc, "Page 4 of 7")
    _heading(doc, "Supporting Analysis: Scoring Breakdown")

    _heading(doc, "Developer Score, by Bucket", size=13, color="404040", space_after=4)
    ds_buckets = cd.rollup_developer_score_buckets(developer_score)
    _simple_table(doc, ["Bucket", "Weight", "Rating", "Coverage"], [
        [b["bucket"], f"{b['weight']}%", b["rating"] or "Not rated", f"{b['scored']} of {b['total']}"]
        for b in ds_buckets
    ])
    _donut_chart(
        doc,
        labels=[b["bucket"] for b in ds_buckets],
        values=[b["weight"] for b in ds_buckets],
        colors=["1F3864", "8FAADC", "BF8F00"],
        title="Developer Score weight by bucket",
    )

    _heading(doc, "Documentation Confidence, by Bucket", size=13, color="404040", space_after=4)
    if dc_stale:
        # Same staleness guard as charter_document.py's section 1: this score
        # cannot be recomputed here (it needs an authenticity_summary built
        # inside _fill_template), so a persisted copy is used, and if none of
        # its criteria map onto the current buckets it is from a superseded
        # scheme -- say so rather than render a figure that no longer means
        # what it says.
        _body(
            doc,
            "Not rated: the stored Documentation Confidence Score is from a superseded scoring scheme and "
            "cannot be mapped onto the current criteria. Regenerate the Charter to score it.",
            italic=True,
        )
    elif dc_buckets:
        _simple_table(doc, ["Bucket", "Weight", "Band", "Coverage"], [
            [b["bucket"], f"{b['weight']}%", b["band"] or "Not rated", f"{b['scored']} of {b['total']}"]
            for b in dc_buckets
        ])
    else:
        _body(doc, "Not rated: no Documentation Confidence Score was recorded this pass.", italic=True)
    _page_footer_note(doc, "Supporting Analysis")


def _page_6_key_data(doc, claim_rows):
    doc.add_page_break()
    _eyebrow(doc, "Page 5 of 7")
    _heading(doc, "Key Data & Evidence")
    _body(
        doc,
        "Every material claim checked this pass, grouped by what it was verified against. A claim with no "
        "recorded source is excluded rather than shown as an empty row."
    )

    status_order = [
        (cd.STATUS_CONFIRMED_INDEPENDENT, cc._TEXT_GREEN),
        (cd.STATUS_CONFIRMED_DOCUMENT, "1F3864"),
        (cd.STATUS_PROMOTER_FILED, cc._TEXT_AMBER),
        (cd.STATUS_DISCREPANCY, cc._TEXT_RED),
        (cd.STATUS_STATED_ONLY, cc._TEXT_AMBER),
        (cd.STATUS_NOT_ESTABLISHED, "808080"),
    ]
    counts = {status: sum(1 for r in claim_rows if r["status"] == status) for status, _ in status_order}
    _donut_chart(
        doc,
        labels=[s for s, _ in status_order],
        values=[counts[s] for s, _ in status_order],
        colors=[c for _, c in status_order],
        title="Material claims by verification status",
    )
    _simple_table(doc, ["Status", "Count"], [[s, str(counts[s])] for s, _ in status_order if counts[s]])
    _page_footer_note(doc, "Key Data & Evidence")


def _page_7_implications(doc, facts, flags, next_steps):
    doc.add_page_break()
    _eyebrow(doc, "Page 6 of 7")
    _heading(doc, "Implications")

    imminent = flags.get("imminent") or []
    structural = flags.get("structural") or []

    _heading(doc, f"What Needs Attention ({len(imminent)} imminent, {len(structural)} structural)",
             size=13, color="404040", space_after=4)
    if imminent or structural:
        rows = [[("Imminent", cc._TEXT_RED), item["text"]] for item in imminent]
        rows += [[("Structural", cc._TEXT_AMBER), item["text"]] for item in structural]
        _simple_table(doc, ["Severity", "Finding"], rows)
    else:
        _body(doc, "No imminent or structural flag was raised this pass.", italic=True)

    green = cd._green_flags(facts)
    if green:
        _heading(doc, "What Checks Out", size=13, color="404040", space_after=4)
        for point in green:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(point)
            _set_run(run, size=10, color=cc._TEXT_GREEN)

    if next_steps:
        _heading(doc, "Recommended Verification Steps", size=13, color="404040", space_after=4)
        for step in next_steps:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(step)
            _set_run(run, size=10)
    _page_footer_note(doc, "Implications")


def _page_8_conclusion(doc, assessment, reg_no):
    doc.add_page_break()
    _eyebrow(doc, "Page 7 of 7")
    _heading(doc, "Conclusion")
    _quote_highlight(doc, assessment["verdict"])
    _body(
        doc,
        "This briefing summarises findings already established in the full Company Charter; it performs no "
        "new analysis of its own. For claim-by-claim sourcing, the complete sub-metric scoring detail, and "
        f"the Harvard-style source list, see the companion Company Charter for registration {reg_no}."
    )
    _page_footer_note(doc, "Conclusion")


# --- entry point -------------------------------------------------------------

def build_executive_briefing(reg_no: str, facts: dict, out_path: str, generated_on: str) -> None:
    """Builds the 8-page unbranded Executive Briefing PDF companion to the
    full Company Charter. `generated_on` is passed in (e.g. "30 July 2026")
    rather than computed here, so the caller controls the date source.

    Always recomputes developer_score: facts["developer_score"] is an output
    of a prior run, never a valid input. charter_report.py's builder makes the
    same guarantee for the same reason."""
    doc = docx.Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    flags = cc._classify_flags(facts)
    developer_score = cc._compute_developer_score(facts, flags)
    claim_rows = cd.build_claim_rows(facts)
    assessment = cd.assess_counterparty(facts, flags, developer_score)
    next_steps = cd._recommended_steps(facts, flags)

    doc_confidence = facts.get("documentation_confidence_score") or {}
    dc_buckets = cd.rollup_doc_confidence_buckets(doc_confidence) if doc_confidence.get("criteria") else []
    dc_stale = bool(dc_buckets) and not any(bucket["scored"] for bucket in dc_buckets)
    if dc_stale:
        dc_buckets = []
    doc_confidence_display = {
        "overall": doc_confidence.get("overall") if dc_buckets else None,
        "band": doc_confidence.get("band") if dc_buckets else None,
        "stale": dc_stale,
    }

    _page_1_cover(doc, facts, reg_no, generated_on)
    _page_2_executive_summary(doc, facts, claim_rows, developer_score, doc_confidence_display, assessment["verdict"])
    _page_3_context(doc, facts)
    _page_4_primary_analysis(doc, assessment)
    _page_5_supporting_analysis(doc, developer_score, dc_buckets, dc_stale)
    _page_6_key_data(doc, claim_rows)
    _page_7_implications(doc, facts, flags, next_steps)
    _page_8_conclusion(doc, assessment, reg_no)

    doc.save(out_path)
