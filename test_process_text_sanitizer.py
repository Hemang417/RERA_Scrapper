"""
Tests for _sanitize_process_text / _sanitize_process_gaps.

CLAUDE.md Section B pulls two ways here, and both have to hold at once:

  * "Internal keeps process failures, External does not" -- so a
    re-verification that could not run STAYS in the Internal document, and
    deleting it is wrong;
  * "Never put a file path, module name, function or parameter name, JSON key,
    or raw exception string into either document" -- so what stays has to be
    rewritten rather than passed through.

Deleting the item would satisfy the second rule and break the first. Every test
below therefore checks a leak is gone AND the information it carried survives.

Run directly: python test_process_text_sanitizer.py
"""

import json
import os
import re

import docx

import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_sanitizer")


def _facts() -> dict:
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        return json.load(f)


def _all_text(variant: str) -> str:
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    cc._fill_template("P51800077150", _facts(), out, doc_variant=variant)
    d = docx.Document(out)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# --- the four rewrites -------------------------------------------------------

def test_a_json_key_label_becomes_a_readable_field_name():
    out = cc._sanitize_process_text("promoter_name: Acme Realty Limited (not re-verified).")
    assert out.startswith("Promoter name: "), out
    assert "promoter_name" not in out
    assert "Acme Realty Limited" in out, "the value itself must survive"
    print("test_a_json_key_label_becomes_a_readable_field_name: PASS")


def test_a_dotted_path_label_is_reduced_to_its_last_segment():
    out = cc._sanitize_process_text("land_identification.total_gross_area: details not found verbatim.")
    assert out.startswith("Total gross area: "), out
    assert "land_identification" not in out
    print("test_a_dotted_path_label_is_reduced_to_its_last_segment: PASS")


def test_known_initialisms_are_not_mangled():
    assert cc._sanitize_process_text("cin_llpin: U70109MH2022PLC385473 (x).").startswith("CIN / LLPIN: ")
    print("test_known_initialisms_are_not_mangled: PASS")


def test_a_quoted_api_exception_becomes_plain_language():
    raw = ('cin_llpin: U123 (NOT independently re-verified this pass -- verification could not run: '
           '"Could not resolve authentication method. Expected one of api_key, auth_token" -- treat as unconfirmed)')
    out = cc._sanitize_process_text(raw)
    assert "Could not resolve authentication" not in out, out
    assert "api_key" not in out and "auth_token" not in out, out
    assert "the verification step could not run" in out, out
    assert "treat as unconfirmed" in out, "the actionable part must survive"
    print("test_a_quoted_api_exception_becomes_plain_language: PASS")


def test_a_run_output_path_is_described_not_printed():
    raw = "CTS land-record lookup: candidates are in output/P51800077150/cts_office_candidates.json. Pick the office."
    out = cc._sanitize_process_text(raw)
    assert "output/P51800077150" not in out, out
    assert ".json" not in out, out
    # The sentence must still end properly -- the path regex used to swallow
    # the full stop, producing "...run output Pick the office."
    assert "run output. Pick the office." in out, out
    print("test_a_run_output_path_is_described_not_printed: PASS")


def test_a_module_name_becomes_the_step_it_performs():
    out = cc._sanitize_process_text("Then run cts_resolve.py to continue.")
    assert ".py" not in out, out
    assert "cts resolve step" in out, out
    print("test_a_module_name_becomes_the_step_it_performs: PASS")


def test_real_prose_labels_are_left_alone():
    """"Cross-corroboration:" and "CTS land-record lookup:" are prose, not keys.
    The key pattern is lowercase-and-underscores only so it cannot eat them."""
    for text in (
        "Cross-corroboration: the 'pricing' topic rests on one source.",
        "CTS land-record lookup: still pending.",
    ):
        assert cc._sanitize_process_text(text).startswith(text.split(":")[0] + ":"), text
    print("test_real_prose_labels_are_left_alone: PASS")


# --- gap-level behaviour -----------------------------------------------------

def test_sanitizing_gaps_is_reversible_and_idempotent():
    facts = _facts()
    before = json.dumps(facts["gaps"])
    original = cc._sanitize_process_gaps(facts)
    assert original is not None, "fixture must contain process-failure gaps"
    assert json.dumps(facts["gaps"]) != before
    assert cc._sanitize_process_gaps(facts) is None, "second pass must be a no-op"
    facts["gaps"] = original
    assert json.dumps(facts["gaps"]) == before, "restore was not lossless"
    print("test_sanitizing_gaps_is_reversible_and_idempotent: PASS")


def test_no_gap_is_lost_or_emptied():
    facts = _facts()
    count = len(facts["gaps"])
    cc._sanitize_process_gaps(facts)
    assert len(facts["gaps"]) == count, "a gap was dropped"
    assert all(g.strip() for g in facts["gaps"]), "a gap was emptied"
    print("test_no_gap_is_lost_or_emptied: PASS")


# --- the rendered documents --------------------------------------------------

def test_neither_document_leaks_internals():
    for variant in ("internal", "external"):
        blob = _all_text(variant)
        assert "output/P51800077150" not in blob, f"{variant}: raw path"
        assert "Could not resolve authentication" not in blob, f"{variant}: raw exception"
        assert not re.search(r"\b\w+\.py\b", blob), f"{variant}: module name"
        assert not re.search(r"^\s*.?\s*[a-z][a-z0-9_]*(\.[a-z_]+)+\s*:", blob, re.M), f"{variant}: key label"
    print("test_neither_document_leaks_internals: PASS")


def test_internal_still_keeps_the_process_failures_themselves():
    """The point of sanitizing rather than deleting: Section B says these stay
    in Internal. A pass that silently removed them would break that rule while
    appearing to satisfy the leak checks."""
    blob = _all_text("internal")
    assert "the verification step could not run" in blob, "process failures were deleted, not sanitized"
    assert "treat as unconfirmed" in blob
    print("test_internal_still_keeps_the_process_failures_themselves: PASS")


if __name__ == "__main__":
    test_a_json_key_label_becomes_a_readable_field_name()
    test_a_dotted_path_label_is_reduced_to_its_last_segment()
    test_known_initialisms_are_not_mangled()
    test_a_quoted_api_exception_becomes_plain_language()
    test_a_run_output_path_is_described_not_printed()
    test_a_module_name_becomes_the_step_it_performs()
    test_real_prose_labels_are_left_alone()
    test_sanitizing_gaps_is_reversible_and_idempotent()
    test_no_gap_is_lost_or_emptied()
    test_neither_document_leaks_internals()
    test_internal_still_keeps_the_process_failures_themselves()
    print("\nAll tests passed.")
