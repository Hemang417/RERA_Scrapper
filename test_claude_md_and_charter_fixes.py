"""
Tests for the CLAUDE.md three-tier loader infrastructure (Task 0) and the
Company Charter content-quality fixes built on top of it (Tasks 1-6):

- Section A/B/C loaders return non-empty, correctly-scoped content, and
  Section A never reaches any constructed Claude API request while Section
  C reaches only external-doc-variant requests.
- The citation-per-bullet-clause fix (the actual root cause behind
  Litigation Status bullets showing zero or all citations).
- Inline CIN/DIN stripping from narrative prose (kept in table cells).
- The multi-column empty-row removal for Group/Affiliated Companies-style
  tables.
- Bullet clause consolidation and capitalization.
- First-use jargon expansion.

Run directly: python test_claude_md_and_charter_fixes.py
"""

import docx
import pytest

import company_charter as cc
import deep_research


# --- Task 0: CLAUDE.md loaders ---------------------------------------------

def test_loaders_return_non_empty_correctly_scoped_content():
    section_a = cc._coding_time_notes()
    section_b = cc._common_content_rules()
    section_c = cc._external_citation_rule()

    assert section_a.strip() and section_b.strip() and section_c.strip()

    # Scoping sanity: each section's own distinguishing rule text shows up
    # only where it should, not smeared across all three.
    assert "charter_document.py" in section_a  # coding-time-only concern
    assert "Deep-dive vs. nothing-found" in section_b
    assert "Numbered citations" in section_c
    assert "charter_document.py" not in section_b
    assert "charter_document.py" not in section_c
    assert "Numbered citations" not in section_a
    assert "Numbered citations" not in section_b
    print("test_loaders_return_non_empty_correctly_scoped_content: PASS")


def _system_texts(system) -> str:
    """Flattens a system param (str, or a list of {"type": "text", "text":
    ...} blocks) into one string, so a test can search it regardless of
    which shape a given call site used."""
    if isinstance(system, str):
        return system
    return "\n".join(block.get("text", "") for block in system)


def test_section_a_never_reaches_any_constructed_api_request(monkeypatch):
    """Section A is coding-time-only documentation -- it must never appear
    in the `system` param built for _run_charter_pass or the citation-
    completeness judge, the two charter-specific Claude API calls this
    file makes."""
    captured = {}

    def fake_run_agentic_pass(user_prompt, system, label="agentic_pass"):
        captured[label] = system
        if label == "charter_pass":
            return {"gaps": []}
        return {"uncited_claims": []}

    monkeypatch.setattr(deep_research, "_run_agentic_pass", fake_run_agentic_pass)

    cc._run_charter_pass("irrelevant test prompt")

    doc = docx.Document()
    doc.add_paragraph("A plain sentence with nothing notable in it.")
    tmp_path = "output/company_charters/_test_scratch_citation_judge.docx"
    import os
    os.makedirs(os.path.dirname(tmp_path), exist_ok=True)
    doc.save(tmp_path)
    cc._llm_verify_citation_completeness(tmp_path)
    os.remove(tmp_path)

    coding_time_marker = "Never edit charter_document.py to change live charter output"
    for label, system in captured.items():
        assert coding_time_marker not in _system_texts(system), (
            f"Section A leaked into the '{label}' API request"
        )
    print("test_section_a_never_reaches_any_constructed_api_request: PASS")


def test_section_c_reaches_only_external_variant_requests(monkeypatch):
    """_run_charter_pass produces variant-agnostic facts (shared by both
    Internal and External renders) and must NEVER get Section C. The
    citation-completeness judge only ever runs against the rendered
    External document and must ALWAYS get Section C."""
    captured = {}

    def fake_run_agentic_pass(user_prompt, system, label="agentic_pass"):
        captured[label] = system
        if label == "charter_pass":
            return {"gaps": []}
        return {"uncited_claims": []}

    monkeypatch.setattr(deep_research, "_run_agentic_pass", fake_run_agentic_pass)

    cc._run_charter_pass("irrelevant test prompt")
    assert "Numbered citations" not in _system_texts(captured["charter_pass"])

    doc = docx.Document()
    doc.add_paragraph("A plain sentence with nothing notable in it.")
    import os
    tmp_path = "output/company_charters/_test_scratch_citation_judge2.docx"
    os.makedirs(os.path.dirname(tmp_path), exist_ok=True)
    doc.save(tmp_path)
    cc._llm_verify_citation_completeness(tmp_path)
    os.remove(tmp_path)
    assert "Numbered citations" in _system_texts(captured["citation_completeness_judge"])
    print("test_section_c_reaches_only_external_variant_requests: PASS")


def test_common_content_rules_are_cache_marked_for_charter_calls():
    """_charter_system_blocks marks Section B as a cacheable prompt prefix
    -- the one thing a plain-string system prompt could never do."""
    blocks = cc._charter_system_blocks(external=False)
    assert blocks[0]["cache_control"] == {"type": "ephemeral"}
    assert blocks[0]["text"] == cc._common_content_rules()
    print("test_common_content_rules_are_cache_marked_for_charter_calls: PASS")


# --- Task 1: citation-per-bullet-clause fix ---------------------------------

def test_citation_attaches_to_every_bullet_clause_not_just_the_last():
    """The actual reported bug: a citation used to be glued once onto the
    whole multi-sentence blob BEFORE splitting, so only the last resulting
    bullet ever carried it. Now every clause gets it."""
    doc = docx.Document(cc.TEMPLATE_PATH)
    paragraph = doc.paragraphs[15]
    facts = {"_deferred_bullets": []}
    text = "First finding happened in 2020. Second finding happened in 2022. Third finding is still pending."
    cc._set_paragraph_as_bullets(facts, paragraph, text, citation="[1]")

    _, clauses = facts["_deferred_bullets"][0]
    assert len(clauses) == 3
    for clause in clauses:
        assert clause.endswith("[1]"), f"clause missing citation: {clause!r}"
    print("test_citation_attaches_to_every_bullet_clause_not_just_the_last: PASS")


def test_cite_marker_returns_bare_marker_not_concatenated_text():
    facts = {
        "_doc_variant": "external",
        "_citation_registry": {"order": [], "index": {}},
        "sources": [{"topic": "distance", "label": "Google Maps", "accessed_date": "2026-01-01"}],
    }
    marker = cc._cite_marker("distance", facts=facts)
    assert marker == "[1]"
    print("test_cite_marker_returns_bare_marker_not_concatenated_text: PASS")


# --- Task 4: inline CIN/DIN stripping ---------------------------------------

def test_strip_inline_cin_din_removes_value_keeps_other_parenthetical_content():
    text = "Bijay Kumar Agarwal (DIN 00448678, appointed 2022-09-07) is a current director."
    cleaned = cc._strip_inline_cin_din(text)
    assert "DIN" not in cleaned
    assert "00448678" not in cleaned
    assert "appointed 2022-09-07" in cleaned
    print("test_strip_inline_cin_din_removes_value_keeps_other_parenthetical_content: PASS")


def test_strip_inline_cin_din_removes_cin_value():
    text = "Vador Properties Private Limited (CIN U70102MH2010PTC210775, found via ZaubaCorp) is unrelated."
    cleaned = cc._strip_inline_cin_din(text)
    assert "CIN" not in cleaned
    assert "U70102MH2010PTC210775" not in cleaned
    assert "found via ZaubaCorp" in cleaned
    print("test_strip_inline_cin_din_removes_cin_value: PASS")


def test_table_cells_keep_cin_din_via_raw_setter():
    """_set_cell must bypass the stripper entirely -- the identity tables
    are the one place CIN/DIN is required."""
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    cc._set_cell(table, 1, 0, "CIN U70102MH2010PTC210775")
    assert "U70102MH2010PTC210775" in table.rows[1].cells[0].text
    print("test_table_cells_keep_cin_din_via_raw_setter: PASS")


# --- Task 5: multi-column empty-row removal ---------------------------------

def test_remove_fully_empty_rows_keeps_row_with_any_real_value():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=3)
    row1 = table.add_row()
    row1.cells[0].text = "Vador Properties Private Limited"
    row1.cells[1].text = "unknown"
    row1.cells[2].text = "shared director: Bijay Kumar Agarwal"
    row2 = table.add_row()
    row2.cells[0].text = "unknown"
    row2.cells[1].text = "unknown"
    row2.cells[2].text = ""

    cc._remove_fully_empty_rows(table)

    assert len(table.rows) == 2  # header + the one real row
    assert table.rows[1].cells[0].text == "Vador Properties Private Limited"
    print("test_remove_fully_empty_rows_keeps_row_with_any_real_value: PASS")


# --- Task 6: bullet consolidation + capitalization --------------------------

def test_consolidate_bullet_clauses_merges_short_trailing_fragment():
    clauses = ["The claim was filed in the Bombay High Court in 2023.", "Confirmed."]
    merged = cc._consolidate_bullet_clauses(clauses)
    assert merged == ["The claim was filed in the Bombay High Court in 2023. Confirmed."]
    print("test_consolidate_bullet_clauses_merges_short_trailing_fragment: PASS")


def test_consolidate_bullet_clauses_keeps_two_standalone_points():
    clauses = [
        "The claim was filed in the Bombay High Court in 2023.",
        "It remains pending as of the most recent cause-list check in 2026.",
    ]
    merged = cc._consolidate_bullet_clauses(clauses)
    assert merged == clauses
    print("test_consolidate_bullet_clauses_keeps_two_standalone_points: PASS")


def test_fix_bullet_capitalization_capitalizes_and_adds_terminal_punctuation():
    assert cc._fix_bullet_capitalization("this is a fragment") == "This is a fragment."
    assert cc._fix_bullet_capitalization("Already fine.") == "Already fine."
    print("test_fix_bullet_capitalization_capitalizes_and_adds_terminal_punctuation: PASS")


# --- Task 3: jargon expansion ------------------------------------------------

def test_expand_jargon_first_use_expands_once_keeps_term():
    text = "The promoter is under CIRP. The CIRP process began in 2024."
    expanded = cc._expand_jargon_first_use(text)
    assert "insolvency proceedings (CIRP)" in expanded
    assert expanded.count("CIRP") == 2  # first expanded, second left bare
    print("test_expand_jargon_first_use_expands_once_keeps_term: PASS")


# --- Task 1 gate-check heuristic ---------------------------------------------

def test_looks_like_uncited_factual_claim_flags_bare_registry_facts():
    assert cc._looks_like_uncited_factual_claim("Registered in 2020 under CIN U70102MH2010PTC210775.")
    assert not cc._looks_like_uncited_factual_claim("Registered in 2020 under CIN U70102MH2010PTC210775. [1]")
    assert not cc._looks_like_uncited_factual_claim("Nothing found.")
    print("test_looks_like_uncited_factual_claim_flags_bare_registry_facts: PASS")


# --- PDF conversion: graceful degradation -----------------------------------

def test_convert_docx_to_pdf_degrades_to_none_on_failure(monkeypatch):
    """No Word installed / a COM error must never crash generation -- the
    .docx stays the only output for that file, and the caller decides how
    to log it."""
    # docx2pdf drives Word via COM and only installs on Windows -- skip
    # rather than fail so the suite still runs on Linux/CI.
    docx2pdf = pytest.importorskip("docx2pdf")

    def fake_convert(*args, **kwargs):
        raise RuntimeError("Word is not installed")

    monkeypatch.setattr(docx2pdf, "convert", fake_convert)
    result = cc._convert_docx_to_pdf("some/path/Charter_External.docx")
    assert result is None
    print("test_convert_docx_to_pdf_degrades_to_none_on_failure: PASS")


def test_convert_docx_to_pdf_returns_sibling_pdf_path_on_success(monkeypatch):
    docx2pdf = pytest.importorskip("docx2pdf")

    captured = {}

    def fake_convert(src, dst):
        captured["src"], captured["dst"] = src, dst

    monkeypatch.setattr(docx2pdf, "convert", fake_convert)
    result = cc._convert_docx_to_pdf("some/path/Charter_External.docx")
    assert result == "some/path/Charter_External.pdf"
    assert captured["dst"] == "some/path/Charter_External.pdf"
    print("test_convert_docx_to_pdf_returns_sibling_pdf_path_on_success: PASS")


if __name__ == "__main__":
    test_loaders_return_non_empty_correctly_scoped_content()
    test_common_content_rules_are_cache_marked_for_charter_calls()
    test_citation_attaches_to_every_bullet_clause_not_just_the_last()
    test_cite_marker_returns_bare_marker_not_concatenated_text()
    test_strip_inline_cin_din_removes_value_keeps_other_parenthetical_content()
    test_strip_inline_cin_din_removes_cin_value()
    test_table_cells_keep_cin_din_via_raw_setter()
    test_remove_fully_empty_rows_keeps_row_with_any_real_value()
    test_consolidate_bullet_clauses_merges_short_trailing_fragment()
    test_consolidate_bullet_clauses_keeps_two_standalone_points()
    test_fix_bullet_capitalization_capitalizes_and_adds_terminal_punctuation()
    test_expand_jargon_first_use_expands_once_keeps_term()
    test_looks_like_uncited_factual_claim_flags_bare_registry_facts()
    print("\nAll non-monkeypatch tests passed (run under pytest for the monkeypatch-based tests).")
