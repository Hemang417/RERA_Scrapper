"""
Tests for CLAUDE.md Section C -- numbered per-clause citations in the External
Company Charter.

Four things are pinned here, all of which shipped broken in the 2026-08
Pranami Bliss charter:
  * factual bullets carry a marker (28% coverage before this work),
  * a marker sits at the end of the clause it supports and never mid-token
    (the rule cites "MahaRERA[10]-registered" by name; that exact string
    was in the shipped document),
  * Sources entries are descriptive citations, never a raw internal filename
    ("Declaration for one registration 18 March 24.pdf" reached a client) and
    never a bare category label ("Project record", "appeals data"),
  * the numbering has no orphans and no dangling markers.

Two deliberate NON-goals, both of which look like missing coverage and are not:
gaps stay uncited, because Section B says never to attach a citation marker to
an absence; and a clause whose supporting source cannot be identified stays
uncited rather than borrowing an adjacent one, because Section C says the
marker must actually support the claim.

Run directly: python test_external_citations.py
"""

import json
import os
import re

import docx

import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_citations")

# Bullets that legitimately carry no marker. Gaps and qualifiers report an
# absence or a limitation, which Section B forbids citing.
_UNCITED_EXEMPT = (
    "rest on a single source", "not independently field-verified",
)


def _gap_prefixes() -> list:
    """First words of every gap, for recognising a gap-derived flag bullet.

    Section B forbids attaching a citation marker to an absence, so gaps are
    exempt from Section C's coverage rule. Two forms have to be recognised: a
    gap material enough to be printed in External carries a "(Gap N)" pointer,
    while a monitor-only gap is rendered in full with no pointer at all (see
    _flag_headline). Matching against the gap list itself catches both without
    depending on how any particular gap happens to be worded."""
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        gaps = json.load(f).get("gaps", []) or []
    return [g.strip()[:55] for g in gaps if g.strip()]


def _is_gap_bullet(text: str, prefixes: list) -> bool:
    body = text.strip().lstrip("• ").strip()
    if re.search(r"\(Gap \d+\)\s*$", body):
        return True
    return any(p and body.startswith(p) for p in prefixes)


def _paragraphs(variant: str) -> list:
    os.makedirs(_SCRATCH, exist_ok=True)
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        facts = json.load(f)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    cc._fill_template("P51800077150", facts, out, doc_variant=variant)
    d = docx.Document(out)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return parts


def _bullets(paras: list) -> list:
    return [p for p in paras if p.strip().startswith("•")]


def _sources(paras: list) -> list:
    starts = [i for i, t in enumerate(paras) if t.strip().lower().startswith("sources")]
    if not starts:
        return []
    return [t.strip() for t in paras[starts[-1]:] if re.match(r"^\s*\[\d+\]", t.strip())]


# --- coverage ----------------------------------------------------------------

def test_external_factual_bullets_are_cited():
    paras = _paragraphs("external")
    prefixes = _gap_prefixes()
    uncited = [
        b for b in _bullets(paras)
        if not re.search(r"\[\d+\]", b)
        and not _is_gap_bullet(b, prefixes)
        and not any(x in b for x in _UNCITED_EXEMPT)
    ]
    assert not uncited, "uncited factual bullets in External:\n" + "\n".join(f"  - {u[:110]}" for u in uncited)
    print("test_external_factual_bullets_are_cited: PASS")


def test_coverage_is_substantially_complete():
    """A floor, so a regression that silently drops citations is caught even if
    it happens to dodge the exemption list above. Was 28% before this work."""
    bullets = _bullets(_paragraphs("external"))
    cited = [b for b in bullets if re.search(r"\[\d+\]", b)]
    ratio = len(cited) / len(bullets)
    assert ratio >= 0.75, f"External citation coverage fell to {ratio:.0%} ({len(cited)}/{len(bullets)})"
    print(f"test_coverage_is_substantially_complete: PASS ({ratio:.0%})")


def test_internal_is_not_given_numbered_markers():
    """Section C applies to External only, and says so explicitly: "Does not
    apply to the Internal document at all"."""
    paras = _paragraphs("internal")
    numbered = [b for b in _bullets(paras) if re.search(r"\[\d+\]", b)]
    assert not numbered, numbered[:3]
    print("test_internal_is_not_given_numbered_markers: PASS")


# --- placement ---------------------------------------------------------------

def test_no_marker_lands_mid_token():
    """"MahaRERA[10]-registered" is the failure the rule names by example."""
    offenders = [p for p in _paragraphs("external") if re.search(r"\[\d+\][A-Za-z-]", p)]
    assert not offenders, offenders[:3]
    print("test_no_marker_lands_mid_token: PASS")


def test_marker_goes_to_the_clause_end_not_next_to_the_keyword():
    text = "Promoter portfolio: 1 MahaRERA-registered project tied to this promoter. More text."
    out = cc._insert_marker_at_clause_end(text, text.index("MahaRERA") + len("MahaRERA"), "[4]")
    assert "MahaRERA[4]-registered" not in out, out
    assert "tied to this promoter[4]." in out, out
    print("test_marker_goes_to_the_clause_end_not_next_to_the_keyword: PASS")


# --- source labels -----------------------------------------------------------

def test_sources_are_descriptive_never_filenames_or_categories():
    entries = _sources(_paragraphs("external"))
    assert entries, "External must render a Sources list"
    bad = [e for e in entries if re.search(r"\.pdf\b|\.json\b|Project record|appeals data", e, re.IGNORECASE)]
    assert not bad, "non-descriptive Sources entries:\n" + "\n".join(f"  - {b}" for b in bad)
    print(f"test_sources_are_descriptive_never_filenames_or_categories: PASS ({len(entries)} entries)")


def test_a_catalogued_document_resolves_to_its_descriptive_ref():
    facts = {"sources": [{
        "label": "Title Report - RERA.pdf",
        "ref": "Legal Title Report and Title Certificate, Adv. Preet J. Chheda, 18 April 2024",
    }]}
    assert cc._external_source_label(facts, "Title Report - RERA.pdf").startswith("Legal Title Report")
    # and still matches once _topic_citation has appended its accessed-date suffix
    assert cc._external_source_label(facts, "Title Report - RERA.pdf, accessed 2026-08-09").startswith("Legal Title Report")
    print("test_a_catalogued_document_resolves_to_its_descriptive_ref: PASS")


def test_an_uncatalogued_document_never_renders_as_project_record():
    assert cc._generic_one_label("6.pdf") != "Project record"
    assert ".pdf" not in cc._generic_one_label("6.pdf")
    print("test_an_uncatalogued_document_never_renders_as_project_record: PASS")


# --- numbering integrity -----------------------------------------------------

def test_no_orphan_sources_and_no_dangling_markers():
    paras = _paragraphs("external")
    listed = {int(m.group(1)) for e in _sources(paras) if (m := re.match(r"^\s*\[(\d+)\]", e))}
    used = {int(n) for p in paras for n in re.findall(r"\[(\d+)\]", p)}
    assert not (listed - used), f"sources listed but never cited: {sorted(listed - used)}"
    assert not (used - listed), f"markers with no Sources entry: {sorted(used - listed)}"
    print(f"test_no_orphan_sources_and_no_dangling_markers: PASS ({len(listed)} sources)")


# --- the marker must support the claim ---------------------------------------

def test_clause_topic_resolution_picks_a_supporting_topic():
    facts = {
        "_doc_variant": "external",
        "_citation_registry": {"order": [], "index": {}},
        "sources": [
            {"label": "t.pdf", "ref": "Title deed", "topic": "land_title"},
            {"label": "c.pdf", "ref": "CRISIL rating rationale", "topic": "credit_rating"},
        ],
    }
    title = cc._clause_topic_citation(facts, "The sale deed conveys the plot to the society.")
    rating = cc._clause_topic_citation(facts, "CRISIL downgraded the group entity.")
    assert title and rating and title != rating, (title, rating)
    print("test_clause_topic_resolution_picks_a_supporting_topic: PASS")


def test_an_unrecognisable_clause_stays_uncited_rather_than_mis_cited():
    """Section C's worked failure: citing the MahaRERA complaints record for a
    sentence beginning "independent web research found...". Both real sources,
    wrong pairing. A missing marker is the better failure."""
    facts = {
        "_doc_variant": "external",
        "_citation_registry": {"order": [], "index": {}},
        "sources": [{"label": "c.pdf", "ref": "CRISIL rating rationale", "topic": "credit_rating"}],
    }
    assert cc._clause_topic_citation(facts, "The weather that quarter was unremarkable.") is None
    print("test_an_unrecognisable_clause_stays_uncited_rather_than_mis_cited: PASS")


# --- External drops process failures -----------------------------------------

def test_external_drops_process_failures_internal_keeps_them():
    """Section B: a re-verification that could not run is a tooling failure,
    not a fact about the project. It also carries raw JSON keys, a file path
    and an API exception string, none of which may appear in either document."""
    ext = "\n".join(_paragraphs("external"))
    internal = "\n".join(_paragraphs("internal"))
    for leak in ("output/P51800077150", "promoter_name:", "Could not resolve authentication"):
        assert leak not in ext, f"External leaked process text: {leak!r}"
    # Internal keeps the process-failure ITEM, but not the raw exception string
    # inside it: Section B forbids that in EITHER document, so it is rewritten
    # rather than passed through (see _sanitize_process_text). This used to
    # assert the raw exception survived in Internal, which was the leak.
    assert "Could not resolve authentication" not in internal, "raw exception text must not survive anywhere"
    assert "the verification step could not run" in internal, "Internal must still keep the process failure itself"
    print("test_external_drops_process_failures_internal_keeps_them: PASS")


def test_single_source_finding_survives_as_one_consolidated_line():
    """The finding inside those gaps -- that a topic rests on one source -- is
    real and must be preserved, consolidated, with the error text dropped."""
    gaps = [
        "Cross-corroboration: the 'credit_rating' topic is backed by only one source (X); auth error text.",
        "Cross-corroboration: the 'pricing' topic is backed by only one source (Y); auth error text.",
        "A genuine unrelated gap.",
    ]
    out = cc._external_gaps(gaps)
    consolidated = [g for g in out if "rest on a single source" in g]
    assert len(consolidated) == 1, out
    assert "credit rating" in consolidated[0] and "pricing" in consolidated[0], consolidated[0]
    assert "A genuine unrelated gap." in out
    print("test_single_source_finding_survives_as_one_consolidated_line: PASS")


if __name__ == "__main__":
    test_external_factual_bullets_are_cited()
    test_coverage_is_substantially_complete()
    test_internal_is_not_given_numbered_markers()
    test_no_marker_lands_mid_token()
    test_marker_goes_to_the_clause_end_not_next_to_the_keyword()
    test_sources_are_descriptive_never_filenames_or_categories()
    test_a_catalogued_document_resolves_to_its_descriptive_ref()
    test_an_uncatalogued_document_never_renders_as_project_record()
    test_no_orphan_sources_and_no_dangling_markers()
    test_clause_topic_resolution_picks_a_supporting_topic()
    test_an_unrecognisable_clause_stays_uncited_rather_than_mis_cited()
    test_external_drops_process_failures_internal_keeps_them()
    test_single_source_finding_survives_as_one_consolidated_line()
    print("\nAll tests passed.")
