"""
Tests for charter_document.py -- the restructured Counterparty + Collateral
Charter builder.

Run directly: python test_charter_document.py
"""

import json
import os
import shutil

import docx

import charter_document as cd
import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_restructure")


def _load_facts():
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        return json.load(f)


# --- bucket rollups -------------------------------------------------------

def test_developer_bucket_rating_uses_only_what_scored():
    """A bucket's rating renormalises over the sub-metrics that actually
    scored, so a half-measured bucket isn't dragged to D by absent data --
    while the composite separately keeps every weight fixed. Both readings
    are honest; they answer different questions, and the unscored names must
    travel with the bucket so they can be flagged (the no-N/A rule)."""
    developer_score = {
        "composite": 25.0, "grade": "C",
        "criteria": {
            # Operational: 2 of 4 scored, both AAA -> bucket reads AAA.
            "team_strength": {"score": None, "weight": 12.5, "display_name": "Team Strength"},
            "area_within_5km": {"score": None, "weight": 12.5, "display_name": "Influence in Micromarket"},
            "past_area_developed": {"score": 100.0, "weight": 12.5, "display_name": "Past Experience - Area"},
            "track_record_years": {"score": 100.0, "weight": 12.5, "display_name": "Track Record"},
            "financial_strength_debt": {"score": None, "weight": 20.0, "display_name": "Financial Strength"},
            "rera_compliance": {"score": 50.0, "weight": 7.5, "display_name": "RERA Compliance"},
            "gst_compliance": {"score": None, "weight": 7.5, "display_name": "GST Compliance"},
            "past_default_count": {"score": 50.0, "weight": 7.5, "display_name": "Cases"},
            "entity_rating": {"score": 50.0, "weight": 7.5, "display_name": "Entity Rating"},
        },
    }
    buckets = {b["bucket"]: b for b in cd.rollup_developer_score_buckets(developer_score)}

    ops = buckets["Operational Strength"]
    assert ops["score"] == 100.0, ops              # renormalised over the 2 that scored
    assert ops["rating"] == "AAA", ops
    assert ops["scored"] == 2 and ops["total"] == 4, ops
    assert ops["unscored_display_names"] == ["Team Strength", "Influence in Micromarket"], ops

    fin = buckets["Financial Strength"]
    assert fin["score"] is None and fin["rating"] is None, fin   # nothing scored -> no fake grade
    assert fin["unscored_display_names"] == ["Financial Strength"], fin

    gov = buckets["Governance Strength"]
    assert gov["score"] == 50.0, gov
    assert gov["unscored_display_names"] == ["GST Compliance"], gov
    print("test_developer_bucket_rating_uses_only_what_scored: PASS")


def test_doc_confidence_buckets_sum_existing_weights_without_reassigning():
    """Bucket weight is just the sum of its criteria's existing weights, so
    the bucketed table and the detailed breakdown can never disagree."""
    doc_confidence = {
        "overall": 42.7, "band": "Limited",
        "criteria": {
            "source_tier_quality": {"score": 59.0, "weight": 26.7},
            "primary_tier_density": {"score": 30.0, "weight": 20.0},
            "cross_corroboration": {"score": 12.5, "weight": 20.0},
            "financial_figures_confirmed": {"score": 25.0, "weight": 13.3},
            "completeness_rate": {"score": 52.6, "weight": 6.7},
            "recency_legal": {"score": 100.0, "weight": 6.7},
            "recency_other": {"score": 75.0, "weight": 6.7},
        },
        "skipped_criteria": ["verification_rate"],
    }
    buckets = {b["bucket"]: b for b in cd.rollup_doc_confidence_buckets(doc_confidence)}
    assert buckets["Source Quality"]["weight"] == 46.7, buckets["Source Quality"]
    assert buckets["Coverage & Recency"]["weight"] == 20.1, buckets["Coverage & Recency"]
    # An explicitly skipped criterion is named, not silently dropped.
    assert buckets["Corroboration & Confirmation"]["unscored_display_names"] == ["Independent Verification Rate"]
    assert buckets["Corroboration & Confirmation"]["scored"] == 2
    total_weight = sum(b["weight"] for b in buckets.values())
    assert abs(total_weight - 100.1) < 0.05, total_weight
    print("test_doc_confidence_buckets_sum_existing_weights_without_reassigning: PASS")


# --- claim classification -------------------------------------------------

def test_maharera_filing_is_promoter_stated_not_independent():
    """The core distinction the whole claim table rests on: MahaRERA is an
    official venue, but what it holds is the promoter's OWN filing. An MCA
    mirror or a government land record is a third party. Collapsing those
    into one "verified" bucket would overstate the evidence."""
    _pos, status, independent = cd.classify_claim_evidence("MahaRERA project record -- https://maharerait.maharashtra.gov.in/x")
    assert status == cd.STATUS_PROMOTER_FILED and independent is False

    _pos, status, independent = cd.classify_claim_evidence("MCA-mirror company registration profile (ZaubaCorp, Tofler)")
    assert status == cd.STATUS_CONFIRMED_INDEPENDENT and independent is True

    _pos, status, independent = cd.classify_claim_evidence("Maha Bhulekh Property Card, CTS 183")
    assert status == cd.STATUS_CONFIRMED_INDEPENDENT and independent is True

    _pos, status, _ = cd.classify_claim_evidence("output/P51800077150/documents/Title Report - RERA.pdf")
    assert status == cd.STATUS_CONFIRMED_DOCUMENT

    _pos, status, _ = cd.classify_claim_evidence("gap -- see Gaps section")
    assert status == cd.STATUS_NOT_ESTABLISHED

    _pos, status, _ = cd.classify_claim_evidence("")
    assert status == cd.STATUS_NOT_ESTABLISHED
    print("test_maharera_filing_is_promoter_stated_not_independent: PASS")


def test_prose_gap_mentioning_a_field_name_does_not_flag_a_discrepancy():
    """Regression for a false positive that fired immediately on real data: a
    model-authored gap describing a PORTFOLIO search ("matched only this one
    project under this exact promoter name") flipped the promoter_name claim
    to Discrepancy despite saying nothing about that claim's verification.
    Only a MACHINE-written per-field gap ("<field_key>: ...") may override a
    status."""
    facts = {
        "corporate_identity": {"promoter_name": {"value": "Test Promoter Ltd", "source": "MahaRERA project record"}},
        "gaps": ["promoter_portfolio's totals are unavailable -- matched only one project under this exact promoter name."],
    }
    row = cd.build_claim_rows(facts)[0]
    assert row["status"] == cd.STATUS_PROMOTER_FILED, row

    # The machine-written shape DOES override.
    facts["gaps"] = ["promoter_name: Test Promoter Ltd (verification: not supported by the cited source)"]
    row = cd.build_claim_rows(facts)[0]
    assert row["status"] == cd.STATUS_DISCREPANCY, row
    print("test_prose_gap_mentioning_a_field_name_does_not_flag_a_discrepancy: PASS")


def test_claim_rows_only_include_sourced_fields():
    """A group can hold plain strings with no source of their own (fsi_metrics
    carries five). A claim with no recorded source cannot honestly be given a
    verification status, so it must not appear in a claim-versus-evidence
    table at all."""
    facts = {
        "fsi_metrics": {
            "mortgage_lender": {"value": "Some Bank", "source": "output/x/documents/deed.pdf"},
            "implied_fsi": "3.0",            # plain string, no source
            "approved_bua": "12,345 sq m",   # plain string, no source
        },
    }
    rows = cd.build_claim_rows(facts)
    assert len(rows) == 1, rows
    assert rows[0]["claim"] == "Mortgage lender", rows
    print("test_claim_rows_only_include_sourced_fields: PASS")


def test_field_display_names_are_not_mangled():
    """Regression: an earlier version applied replacements then .capitalize(),
    which lowercased them again and rendered "CIN / LLPIN" as "Cin / llpin"."""
    assert cd._field_display_name("cin_llpin") == "CIN / LLPIN"
    assert cd._field_display_name("survey_cts_plot_numbers") == "Survey / CTS / plot numbers"
    assert cd._field_display_name("implied_fsi") == "Implied FSI"
    assert cd._field_display_name("village_locality") == "Village locality"  # plain fallback
    print("test_field_display_names_are_not_mangled: PASS")


# --- External-variant quality-gate fixes -----------------------------------

def test_prose_generalises_document_library_to_document_list_in_external():
    """Regression: real model-authored prose from the fixture refers to the
    OLD section name ("...confirmed present in the Document Library but not
    individually opened...") -- section 11 is "Document & Diligence Trail"
    now, and the phrase reads as internal-process naming, exactly what the
    External quality gate's Document-Library check exists to catch. Must be
    generalised (case-insensitive, word-boundary) rather than pattern-matched
    against one project's exact sentence. Internal keeps the wording as-is."""
    b = cd._Builder(docx.Document(), {}, {"_doc_variant": "external"})
    out = b.prose("Confirmed present in the Document Library but not individually opened.")
    assert "Document Library" not in out, out
    assert "document list" in out, out

    b_internal = cd._Builder(docx.Document(), {}, {"_doc_variant": "internal"})
    out_internal = b_internal.prose("Confirmed present in the Document Library but not individually opened.")
    assert "Document Library" in out_internal, out_internal
    print("test_prose_generalises_document_library_to_document_list_in_external: PASS")


def test_prose_replaces_dashes_with_semicolons_in_external_only():
    """Last-resort fallback for arbitrary model-authored prose the curated
    _externalize_prose dictionary doesn't cover -- the External quality gate
    hard-forbids " -- " / " — " / "—", so anything the dictionary misses must
    still be caught here. Internal is untouched."""
    b = cd._Builder(docx.Document(), {}, {"_doc_variant": "external"})
    out = b.prose("Two figures were found -- the second could not be reconciled.")
    assert " -- " not in out and "; " in out, out

    b_internal = cd._Builder(docx.Document(), {}, {"_doc_variant": "internal"})
    out_internal = b_internal.prose("Two figures were found -- the second could not be reconciled.")
    assert " -- " in out_internal, out_internal
    print("test_prose_replaces_dashes_with_semicolons_in_external_only: PASS")


def test_green_flag_color_is_allowed_in_external():
    """_section_12_consolidated_flags renders "What Checks Out" bullets in
    _TEXT_GREEN (375623) to complete the red/amber/green traffic-light
    convention alongside the pre-existing red/amber. The External quality
    gate's run-color allowlist must include it, or every green bullet in a
    real document is flagged as an unexpected color -- as it was the first
    time this table was built (7 violations)."""
    assert "375623" in cc._EXTERNAL_ALLOWED_RUN_COLORS
    print("test_green_flag_color_is_allowed_in_external: PASS")


def test_developer_score_table_drops_weight_column_in_external_only():
    """Regression: _section_14_scoring_detail's Developer Score table
    originally rendered a Weight column unconditionally, the 9th and last of
    the External quality-gate violations found when this section was first
    built -- every other weighted table in the document already drops Weight
    for External (a numeric weight per row reads as internal methodology,
    not a finding), and this one had been missed. The Documentation
    Confidence table right below it is a deliberate exception and keeps
    Weight in both variants."""
    facts = _load_facts()
    flags = cc._classify_flags(facts)
    developer_score = cc._compute_developer_score(facts, flags)

    doc_internal = docx.Document()
    b_internal = cd._Builder(doc_internal, cd._capture_styles(docx.Document()), {"_doc_variant": "internal"})
    cd._section_14_scoring_detail(b_internal, developer_score)
    doc_external = docx.Document()
    b_external = cd._Builder(doc_external, cd._capture_styles(docx.Document()), {"_doc_variant": "external"})
    cd._section_14_scoring_detail(b_external, developer_score)

    def dev_score_header(doc):
        for t in doc.tables:
            header = [c.text.strip() for c in t.rows[0].cells]
            if header[:2] == ["Bucket", "Sub-metric"]:
                return header
        raise AssertionError("Developer Score table not found")

    assert "Weight" in dev_score_header(doc_internal)
    assert "Weight" not in dev_score_header(doc_external)
    print("test_developer_score_table_drops_weight_column_in_external_only: PASS")


# --- document build -------------------------------------------------------

def test_builds_both_variants_and_external_passes_the_quality_gate():
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    for variant in ("internal", "external"):
        out = os.path.join(_SCRATCH, f"{variant}.docx")
        cd.build_charter_document("P51800077150", dict(facts), out, doc_variant=variant)
        assert os.path.exists(out)
        doc = docx.Document(out)
        # sectPr carries page size/orientation/margins; clearing the body must
        # never remove it or Word silently reverts to its own defaults.
        assert any(c.tag.endswith("}sectPr") for c in doc.element.body), f"{variant} lost its sectPr"
        assert len(doc.tables) >= 3, f"{variant} should have the claim table plus both rating tables"

    violations = cc._verify_external_document_quality(os.path.join(_SCRATCH, "external.docx"))
    assert violations == [], violations
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_builds_both_variants_and_external_passes_the_quality_gate: PASS")


def test_never_trusts_a_stale_persisted_developer_score():
    """The bug this guards against was live in the real fixture: its persisted
    developer_score was from the superseded flat 7-criteria scheme (composite
    100.0 / grade A / weights 33.3, no buckets, no rera_compliance or
    gst_compliance), so reading it instead of recomputing rendered 100.0/A
    against a true 39.6/C and reported scored sub-metrics as "Not rated".
    facts["developer_score"] is an OUTPUT of a prior run, never an input."""
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    facts["developer_score"] = {
        "composite": 100.0, "grade": "A",
        "criteria": {"track_record_years": {"score": 100.0, "tier": "AAA", "weight": 33.3}},
    }
    out = os.path.join(_SCRATCH, "stale.docx")
    cd.build_charter_document("P51800077150", facts, out, doc_variant="external")

    text = "\n".join(p.text for p in docx.Document(out).paragraphs)
    assert "Composite 100.0/100" not in text, "stale persisted composite was rendered"
    assert "grade A" not in text, "stale persisted grade was rendered"
    assert "Composite 39.6/100, grade C" in text, text[:400]
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_never_trusts_a_stale_persisted_developer_score: PASS")


def test_stale_doc_confidence_is_flagged_not_rendered():
    """The Documentation Confidence Score cannot be recomputed here (it needs
    an authenticity_summary built inside _fill_template), so a persisted copy
    is used. If none of its criteria map onto the current buckets it is from a
    superseded scheme, and saying so beats printing a figure that no longer
    means what it says."""
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    os.makedirs(_SCRATCH, exist_ok=True)
    facts = _load_facts()
    facts["documentation_confidence_score"] = {
        "overall": 88.0, "band": "High",
        "criteria": {"some_retired_criterion": {"score": 88.0, "weight": 100.0}},
    }
    out = os.path.join(_SCRATCH, "stale_dc.docx")
    cd.build_charter_document("P51800077150", facts, out, doc_variant="external")

    text = "\n".join(p.text for p in docx.Document(out).paragraphs)
    assert "88.0" not in text, "a superseded Documentation Confidence figure was rendered"
    assert "superseded scoring scheme" in text, text[:600]
    shutil.rmtree(_SCRATCH, ignore_errors=True)
    print("test_stale_doc_confidence_is_flagged_not_rendered: PASS")


if __name__ == "__main__":
    test_developer_bucket_rating_uses_only_what_scored()
    test_doc_confidence_buckets_sum_existing_weights_without_reassigning()
    test_maharera_filing_is_promoter_stated_not_independent()
    test_prose_gap_mentioning_a_field_name_does_not_flag_a_discrepancy()
    test_claim_rows_only_include_sourced_fields()
    test_field_display_names_are_not_mangled()
    test_prose_generalises_document_library_to_document_list_in_external()
    test_prose_replaces_dashes_with_semicolons_in_external_only()
    test_green_flag_color_is_allowed_in_external()
    test_developer_score_table_drops_weight_column_in_external_only()
    test_builds_both_variants_and_external_passes_the_quality_gate()
    test_never_trusts_a_stale_persisted_developer_score()
    test_stale_doc_confidence_is_flagged_not_rendered()
    print("\nAll tests passed.")
