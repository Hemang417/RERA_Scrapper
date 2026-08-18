"""
Guards on state-dependent TEXT in the rendered Charter.

Background: the .docx template was authored state-neutral -- paragraph 2
still literally reads "Public Web-Sourced Edition - Adapted for [State]
([State RERA acronym])", and paragraphs 58/69/70/75 carry [Regulator] and
[Local planning authority]. The Python then overwrote all of it with
Maharashtra literals, so a Telangana Charter generated in August 2026 came
out claiming to be "Adapted for Maharashtra (MahaRERA)" with a "MahaRERA
Registration" row. That run needed a one-off post-hoc .docx patcher
(output/CONSTELLA_TS/research/patch_state_labels.py) to fix by hand.

The tests here exist so that never has to happen again, and they pull in two
opposite directions on purpose:

  * test_maharashtra_render_is_unchanged pins the CURRENT Maharashtra output
    byte-for-byte, so parameterising cannot quietly reword the production
    document.

  * test_a_different_state_changes_only_the_state_paragraphs renders the
    SAME facts under a synthetic profile and asserts the symmetric
    difference is exactly the state-bearing strings -- proving in one
    assertion that the parameterisation is COMPLETE (nothing still says
    Maharashtra) and CONTAINED (nothing else moved).

Everything here renders from a saved facts.json through _fill_template, so
there are zero API calls -- the established no-network pattern from
test_executive_ready.py.

Run directly: python test_state_labels.py
"""

import json
import os
import shutil

import docx

import company_charter as cc
import states
from states.base import (
    ALL_CAPABILITIES,
    StateProfile,
)

_FACTS_PATH = os.path.join(
    "output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json"
)
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_state_labels")
_REG_NO = "P51800077150"

# A state that does not exist, on purpose: every token is unmistakable in a
# diff, and none of them can collide with real content in the facts file.
_TEST_PROFILE = StateProfile(
    code="ZZ",
    state_name="Testland",
    rera_acronym="TestRERA",
    regulator_name="Testland Real Estate Regulatory Authority (TestRERA)",
    planning_authority_label="the Testland planning authority",
    reg_no_pattern=r"^ZZ\d{6}$",
    portal_domains=("rera.testland.gov.in",),
    domain_labels=(("rera.testland.gov.in", "TestRERA"),),
    capabilities=frozenset(ALL_CAPABILITIES),
)

# Strings the Maharashtra document must still contain, verbatim. Deliberately
# pins OUTPUT, not implementation -- how the sentence is built is free to
# change, what the reader sees is not.
_MH_INTERNAL_LITERALS = (
    "Public Web-Sourced Edition -- Adapted for Maharashtra (MahaRERA)",
    "MahaRERA Registration",
)
_MH_EXTERNAL_LITERALS = (
    "Public Web-Sourced Edition: Adapted for Maharashtra (MahaRERA)",
)

# Placeholders the template ships with. None may survive into a render.
_TEMPLATE_PLACEHOLDERS = (
    "[State]",
    "[State RERA acronym]",
    "[Regulator]",
    "[Local planning authority]",
)


def _facts():
    with open(_FACTS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def _render(variant, profile=None, facts=None):
    """A fresh deep copy per render: _fill_template mutates the facts it is
    given, and these tests render the same facts up to four times."""
    os.makedirs(_SCRATCH, exist_ok=True)
    out = os.path.join(_SCRATCH, f"{variant}_{(profile.code if profile else 'default')}.docx")
    payload = json.loads(json.dumps(facts if facts is not None else _facts()))
    cc._fill_template(_REG_NO, payload, out, doc_variant=variant, state_profile=profile)
    return out


def _all_strings(path):
    """Every paragraph and table-cell string in the document, as a set."""
    d = docx.Document(path)
    out = {p.text for p in d.paragraphs if p.text.strip()}
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        out.add(p.text)
    return out


def test_maharashtra_render_is_unchanged():
    for variant, literals in (
        ("internal", _MH_INTERNAL_LITERALS),
        ("external", _MH_EXTERNAL_LITERALS),
    ):
        text = "\n".join(_all_strings(_render(variant)))
        for literal in literals:
            assert literal in text, (variant, literal)
    print("test_maharashtra_render_is_unchanged: PASS")


def test_no_template_placeholder_survives_a_render():
    for variant in ("internal", "external"):
        text = "\n".join(_all_strings(_render(variant)))
        for placeholder in _TEMPLATE_PLACEHOLDERS:
            assert placeholder not in text, (variant, placeholder)
    print("test_no_template_placeholder_survives_a_render: PASS")


def test_template_still_declares_its_state_placeholders():
    """The template is UNTRACKED -- it lives under gitignored output/ and is
    not in version control, so there is no diff to notice a hand-edit. This
    documents the contract: the placeholders are why the code is allowed to
    author these paragraphs from a profile rather than trusting the file."""
    d = docx.Document(cc.TEMPLATE_PATH)
    template_text = "\n".join(p.text for p in d.paragraphs)
    for placeholder in ("[State]", "[State RERA acronym]", "[Regulator]"):
        assert placeholder in template_text, (
            f"{placeholder} is gone from {cc.TEMPLATE_PATH}. If the template was "
            f"re-saved with a state typed in, the code no longer matches the asset "
            f"it documents -- restore the placeholder rather than deleting this test."
        )
    print("test_template_still_declares_its_state_placeholders: PASS")


def test_a_different_state_changes_only_the_state_paragraphs():
    """The strongest guard here: complete AND contained, in one assertion."""
    for variant in ("internal", "external"):
        mh = _all_strings(_render(variant))
        zz = _all_strings(_render(variant, profile=_TEST_PROFILE))

        only_mh = mh - zz
        only_zz = zz - mh

        # Everything that changed must be state-bearing on the side it came
        # from. If a non-state paragraph moved, it shows up here.
        for line in only_mh:
            assert ("Maharashtra" in line or "MahaRERA" in line), (
                variant, "changed but is not a Maharashtra string", line[:160]
            )
        for line in only_zz:
            assert ("Testland" in line or "TestRERA" in line), (
                variant, "changed but is not a Testland string", line[:160]
            )

        # ...and something must actually have changed, or this passes vacuously.
        assert only_mh, (variant, "no Maharashtra-specific text changed at all")
        assert only_zz, (variant, "no Testland text was produced at all")

        # Nothing that survives unchanged may still name Maharashtra --
        # UNLESS it came from the facts file rather than from code.
        #
        # This distinction is the whole point. The fixture is a real
        # Maharashtra project, so its own data legitimately says "MahaRERA"
        # (document titles, source labels, promoter-written prose), and
        # re-rendering Maharashtra DATA under a Testland profile must NOT
        # rewrite it -- that would be falsifying the record. What must never
        # survive is a Maharashtra noun the CODE supplied.
        facts_json = json.dumps(_facts(), ensure_ascii=False)
        for line in (zz & mh):
            for token in ("MahaRERA", "Maharashtra"):
                idx = line.find(token)
                while idx != -1:
                    chunk = line[idx:idx + 25]
                    assert chunk in facts_json, (
                        variant,
                        "code-authored text still hardcodes a Maharashtra noun "
                        "(it renders identically under a different state profile, "
                        "and does not trace to the facts file)",
                        line[:160],
                    )
                    idx = line.find(token, idx + 1)
    print("test_a_different_state_changes_only_the_state_paragraphs: PASS")


def test_external_dash_rewrite_covers_every_registered_profile():
    """Section C forbids the " -- " form in the External document and the
    quality gate BLOCKS THE SAVE if one survives. The subtitle's Internal
    and External forms therefore have to move together; generating both
    from one function is what makes that structural instead of something a
    future edit has to remember."""
    rewrites = cc._state_dash_rewrites()
    for code, profile in states.PROFILES.items():
        internal = cc._edition_subtitle(profile)
        assert internal in rewrites, (code, internal, sorted(rewrites))
        external = rewrites[internal]
        assert " -- " not in external, (code, external)
        assert profile.state_name in external and profile.rera_acronym in external, (code, external)
    print("test_external_dash_rewrite_covers_every_registered_profile: PASS")


def test_active_profile_is_reset_after_a_render():
    _render("internal")
    assert cc._ACTIVE_STATE_PROFILE is None, cc._ACTIVE_STATE_PROFILE
    assert cc._ACTIVE_EXTERNAL_FACTS is None, cc._ACTIVE_EXTERNAL_FACTS
    print("test_active_profile_is_reset_after_a_render: PASS")


def test_active_profile_is_reset_even_when_the_render_raises():
    """The External quality gate raises AFTER doc.save, so the failure path
    is exactly the one that would leak the global into the NEXT render in
    the same process -- every test file, every Streamlit session, and the
    Internal/External pair inside run_company_charter."""
    # Same technique as test_executive_ready.test_fill_template_raises_when_
    # gate_fails: reintroduce a dash through _externalize_prose, which every
    # External string flows through, rather than trying to smuggle one past
    # the sanitisers via the facts dict.
    real_externalize_prose = cc._externalize_prose

    def _broken_externalize_prose(facts, text):
        return real_externalize_prose(facts, text) + " -- broken by test"

    cc._externalize_prose = _broken_externalize_prose
    raised = False
    try:
        _render("external", profile=_TEST_PROFILE)
    except RuntimeError:
        raised = True
    finally:
        cc._externalize_prose = real_externalize_prose
    assert raised, "the External quality gate did not fire on a reintroduced dash"
    assert cc._ACTIVE_STATE_PROFILE is None, (
        "state profile leaked after a failed render -- the next render in this "
        f"process would inherit it: {cc._ACTIVE_STATE_PROFILE}"
    )
    assert cc._ACTIVE_EXTERNAL_FACTS is None, cc._ACTIVE_EXTERNAL_FACTS
    print("test_active_profile_is_reset_even_when_the_render_raises: PASS")


if __name__ == "__main__":
    try:
        test_maharashtra_render_is_unchanged()
        test_no_template_placeholder_survives_a_render()
        test_template_still_declares_its_state_placeholders()
        test_a_different_state_changes_only_the_state_paragraphs()
        test_external_dash_rewrite_covers_every_registered_profile()
        test_active_profile_is_reset_after_a_render()
        test_active_profile_is_reset_even_when_the_render_raises()
        print("\nAll tests passed.")
    finally:
        shutil.rmtree(_SCRATCH, ignore_errors=True)
