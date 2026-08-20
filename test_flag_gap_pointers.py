"""
Tests for CLAUDE.md Section B's "Flags summarise, Gaps explain" rule as
implemented by _classify_flags (stable gap numbers), _rendered_gap_numbers
(which gaps a variant prints), _flag_headline (the "(Gap N)" pointer), and
_fill_template's numbered Gaps & Sources list.

The contract in one line: a "(Gap N)" pointer must always resolve to a
"Gap N." entry printed in the SAME document, and a gap's body must not be
restated in full in the flag list when the Gaps entry already carries it.

Run directly: python test_flag_gap_pointers.py
"""

import json
import os
import re

import docx

import company_charter as cc

_PRANAMI_FACTS = os.path.join("output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json")
_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_flag_gaps")


def _all_text(path: str) -> str:
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _printed_gap_numbers(blob: str) -> set:
    """The gap numbers actually printed under Gaps & Sources. Both variants
    render the list inside one paragraph with line breaks between entries,
    which python-docx surfaces as newlines."""
    return {int(n) for n in re.findall(r"^Gap (\d+)\.", blob, re.M)}


def _build(variant: str) -> str:
    os.makedirs(_SCRATCH, exist_ok=True)
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        facts = json.load(f)
    out = os.path.join(_SCRATCH, f"{variant}.docx")
    # A fresh deep copy per build: _fill_template mutates the facts it is given.
    cc._fill_template("P51800077150", json.loads(json.dumps(facts)), out, doc_variant=variant)
    return out


# --- unit level --------------------------------------------------------------

def test_gap_flags_carry_a_stable_one_based_number():
    """The number indexes facts["gaps"] and is assigned once, at classification
    time, so both renderers agree without re-deriving it."""
    facts = {"gaps": ["First gap.", "Second gap.", "Third gap."]}
    flags = cc._classify_flags(facts)
    by_text = {i["text"]: i for tier in flags.values() for i in tier}
    assert by_text["First gap."]["gap_number"] == 1
    assert by_text["Second gap."]["gap_number"] == 2
    assert by_text["Third gap."]["gap_number"] == 3
    print("test_gap_flags_carry_a_stable_one_based_number: PASS")


def test_non_gap_flags_carry_no_gap_number():
    """A flag with no Gaps entry to point at must not get a number -- there
    would be nothing for the pointer to resolve to."""
    facts = {"cts_mismatch_note": "The CTS number on record disagrees with the carried-over intake."}
    flags = cc._classify_flags(facts)
    items = [i for tier in flags.values() for i in tier]
    assert items, "fixture must produce at least one flag"
    assert all("gap_number" not in i for i in items), items
    print("test_non_gap_flags_carry_no_gap_number: PASS")


def _gap_item(facts: dict, number: int = 1) -> dict:
    """The flag carrying gap `number`. Selected by number, not by position:
    _classify_flags also emits non-gap flags for a sparse fixture (no
    registration number, no CTS, no CIN), so indexing [0] picks one of those."""
    for tier in cc._classify_flags(facts).values():
        for item in tier:
            if item.get("gap_number") == number:
                return item
    raise AssertionError(f"no flag carrying gap_number={number}")


def test_headline_takes_the_first_sentence_and_appends_the_pointer():
    facts = {"gaps": ["The first sentence states the finding. The second sentence adds detail."]}
    text, points_at_a_gap = cc._flag_headline(facts, _gap_item(facts))
    assert points_at_a_gap is True
    assert text == "The first sentence states the finding. (Gap 1)", text
    print("test_headline_takes_the_first_sentence_and_appends_the_pointer: PASS")


def test_a_flag_with_no_gap_keeps_its_full_text():
    facts = {"cts_mismatch_note": "The CTS number on record disagrees with the carried-over intake."}
    item = [i for tier in cc._classify_flags(facts).values() for i in tier][0]
    text, points_at_a_gap = cc._flag_headline(facts, item)
    assert points_at_a_gap is False
    assert text == item["text"], text
    print("test_a_flag_with_no_gap_keeps_its_full_text: PASS")


def test_external_prints_monitor_only_gaps_too():
    """rules.md Section B: "Gaps & Sources keeps every item with its full
    explanation, including permanent ones."

    External used to print only gaps that had also earned an Imminent or
    Structural flag. On a project whose gaps were ALL monitor severity that
    subset came out empty, and the section then asserted "No additional
    material gaps identified" directly beneath a dozen unresolved gaps printed
    in full as flags: false, and self-contradicting. Both variants now print
    every gap, so a monitor-only gap gets a pointer in External exactly as it
    does in Internal."""
    monitor_only = "Some minor detail could not be confirmed. It is not material to this project."
    facts = {"gaps": [monitor_only], "_doc_variant": "external"}
    assert cc._rendered_gap_numbers(facts) == {1}, "External must print a monitor-only gap"

    item = _gap_item(facts)
    text, points_at_a_gap = cc._flag_headline(facts, item)
    assert points_at_a_gap is True
    assert text.endswith("(Gap 1)"), text

    # Internal is unchanged, and the two now agree.
    internal = dict(facts, _doc_variant="internal")
    assert cc._rendered_gap_numbers(internal) == {1}
    text_i, points_i = cc._flag_headline(internal, item)
    assert points_i is True and text_i.endswith("(Gap 1)"), text_i
    print("test_external_prints_monitor_only_gaps_too: PASS")


def test_external_never_points_at_a_gap_it_does_not_print():
    """The invariant that survives the change above: a "(Gap N)" a reader
    cannot follow is worse than no pointer at all.

    The one gap External still does not print is one that externalizes to
    nothing, meaning it was pure internal-process phrasing with no finding
    left once Section B's "Internal keeps process failures, External does not"
    rule has been applied. Such a gap must keep its full text rather than emit
    a pointer into a list that will not contain it."""
    vanishes = ("Specific schools/hospitals within a fixed radius were not "
                "independently verified this pass.")
    facts = {"gaps": [vanishes], "_doc_variant": "external"}
    assert cc._externalize_prose(facts, vanishes).strip() == "", "premise: this gap externalizes away"
    assert cc._rendered_gap_numbers(facts) == set()

    item = _gap_item(facts)
    text, points_at_a_gap = cc._flag_headline(facts, item)
    assert points_at_a_gap is False
    assert text == vanishes, text

    # Internal prints it, so Internal still gets its pointer.
    internal = dict(facts, _doc_variant="internal")
    assert cc._rendered_gap_numbers(internal) == {1}
    _, points_i = cc._flag_headline(internal, item)
    assert points_i is True
    print("test_external_never_points_at_a_gap_it_does_not_print: PASS")


# --- document level ----------------------------------------------------------

def test_every_pointer_resolves_in_both_variants():
    """The core contract. A "(Gap N)" a reader cannot follow is worse than no
    pointer at all, so this is checked against the real fixture, in both
    variants, over body paragraphs AND table cells."""
    for variant in ("internal", "external"):
        blob = _all_text(_build(variant))
        pointers = {int(n) for n in re.findall(r"\(Gap (\d+)\)", blob)}
        printed = {int(n) for n in re.findall(r"^Gap (\d+)\.", blob, re.M)}
        dangling = sorted(pointers - printed)
        assert not dangling, f"{variant}: pointers with no matching entry: {dangling}"
        assert pointers, f"{variant}: fixture must produce at least one pointer"
    print("test_every_pointer_resolves_in_both_variants: PASS")


def test_internal_numbers_every_gap_contiguously():
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        expected = len(json.load(f).get("gaps", []))
    blob = _all_text(_build("internal"))
    printed = sorted({int(n) for n in re.findall(r"^Gap (\d+)\.", blob, re.M)})
    assert printed == list(range(1, expected + 1)), printed
    print(f"test_internal_numbers_every_gap_contiguously: PASS ({expected} gaps)")


def test_external_prints_every_gap_internal_does():
    """rules.md Section B: "Gaps & Sources keeps every item with its full
    explanation, including permanent ones."

    Regression for a real defect: External printed only Imminent/Structural
    gaps, so on a project whose gaps were all monitor severity the section came
    out empty. The only gaps External may now omit are ones that externalize to
    nothing, which carry no finding to report."""
    internal = _printed_gap_numbers(_all_text(_build("internal")))
    external = _printed_gap_numbers(_all_text(_build("external")))
    assert internal, "fixture must produce gaps at all"
    assert external, "External printed no gaps at all"
    assert external <= internal, f"External invented gap numbers: {sorted(external - internal)}"

    # The regression itself: monitor-severity gaps used to be excluded from
    # External outright, which is what let the section collapse to empty. At
    # least one must now survive. (Some gaps legitimately drop out here, being
    # pure process detail that externalizes to nothing, so this asserts the
    # rule rather than an exact count.)
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        flags = cc._classify_flags({"gaps": json.load(f).get("gaps", [])})
    escalated = {i["gap_number"] for tier in ("imminent", "structural")
                 for i in flags.get(tier, []) if i.get("gap_number")}
    monitor_printed = external - escalated
    assert monitor_printed, (
        "External printed only escalated gaps; monitor-severity gaps are still being dropped"
    )
    print(f"test_external_prints_every_gap_internal_does: PASS "
          f"({len(external)} printed, {len(monitor_printed)} of them monitor-severity)")


def test_external_never_claims_no_gaps_while_gaps_exist():
    """The exact false sentence this change removes: "No additional material
    gaps identified beyond the standing gap below.", printed directly beneath a
    list of unresolved gaps."""
    blob = _all_text(_build("external"))
    printed = re.findall(r"^Gap (\d+)\.", blob, re.M)
    assert printed, "fixture must print gaps in External"
    assert "No additional material gaps identified" not in blob, (
        "External claimed there are no material gaps while printing "
        f"{len(printed)} of them"
    )
    print("test_external_never_claims_no_gaps_while_gaps_exist: PASS")


# --- the governing-act line, and who it is generic for -----------------------

def test_project_is_rera_registered_reads_the_registration_number():
    reg = lambda v: {"rera_core_fields": {"registration_number": v}}
    assert cc._project_is_rera_registered(reg("P51800077150")) is True
    assert cc._project_is_rera_registered(reg("")) is False
    assert cc._project_is_rera_registered({}) is False
    for prose in ("Not registered with the Maharashtra Real Estate Regulatory Authority, "
                  "per the client's project record.",
                  "No registration on record.", "Unregistered.", "N/A", "None", "Not available"):
        assert cc._project_is_rera_registered(reg(prose)) is False, prose
    print("test_project_is_rera_registered_reads_the_registration_number: PASS")


def test_governing_act_is_dropped_from_external_only_when_registered():
    """It is generic-by-definition for a RERA-registered project, so External
    drops it there. For an UNREGISTERED project the governing legislation is
    different, and its consent thresholds and developer security requirements
    are among the most decision-relevant facts in the document, so it stays.

    Regression: this suppression used to be unconditional, which silently
    stripped that statutory basis out of every unregistered project's External
    document."""
    os.makedirs(_SCRATCH, exist_ok=True)
    with open(_PRANAMI_FACTS, encoding="utf-8") as f:
        base = json.load(f)
    act = (base.get("rules_statutory") or {}).get("governing_act", "")
    assert act, "fixture must carry a governing act"
    probe = act.split(".")[0][:60]

    def build(reg_value, name):
        facts = json.loads(json.dumps(base))
        facts.setdefault("rera_core_fields", {})["registration_number"] = reg_value
        out = os.path.join(_SCRATCH, name)
        cc._fill_template("P51800077150", facts, out, doc_variant="external")
        return _all_text(out)

    registered = build("P51800077150", "gov_registered.docx")
    unregistered = build("Not registered with MahaRERA, per the project record.",
                         "gov_unregistered.docx")

    assert probe not in registered, "registered project: generic act line should be dropped"
    assert probe in unregistered, (
        "unregistered project: the governing legislation is a project-specific fact "
        "and must survive into External"
    )
    print("test_governing_act_is_dropped_from_external_only_when_registered: PASS")


def test_gap_pointers_replace_the_raw_facts_path_annotation():
    """CLAUDE.md Section B forbids a JSON key in EITHER document. A flag that
    points at a gap must render "(Gap N)", never "(see gaps[3])"."""
    for variant in ("internal", "external"):
        blob = _all_text(_build(variant))
        assert "see gaps[" not in blob, f"{variant} still leaks a raw gaps[N] path"
    print("test_gap_pointers_replace_the_raw_facts_path_annotation: PASS")


def test_an_empty_flag_tier_still_collapses_to_nothing_found():
    """The `not items` branch is deliberately untouched by this change --
    CLAUDE.md: an empty section keeps its heading and one bare line."""
    blob = _all_text(_build("internal"))
    assert re.search(r"Imminent Red Flags[^\n]*\(0\)\n+Nothing found\.", blob), \
        "expected an empty flag tier to render its heading plus a bare 'Nothing found.'"
    print("test_an_empty_flag_tier_still_collapses_to_nothing_found: PASS")


if __name__ == "__main__":
    test_gap_flags_carry_a_stable_one_based_number()
    test_non_gap_flags_carry_no_gap_number()
    test_headline_takes_the_first_sentence_and_appends_the_pointer()
    test_a_flag_with_no_gap_keeps_its_full_text()
    test_external_prints_monitor_only_gaps_too()
    test_external_never_points_at_a_gap_it_does_not_print()
    test_every_pointer_resolves_in_both_variants()
    test_internal_numbers_every_gap_contiguously()
    test_external_prints_every_gap_internal_does()
    test_external_never_claims_no_gaps_while_gaps_exist()
    test_project_is_rera_registered_reads_the_registration_number()
    test_governing_act_is_dropped_from_external_only_when_registered()
    test_gap_pointers_replace_the_raw_facts_path_annotation()
    test_an_empty_flag_tier_still_collapses_to_nothing_found()
    print("\nAll tests passed.")
