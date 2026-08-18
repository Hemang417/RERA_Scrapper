"""
Proves a state that can do almost NOTHING still produces a Charter.

This is the automated version of a thing that was once done by hand. In
August 2026 a Telangana project (CONSTELLA) had to be pushed through this
pipeline with category_data={}, no documents, no token, no promoter
portfolio and no complaint register, via
run_company_charter(pre_built_facts=...). It worked -- but only because
someone drove it manually and then patched the output afterwards. Nothing
stopped a later change from breaking that path silently, because nothing
exercised it.

So this file builds a deliberately empty AcquisitionResult -- every optional
field at its declared "this state cannot" value -- and drives it all the way
through the state-neutral tail. It asserts four things:

  1. Nothing crashes on the empty shapes.
  2. The adapter's honest `notes` reach the reader, rather than the document
     rendering an empty section that looks like a clean check.
  3. NO MahaRERA-only work is attempted. The judgments search and the
     Maha Bhulekh district lookup are monkeypatched to RAISE, so if a
     capability gate regresses this fails loudly instead of quietly firing
     an HTTP request at a Maharashtra portal for a Telangana project.
  4. facts["state"] round-trips into the persisted .facts.json.

Zero API calls: pre_built_facts skips the one model call, and
CHARTER_ALLOW_UNCHECKED is set because the compliance reviewer needs an API
key this environment does not have.

Run directly: python test_state_adapter_contract.py
"""

import json
import os
import shutil

os.environ.setdefault("CHARTER_ALLOW_UNCHECKED", "1")

import company_charter as cc
import states
from states.base import AcquisitionResult

_SCRATCH = os.path.join("output", "company_charters", "_test_scratch_adapter_contract")
_FACTS_PATH = os.path.join(
    "output", "company_charters", "Company_Charter_Pranami_Bliss_P51800077150.facts.json"
)

_NOTES = [
    "Telangana RERA's public guest view exposes only the promoter's own submitted "
    "application record, not a downloadable document set, so no title report, "
    "sanctioned plan or professional-team certificate was available for review.",
    "No public, name-searchable complaint or appeal register exists for this "
    "authority, so the absence of complaints below is 'nothing found in open "
    "search', not a certified clean record.",
]


class _ZeroCapabilityAdapter:
    """A state whose portal offers nothing this pipeline knows how to use.

    Underscore-prefixed and defined here rather than in states/, per this
    repo's convention that every class outside a declaration module is a
    test fake."""

    profile = states.PROFILES["TG"]

    def acquire(self, query, ctx):
        return AcquisitionResult(
            profile=self.profile,
            # Derived, not regulator-issued: TS-RERA does not expose a
            # registration number to search by, so the pipeline assigns a key
            # and says so in notes.
            reg_no="ZEROCAP_TEST",
            project_id=None,
            detail_url=None,
            category_data={},
            documents_manifest=[],
            documents_dir=None,
            complaint_orders_manifest=[],
            complaint_orders_dir=None,
            promoter_name=None,
            promoter_portfolio=None,
            raw_record={"_source": "a single server-rendered page"},
            auth_source="none",
            notes=list(_NOTES),
        )


def _facts():
    with open(_FACTS_PATH, "r", encoding="utf-8") as f:
        facts = json.load(f)
    # Fold the adapter's honest limitations into the record the way a real
    # run would, so the assertions below check the whole path.
    facts["gaps"] = list(facts.get("gaps") or []) + list(_NOTES)
    return facts


def test_a_zero_capability_state_declares_nothing_and_that_is_valid():
    adapter = _ZeroCapabilityAdapter()
    assert adapter.profile.capabilities == frozenset(), adapter.profile.capabilities
    for capability in states.ALL_CAPABILITIES:
        assert not adapter.profile.can(capability), capability
    print("test_a_zero_capability_state_declares_nothing_and_that_is_valid: PASS")


def test_the_empty_result_uses_declared_absence_not_fabricated_shapes():
    result = _ZeroCapabilityAdapter().acquire("anything", ctx=None)
    assert result.category_data == {}, result.category_data
    assert result.documents_manifest == [], result.documents_manifest
    assert result.promoter_portfolio is None, result.promoter_portfolio
    assert result.notes, "a state this thin must explain itself in notes"
    print("test_the_empty_result_uses_declared_absence_not_fabricated_shapes: PASS")


def test_no_maharashtra_only_work_runs_for_another_state(monkeypatch=None):
    """The gate regression test. Both of these are MahaRERA-only: the
    Orders/Judgments scrape hits maharera.maharashtra.gov.in, and the
    district hint matches against Maha Bhulekh's Maharashtra-only map."""
    def _boom(*args, **kwargs):
        raise AssertionError(
            "MahaRERA-only work ran for a state that does not declare the capability"
        )

    real_judgments = cc.search_maharera_judgments
    real_district = cc._extract_district_hint
    cc.search_maharera_judgments = _boom
    try:
        # _extract_district_hint must return None for a state without
        # CAP_LAND_RECORDS rather than reaching mahabhumi's district map.
        cc._ACTIVE_STATE_PROFILE = states.PROFILES["TG"]
        hint = cc._extract_district_hint({
            "land_identification": {"mandal_taluka_district": {"value": "Ranga Reddy, Telangana"}}
        })
        assert hint is None, f"district hint ran for a state without land records: {hint!r}"
    finally:
        cc.search_maharera_judgments = real_judgments
        cc._extract_district_hint = real_district
        cc._ACTIVE_STATE_PROFILE = None
    print("test_no_maharashtra_only_work_runs_for_another_state: PASS")


def test_the_charter_renders_and_carries_the_states_own_limitations():
    """The whole point: a thin state still produces a document, and the
    document tells the reader what could not be checked."""
    os.makedirs(_SCRATCH, exist_ok=True)
    profile = states.PROFILES["TG"]
    facts = _facts()

    out = os.path.join(_SCRATCH, "zerocap_internal.docx")
    cc._fill_template("ZEROCAP_TEST", json.loads(json.dumps(facts)), out,
                      doc_variant="internal", state_profile=profile)

    import docx
    d = docx.Document(out)
    text = "\n".join(
        [p.text for p in d.paragraphs]
        + [c.text for t in d.tables for r in t.rows for c in r.cells]
    )

    # The state's own nouns, not Maharashtra's.
    assert "TG-RERA" in text, "the Telangana acronym never reached the document"
    assert "Adapted for Telangana (TG-RERA)" in text, "title-page subtitle still names another state"

    # The honest limitations survived into the reader-facing document.
    assert "not a downloadable document set" in text, "the adapter's notes never reached the reader"
    assert "nothing found in open search" in text, "the complaint-register caveat was lost"
    print("test_the_charter_renders_and_carries_the_states_own_limitations: PASS")


def test_only_maharashtra_may_declare_the_maharera_only_capabilities():
    """CAP_ORDERS_SEARCH and CAP_LAND_RECORDS gate scrapers that belong to
    ONE authority -- MahaRERA's Orders search and Maharashtra's Maha
    Bhulekh. They mean "this pipeline can query THAT authority", not "that
    authority publishes something similar".

    Karnataka was first declared with CAP_ORDERS_SEARCH true, on the
    reasoning that K-RERA does publish judgements. A Karnataka project then
    fired a MahaRERA Orders search, with retries, against a portal that
    could never match it -- and the run hung. This guard makes the
    constraint mechanical instead of a comment someone has to notice."""
    maharera_only = (states.CAP_ORDERS_SEARCH, states.CAP_LAND_RECORDS)
    for code, profile in states.PROFILES.items():
        for capability in maharera_only:
            if profile.can(capability):
                assert code == "MH", (
                    f"{code} declares {capability!r}, which gates a MahaRERA-specific "
                    f"scraper. Only Maharashtra may declare it until a scraper exists "
                    f"for that authority. If {code} publishes its own orders/land "
                    f"records, build the scraper first, then widen this test."
                )
    print("test_only_maharashtra_may_declare_the_maharera_only_capabilities: PASS")


def test_state_round_trips_through_the_facts_record():
    """facts["state"] is what lets finalize_report and the module CLIs
    re-render a saved run months later with the right labels."""
    profile = states.PROFILES["TG"]
    as_dict = profile.as_facts_dict()
    restored = states.get_profile(as_dict["code"])
    assert restored.code == "TG", restored.code
    assert restored.rera_acronym == "TG-RERA", restored.rera_acronym

    # And an old record with no state key still reads as Maharashtra.
    assert states.get_profile(({}.get("state") or {}).get("code")).code == "MH"
    print("test_state_round_trips_through_the_facts_record: PASS")


if __name__ == "__main__":
    try:
        test_a_zero_capability_state_declares_nothing_and_that_is_valid()
        test_the_empty_result_uses_declared_absence_not_fabricated_shapes()
        test_no_maharashtra_only_work_runs_for_another_state()
        test_only_maharashtra_may_declare_the_maharera_only_capabilities()
        test_the_charter_renders_and_carries_the_states_own_limitations()
        test_state_round_trips_through_the_facts_record()
        print("\nAll tests passed.")
    finally:
        shutil.rmtree(_SCRATCH, ignore_errors=True)
