"""
Company Charter -- Company / Promoters / Collateral report, built to the
external spec supplied by the user (the "Company Charter Generator" prompt,
revised after a first-draft review -- see COMPANY_CHARTER_SPEC.md for the
current, authoritative version of that spec; this module implements it).

Three subjects, each with its OWN litigation screen, kept structurally
separate rather than folded together:
  1. The Company    -- the legal entity: identity, ratings, portfolio,
     company-level litigation only.
  2. The Promoters   -- the people who control it, one subsection per
     person: biographical research, THEIR OWN related-entity table, and a
     litigation flag inside their own subsection if (and only if) something
     was found against them personally.
  3. The Collateral  -- the asset: land, approvals, RERA compliance,
     project-level litigation only.

Editorial policy, applied throughout (locked in after the first-draft
review): report findings, not the absence of findings. A clean check
("no disagreement," "no adverse coverage," "no complaint on record") is
never rendered as its own sentence or bullet -- it is simply not mentioned.
Only two things earn a line: (1) core data the reader came for (a rating, a
filed record, a chronology), and (2) something that needs the reader's
attention (a discrepancy, a gap, an open item). The second category is
always consolidated into a section's own "Needs Attention" subsection,
never scattered as asides under whatever table happened to raise it. There
is no "What Checks Out" heading anywhere in this document.

WHY A SEPARATE MODULE from charter_document.py: that module's 15-section
layout treats leadership as one part of the Company half. This spec
elevates "The Promoters" to a full top-level subject, adds real
biographical research on named individuals, and applies the above
find-first editorial policy -- a different organising principle, not a
rename of the same sections. Both modules read the same underlying
facts.json; nothing here duplicates data collection, only re-presents it
under the new structure and a different visual design (Calibri, navy
#1B2A4A headings, Table Grid with alternating row shading, no-dash prose,
numbered citations throughout, tick/cross/dot status glyphs in place of
"Confirmed" / "Not established" text).

Reuses computation directly from charter_document.py (rollups, claim
classification, the closing-verdict rule engine, related-entity link typing)
and company_charter.py (developer score, flag classification, entity-name
normalisation) so numbers can never disagree between documents built from
the same facts.json.
"""

from __future__ import annotations

import re
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import charter_document as cd
import company_charter as cc

# ---------------------------------------------------------------------------
# Visual design (per spec: Calibri 10.5pt body, navy #1B2A4A headings,
# Table Grid with navy header / alternating grey-white body rows, ~2.1cm
# margins, explicit column widths).
# ---------------------------------------------------------------------------
_NAVY = "1B2A4A"
_LIGHT_GREY = "F2F4F8"
_WHITE = "FFFFFF"
_GREY_TEXT = "595959"
_MARGIN_CM = 2.1
_PAGE_WIDTH_CM = 21.0  # A4
_CONTENT_WIDTH_CM = _PAGE_WIDTH_CM - 2 * _MARGIN_CM

_TICK = "✓"    # confirmed
_CROSS = "✗"   # not established / discrepancy
_DOT = "●"     # stated but not independently confirmed

_STATUS_GLYPH = {
    cd.STATUS_CONFIRMED_INDEPENDENT: (_TICK, cc._TEXT_GREEN),
    cd.STATUS_CONFIRMED_DOCUMENT: (_TICK, cc._TEXT_GREEN),
    cd.STATUS_PROMOTER_FILED: (_DOT, cc._TEXT_AMBER),
    cd.STATUS_STATED_ONLY: (_DOT, cc._TEXT_AMBER),
    cd.STATUS_DISCREPANCY: (_CROSS, cc._TEXT_RED),
    cd.STATUS_NOT_ESTABLISHED: (_CROSS, cc._TEXT_RED),
}


def _configure_base_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for level, size in ((1, 19), (2, 14.5), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(_NAVY)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Cm(_MARGIN_CM)
        section.bottom_margin = Cm(_MARGIN_CM)
        section.left_margin = Cm(_MARGIN_CM)
        section.right_margin = Cm(_MARGIN_CM)


def _add_field(paragraph, field_code: str):
    """Inserts a live Word field (e.g. "PAGE", "NUMPAGES") into `paragraph`
    -- python-docx has no high-level API for this, so it's built directly
    from the OOXML field-character sequence (begin/instrText/separate/end)
    Word itself uses. Word resolves the actual number when it opens or
    prints the document (including on a Word-COM SaveAs to PDF), so this
    correctly renders as "3" rather than the literal field code. Returns
    the run carrying the field, so the caller can style it like any other
    run (font.size, font.color, ...)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, end):
        run._r.append(element)
    return run


def _add_page_numbers(doc) -> None:
    """"Page X of Y", centered, in the footer of every page. Set once on
    the document's only section -- charter_report.py never inserts a
    section break (add_page_break() is a plain page break, not a new
    section), so one footer covers the whole document."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    runs = [paragraph.add_run("Page ")]
    runs.append(_add_field(paragraph, "PAGE"))
    runs.append(paragraph.add_run(" of "))
    runs.append(_add_field(paragraph, "NUMPAGES"))
    for run in runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(_GREY_TEXT)


# ---------------------------------------------------------------------------
# No-dash rule: em dash / en dash / this codebase's own " -- " convention are
# all normalised out. A dash directly between two numbers is a range and
# becomes a plain hyphen ("28,300-35,250"); anywhere else a dash is joining
# or introducing a clause and becomes a semicolon, which is grammatical
# wherever a dash was and (unlike a comma) cannot create a comma splice.
# ---------------------------------------------------------------------------
_NUMERIC_DASH_RE = re.compile(r"(?<=[\d,])\s*(?:--|[–—])\s*(?=[\d])")
_WORD_DASH_RE = re.compile(r"\s*(?:--|[–—])\s*")
_MULTI_SPACE_RE = re.compile(r"[ \t]{2,}")
_EMPTY_COMMA_RE = re.compile(r"(?:,\s*){2,}")


def _clean_text(text) -> str:
    """Dash normalisation plus purely typographic whitespace/punctuation
    tidy-up (collapsing "X,   , Y" scraped-address artifacts to "X, Y") --
    never changes a value, only how its existing characters are spaced."""
    if not text:
        return ""
    text = str(text)
    text = _NUMERIC_DASH_RE.sub("-", text)
    text = _WORD_DASH_RE.sub("; ", text)
    text = _EMPTY_COMMA_RE.sub(", ", text)
    text = _MULTI_SPACE_RE.sub(" ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Redundant-attribution stripping: this codebase's own field values often
# spell out their own provenance inline ("(per ZaubaCorp/InstaFinancials --
# Company limited by shares, Non-government)", "value -- confirmed
# identically across ZaubaCorp, Tofler, and InstaFinancials"). That was the
# only attribution available before this document had numbered citations;
# now that a value carrying one of these gets a bracketed [n] instead, the
# inline mention just repeats what the citation already says. Only a
# RECOGNISED attribution shape is stripped -- anything not matching either
# pattern is returned unchanged, so this never eats content it doesn't
# understand (e.g. the "Company limited by shares, Non-government"
# classification detail inside the same parenthetical survives).
# ---------------------------------------------------------------------------
_PER_PARENTHETICAL_WITH_DETAIL_RE = re.compile(r"\(per\s+.*?(?:--|;)\s*([^)]*)\)", re.IGNORECASE)
_PER_PARENTHETICAL_ONLY_RE = re.compile(r"\s*\(per\s+[^)]+\)", re.IGNORECASE)
_CONFIRMED_ACROSS_SUFFIX_RE = re.compile(
    r"\s*(?:--|;)\s*confirmed\s+identically\s+(?:across|by)\s+[^.]*$", re.IGNORECASE
)


def _strip_redundant_attribution(text: str) -> str:
    text = str(text or "")
    text = _PER_PARENTHETICAL_WITH_DETAIL_RE.sub(r"(\1)", text)
    text = _PER_PARENTHETICAL_ONLY_RE.sub("", text)
    text = _CONFIRMED_ACROSS_SUFFIX_RE.sub("", text)
    return text.strip()


_BARE_DOMAIN_MENTION_RE = re.compile(
    r"\((?:www\.)?([a-z0-9-]+\.[a-z]{2,}(?:/[^\s)]*)?)\)", re.IGNORECASE
)


def _extract_bare_domain_citation(text: str) -> tuple:
    """Some free-text fields cite themselves inline with a bare domain
    mention, e.g. "own website (pranamigroup.com/about)", rather than a
    proper URL in a separate source attribute. Pulls the first such mention
    out as a real https:// source and drops the now-redundant parenthetical
    from the display text (a citation number will do that job instead).
    Returns (cleaned_text, source_url_or_None); returns the text unchanged
    with None if no bare-domain mention is found."""
    match = _BARE_DOMAIN_MENTION_RE.search(text or "")
    if not match:
        return text, None
    url = f"https://{match.group(1)}"
    cleaned = (text[:match.start()] + text[match.end():]).strip()
    cleaned = re.sub(r"\s+,", ",", cleaned)  # "website , the group" -> "website, the group"
    return cleaned, url


_LABELED_CLAUSE_RE = re.compile(r"([A-Z][A-Za-z /]{2,40}:)\s*")


def split_labeled_narrative(text: str) -> list:
    """Splits a machine-authored "Label one: clause. Label two: clause."
    narrative into (label, clause) pairs -- generic to any field using this
    codebase's own labeled-clause convention (seen on corporate_identity and
    similar fields), NOT hardcoded to one project's entity names. Falls
    back to a single unlabeled pointer if the text doesn't carry two or
    more such labels."""
    text = str(text or "").strip()
    matches = list(_LABELED_CLAUSE_RE.finditer(text))
    if len(matches) < 2:
        return [(None, text)]
    pointers = []
    for i, m in enumerate(matches):
        label = m.group(1).rstrip(":").strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        clause = text[start:end].strip().rstrip(".").strip()
        pointers.append((label, clause))
    return pointers


def reframe_arrow_chain(text: str) -> list:
    """Turns a machine-authored "Intro label: A -> B -> C." arrow chain into
    an ordered list of clean steps -- generic to any project's approval or
    event chronology, not hardcoded to specific milestone names. Strips a
    leading "Label:" lead-in (if present) since that belongs in the calling
    heading, not step 1. Only the LAST step's trailing period is stripped
    (the chain's own sentence-ending full stop) -- an intermediate step's
    trailing period is an abbreviation (e.g. "...Pvt. Ltd." before the next
    arrow) and must survive, or "Ltd." silently loses its own period."""
    text = str(text or "").strip()
    lead_match = re.match(r"^[A-Z][A-Za-z /]{2,60}:\s*", text)
    if lead_match:
        text = text[lead_match.end():]
    raw_steps = [s.strip() for s in text.split("->") if s.strip()]
    steps = []
    for i, s in enumerate(raw_steps):
        if i == len(raw_steps) - 1:
            s = s.rstrip(".").strip()
        steps.append(s[0].upper() + s[1:] if s else s)
    return steps


# ---------------------------------------------------------------------------
# Citation registry -- numbered [n] markers matching the References list.
# Internal keeps raw local file paths as the citation text (useful to
# relocate the file); External replaces a raw path with a clean description
# of the same document, since a filesystem path is internal-process detail.
# Both are deduped to the SAME key (the file path or its filename) so the
# same underlying document never fragments into two citation numbers.
# ---------------------------------------------------------------------------
_FILE_PATH_RE = re.compile(r"(output/[^;]+?\.pdf)", re.IGNORECASE)


def _polish_source(source, doc_variant) -> str:
    source = (source or "").strip()
    if not source or source.lower().startswith("gap"):
        return ""
    match = _FILE_PATH_RE.search(source)
    if not match:
        # A trailing slash on an otherwise-bare URL ("https://x.com/about/"
        # vs "https://x.com/about") is not a different source -- normalise
        # it away so the same domain always dedupes to one citation number
        # regardless of which call site supplied the trailing slash.
        if re.fullmatch(r"https?://\S+", source):
            source = source.rstrip("/")
        return source
    path = match.group(1)
    if doc_variant == "external":
        filename = path.rsplit("/", 1)[-1]
        return f"{filename} (project document on file with the promoter)"
    return path


def _inline_citation_label(source: str) -> str:
    """Short, readable inline-citation text for the Internal variant, e.g.
    "(per ZaubaCorp company profile)" or "(per zaubacorp.com)" -- Internal
    uses plain readable attribution in place of External's numbered [n] /
    References-list mechanism (see _Builder.cite). Prefers a description
    already present before a URL (e.g. "ZaubaCorp company profile --
    https://...") over the bare domain, same preference order _harvardish
    uses for the References-list line."""
    source = (source or "").strip()
    if not source or source.lower().startswith("gap"):
        return ""
    match = _FILE_PATH_RE.search(source)
    if match:
        filename = match.group(1).rsplit("/", 1)[-1]
        return f"per {filename}, project document on file"
    url_match = re.search(r"https?://\S+", source)
    if url_match:
        description = source[:url_match.start()].strip(" ;-")
        if description:
            return f"per {description}"
        domain = urlparse(url_match.group(0)).netloc
        domain = domain[4:] if domain.startswith("www.") else domain
        return f"per {domain}"
    return f"per {source}"


class _CitationRegistry:
    def __init__(self):
        self.order = []
        self.index = {}

    def cite(self, source, doc_variant) -> str:
        text = _polish_source(source, doc_variant)
        if not text:
            return ""
        if text not in self.index:
            self.index[text] = len(self.order) + 1
            self.order.append(text)
        return f"[{self.index[text]}]"


def _harvardish(source_text: str, generated_on: str) -> str:
    """Renders one References-list line. If the source string carries a URL,
    everything before it is the description and the URL is appended in the
    Harvard "Available at" form, falling back to the URL's own domain when
    no description text precedes it. A bare file path or plain description
    is rendered as-is."""
    url_match = re.search(r"https?://\S+", source_text)
    if url_match:
        url = url_match.group(0).rstrip(").,;")
        description = source_text[:url_match.start()].strip(" ;-")
        if not description:
            description = urlparse(url).netloc or "Source"
        return f"{description}. Available at: {url} (Accessed: {generated_on})."
    if source_text.startswith("output/"):
        return f"{source_text} (project document on file, accessed {generated_on})."
    return f"{source_text}."


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------
class _Builder:
    def __init__(self, doc, facts, doc_variant):
        self.doc = doc
        self.facts = facts
        self.doc_variant = doc_variant
        self.citations = _CitationRegistry()

    def cite(self, source) -> str:
        if not source:
            return ""
        if self.doc_variant == "internal":
            # Internal reads inline and self-descriptive ("per X"); only
            # External uses numbered [n] markers resolving to a References
            # list -- Internal has no such list (see build_charter_report,
            # which skips references_section for this variant).
            label = _inline_citation_label(source)
            return f"({label})" if label else ""
        return self.citations.cite(source, self.doc_variant)

    def _with_citation(self, text, source) -> str:
        cleaned = _clean_text(text)
        marker = self.cite(source) if source else ""
        return f"{cleaned} {marker}".rstrip() if marker else cleaned

    def heading(self, text, level=1, page_break_before=False):
        """`page_break_before=True` sets the OOXML page-break-before
        paragraph property instead of inserting a standalone page-break
        character. This is deliberate, not a style preference: an explicit
        page-break character always forces a new page, even when the
        previous content already happened to end exactly at a page
        boundary -- producing a genuinely blank page (confirmed live: the
        Internal variant's Company/Promoters transition did exactly this
        once its content grew past a certain length). page-break-before
        tells Word "start this heading on a fresh page, but only if it
        isn't already on one," which cannot produce a blank page no matter
        how the surrounding content's length happens to land."""
        heading = self.doc.add_heading(_clean_text(text), level=level)
        if page_break_before:
            heading.paragraph_format.page_break_before = True
        return heading

    def para(self, text="", source=None, bold=False, italic=False, color=None, align=None):
        content = self._with_citation(text, source) if text else ""
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        if content:
            run = p.add_run(content)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        return p

    def lead_bold_para(self, lead, rest, source=None):
        """A paragraph shaped "**Lead:** rest of the sentence [n]." -- used
        for a headline statistic that then explains itself in the same
        breath (e.g. the Documentation Confidence overall figure)."""
        p = self.doc.add_paragraph()
        lead_run = p.add_run(_clean_text(lead) + ": ")
        lead_run.bold = True
        p.add_run(self._with_citation(rest, source))
        return p

    def bullet(self, text, source=None):
        p = self.doc.add_paragraph(style="List Bullet")
        p.add_run(self._with_citation(text, source))
        return p

    def flag_line(self, text, source=None):
        return self.para(text, source=source, italic=True, color=_GREY_TEXT)

    def _set_col_widths(self, table, widths_cm):
        table.autofit = False
        for row in table.rows:
            for cell, width in zip(row.cells, widths_cm):
                cell.width = Cm(width)

    def table(self, headers, rows, col_widths=None):
        """rows: list of lists; each cell is plain text, (text, source) for
        an auto-cited cell, or (text, source, color) to also colour the run
        (used for the tick/cross/dot status glyphs)."""
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            header_cells[i].text = ""
            cc._shade_cell(header_cells[i], _NAVY)
            run = header_cells[i].paragraphs[0].add_run(_clean_text(h))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(_WHITE)

        for ridx, row_values in enumerate(rows):
            row = t.add_row()
            shade = _LIGHT_GREY if ridx % 2 == 0 else None
            for i, value in enumerate(row_values):
                if isinstance(value, tuple):
                    if len(value) == 3:
                        text, source, color = value
                    else:
                        text, source = value
                        color = None
                else:
                    text, source, color = value, None, None
                cell = row.cells[i]
                cell.text = ""
                if shade:
                    cc._shade_cell(cell, shade)
                run = cell.paragraphs[0].add_run(self._with_citation(str(text or ""), source))
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)

        widths = col_widths or [_CONTENT_WIDTH_CM / len(headers)] * len(headers)
        self._set_col_widths(t, widths)
        return t

    def references_section(self, generated_on):
        self.heading("References", page_break_before=True)
        if not self.citations.order:
            self.para("No source citations were registered in this pass.")
            return
        for i, source_text in enumerate(self.citations.order, start=1):
            p = self.doc.add_paragraph()
            p.add_run(f"[{i}] ").bold = True
            p.add_run(_clean_text(_harvardish(source_text, generated_on)))


def _status_glyph_cell(status: str, source) -> tuple:
    glyph, color = _STATUS_GLYPH.get(status, ("?", None))
    return (glyph, source, color)


def _needs_attention(b: _Builder, label: str, items: list, empty_text: str):
    """Consolidated closing subsection for a major part of the document.
    Renders bullets ONLY for items actually found; an empty list renders a
    single clean-state line instead -- there is no "What Checks Out"
    counterpart anywhere in this document."""
    b.para()
    b.heading(f"Needs Attention, {label}", level=2)
    if items:
        for item in items:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.bullet(text, source=source)
    else:
        b.flag_line(empty_text)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def _cover_page(b: _Builder, reg_no, generated_on, composite_line):
    core = b.facts.get("rera_core_fields") or {}
    corp = b.facts.get("corporate_identity") or {}
    promoter = cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", "")) or "[Unknown]"
    project = core.get("project_name", "[Unknown]")
    location = (b.facts.get("land_identification") or {}).get("village_locality", {})
    location_value = location.get("value") if isinstance(location, dict) else location

    for _ in range(3):
        b.para()
    b.para("COMPANY CHARTER", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=_NAVY)
    title_run = b.doc.paragraphs[-1].runs[0]
    title_run.font.size = Pt(28)

    b.para(
        "Company, Promoters and Collateral Due-Diligence Note",
        italic=True, color=_GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    b.doc.paragraphs[-1].runs[0].font.size = Pt(13)

    b.para()
    b.para(f"Company: {_clean_text(promoter)}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    b.para(f"Collateral: {_clean_text(project)} ({reg_no})", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if location_value:
        b.para(f"Location: {_clean_text(location_value)}", align=WD_ALIGN_PARAGRAPH.CENTER)
    b.para(f"Date of review: {generated_on}", align=WD_ALIGN_PARAGRAPH.CENTER)
    classification = "Internal -- Integrow Asset Management" if b.doc_variant == "internal" else "Strictly Private and Confidential"
    b.para(classification, align=WD_ALIGN_PARAGRAPH.CENTER, color=_GREY_TEXT)

    for _ in range(4):
        b.para()
    if composite_line:
        b.para(composite_line, italic=True, color=_GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# 1. About the Company -- prose only, no tables.
# ---------------------------------------------------------------------------
def _about_the_company(b: _Builder):
    b.heading("1. About the Company", page_break_before=True)

    corp = b.facts.get("corporate_identity") or {}
    track = b.facts.get("developer_track_record") or {}
    profile = b.facts.get("company_profile_check") or {}
    group = b.facts.get("group_companies_check") or {}
    promoter = cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", ""))
    org_type = (corp.get("organization_type") or {}).get("value", "")
    registry_source = (corp.get("organization_type") or {}).get("source")

    org_clause = _clean_text(_strip_redundant_attribution(org_type)) or "a company"
    org_cite = f" {b.cite(registry_source)}" if registry_source else ""
    b.para(
        f"{promoter} is {org_clause}{org_cite}, incorporated on "
        f"{profile.get('incorporation_date', 'an unconfirmed date')}, acting as the registered promoter for the "
        f"project reviewed in this Charter. It is the Mumbai-market vehicle for the wider group described below, "
        f"not a standalone developer with its own independent history."
    )

    if track.get("years_in_industry"):
        basis_text, basis_source = _extract_bare_domain_citation(str(track.get("years_in_industry_basis") or ""))
        b.para(basis_text, source=basis_source)

    directors = profile.get("current_directors") or []
    if directors:
        names = ", ".join(cc._normalise_entity_name(d.get("Director Name") or "") for d in directors)
        b.para(f"Corporate control sits with {len(directors)} current {cc._role_word(b.facts, len(directors))}: {names}.", source=registry_source)

    company_count = len(group.get("companies") or [])
    if company_count:
        b.para(
            f"The group's footprint extends to at least {company_count} related entities identified through "
            f"shared {cc._role_word(b.facts, 2)} or a shared registered office, discussed in full under The Promoters.",
            source=group.get("url"),
        )


# ---------------------------------------------------------------------------
# 2. The Company
# ---------------------------------------------------------------------------
def _identity_and_registration(b: _Builder, needs_attention: list):
    b.heading("Identity and Registration", level=2)
    corp = b.facts.get("corporate_identity") or {}
    profile = b.facts.get("company_profile_check") or {}
    registry_source = (corp.get("organization_type") or {}).get("source")

    rows = []

    def add(label, value, source=None):
        if value:
            rows.append([label, (str(value), source or registry_source)])

    add("Legal name", cc._normalise_entity_name((corp.get("promoter_name") or {}).get("value", "")),
        (corp.get("promoter_name") or {}).get("source"))
    add("CIN / LLPIN", profile.get("cin"), (corp.get("cin_llpin") or {}).get("source"))
    add("Status", profile.get("status"))
    add("Class / category", f"{profile.get('class_of_company', '')} / {profile.get('company_category', '')}".strip(" /"))
    add("Registrar", profile.get("roc"))
    add("Incorporation date", profile.get("incorporation_date"))
    add("Registered office", profile.get("registered_address"), (corp.get("registered_office_main") or {}).get("source"))
    add("Authorised capital", profile.get("authorised_capital"))
    add("Paid-up capital", profile.get("paid_up_capital"))
    b.table(["Field", "As recorded"], rows, col_widths=(6.0, _CONTENT_WIDTH_CM - 6.0))

    conflicts = profile.get("roster_conflicts") or []
    if conflicts:
        needs_attention.append(
            f"{len(conflicts)} disagreement(s) were found across registry mirrors on the {cc._role_word(b.facts, 2)} roster."
        )


# Fields deliberately excluded from the Key Items table: the two registered-
# office variants add nothing beyond the "Registered office" row already in
# Identity and Registration; partners_directors duplicates the Director /
# Partner Profiles subsections; landowner_investor is a chronology, not a
# single verifiable claim, and gets its own subsection under The Collateral.
_KEY_ITEMS_EXCLUDED_FIELDS = {
    "registered_office_board_resolution",
    "registered_office_planning_stage",
    "partners_directors",
    "landowner_investor",
}


def _curate_key_items(facts) -> list:
    """Legal-identity items, financial-exposure items, and anything with a
    discrepancy or gap only -- a deliberate curation, not the full working
    (that stays in the internal facts.json / scoring appendix)."""
    claim_rows = cd.build_claim_rows(facts)
    rows = []
    for r in claim_rows:
        if r.get("field_key") in _KEY_ITEMS_EXCLUDED_FIELDS:
            continue
        if r["group"] == "Counterparty identity" or r["status"] in (cd.STATUS_DISCREPANCY, cd.STATUS_NOT_ESTABLISHED):
            stated = _strip_redundant_attribution(r["stated"])
            rows.append([r["claim"], stated, _status_glyph_cell(r["status"], r.get("source"))])

    profile = facts.get("company_profile_check") or {}
    corp = facts.get("corporate_identity") or {}
    registry_source = (corp.get("organization_type") or {}).get("source")
    if profile.get("authorised_capital") or profile.get("paid_up_capital"):
        rows.append([
            "Capital structure",
            f"Authorised {profile.get('authorised_capital', 'not stated')}; paid-up {profile.get('paid_up_capital', 'not stated')}",
            _status_glyph_cell(cd.STATUS_CONFIRMED_INDEPENDENT, registry_source),
        ])
    return rows


def _verification_summary(b: _Builder, developer_score, doc_confidence_display, dc_buckets, dc_stale, needs_attention):
    b.para()
    b.heading("Verification Summary", level=2)

    b.heading("Overall Rating", level=3)
    ds_buckets = cd.rollup_developer_score_buckets(developer_score)
    ds_rows = []
    for bucket in ds_buckets:
        row = [bucket["bucket"], f"{bucket['weight']}%", bucket["rating"] or "Not rated", f"{bucket['scored']} of {bucket['total']}"]
        ds_rows.append(row)
    headers = ["Bucket", "Weight", "Rating", "Coverage"] if b.doc_variant == "internal" else ["Bucket", "Rating", "Coverage"]
    if b.doc_variant != "internal":
        ds_rows = [[r[0], r[2], r[3]] for r in ds_rows]
    b.table(headers, ds_rows)
    b.para(f"Composite Developer Score: {developer_score.get('composite')}/100, grade {developer_score.get('grade')}.", bold=True)
    for bucket in ds_buckets:
        for name in bucket["unscored_display_names"]:
            needs_attention.append(f"Not rated in {bucket['bucket']}: {name}. No public data source established this for this pass.")

    b.para()
    b.heading("Documentation Confidence", level=3)
    if dc_stale:
        needs_attention.append(
            "The stored Documentation Confidence Score is from a superseded scoring scheme and cannot be "
            "mapped onto the current criteria. Regenerate the Charter to score it."
        )
    elif dc_buckets:
        dc_rows = [[b_["bucket"], f"{b_['weight']}%", b_["band"] or "Not rated", f"{b_['scored']} of {b_['total']}"] for b_ in dc_buckets]
        dc_headers = ["Bucket", "Weight", "Band", "Coverage"] if b.doc_variant == "internal" else ["Bucket", "Band", "Coverage"]
        if b.doc_variant != "internal":
            dc_rows = [[r[0], r[2], r[3]] for r in dc_rows]
        b.table(dc_headers, dc_rows)
        b.lead_bold_para(
            f"Overall {doc_confidence_display['overall']}/100, band {doc_confidence_display['band']}",
            "This rates how well evidenced this Charter's own claims are, not the project's quality.",
        )
    else:
        needs_attention.append("No Documentation Confidence Score was recorded this pass.")

    b.para()
    b.heading("Key Items to Verify Before Engaging the Developer", level=3)
    key_items = _curate_key_items(b.facts)
    if key_items:
        b.table(["Item", "Stated position", "Verified"], key_items, col_widths=(4.5, _CONTENT_WIDTH_CM - 6.5, 2.0))
    else:
        b.flag_line("No legal-identity, financial-exposure or discrepancy items were flagged this pass.")


def _is_unconfirmed_read(read: str) -> bool:
    """True for a headline-claim "read" that could not actually be
    confirmed (charter_document._headline_claim_rows's "Not established"
    and "Possibly related, unconfirmed" outcomes) -- these move to Needs
    Attention rather than the main narrative. "Promoter-declared, not
    independently verified" and "Consistent" both describe a check that
    WAS completed (just resting on a lower evidence tier in the first
    case), so they stay in the body with their caveat stated plainly, the
    same treatment self-reported data gets everywhere else in this
    document (the amber-dot tier, not the red-cross one)."""
    read_lower = (read or "").lower()
    if "not established" in read_lower:
        return True
    return "unconfirmed" in read_lower and "not independently verified" not in read_lower


def _portfolio_and_track_record(b: _Builder, group_research, needs_attention: list):
    b.para()
    b.heading("Portfolio and Track Record", level=2)
    portfolio = b.facts.get("promoter_portfolio") or {}
    totals = portfolio.get("totals") or {}

    b.para(
        f"{totals.get('total_projects', 0)} MahaRERA registration(s) are held under this promoter's name, "
        f"with {totals.get('total_experience_entries_found', 0)} prior delivery or deliveries declared to the "
        f"regulator; these figures are promoter-reported and have not been independently re-verified project "
        f"by project."
    )

    # Each confirmed headline claim (years in the industry, area delivered)
    # gets its own paragraph rather than being run together, so a reader can
    # take in one claim, its basis and its caveat before moving to the next.
    for claim, shows, read in cd._headline_claim_rows(b.facts, totals):
        if _is_unconfirmed_read(read):
            needs_attention.append(f"{claim}: {shows}")
            continue
        lead = shows if shows.endswith((".", "!", "?")) else f"{shows}."
        if "years in the industry" in claim.lower():
            # The basis text for this specific claim is already a complete,
            # self-contained account (who says so, when, on what authority)
            # -- restating the short claim label first would just repeat it.
            lead, lead_source = _extract_bare_domain_citation(lead)
            b.para(lead, source=lead_source)
        else:
            caveat = "" if read.lower() == "consistent" else f" ({read}.)"
            b.para(f"{claim}, {lead[0].lower()}{lead[1:]}{caveat}")

    corroboration_sentences = []
    for item in group_research.get("track_record_corroboration", []):
        text, source = item if isinstance(item, tuple) else (item, None)
        marker = f" {b.cite(source)}" if source else ""
        corroboration_sentences.append(f"{text}{marker}")
    if corroboration_sentences:
        b.para(" ".join(corroboration_sentences))

    for item in group_research.get("track_record_gaps", []):
        needs_attention.append(item if isinstance(item, tuple) else (item, None))


def _company_litigation(b: _Builder, facts, group_research):
    b.para()
    b.heading("Litigation and Regulatory Screening, Company Level", level=2)
    ibbi = facts.get("ibbi_insolvency_check") or {}
    credit = facts.get("credit_rating_check") or {}

    findings = []
    if ibbi.get("found_process") is True:
        findings.append((str(ibbi.get("status_text") or "An active insolvency process is on record against this entity in IBBI's register."), ibbi.get("url")))
    if credit.get("found") is True:
        findings.append((str(credit.get("note") or "A credit rating was found for this entity."), None))
    for item in group_research.get("corporate_litigation", []):
        text, source = item if isinstance(item, tuple) else (item, None)
        findings.append((text, source))

    if findings:
        for text, source in findings:
            b.bullet(text, source=source)
    else:
        b.flag_line("No company-level litigation or regulatory concern was identified this pass.")


def _the_company(b: _Builder, flags, developer_score, doc_confidence_display, dc_buckets, dc_stale, group_research):
    b.heading("2. The Company", page_break_before=True)
    needs_attention = []
    _identity_and_registration(b, needs_attention)
    _verification_summary(b, developer_score, doc_confidence_display, dc_buckets, dc_stale, needs_attention)
    _portfolio_and_track_record(b, group_research, needs_attention)
    _company_litigation(b, b.facts, group_research)

    company_prefixes = ("rera_core_fields", "corporate_identity", "company_profile_check")
    for item in (flags.get("imminent") or []):
        if item.get("field", "").split(".")[0] in company_prefixes:
            needs_attention.append(f"Imminent: {item['text']}")
    for item in (flags.get("structural") or []):
        if item.get("field", "").split(".")[0] in company_prefixes:
            needs_attention.append(f"Structural: {item['text']}")

    _needs_attention(b, "Company", needs_attention, "No company-level flag was raised this pass.")


# ---------------------------------------------------------------------------
# 3. The Promoters -- one subsection per current director, plus one for any
# past director whose related-entity footprint is material enough to matter.
# ---------------------------------------------------------------------------
_NOT_FOUND = "Not found in public sources searched. Placeholder for further enrichment."


def _person_bio_bullets(b: _Builder, bio: dict):
    if bio.get("education"):
        for item in bio["education"]:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.bullet(text, source=source)
    else:
        b.bullet(f"Education: {_NOT_FOUND}")
    if bio.get("career"):
        for item in bio["career"]:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.bullet(text, source=source)
    else:
        b.bullet(f"Career history and core competencies: {_NOT_FOUND}")
    for item in bio.get("positive", []):
        text, source = item if isinstance(item, tuple) else (item, None)
        b.bullet(text, source=source)
    if b.doc_variant == "internal":
        for item in bio.get("identity_notes", []):
            text, source = item if isinstance(item, tuple) else (item, None)
            b.bullet(text, source=source)


def _director_subsection(b: _Builder, director: dict, research: dict):
    name = cc._normalise_entity_name(director.get("Director Name") or "")
    bio = (research.get("directors") or {}).get(name, {})

    b.heading(name, level=2)
    b.bullet(f"Appointed {director.get('Appointment Date', 'not disclosed')}, per MCA-mirror registry crosswalk.")
    other = bio.get("other_directorships_count")
    if other:
        b.bullet(
            f"Named on {other} other {'entity' if other == 1 else 'entities'} via shared-{cc._role_word(b.facts)} "
            f"registry data; see the Annexure for the full related-entity mapping."
        )
    _person_bio_bullets(b, bio)

    adverse = bio.get("adverse") or []
    if adverse:
        b.para()
        for item in adverse:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.flag_line(f"Needs attention: {text}", source=source)


def _past_director_subsection(b: _Builder, research: dict):
    summary = research.get("past_director_summary")
    bio = research.get("past_director_bio") or {}
    findings = research.get("past_director_findings") or []
    name = research.get("past_director_name")
    if not name:
        return
    b.para()
    b.heading(name, level=2)
    if summary:
        b.para(summary)
    # Was previously skipped entirely for past directors -- their own
    # research-agent response (education/career/positive/adverse) rendered
    # nowhere, even though a full per-person research call was made for
    # them; only the group agent's brief past_director_findings note ever
    # reached the document. Now matches a current director's treatment.
    _person_bio_bullets(b, bio)
    adverse = bio.get("adverse") or []
    if adverse:
        b.para()
        for item in adverse:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.flag_line(f"Needs attention: {text}", source=source)
    if findings:
        b.para()
        for item in findings:
            text, source = item if isinstance(item, tuple) else (item, None)
            b.flag_line(f"Needs attention: {text}", source=source)


def _the_promoters(b: _Builder, research):
    b.heading("3. The Promoters", page_break_before=True)

    profile = b.facts.get("company_profile_check") or {}
    directors = profile.get("current_directors") or []

    director_names = ", ".join(cc._normalise_entity_name(d.get("Director Name") or "") for d in directors)
    past_name = research.get("past_director_name")
    past_clause = (
        f", followed by a subsection on {past_name}, a past {cc._role_word(b.facts)} whose related-entity footprint is "
        f"large enough to warrant coverage of their own"
        if past_name else ""
    )
    b.para(
        f"This section profiles the {len(directors)} current {cc._role_word(b.facts, len(directors))} of this company: "
        f"{director_names}{past_clause}. Each profile sets out registry facts (appointment date, "
        f"other {cc._role_word(b.facts, 1)}ships), biographical detail actually found in public sources (education, career "
        f"history, notable coverage); a biographical field not found in public sources is marked as such rather "
        f"than invented. A litigation or adverse-coverage finding against a specific individual is stated directly "
        f"under their own name, rather than in a section shared by everyone. Every promoter's related-entity "
        f"connections are combined into a single table in the Annexure, rather than repeated per person."
    )

    for d in directors:
        b.para()
        _director_subsection(b, d, research)

    _past_director_subsection(b, research)

    all_bios = list((research.get("directors") or {}).values())
    if research.get("past_director_name"):
        all_bios.append(research.get("past_director_bio") or {})
    needs_attention = []
    if any(not bio.get("education") or not bio.get("career") for bio in all_bios):
        needs_attention.append(f"Education and career-history detail remains unconfirmed for at least one {cc._role_word(b.facts)}; see Recommended Verification Steps.")
    _needs_attention(b, "Promoters", needs_attention, "No promoter-level flag was raised this pass.")


# ---------------------------------------------------------------------------
# Annexure -- one combined related-entity table covering every current and
# past promoter, replacing a separate table repeated under each person's own
# subsection (confirmed the user wants ONE shared table here instead).
# ---------------------------------------------------------------------------
def _annexure_related_entities(b: _Builder):
    b.heading("Annexure: Related Entity Mapping", page_break_before=True)

    group = b.facts.get("group_companies_check") or {}
    companies = group.get("companies") or []
    if not companies:
        b.flag_line("No related entities were identified from the registry crosswalk this pass.")
        return

    b.para(
        f"{len(companies)} entities share at least one concrete, named link (a shared {cc._role_word(b.facts)}, "
        f"a shared registered office, or a filed subsidiary/associate/JV relationship) with this company's "
        f"current or past {cc._role_word(b.facts, 2)}, combined into the single table below rather than "
        f"repeated once per person."
    )
    rows = []
    for c in sorted(companies, key=lambda c: c.get("name") or ""):
        basis_list = c.get("basis") or []
        linked_via = cd._related_entity_linked_via(basis_list)
        if not linked_via:
            # No shared-director basis on this entity (e.g. office-only or a
            # subsidiary/associate/JV link) -- show the link type itself
            # rather than leaving the column blank.
            linked_via = "; ".join(sorted(set(basis_list))) or "Not specified"
        rows.append([c.get("name") or "", linked_via, cd._related_entity_implication(basis_list)])

    b.table(
        ["Related entity", "Linked via", "Nature and implication"],
        rows,
        col_widths=(5.0, 5.0, _CONTENT_WIDTH_CM - 10.0),
    )

    for relationship, count in sorted((group.get("undisclosed_relationship_counts") or {}).items()):
        b.flag_line(
            f"{count} {relationship} relationship(s) appear on the registry record but their counterparty "
            f"identities sit behind that source's paid tier. Their existence is confirmed; who they are is not."
        )


# ---------------------------------------------------------------------------
# 4. The Collateral
# ---------------------------------------------------------------------------
def _cited_field_rows(group: dict) -> list:
    """Same idea as charter_document.py's _field_table_rows, but returns
    rows shaped for THIS module's (label, (value, source)) citation
    convention instead of a separate "checked against" text column. A
    local function rather than a change to the shared one, so
    charter_document.py's own tables (3 columns, no bracket citations)
    keep rendering exactly as before."""
    rows = []
    for key, field in (group or {}).items():
        if isinstance(field, dict):
            value = str(field.get("value") or "").strip()
            if not value:
                continue
            rows.append([cd._field_display_name(key), (value, field.get("source"))])
        elif field:
            rows.append([cd._field_display_name(key), str(field)])
    return rows


def _land_ownership_history(b: _Builder):
    corp = b.facts.get("corporate_identity") or {}
    entry = corp.get("landowner_investor") or {}
    value = entry.get("value")
    if not value:
        return
    source = entry.get("source")
    b.para()
    b.heading("Land Ownership and Development Rights History", level=2)
    for label, clause in split_labeled_narrative(value):
        if label:
            b.bullet(f"{label}: {clause}", source=source)
        else:
            b.bullet(clause, source=source)


def _asset_identity(b: _Builder):
    b.heading("Asset Identity and Land Record", level=2)
    land_rows = _cited_field_rows(b.facts.get("land_identification") or {})
    if land_rows:
        b.table(["Field", "Value"], land_rows, col_widths=(5.0, _CONTENT_WIDTH_CM - 5.0))
    else:
        b.flag_line("No land identification fields were established this pass.")

    _land_ownership_history(b)

    neighbourhood = b.facts.get("neighbourhood") or {}
    if any(neighbourhood.values()):
        b.para()
        b.table(
            ["Direction", "Abutting"],
            [[d.title(), str(neighbourhood.get(d) or "Not stated")] for d in ("north", "east", "south", "west")],
            col_widths=(4.0, _CONTENT_WIDTH_CM - 4.0),
        )

    cts_check = b.facts.get("cts_land_record_check") or {}
    if cts_check.get("found"):
        b.para()
        b.para(
            "A Property Card was retrieved from Maha Bhulekh, Maharashtra's official land-records portal, by "
            "exact CTS number.",
            source=cts_check.get("url"),
        )
        fields = cts_check.get("fields") or {}
        if fields:
            b.table(["Field", "Value"], [[k, v] for k, v in fields.items()], col_widths=(6.0, _CONTENT_WIDTH_CM - 6.0))


def _approvals_fsi_title(b: _Builder):
    b.para()
    b.heading("Approvals, FSI and Title", level=2)
    fsi_rows = _cited_field_rows(b.facts.get("fsi_metrics") or {})
    if fsi_rows:
        b.table(["Metric", "Value"], fsi_rows, col_widths=(5.0, _CONTENT_WIDTH_CM - 5.0))
    for key in ("fsi_governing_framework", "fsi_interpretation"):
        if b.facts.get(key):
            b.para(str(b.facts[key]))
    rules = b.facts.get("rules_statutory") or {}
    if rules.get("planning_approval_sequence"):
        b.para()
        b.heading("Planning Approval Sequence", level=3)
        for step in reframe_arrow_chain(str(rules["planning_approval_sequence"])):
            b.bullet(step)
    if rules.get("governing_act"):
        b.para(f"Governing statute: {rules['governing_act']}")


def _rera_compliance_escrow(b: _Builder):
    b.para()
    b.heading("RERA Compliance and Escrow", level=2)
    compliance = b.facts.get("rera_compliance") or {}
    for key, title in (
        ("registration_summary", "Registration"),
        ("collection_account", "Collection account"),
        ("escrow_subaccounts", "Escrow sub-accounts"),
        ("litigations_complaints_appeals", "Complaints and appeals as filed"),
        ("statutory_declaration", "Statutory declaration"),
        ("construction_progress", "Construction progress"),
    ):
        if compliance.get(key):
            b.bullet(f"{title}: {compliance[key]}")

    core = b.facts.get("rera_core_fields") or {}
    core_keys = [
        ("registration_number", "Registration number"), ("project_status", "Project status"),
        ("approved_date", "Approved on"), ("proposed_completion_date", "Proposed completion"),
        ("plan_approval_number", "Plan approval number"), ("authority", "Authority"),
    ]
    rows = [[title, str(core[key])] for key, title in core_keys if core.get(key)]
    if rows:
        b.para()
        b.heading("RERA Record as Filed", level=3)
        b.table(["Field", "As filed"], rows, col_widths=(6.0, _CONTENT_WIDTH_CM - 6.0))

    b.para()
    b.heading("GST Filing Compliance", level=3)
    gst = b.facts.get("gst_compliance_check") or {}
    if gst.get("found"):
        summary = gst.get("summary") or {}
        b.bullet(
            f"GSTIN {gst.get('gstin', '')}: {summary.get('late_pct', 0)}% of rated return periods were filed "
            f"late, worst delay {summary.get('worst_delay_days') or 0} day(s), with "
            f"{summary.get('delays_last_12_months', 0)} delayed or unfiled period(s) falling due in the "
            f"trailing 12 months."
        )
    else:
        b.flag_line(
            "GST filing compliance is not assessed: it requires the developer's GSTIN and their filing "
            "history, and the GST portal gates that history behind a CAPTCHA that must be solved per "
            "lookup, so it cannot be collected unattended."
        )


def _project_professionals(b: _Builder):
    b.para()
    b.heading("Project Professionals", level=2)
    b.para(
        "These are project-level consultants named on the regulatory filing, not the promoter's internal "
        "management; see The Promoters for corporate control."
    )
    team = b.facts.get("professional_team") or []
    if team:
        b.table(
            ["Name", "Designation", "Role"],
            [[p.get("name") or "", p.get("role") or "", f"{p['registration_label']}: {p['registration_number']}" if p.get("registration_number") else "None filed"] for p in team],
            col_widths=(6.0, 4.0, _CONTENT_WIDTH_CM - 10.0),
        )
    else:
        b.flag_line("No professionals are named on this project's MahaRERA record.")


def _location_and_market(b: _Builder):
    b.para()
    b.heading("Location, Connectivity and Market Read", level=2)
    connectivity = b.facts.get("connectivity") or {}
    if any(connectivity.values()):
        b.table(
            ["Mode", "Position"],
            [[m.title(), str(connectivity.get(m) or "Not stated")] for m in ("road", "rail", "metro", "air")],
            col_widths=(4.0, _CONTENT_WIDTH_CM - 4.0),
        )
    distances = b.facts.get("distances") or []
    if distances:
        b.para()
        b.table(
            ["Landmark", "Distance / time", "Basis"],
            [[x.get("landmark", ""), x.get("distance_time", ""), x.get("route_note", "")] for x in distances],
            col_widths=(5.0, 4.0, _CONTENT_WIDTH_CM - 9.0),
        )
    if b.facts.get("micro_market_overview"):
        b.para()
        b.para(str(b.facts["micro_market_overview"]))
    comparables = b.facts.get("comparables") or []
    if comparables:
        b.para()
        b.table(
            ["Comparable project", "Distance", "Configuration", "Pricing"],
            [[c.get("project", ""), f"{c['distance_km']} km" if c.get("distance_km") else "Not stated", c.get("configuration", ""), c.get("pricing", "")] for c in comparables],
            col_widths=(4.5, 2.5, 3.5, _CONTENT_WIDTH_CM - 10.5),
        )


def _document_trail(b: _Builder):
    b.para()
    b.heading("Document and Diligence Trail", level=2)
    library = b.facts.get("document_library") or []
    if library:
        opened = sum(1 for d in library if "download" in str(d.get("status", "")).lower() or "reused" in str(d.get("status", "")).lower())
        b.para(f"{len(library)} documents are listed on this project's MahaRERA record, of which {opened} were retrieved for this review.")
    else:
        b.flag_line("No document library was recorded for this project.")
    for key, title in (("documents_reviewed_note", "Documents read in full"), ("documents_absent_note", "Documents absent from the registry library")):
        if b.facts.get(key):
            b.para(f"{title}: {b.facts[key]}")


def _collateral_litigation(b: _Builder):
    b.para()
    b.heading("Litigation and Regulatory Screening, Collateral Level", level=2)
    litigation = b.facts.get("litigation_status")
    if isinstance(litigation, dict) and litigation.get("value"):
        b.para(str(litigation["value"]), source=litigation.get("source"))
    else:
        b.flag_line("No collateral-level litigation finding was recorded this pass.")


def _the_collateral(b: _Builder, flags, project_name, group_research):
    b.heading(f"4. The Collateral: {project_name}", page_break_before=True)
    _asset_identity(b)
    _approvals_fsi_title(b)
    _rera_compliance_escrow(b)
    _project_professionals(b)
    _location_and_market(b)
    _document_trail(b)
    _collateral_litigation(b)

    needs_attention = []
    collateral_prefixes = ("land_identification", "fsi_metrics", "rera_compliance", "rules_statutory")
    for item in (flags.get("imminent") or []):
        if item.get("field", "").split(".")[0] in collateral_prefixes:
            needs_attention.append(f"Imminent: {item['text']}")
    for item in (flags.get("structural") or []):
        if item.get("field", "").split(".")[0] in collateral_prefixes:
            needs_attention.append(f"Structural: {item['text']}")
    for item in group_research.get("collateral_discrepancies", []):
        needs_attention.append(item if isinstance(item, tuple) else (item, None))

    _needs_attention(b, "Collateral", needs_attention, "No collateral-level flag was raised this pass.")


# ---------------------------------------------------------------------------
# 5. Closing Read
# ---------------------------------------------------------------------------
def _closing_read(b: _Builder, flags, developer_score, next_steps):
    b.heading("5. Closing Read", page_break_before=True)
    assessment = cd.assess_counterparty(b.facts, flags, developer_score)
    b.table(["Signal", "Finding"], [[label, finding] for label, finding in assessment["signals"]], col_widths=(5.0, _CONTENT_WIDTH_CM - 5.0))
    b.para()
    b.para(assessment["verdict"], bold=True)
    b.para(
        "This read is derived by rule from the signals tabled above, not written as an opinion. Each signal is a "
        "fact recorded elsewhere in this Charter and can be traced back to its source."
    )

    b.para()
    b.heading("Recommended Verification Steps", level=2)
    for step in next_steps:
        b.bullet(step)
    return assessment


# ---------------------------------------------------------------------------
# 6. Scoring Detail (Appendix)
# ---------------------------------------------------------------------------
def _scoring_detail(b: _Builder, developer_score, dc_buckets, dc_stale, doc_confidence_display):
    b.heading("6. Scoring Detail (Appendix)", page_break_before=True)
    b.para(
        "Full sub-metric working behind the two summary tables in section 2, so any rating there can be traced "
        "back to what produced it."
    )
    show_weight = b.doc_variant == "internal"

    b.heading("Developer Score, by Sub-Metric", level=2)
    criteria = developer_score.get("criteria") or {}
    rows = []
    for bucket_name, _w, metrics in cc._DEVELOPER_SCORE_STRUCTURE:
        for key, _display, _fn in metrics:
            criterion = criteria.get(key) or {}
            row = [bucket_name, criterion.get("display_name") or key.replace("_", " ").title()]
            if show_weight:
                row.append(f"{criterion.get('weight', '')}%")
            row += [criterion.get("tier") or "Not rated", str(criterion.get("score")) if criterion.get("score") is not None else "", criterion.get("note") or criterion.get("reason") or ""]
            rows.append(row)
    headers = ["Bucket", "Sub-metric"] + (["Weight"] if show_weight else []) + ["Tier", "Score", "Basis"]
    widths = [3.0, 3.5] + ([1.5] if show_weight else []) + [1.8, 1.5, _CONTENT_WIDTH_CM - (3.0 + 3.5 + (1.5 if show_weight else 0) + 1.8 + 1.5)]
    b.table(headers, rows, col_widths=widths)

    b.para()
    b.heading("Documentation Confidence, by Criterion", level=2)
    if dc_stale:
        b.flag_line("Not rated: the stored Documentation Confidence Score is from a superseded scoring scheme.")
    elif dc_buckets:
        dc_rows = [[bk["bucket"], f"{bk['weight']}%", bk["band"] or "Not rated", f"{bk['scored']} of {bk['total']}"] for bk in dc_buckets]
        dc_headers = ["Bucket", "Weight", "Band", "Coverage"] if show_weight else ["Bucket", "Band", "Coverage"]
        if not show_weight:
            dc_rows = [[r[0], r[2], r[3]] for r in dc_rows]
        b.table(dc_headers, dc_rows)
    else:
        b.flag_line("Not rated: no Documentation Confidence Score was recorded this pass.")


# ---------------------------------------------------------------------------
# Quality gate -- re-opens a just-saved Charter and checks for the exact
# regressions found and fixed by hand during review (dash artifacts, a
# redundant inline attribution surviving alongside its own citation, a
# bare-domain mention that should have become one, a manual page break
# reintroducing the blank-page bug, missing page-number fields, and a
# citation numbering that doesn't match the References list one-for-one).
# Not a general style linter -- a guard against these specific bugs coming
# back, following the exact pattern company_charter.py's own
# _verify_external_document_quality already established for the older
# document structure. Runs for BOTH variants (unlike that older gate, which
# is External-only) since none of these checks are variant-specific.
# ---------------------------------------------------------------------------
_CITATION_MARKER_RE = re.compile(r"\[\d+\]")
_REDUNDANT_ATTRIBUTION_MARKERS_RE = re.compile(r"\(per\s+[^)]+\)|confirmed\s+identically\s+(?:across|by)", re.IGNORECASE)


def _iter_all_paragraphs(doc):
    """Yields every paragraph in the document body plus every table cell
    (recursively, in case a cell itself contains a nested table). Mirrors
    company_charter.py's own helper of the same name, duplicated rather
    than imported since charter_report.py is deliberately independent of
    _fill_template's internals."""
    def _from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from _from_tables(cell.tables)

    yield from doc.paragraphs
    yield from _from_tables(doc.tables)


def verify_charter_report_quality(docx_path: str) -> list:
    """Returns a list of violation strings; empty means clean. Raises
    nothing itself -- the caller (build_charter_report) decides how loudly
    to fail."""
    from docx import Document as _Document
    from docx.oxml.ns import qn

    doc = _Document(docx_path)
    violations = []
    all_text_pieces = []

    for para in _iter_all_paragraphs(doc):
        text = para.text
        all_text_pieces.append(text)
        if " -- " in text or "–" in text or "—" in text:
            violations.append(f"dash artifact survived _clean_text: {text[:80]!r}")

        has_citation = bool(_CITATION_MARKER_RE.search(text))
        has_attribution_marker = bool(_REDUNDANT_ATTRIBUTION_MARKERS_RE.search(text))
        if has_citation and has_attribution_marker:
            # Only a violation when BOTH are present in the same text: the
            # "(per X)" / "confirmed identically across..." phrasing is
            # legitimate and necessary on a value that carries no citation
            # at all (its only attribution) -- e.g. a plain-string
            # fsi_metrics field with no separate source. It is redundant
            # only once a bracketed citation is ALSO doing that job.
            violations.append(f"redundant inline attribution survived alongside its own citation: {text[:80]!r}")

        if _BARE_DOMAIN_MENTION_RE.search(text):
            violations.append(f"bare-domain mention was never converted to a citation: {text[:80]!r}")

    # No manual page-break character anywhere -- every major-section heading
    # uses page-break-before paragraph formatting instead (see
    # _Builder.heading's own note on why), the only approach proven immune
    # to the blank-page bug a manual page break produced once a variant's
    # content grew past a certain length. A manual page break reappearing
    # is exactly how that bug would come back.
    for br in doc.element.body.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            violations.append("a manual page-break character was found; use heading(text, page_break_before=True) instead")

    # Page numbers: the footer must carry both a PAGE and a NUMPAGES field.
    footer_field_codes = set()
    for section in doc.sections:
        for para in section.footer.paragraphs:
            for instr in para._p.iter(qn("w:instrText")):
                footer_field_codes.add((instr.text or "").strip())
    if "PAGE" not in footer_field_codes:
        violations.append("footer is missing the PAGE field -- page numbers will not render")
    if "NUMPAGES" not in footer_field_codes:
        violations.append('footer is missing the NUMPAGES field -- "of Y" will not render')

    # Citation integrity: every [n] used in the body has a matching
    # References entry, numbered contiguously from 1 with no gaps or
    # orphans. True by construction given how _CitationRegistry assigns
    # numbers, but checked defensively in case a future edit breaks that
    # invariant.
    used_numbers = set()
    for text in all_text_pieces:
        for m in re.finditer(r"\[(\d+)\]", text):
            used_numbers.add(int(m.group(1)))
    reference_numbers = set()
    for para in doc.paragraphs:
        m = re.match(r"^\[(\d+)\]\s+\S", para.text.strip())
        if m:
            reference_numbers.add(int(m.group(1)))
    if used_numbers or reference_numbers:
        missing_refs = used_numbers - reference_numbers
        orphan_refs = reference_numbers - used_numbers
        if missing_refs:
            violations.append(f"citation number(s) {sorted(missing_refs)} used in text but missing from References")
        if orphan_refs:
            violations.append(f"References entry number(s) {sorted(orphan_refs)} never cited in the body")

    return violations


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def build_charter_report(reg_no: str, facts: dict, research: dict, out_path: str, doc_variant: str, generated_on: str) -> None:
    """Builds one variant (doc_variant in {"internal", "external"}) of the
    Company Charter following the Cover / About / Company / Promoters /
    Collateral / Closing / Appendix / References structure.

    `research` carries the biographical and group-background findings
    gathered separately (director bios, adverse-coverage screening, group
    press coverage) -- see COMPANY_CHARTER_SPEC.md for its documented shape.
    Always recomputes developer_score, never trusts facts["developer_score"]
    (an output of a prior run, see charter_document.py's own note on this)."""
    if doc_variant == "external":
        facts = cc._externalized_facts_copy(facts)
    facts["_doc_variant"] = doc_variant

    doc = Document()
    _configure_base_styles(doc)
    _add_page_numbers(doc)
    b = _Builder(doc, facts, doc_variant)

    flags = cc._classify_flags(facts)
    developer_score = cc._compute_developer_score(facts, flags)
    facts["developer_score"] = developer_score
    next_steps = cd._recommended_steps(facts, flags)
    all_promoter_bios = list((research.get("directors") or {}).values())
    if research.get("past_director_name"):
        all_promoter_bios.append(research.get("past_director_bio") or {})
    if any(not bio.get("education") or not bio.get("career") for bio in all_promoter_bios):
        next_steps.append("Commission a formal biographical/KYC check on the promoters to close the education and career-history gaps noted in The Promoters.")
    next_steps.extend(research.get("additional_next_steps", []))

    doc_confidence = facts.get("documentation_confidence_score") or {}
    dc_buckets = cd.rollup_doc_confidence_buckets(doc_confidence) if doc_confidence.get("criteria") else []
    dc_stale = bool(dc_buckets) and not any(bucket["scored"] for bucket in dc_buckets)
    if dc_stale:
        dc_buckets = []
    doc_confidence_display = {
        "overall": doc_confidence.get("overall") if dc_buckets else None,
        "band": doc_confidence.get("band") if dc_buckets else None,
    }

    composite_line = f"Composite Developer Score: {developer_score.get('composite')}/100 (grade {developer_score.get('grade')})"
    _cover_page(b, reg_no, generated_on, composite_line)

    core = facts.get("rera_core_fields") or {}
    _about_the_company(b)
    _the_company(b, flags, developer_score, doc_confidence_display, dc_buckets, dc_stale, research.get("group", {}))
    _the_promoters(b, research)
    _the_collateral(b, flags, core.get("project_name", "the project"), research.get("group", {}))
    _closing_read(b, flags, developer_score, next_steps)
    _scoring_detail(b, developer_score, dc_buckets, dc_stale, doc_confidence_display)
    _annexure_related_entities(b)
    if doc_variant == "external":
        # Internal's citations are already self-descriptive inline "(per X)"
        # text (see _Builder.cite) -- a References list would just repeat
        # them with no numbers to resolve, so only External gets one.
        b.references_section(generated_on)

    doc.save(out_path)

    # Re-opens the file just saved (not the in-memory `doc`) so this checks
    # exactly what a reader will actually open -- catches a future code
    # change silently reintroducing any of the specific bugs this pipeline
    # found and fixed by hand (see verify_charter_report_quality's own
    # docstring). Fails loudly rather than silently shipping a regressed
    # document, for both variants.
    violations = verify_charter_report_quality(out_path)
    if violations:
        raise RuntimeError(
            f"Company Charter quality gate failed for {out_path} "
            f"({len(violations)} violation(s)):\n" + "\n".join(f"  - {v}" for v in violations)
        )
