"""
Tests for _verify_external_document_quality, the quality gate wired into
_fill_template right after doc.save(out_path) for doc_variant="external".

The gate's job is to catch a FUTURE code change silently reintroducing one
of the specific bugs this session found and fixed by hand in the External
Company Charter: grey/italic body text (from _set_paragraph_text reusing
template placeholder run formatting), hyphen-pair/em-dashes, a stray Weight
column on the Developer Score table, a resurrected Document Library
section, a citation entry that lost its bullet numbering, or a dangling
unmatched paren in a Sources entry.

Two kinds of test here:
  1. The real Pranami Bliss / IRA Insignia External fixtures must pass
     cleanly today (a positive control -- proves the gate isn't so strict
     it flags legitimate content like heading colors or the deliberately
     red+italic Standing Gap paragraph).
  2. Each specific violation, deliberately reintroduced into a COPY of a
     real fixture, must actually be caught (negative controls -- proves
     the gate fires, not just silently passes everything).

Run directly: python test_executive_ready.py
"""

import copy
import json
import os
import shutil

import docx
from docx.shared import RGBColor

import company_charter as cc

_FIXTURES = [
    ("P51800077150", os.path.join("output", "company_charters", "Company_Charter_PranamiBliss_P51800077150.facts.json")),
    ("P51700031409", os.path.join("output", "company_charters", "Company_Charter_IRAInsignia_P51700031409.facts.json")),
]
_SCRATCH_DIR = os.path.join("output", "company_charters", "_test_scratch")


def _regenerate_external(reg_no: str, facts_path: str, out_name: str) -> str:
    with open(facts_path, encoding="utf-8") as f:
        facts = json.load(f)
    os.makedirs(_SCRATCH_DIR, exist_ok=True)
    out_path = os.path.join(_SCRATCH_DIR, out_name)
    cc._fill_template(reg_no, facts, out_path, doc_variant="external")
    return out_path


def test_real_fixtures_pass_the_quality_gate():
    """Both real External Charters must already be clean -- _fill_template
    itself would have raised RuntimeError if not, but this also confirms
    _verify_external_document_quality's return value directly."""
    for reg_no, facts_path in _FIXTURES:
        out_path = _regenerate_external(reg_no, facts_path, f"{reg_no}_External.docx")
        violations = cc._verify_external_document_quality(out_path)
        assert violations == [], f"{reg_no}: unexpected violations: {violations}"
    print("test_real_fixtures_pass_the_quality_gate: PASS")


def test_standing_gap_red_italic_paragraph_is_not_falsely_flagged():
    """The one deliberate red+italic paragraph in the real document must
    NOT trip the italic check -- proves the exception is scoped to that
    specific color pairing, not a blanket italic allowance."""
    out_path = os.path.join(_SCRATCH_DIR, "P51800077150_External.docx")
    assert os.path.exists(out_path), "run test_real_fixtures_pass_the_quality_gate first"
    doc = docx.Document(out_path)
    found = False
    for para in doc.paragraphs:
        if para.text.strip().startswith("Standing gap (permanent"):
            found = True
            assert any(r.italic for r in para.runs), "fixture no longer has an italic Standing Gap run to test against"
    assert found, "fixture no longer contains a Standing Gap paragraph to test against"
    violations = cc._verify_external_document_quality(out_path)
    assert not any("Standing gap" in v for v in violations)
    print("test_standing_gap_red_italic_paragraph_is_not_falsely_flagged: PASS")


def _first_nonempty_run(doc):
    """The External document's paragraph indices shift depending on which
    sections a given fixture's facts happen to populate, so tests mutate
    whichever paragraph actually has real run text instead of a hardcoded
    index -- more robust, and it's the run's own text/style being tested,
    not its position."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                return run
    raise AssertionError("no paragraph with non-empty run text found in fixture")


def _reopen_mutate_and_check(src_path: str, mutate_fn, expect_substring: str):
    """Copies src_path, applies mutate_fn(doc) to the copy, saves it, runs
    the gate, and asserts at least one violation contains expect_substring."""
    mutated_path = src_path.replace(".docx", "_mutated.docx")
    shutil.copy2(src_path, mutated_path)
    doc = docx.Document(mutated_path)
    mutate_fn(doc)
    doc.save(mutated_path)
    violations = cc._verify_external_document_quality(mutated_path)
    assert any(expect_substring in v for v in violations), (
        f"expected a violation containing {expect_substring!r}, got: {violations}"
    )
    os.remove(mutated_path)


def test_quality_gate_detects_reintroduced_hyphen_dash():
    def mutate(doc):
        run = _first_nonempty_run(doc)
        run.text = "This is a test -- reintroducing a dash."
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "hyphen-pair dash"
    )
    print("test_quality_gate_detects_reintroduced_hyphen_dash: PASS")


def test_quality_gate_detects_reintroduced_grey_italic_run():
    def mutate(doc):
        run = _first_nonempty_run(doc)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string("808080")
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "italic run"
    )
    print("test_quality_gate_detects_reintroduced_grey_italic_run: PASS")


def test_quality_gate_detects_document_library_regression():
    def mutate(doc):
        run = _first_nonempty_run(doc)
        run.text = "Document Library reappeared by accident."
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "Document Library"
    )
    print("test_quality_gate_detects_document_library_regression: PASS")


def test_quality_gate_detects_weight_column_regression():
    def mutate(doc):
        for table in doc.tables:
            header_texts = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
            if header_texts[:2] == ["Bucket", "Sub-metric"]:
                # The gate only checks for the literal string "Weight" among
                # header texts, so relabeling one existing header cell is
                # enough to simulate the regression without rebuilding the
                # table's column structure.
                table.rows[0].cells[2].text = "Weight"
                return
        raise AssertionError("Developer Score table not found in fixture")
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "Weight column"
    )
    print("test_quality_gate_detects_weight_column_regression: PASS")


def test_quality_gate_detects_missing_citation_bullet():
    def mutate(doc):
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped.startswith("[1] ") or stripped.startswith("[2] "):
                p_pr = para._p.get_or_add_pPr()
                if p_pr.numPr is not None:
                    p_pr.remove(p_pr.numPr)
                return
        raise AssertionError("no numbered citation entry found in fixture")
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "lost its bullet numbering"
    )
    print("test_quality_gate_detects_missing_citation_bullet: PASS")


def test_quality_gate_detects_dangling_paren_citation():
    def mutate(doc):
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped.startswith("[1] ") or stripped.startswith("[2] "):
                para.runs[0].text = stripped + " (dangling annotation"
                return
        raise AssertionError("no numbered citation entry found in fixture")
    _reopen_mutate_and_check(
        os.path.join(_SCRATCH_DIR, "P51800077150_External.docx"), mutate, "unbalanced parentheses"
    )
    print("test_quality_gate_detects_dangling_paren_citation: PASS")


def test_fill_template_raises_when_gate_fails():
    """Confirms the wiring in _fill_template itself (not just the standalone
    checker function) actually raises -- monkeypatches _externalize_prose
    to reintroduce a dash into whatever text flows through it, then asserts
    _fill_template raises RuntimeError for doc_variant='external'."""
    reg_no, facts_path = _FIXTURES[0]
    with open(facts_path, encoding="utf-8") as f:
        facts = json.load(f)

    real_externalize_prose = cc._externalize_prose

    def _broken_externalize_prose(facts, text):
        result = real_externalize_prose(facts, text)
        return result + " -- broken by test"

    cc._externalize_prose = _broken_externalize_prose
    try:
        out_path = os.path.join(_SCRATCH_DIR, "broken_External.docx")
        try:
            cc._fill_template(reg_no, facts, out_path, doc_variant="external")
            raised = False
        except RuntimeError as e:
            raised = True
            assert "hyphen-pair dash" in str(e)
        assert raised, "_fill_template did not raise despite a reintroduced dash"
    finally:
        cc._externalize_prose = real_externalize_prose
    print("test_fill_template_raises_when_gate_fails: PASS")


if __name__ == "__main__":
    test_real_fixtures_pass_the_quality_gate()
    test_standing_gap_red_italic_paragraph_is_not_falsely_flagged()
    test_quality_gate_detects_reintroduced_hyphen_dash()
    test_quality_gate_detects_reintroduced_grey_italic_run()
    test_quality_gate_detects_document_library_regression()
    test_quality_gate_detects_weight_column_regression()
    test_quality_gate_detects_missing_citation_bullet()
    test_quality_gate_detects_dangling_paren_citation()
    test_fill_template_raises_when_gate_fails()
    shutil.rmtree(_SCRATCH_DIR, ignore_errors=True)
    print("\nAll tests passed.")
